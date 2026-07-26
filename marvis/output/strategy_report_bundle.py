"""Deterministic projections of a canonical StrategyReportBundle V2.

This module is intentionally a renderer, not an analysis layer.  It validates
the self-authenticating bundle and projects the exact same facts to canonical
JSON, Markdown, a formula-free XLSX workbook, and a macro-free DOCX review
brief.  Missing values remain blank; their typed availability and reason stay
adjacent to the blank presentation cell and in the evidence index.
"""

from __future__ import annotations

from collections.abc import Iterable, Mapping, Sequence
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from itertools import islice
import json
import re
import unicodedata
from typing import Any
from xml.etree import ElementTree
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from marvis.packs.strategy.errors import StrategyError
from marvis.packs.strategy.report_bundle import (
    REPORT_CORE_SHEET_KEYS,
    canonical_strategy_report_bundle_json,
    validate_strategy_report_bundle,
)


STRATEGY_REPORT_OUTPUT_SCHEMA_VERSION = "strategy.report-output.v2"
STRATEGY_REPORT_SHEET_TITLES = {
    "00_summary": "结论与变更摘要",
    "01_current_state": "项目现状",
    "02_history": "历史策略",
    "03_sample": "样本与口径",
    "04_univariate_model": "单变量与模型",
    "05_candidates": "候选组合",
    "06_strategy": "最终策略明细",
    "07_waterfall_swap": "Waterfall 与 Swap",
    "08_impact": "逐月与分群影响",
    "09_economics": "收益与数据成本",
    "10_validation": "验证与稳定性",
    "11_evidence": "证据与版本",
}

_SECTION_PRIMARY_SHEETS = {
    "current_project": "01_current_state",
    "historical_versions": "02_history",
    "sample_design": "03_sample",
    "univariate_and_models": "04_univariate_model",
    "candidate_combinations": "05_candidates",
    "impact_assessment": "08_impact",
    "final_document": "06_strategy",
}
_SHEET_SECTION_KEYS = {
    "01_current_state": ("current_project",),
    "02_history": ("historical_versions",),
    "03_sample": ("sample_design",),
    "04_univariate_model": ("univariate_and_models",),
    "05_candidates": ("candidate_combinations",),
    "06_strategy": ("candidate_combinations", "final_document"),
    "07_waterfall_swap": ("impact_assessment",),
    "08_impact": ("impact_assessment",),
    "09_economics": ("impact_assessment",),
    "10_validation": ("impact_assessment", "final_document"),
}
_AVAILABILITY_LABELS = {
    "present": "已有可信值",
    "unavailable": "暂缺",
    "not_applicable": "本次不涉及",
    "not_matured": "尚未成熟",
}
_EFFECT_STAGE_LABELS = {
    "estimated": "预估",
    "backtested": "开发回测",
    "oot_validated": "OOT",
    "post_launch_observed": "上线后观察",
}
_REPORT_STATUS_LABELS = {
    "draft": "草稿",
    "partial": "部分完成",
    "final": "当前允许范围内已完成",
}

_FIXED_WORKBOOK_DATETIME = datetime(2000, 1, 1)
_FIXED_DOCX_DATETIME = datetime(2000, 1, 1)
_FIXED_ZIP_DATETIME = (1980, 1, 1, 0, 0, 0)
_FIXED_ZIP_EXTERNAL_ATTR = 0o600 << 16
_FORMULA_PREFIXES = frozenset("=+-@")
_ILLEGAL_PRESENTATION_CONTROL = re.compile(
    r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f\ud800-\udfff\ufffe\uffff]"
)
_MARKDOWN_INLINE_META = re.compile(r"([`*_{}\[\]()!])")
_DOCX_CJK_TEXT = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]")
_CORE_MODIFIED_TIMESTAMP = re.compile(
    rb"(<dcterms:modified\b[^>]*>)[^<]*(</dcterms:modified>)"
)
_MAX_XLSX_CELL_CHARACTERS = 32_767
_MAX_XLSX_ROWS = 1_048_576
_MAX_XLSX_COLUMNS = 16_384

_DOCX_TABLE_WIDTH_DXA = 9_360
_DOCX_TABLE_INDENT_DXA = 120
_DOCX_CJK_FONT = "Microsoft YaHei"
_DOCX_CELL_MARGINS_DXA = {
    "top": 80,
    "bottom": 80,
    "start": 120,
    "end": 120,
}
_DOCX_RELATIONSHIP_NAMESPACE = (
    "http://schemas.openxmlformats.org/package/2006/relationships"
)
_DOCX_WORDPROCESSINGML_NAMESPACE = (
    "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
)
_DOCX_FORBIDDEN_FIELD_LOCAL_NAMES = frozenset(
    {"fldChar", "fldSimple", "instrText"}
)
_DOCX_MAX_CELL_CHARACTERS = 2_048
_DOCX_MAX_REFS_PER_CELL = 16
_DOCX_MAX_SUMMARY_FIELDS_PER_SECTION = 60
_DOCX_MAX_SUMMARY_FIELDS_TOTAL = 200
_DOCX_MAX_STAGE_ROWS_PER_SECTION = 30
_DOCX_MAX_STAGE_ROWS_TOTAL = 50
_DOCX_MAX_REPORT_TABLES = 20
_DOCX_MAX_FACT_ROWS_PER_TABLE = 120
_DOCX_MAX_FACT_ROWS_TOTAL = 360
_DOCX_VOTING_SEARCH_TABLE_ID = "voting_candidate_search_combinations"
_DOCX_MAX_VOTING_SEARCH_ROWS = 20
_DOCX_MAX_RED_FLAGS_PER_SECTION = 30
_DOCX_MAX_RED_FLAGS_TOTAL = 80
_DOCX_MAX_MISSING_INFORMATION_ROWS = 100
_DOCX_MAX_EVIDENCE_ROWS = 300
_DOCX_TRUNCATION_SUFFIX = "…[内容已截断；完整内容见 JSON/XLSX]"
_DOCX_SUMMARY_WIDTHS = (1_600, 1_800, 1_200, 2_000, 2_760)
_DOCX_STAGE_WIDTHS = (1_450, 1_300, 1_300, 5_310)
_DOCX_FACT_WIDTHS = (1_200, 1_600, 1_400, 1_300, 2_100, 1_760)
_DOCX_VOTING_SEARCH_WIDTHS = (
    1_200,
    1_360,
    1_460,
    420,
    620,
    1_100,
    1_360,
    1_840,
)
_DOCX_FLAG_WIDTHS = (1_150, 1_700, 4_150, 2_360)
_DOCX_COMPLETENESS_WIDTHS = (2_100, 4_000, 3_260)
_DOCX_MISSING_WIDTHS = (1_850, 1_150, 1_150, 4_050, 1_160)
_DOCX_EVIDENCE_WIDTHS = (2_500, 1_200, 2_200, 3_460)

_NAVY = "18324A"
_BLUE = "235A7A"
_PALE_BLUE = "EAF2F8"
_PALE_AMBER = "FFF3CD"
_PALE_RED = "F8D7DA"
_PALE_GREEN = "DDEFE3"
_WHITE = "FFFFFF"
_GRID = Side(style="thin", color="D9E1E8")
_BORDER = Border(left=_GRID, right=_GRID, top=_GRID, bottom=_GRID)


@dataclass
class _DocxRenderBudget:
    summary_fields_remaining: int = _DOCX_MAX_SUMMARY_FIELDS_TOTAL
    stage_rows_remaining: int = _DOCX_MAX_STAGE_ROWS_TOTAL
    report_tables_remaining: int = _DOCX_MAX_REPORT_TABLES
    fact_rows_remaining: int = _DOCX_MAX_FACT_ROWS_TOTAL
    red_flags_remaining: int = _DOCX_MAX_RED_FLAGS_TOTAL


class StrategyReportOutputError(StrategyError):
    """A valid report bundle cannot be safely projected to an output format."""


def render_strategy_report_bundle(
    bundle: Mapping[str, Any],
) -> dict[str, bytes]:
    """Render all four deterministic projections of one validated bundle."""

    canonical = validate_strategy_report_bundle(bundle)
    return {
        "json": canonical_strategy_report_bundle_json(canonical).encode("utf-8"),
        "markdown": _render_markdown(canonical).encode("utf-8"),
        "xlsx": _render_xlsx(canonical),
        "docx": _render_docx(canonical),
    }


def render_strategy_report_bundle_json(bundle: Mapping[str, Any]) -> bytes:
    """Return the canonical machine-readable report manifest."""

    return canonical_strategy_report_bundle_json(bundle).encode("utf-8")


def render_strategy_report_bundle_markdown(bundle: Mapping[str, Any]) -> bytes:
    """Return the deterministic Markdown executive report."""

    return _render_markdown(validate_strategy_report_bundle(bundle)).encode("utf-8")


def render_strategy_report_bundle_xlsx(bundle: Mapping[str, Any]) -> bytes:
    """Return a deterministic, formula-free multi-sheet workbook."""

    return _render_xlsx(validate_strategy_report_bundle(bundle))


def render_strategy_report_bundle_docx(bundle: Mapping[str, Any]) -> bytes:
    """Return a deterministic, macro-free formal strategy review brief."""

    return _render_docx(validate_strategy_report_bundle(bundle))


def _render_markdown(bundle: Mapping[str, Any]) -> str:
    title = _field_display_value(bundle["title"])
    lines = [
        f"# {_markdown_cell(title)}",
        "",
        "> 本报告是受治理结构化证据的投影。空白不等于 0，效果阶段不自动升级。",
        "",
        "| 报告属性 | 值 |",
        "|---|---|",
        f"| 报告 ID | {_markdown_cell(bundle['report_id'])} |",
        f"| 修订 | {bundle['report_revision']} |",
        f"| 状态 | {_markdown_cell(_REPORT_STATUS_LABELS[bundle['status']])} |",
        f"| 策略 ID | {_markdown_cell(bundle['strategy_id'] or '')} |",
        f"| 策略版本 | {_markdown_cell(bundle['strategy_version'] or '')} |",
        f"| 策略类型 | {_markdown_cell(bundle['strategy_type'] or '')} |",
        f"| 效果阶段 | {_markdown_cell(_stage_labels(bundle['effect_stages']))} |",
        f"| 生成时间 | {_markdown_cell(bundle['generated_at'])} |",
        f"| 数据分级 | {_markdown_cell(bundle['data_classification'])} |",
        f"| 内容 SHA-256 | {_markdown_cell(bundle['content_sha256'])} |",
        "",
    ]
    for index, section in enumerate(bundle["sections"], start=1):
        lines.extend(_markdown_section(index, section))

    lines.extend(
        [
            "## 证据、缺失信息与完整度",
            "",
            "### 完整度",
            "",
            "| 类别 | 状态 | 数量 |",
            "|---|---|---:|",
        ]
    )
    completeness = bundle["completeness_summary"]
    for key, count in completeness["field_counts"].items():
        lines.append(f"| 字段 | {_markdown_cell(key)} | {count} |")
    for key, count in completeness["missing_information_counts"].items():
        lines.append(f"| 缺失信息 | {_markdown_cell(key)} | {count} |")
    for key, count in completeness["blocking_counts"].items():
        lines.append(f"| 阻塞 | {_markdown_cell(key)} | {count} |")

    lines.extend(
        [
            "",
            "### 缺失信息",
            "",
            "| 字段 | 阻塞级别 | 状态 | 原因 | 已询问次数 |",
            "|---|---|---|---|---:|",
        ]
    )
    if bundle["missing_information"]:
        for item in bundle["missing_information"]:
            lines.append(
                "| "
                + " | ".join(
                    (
                        _markdown_cell(item["field_path"]),
                        _markdown_cell(item["blocking"]),
                        _markdown_cell(item["status"]),
                        _markdown_cell(item["reason"]),
                        str(item["asked_count"]),
                    )
                )
                + " |"
            )
    else:
        lines.append("|  |  |  |  | 0 |")

    lines.extend(
        [
            "",
            "### 来源索引",
            "",
            "| 位置 | 类型 | 引用 ID | 内容 SHA-256 |",
            "|---|---|---|---|",
        ]
    )
    for location, ref in _source_locations(bundle):
        lines.append(
            "| "
            + " | ".join(
                (
                    _markdown_cell(location),
                    _markdown_cell(ref["kind"]),
                    _markdown_cell(ref["ref_id"]),
                    _markdown_cell(ref["content_hash"]),
                )
            )
            + " |"
        )

    # A final newline gives stable CLI/file rendering and is part of the
    # deterministic projection contract.
    return "\n".join(lines).rstrip() + "\n"


def _markdown_section(index: int, section: Mapping[str, Any]) -> list[str]:
    lines = [
        f"## {index}. {_markdown_cell(section['title'])}",
        "",
        f"- 状态：{_markdown_cell(_AVAILABILITY_LABELS[section['availability']])}",
    ]
    stages = sorted(
        {item["effect_stage"] for item in section["stage_evidence"]},
        key=("estimated", "backtested", "oot_validated", "post_launch_observed").index,
    )
    if stages:
        lines.append(f"- 效果阶段：{_markdown_cell(_stage_labels(stages))}")
    lines.append("")
    if section["summary_fields"]:
        lines.extend(
            [
                "| 字段 | 值 | 状态 | 说明 | 来源 |",
                "|---|---|---|---|---|",
            ]
        )
        for item in section["summary_fields"]:
            field = item["field"]
            lines.append(
                "| "
                + " | ".join(
                    (
                        _markdown_cell(item["label"]),
                        _markdown_cell(_field_display_value(field)),
                        _markdown_cell(_AVAILABILITY_LABELS[field["availability"]]),
                        _markdown_cell(field["note"] or ""),
                        _markdown_cell(_compact_refs(field["source_refs"])),
                    )
                )
                + " |"
            )
        lines.append("")
    for table in section["tables"]:
        lines.extend(_markdown_table(table))
    if section["red_flags"]:
        lines.extend(
            [
                "### 红旗",
                "",
                "| 级别 | 编码 | 说明 |",
                "|---|---|---|",
            ]
        )
        for flag in section["red_flags"]:
            lines.append(
                "| "
                + " | ".join(
                    (
                        _markdown_cell(flag["level"]),
                        _markdown_cell(flag["code"]),
                        _markdown_cell(flag["message"]),
                    )
                )
                + " |"
            )
        lines.append("")
    if (
        not section["summary_fields"]
        and not section["tables"]
        and not section["red_flags"]
    ):
        lines.extend(["该部分当前没有可展示的结构化事实。", ""])
    return lines


def _markdown_table(table: Mapping[str, Any]) -> list[str]:
    stage = (
        ""
        if table["effect_stage"] is None
        else f"（{_EFFECT_STAGE_LABELS[table['effect_stage']]}）"
    )
    headers: list[str] = []
    separators: list[str] = []
    for column in table["columns"]:
        label = column["label"]
        if column["unit"] is not None:
            label = f"{label} [{column['unit']}]"
        headers.extend((label, f"{label}状态", f"{label}说明"))
        separators.extend(("---", "---", "---"))
    lines = [
        f"### {_markdown_cell(table['title'])}{stage}",
        "",
        "| " + " | ".join(_markdown_cell(item) for item in headers) + " |",
        "|" + "|".join(separators) + "|",
    ]
    for row in table["rows"]:
        values: list[str] = []
        for column in table["columns"]:
            field = row["cells"][column["key"]]
            values.extend(
                (
                    _markdown_cell(
                        _table_field_display_value(field, column=column)
                    ),
                    _markdown_cell(_AVAILABILITY_LABELS[field["availability"]]),
                    _markdown_cell(field["note"] or ""),
                )
            )
        lines.append("| " + " | ".join(values) + " |")
    if not table["rows"]:
        lines.append("| " + " | ".join("" for _ in headers) + " |")
    lines.extend(
        [
            "",
            f"来源：{_markdown_cell(_compact_refs(table['source_refs']))}",
            "",
        ]
    )
    return lines


def _render_docx(bundle: Mapping[str, Any]) -> bytes:
    document = Document()
    budget = _DocxRenderBudget()
    _configure_docx_styles(document)
    _configure_docx_section(document, bundle=bundle)
    _configure_docx_properties(document, bundle=bundle)
    _write_docx_masthead(document, bundle=bundle)

    for index, section in enumerate(bundle["sections"], start=1):
        _write_docx_section(
            document,
            index=index,
            section=section,
            budget=budget,
        )
    _write_docx_evidence_appendix(document, bundle=bundle)

    raw = BytesIO()
    document.save(raw)
    canonical = _canonicalize_docx_bytes(raw.getvalue())
    _assert_safe_docx_package(canonical)
    return canonical


def _configure_docx_styles(document: Any) -> None:
    styles = document.styles
    normal = styles["Normal"]
    _configure_docx_style(
        normal,
        font_name="Calibri",
        size=11,
        color="000000",
        before=0,
        after=6,
        line_spacing=1.10,
    )
    _configure_docx_style(
        styles["Title"],
        font_name="Calibri",
        size=23,
        color="000000",
        before=0,
        after=4,
        line_spacing=1.0,
        bold=True,
    )
    _configure_docx_style(
        styles["Subtitle"],
        font_name="Calibri",
        size=14,
        color="373737",
        before=0,
        after=16,
        line_spacing=1.0,
    )
    _configure_docx_style(
        styles["Heading 1"],
        font_name="Calibri",
        size=16,
        color="2E74B5",
        before=16,
        after=8,
        line_spacing=1.0,
        bold=True,
        keep_with_next=True,
    )
    _configure_docx_style(
        styles["Heading 2"],
        font_name="Calibri",
        size=13,
        color="2E74B5",
        before=12,
        after=6,
        line_spacing=1.0,
        bold=True,
        keep_with_next=True,
    )
    _configure_docx_style(
        styles["Heading 3"],
        font_name="Calibri",
        size=12,
        color="1F4D78",
        before=8,
        after=4,
        line_spacing=1.0,
        bold=True,
        keep_with_next=True,
    )
    _configure_docx_style(
        styles["Header"],
        font_name="Calibri",
        size=9,
        color="536878",
        before=0,
        after=0,
        line_spacing=1.0,
    )
    _configure_docx_style(
        styles["Footer"],
        font_name="Calibri",
        size=9,
        color="536878",
        before=0,
        after=0,
        line_spacing=1.0,
    )


def _configure_docx_style(
    style: Any,
    *,
    font_name: str,
    size: float,
    color: str,
    before: float,
    after: float,
    line_spacing: float,
    bold: bool = False,
    keep_with_next: bool = False,
) -> None:
    style.font.name = font_name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    _set_docx_font_family(style.element, font_name)
    paragraph = style.paragraph_format
    paragraph.space_before = Pt(before)
    paragraph.space_after = Pt(after)
    paragraph.line_spacing = line_spacing
    paragraph.keep_with_next = keep_with_next


def _set_docx_font_family(element: Any, font_name: str) -> None:
    run_properties = element.get_or_add_rPr()
    fonts = run_properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, fonts)
    for attribute in ("ascii", "hAnsi", "cs"):
        fonts.set(qn(f"w:{attribute}"), font_name)
    # Named glyph-coverage override: Calibri remains the prescribed Latin
    # business font while the explicit CJK face prevents Chinese report labels
    # from degrading to tofu boxes in Word/LibreOffice.
    fonts.set(qn("w:eastAsia"), _DOCX_CJK_FONT)


def _configure_docx_section(document: Any, *, bundle: Mapping[str, Any]) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.style = document.styles["Header"]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_docx_run(
        header,
        "MARVIS | 策略评审 | "
        f"{bundle['report_id']} | 修订 {bundle['report_revision']}",
    )

    footer = section.footer.paragraphs[0]
    footer.style = document.styles["Footer"]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_docx_run(
        footer,
        f"{bundle['data_classification']} | {bundle['generated_at']}",
    )


def _configure_docx_properties(
    document: Any,
    *,
    bundle: Mapping[str, Any],
) -> None:
    properties = document.core_properties
    properties.title = _docx_text(_field_display_value(bundle["title"]))
    properties.subject = "MARVIS 策略迭代评审"
    properties.author = "MARVIS"
    properties.last_modified_by = "MARVIS"
    properties.keywords = "MARVIS;策略评审;受治理证据"
    properties.comments = STRATEGY_REPORT_OUTPUT_SCHEMA_VERSION
    properties.category = "Strategy Review"
    properties.identifier = bundle["report_id"]
    properties.created = _FIXED_DOCX_DATETIME
    properties.modified = _FIXED_DOCX_DATETIME
    properties.revision = 1


def _write_docx_masthead(document: Any, *, bundle: Mapping[str, Any]) -> None:
    title = document.add_paragraph(style="Title")
    _add_docx_run(title, _field_display_value(bundle["title"]))

    subtitle = document.add_paragraph(style="Subtitle")
    _add_docx_run(subtitle, "策略迭代评审报告 | 受治理证据投影")

    metadata = (
        ("报告 ID", bundle["report_id"]),
        ("修订", bundle["report_revision"]),
        ("前一修订", bundle["previous_report_id"]),
        ("状态", _REPORT_STATUS_LABELS[bundle["status"]]),
        ("策略 ID", bundle["strategy_id"]),
        ("策略版本", bundle["strategy_version"]),
        ("策略类型", bundle["strategy_type"]),
        ("效果阶段", _stage_labels(bundle["effect_stages"])),
        ("生成时间", bundle["generated_at"]),
        ("数据分级", bundle["data_classification"]),
        ("内容 SHA-256", bundle["content_sha256"]),
    )
    for label, value in metadata:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        _add_docx_run(paragraph, f"{label}: ", bold=True)
        _add_docx_run(paragraph, value)

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(10)
    _set_docx_paragraph_bottom_border(rule, color="2E74B5", size=10)

    notice = document.add_paragraph()
    _add_docx_run(
        notice,
        "本报告是受治理结构化证据的投影。"
        "空白不等于 0，效果阶段不自动升级。",
        italic=True,
        color="536878",
    )


def _set_docx_paragraph_bottom_border(
    paragraph: Any,
    *,
    color: str,
    size: int,
) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    borders = paragraph_properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        paragraph_properties.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def _write_docx_voting_search_table(
    document: Any,
    table: Mapping[str, Any],
    *,
    maximum_rows: int,
) -> tuple[int, int]:
    """Render the reserved Voting table within local and global row budgets."""

    expected_keys = (
        "search_id",
        "combo_id",
        "member_ids",
        "n",
        "eligible",
        "objective_metric",
        "objective_direction",
        "objective_value",
        "constraint_failures",
        "metrics",
    )
    columns = {column["key"]: column for column in table["columns"]}
    if set(columns) != set(expected_keys):
        raise StrategyReportOutputError(
            "Voting search report table columns are invalid"
        )
    total = len(table["rows"])
    shown_rows = list(
        islice(
            table["rows"],
            min(
                _DOCX_MAX_VOTING_SEARCH_ROWS,
                max(0, maximum_rows),
            ),
        )
    )
    rendered_rows = []
    for row in shown_rows:
        cells = row["cells"]

        def display(key: str) -> str:
            return _docx_text(
                _table_field_display_value(
                    cells[key],
                    column=columns[key],
                )
            )

        rendered_rows.append(
            (
                display("search_id"),
                display("combo_id"),
                display("member_ids"),
                display("n"),
                display("eligible"),
                " / ".join(
                    (
                        display("objective_metric"),
                        display("objective_direction"),
                        display("objective_value"),
                    )
                ),
                display("constraint_failures"),
                display("metrics"),
            )
        )
    _add_docx_table(
        document,
        headers=(
            "搜索ID",
            "组合ID",
            "成员规则ID",
            "n",
            "约束通过",
            "目标指标 / 方向 / 值",
            "约束未通过明细",
            "完整指标",
        ),
        rows=rendered_rows,
        widths=_DOCX_VOTING_SEARCH_WIDTHS,
        centered_columns=frozenset({3, 4}),
    )
    return len(rendered_rows), total


def _write_docx_section(
    document: Any,
    *,
    index: int,
    section: Mapping[str, Any],
    budget: _DocxRenderBudget,
) -> None:
    _add_docx_heading(
        document,
        f"{index}. {section['title']}",
        level=1,
    )
    stages = sorted(
        {item["effect_stage"] for item in section["stage_evidence"]},
        key=("estimated", "backtested", "oot_validated", "post_launch_observed").index,
    )
    state = document.add_paragraph()
    _add_docx_run(state, "状态: ", bold=True)
    _add_docx_run(
        state,
        _AVAILABILITY_LABELS[section["availability"]],
    )
    if stages:
        _add_docx_run(state, "    效果阶段: ", bold=True)
        _add_docx_run(state, _stage_labels(stages))

    summary_total = len(section["summary_fields"])
    summary_shown = min(
        summary_total,
        _DOCX_MAX_SUMMARY_FIELDS_PER_SECTION,
        budget.summary_fields_remaining,
    )
    if summary_total:
        _add_docx_heading(document, "摘要字段", level=2)
        rows = []
        for item in islice(section["summary_fields"], summary_shown):
            field = item["field"]
            rows.append(
                (
                    item["label"],
                    _field_display_value(field),
                    _AVAILABILITY_LABELS[field["availability"]],
                    field["note"],
                    _docx_compact_refs(field["source_refs"]),
                )
            )
        if rows:
            _add_docx_table(
                document,
                headers=("字段", "值", "状态", "说明", "来源"),
                rows=rows,
                widths=_DOCX_SUMMARY_WIDTHS,
                centered_columns=frozenset({2}),
            )
        budget.summary_fields_remaining -= summary_shown
        if summary_shown < summary_total:
            _write_docx_truncation_notice(
                document,
                scope=f"{section['title']}摘要字段",
                shown=summary_shown,
                total=summary_total,
            )

    stage_total = len(section["stage_evidence"])
    stage_shown = min(
        stage_total,
        _DOCX_MAX_STAGE_ROWS_PER_SECTION,
        budget.stage_rows_remaining,
    )
    if stage_total:
        _add_docx_heading(document, "效果证据绑定", level=2)
        rows = [
            (
                _EFFECT_STAGE_LABELS[item["effect_stage"]],
                item["population"],
                item["partition"],
                item["binding"],
            )
            for item in islice(section["stage_evidence"], stage_shown)
        ]
        if rows:
            _add_docx_table(
                document,
                headers=("效果阶段", "样本口径", "分区", "绑定"),
                rows=rows,
                widths=_DOCX_STAGE_WIDTHS,
                centered_columns=frozenset({0, 1, 2}),
            )
        budget.stage_rows_remaining -= stage_shown
        if stage_shown < stage_total:
            _write_docx_truncation_notice(
                document,
                scope=f"{section['title']}效果证据",
                shown=stage_shown,
                total=stage_total,
            )

    tables_total = len(section["tables"])
    tables_shown = min(tables_total, budget.report_tables_remaining)
    for table in islice(section["tables"], tables_shown):
        stage = (
            ""
            if table["effect_stage"] is None
            else f"（{_EFFECT_STAGE_LABELS[table['effect_stage']]}）"
        )
        _add_docx_heading(
            document,
            f"{table['title']}{stage}",
            level=2,
        )
        metadata = document.add_paragraph()
        metadata.paragraph_format.space_before = Pt(4)
        metadata.paragraph_format.space_after = Pt(4)
        _add_docx_run(metadata, "Table ID: ", bold=True)
        _add_docx_run(metadata, table["table_id"])
        _add_docx_run(metadata, "    粒度: ", bold=True)
        _add_docx_run(metadata, table["granularity"])
        _add_docx_run(metadata, "    内容类型: ", bold=True)
        _add_docx_run(metadata, table["content_class"])

        if table["table_id"] == _DOCX_VOTING_SEARCH_TABLE_ID:
            voting_shown, voting_total = _write_docx_voting_search_table(
                document,
                table,
                maximum_rows=budget.fact_rows_remaining,
            )
            budget.fact_rows_remaining -= voting_shown
            if voting_shown < voting_total:
                _write_docx_truncation_notice(
                    document,
                    scope=f"表 {table['table_id']} 组合行",
                    shown=voting_shown,
                    total=voting_total,
                )
        else:
            fact_total = len(table["rows"]) * len(table["columns"])
            fact_limit = min(
                fact_total,
                _DOCX_MAX_FACT_ROWS_PER_TABLE,
                budget.fact_rows_remaining,
            )
            fact_rows = []
            for row in table["rows"]:
                for column in table["columns"]:
                    if len(fact_rows) >= fact_limit:
                        break
                    field = row["cells"][column["key"]]
                    fact_rows.append(
                        (
                            row["row_id"],
                            column["label"],
                            _table_field_display_value(field, column=column),
                            _AVAILABILITY_LABELS[field["availability"]],
                            field["note"],
                            _docx_compact_refs(field["source_refs"]),
                        )
                    )
                if len(fact_rows) >= fact_limit:
                    break
            _add_docx_table(
                document,
                headers=("行 ID", "字段", "值", "状态", "说明", "来源"),
                rows=fact_rows,
                widths=_DOCX_FACT_WIDTHS,
                centered_columns=frozenset({2, 3}),
            )
            budget.fact_rows_remaining -= len(fact_rows)
            if len(fact_rows) < fact_total:
                _write_docx_truncation_notice(
                    document,
                    scope=f"表 {table['table_id']} 事实行",
                    shown=len(fact_rows),
                    total=fact_total,
                )
        source = document.add_paragraph()
        source.paragraph_format.space_before = Pt(4)
        source.paragraph_format.space_after = Pt(4)
        _add_docx_run(source, "来源: ", bold=True)
        _add_docx_run(source, _docx_compact_refs(table["source_refs"]))
    budget.report_tables_remaining -= tables_shown
    if tables_shown < tables_total:
        _write_docx_truncation_notice(
            document,
            scope=f"{section['title']}结构化表",
            shown=tables_shown,
            total=tables_total,
        )

    flags_total = len(section["red_flags"])
    flags_shown = min(
        flags_total,
        _DOCX_MAX_RED_FLAGS_PER_SECTION,
        budget.red_flags_remaining,
    )
    if flags_total:
        _add_docx_heading(document, "红旗与限制", level=2)
        rows = [
            (
                flag["level"],
                flag["code"],
                flag["message"],
                _docx_compact_refs(flag["source_refs"]),
            )
            for flag in islice(section["red_flags"], flags_shown)
        ]
        if rows:
            _add_docx_table(
                document,
                headers=("级别", "编码", "说明", "来源"),
                rows=rows,
                widths=_DOCX_FLAG_WIDTHS,
                centered_columns=frozenset({0}),
            )
        budget.red_flags_remaining -= flags_shown
        if flags_shown < flags_total:
            _write_docx_truncation_notice(
                document,
                scope=f"{section['title']}红旗",
                shown=flags_shown,
                total=flags_total,
            )

    if (
        not section["summary_fields"]
        and not section["tables"]
        and not section["stage_evidence"]
        and not section["red_flags"]
    ):
        paragraph = document.add_paragraph()
        _add_docx_run(paragraph, "该部分当前没有可展示的结构化事实。")


def _write_docx_evidence_appendix(
    document: Any,
    *,
    bundle: Mapping[str, Any],
) -> None:
    _add_docx_heading(document, "证据、缺失信息与完整度", level=1)
    _add_docx_heading(document, "完整度", level=2)
    completeness = bundle["completeness_summary"]
    rows = []
    for category, counts in (
        ("字段", completeness["field_counts"]),
        ("Section", completeness["section_counts"]),
        ("缺失信息", completeness["missing_information_counts"]),
        ("阻塞", completeness["blocking_counts"]),
    ):
        rows.extend((category, key, count) for key, count in counts.items())
    _add_docx_table(
        document,
        headers=("类别", "状态", "数量"),
        rows=rows,
        widths=_DOCX_COMPLETENESS_WIDTHS,
        centered_columns=frozenset({2}),
    )

    _add_docx_heading(document, "缺失信息", level=2)
    missing_total = len(bundle["missing_information"])
    missing_shown = min(
        missing_total,
        _DOCX_MAX_MISSING_INFORMATION_ROWS,
    )
    missing_rows = [
        (
            item["field_path"],
            item["blocking"],
            item["status"],
            item["reason"],
            item["asked_count"],
        )
        for item in islice(bundle["missing_information"], missing_shown)
    ]
    _add_docx_table(
        document,
        headers=("字段", "阻塞级别", "状态", "原因", "已询问次数"),
        rows=missing_rows,
        widths=_DOCX_MISSING_WIDTHS,
        centered_columns=frozenset({1, 2, 4}),
    )
    if not missing_rows:
        paragraph = document.add_paragraph()
        _add_docx_run(paragraph, "当前没有缺失信息记录。")
    if missing_shown < missing_total:
        _write_docx_truncation_notice(
            document,
            scope="缺失信息",
            shown=missing_shown,
            total=missing_total,
        )

    _add_docx_heading(document, "来源索引", level=2)
    bounded_sources = list(
        islice(
            _source_locations(bundle),
            _DOCX_MAX_EVIDENCE_ROWS + 1,
        )
    )
    evidence_was_truncated = len(bounded_sources) > _DOCX_MAX_EVIDENCE_ROWS
    evidence_rows = [
        (
            location,
            ref["kind"],
            ref["ref_id"],
            ref["content_hash"],
        )
        for location, ref in bounded_sources[:_DOCX_MAX_EVIDENCE_ROWS]
    ]
    _add_docx_table(
        document,
        headers=("位置", "类型", "引用 ID", "内容 SHA-256"),
        rows=evidence_rows,
        widths=_DOCX_EVIDENCE_WIDTHS,
    )
    if evidence_was_truncated:
        _write_docx_truncation_notice(
            document,
            scope="来源索引",
            shown=_DOCX_MAX_EVIDENCE_ROWS,
            total=None,
        )


def _write_docx_truncation_notice(
    document: Any,
    *,
    scope: object,
    shown: int,
    total: int | None,
) -> None:
    paragraph = document.add_paragraph()
    if total is None:
        message = (
            f"已截断：{scope}仅展示前 {shown} 项，仍有更多内容。"
            "完整内容见同一报告修订的 JSON/XLSX 输出。"
        )
    else:
        message = (
            f"已截断：{scope}仅展示前 {shown} 项，共 {total} 项。"
            "完整内容见同一报告修订的 JSON/XLSX 输出。"
        )
    _add_docx_run(
        paragraph,
        message,
        bold=True,
        color="7A5A00",
    )


def _add_docx_table(
    document: Any,
    *,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
    widths: Sequence[int],
    centered_columns: frozenset[int] = frozenset(),
) -> Any:
    if len(headers) != len(widths):
        raise StrategyReportOutputError(
            "DOCX table header and geometry column counts differ"
        )
    if sum(widths) != _DOCX_TABLE_WIDTH_DXA:
        raise StrategyReportOutputError(
            "DOCX table geometry does not equal 9360 DXA"
        )
    if any(len(row) != len(headers) for row in rows):
        raise StrategyReportOutputError(
            "DOCX table row and header column counts differ"
        )

    table = document.add_table(rows=1, cols=len(headers))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_docx_table_geometry(table, widths=widths)
    _set_docx_header_repeat(table.rows[0])

    for column, value in enumerate(headers):
        cell = table.rows[0].cells[column]
        _set_docx_cell_text(
            cell,
            value,
            bold=True,
            fill="F2F4F7",
            centered=True,
        )
    for values in rows:
        cells = table.add_row().cells
        for column, value in enumerate(values):
            _set_docx_cell_width(cells[column], widths[column])
            _set_docx_cell_text(
                cells[column],
                value,
                centered=column in centered_columns,
            )
    return table


def _set_docx_table_geometry(
    table: Any,
    *,
    widths: Sequence[int],
) -> None:
    table_properties = table._tbl.tblPr
    _set_docx_width_element(
        table_properties,
        "w:tblW",
        _DOCX_TABLE_WIDTH_DXA,
    )
    _set_docx_width_element(
        table_properties,
        "w:tblInd",
        _DOCX_TABLE_INDENT_DXA,
    )
    layout = table_properties.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_properties.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = table_properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "D9E1E8")

    margins = table_properties.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        table_properties.append(margins)
    for side, width in _DOCX_CELL_MARGINS_DXA.items():
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(width))
        element.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in tuple(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            _set_docx_cell_width(cell, width)


def _set_docx_width_element(
    parent: Any,
    tag: str,
    width: int,
) -> None:
    element = parent.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        parent.append(element)
    element.set(qn("w:w"), str(width))
    element.set(qn("w:type"), "dxa")


def _set_docx_cell_width(cell: Any, width: int) -> None:
    cell.width = Twips(width)
    properties = cell._tc.get_or_add_tcPr()
    _set_docx_width_element(properties, "w:tcW", width)


def _set_docx_header_repeat(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = properties.find(qn("w:tblHeader"))
    if repeat is None:
        repeat = OxmlElement("w:tblHeader")
        properties.append(repeat)
    repeat.set(qn("w:val"), "true")


def _set_docx_cell_text(
    cell: Any,
    value: object,
    *,
    bold: bool = False,
    fill: str | None = None,
    centered: bool = False,
) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.style = "Normal"
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
        if centered
        else WD_ALIGN_PARAGRAPH.LEFT
    )
    _add_docx_run(paragraph, value, bold=bold)
    if fill is not None:
        properties = cell._tc.get_or_add_tcPr()
        shading = properties.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            properties.append(shading)
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), fill)


def _docx_text(value: object) -> str:
    presentation = _presentation_value(value)
    text = _safe_text_projection("" if presentation is None else presentation)
    if len(text) <= _DOCX_MAX_CELL_CHARACTERS:
        return text
    prefix_characters = (
        _DOCX_MAX_CELL_CHARACTERS - len(_DOCX_TRUNCATION_SUFFIX)
    )
    return text[:prefix_characters] + _DOCX_TRUNCATION_SUFFIX


def _add_docx_heading(
    document: Any,
    value: object,
    *,
    level: int,
) -> Any:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    _add_docx_run(paragraph, value)
    return paragraph


def _add_docx_run(
    paragraph: Any,
    value: object,
    *,
    bold: bool = False,
    italic: bool = False,
    color: str | None = None,
) -> Any:
    text = _docx_text(value)
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if _DOCX_CJK_TEXT.search(text):
        run.font.name = _DOCX_CJK_FONT
        properties = run._element.get_or_add_rPr()
        fonts = properties.rFonts
        if fonts is None:
            fonts = OxmlElement("w:rFonts")
            properties.insert(0, fonts)
        for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts.set(qn(f"w:{attribute}"), _DOCX_CJK_FONT)
        language = properties.find(qn("w:lang"))
        if language is None:
            language = OxmlElement("w:lang")
            properties.append(language)
        language.set(qn("w:val"), "zh-CN")
        language.set(qn("w:eastAsia"), "zh-CN")
    return run


def _canonicalize_docx_bytes(raw: bytes) -> bytes:
    _assert_safe_docx_package(raw)
    source = BytesIO(raw)
    destination = BytesIO()
    with ZipFile(source, "r") as input_archive:
        members = sorted(input_archive.infolist(), key=lambda item: item.filename)
        if len({item.filename for item in members}) != len(members):
            raise StrategyReportOutputError(
                "generated DOCX contains duplicate package members"
            )
        with ZipFile(
            destination,
            "w",
            compression=ZIP_DEFLATED,
            compresslevel=9,
            allowZip64=True,
        ) as output_archive:
            output_archive.comment = b""
            for source_member in members:
                member = ZipInfo(
                    source_member.filename,
                    date_time=_FIXED_ZIP_DATETIME,
                )
                member.compress_type = ZIP_DEFLATED
                member.create_system = 0
                member.create_version = 20
                member.extract_version = 20
                member.external_attr = _FIXED_ZIP_EXTERNAL_ATTR
                member.internal_attr = 0
                member.flag_bits = 0
                member.volume = 0
                member.comment = b""
                member.extra = b""
                output_archive.writestr(
                    member,
                    input_archive.read(source_member.filename),
                    compress_type=ZIP_DEFLATED,
                    compresslevel=9,
                )
    return destination.getvalue()


def _assert_safe_docx_package(raw: bytes) -> None:
    try:
        archive = ZipFile(BytesIO(raw), "r")
    except Exception as exc:
        raise StrategyReportOutputError(
            "generated DOCX is not a readable OOXML package"
        ) from exc
    with archive:
        names = archive.namelist()
        if len(set(names)) != len(names):
            raise StrategyReportOutputError(
                "generated DOCX contains duplicate package members"
            )
        for name in names:
            pieces = name.replace("\\", "/").split("/")
            if (
                name.startswith(("/", "\\"))
                or "\\" in name
                or any(piece in {"", ".", ".."} for piece in pieces)
            ):
                raise StrategyReportOutputError(
                    "generated DOCX contains an unsafe package member path"
                )
            lowered = name.lower()
            if (
                lowered.endswith("vbaproject.bin")
                or "macros" in lowered
                or lowered.startswith("word/media/")
            ):
                raise StrategyReportOutputError(
                    "generated DOCX contains forbidden active or image content"
                )

        content_types = archive.read("[Content_Types].xml").lower()
        if b"macroenabled" in content_types or b"vbaproject" in content_types:
            raise StrategyReportOutputError(
                "generated DOCX contains a macro-enabled content type"
            )

        for name in names:
            payload = archive.read(name)
            if name.endswith(".rels"):
                try:
                    relationships = ElementTree.fromstring(payload)
                except ElementTree.ParseError as exc:
                    raise StrategyReportOutputError(
                        "generated DOCX contains invalid relationships XML"
                    ) from exc
                for relationship in relationships.findall(
                    f"{{{_DOCX_RELATIONSHIP_NAMESPACE}}}Relationship"
                ):
                    if (
                        relationship.attrib.get("TargetMode", "").lower()
                        == "external"
                    ):
                        raise StrategyReportOutputError(
                            "generated DOCX contains an external relationship"
                        )
            if name.endswith(".xml"):
                try:
                    root = ElementTree.fromstring(payload)
                except ElementTree.ParseError as exc:
                    raise StrategyReportOutputError(
                        "generated DOCX contains invalid OOXML"
                    ) from exc
                for element in root.iter():
                    tag = element.tag
                    if not isinstance(tag, str) or not tag.startswith("{"):
                        continue
                    namespace, _, local_name = tag[1:].partition("}")
                    if (
                        namespace == _DOCX_WORDPROCESSINGML_NAMESPACE
                        and local_name
                        in _DOCX_FORBIDDEN_FIELD_LOCAL_NAMES
                    ):
                        raise StrategyReportOutputError(
                            "generated DOCX contains a forbidden field code"
                        )


def _render_xlsx(bundle: Mapping[str, Any]) -> bytes:
    workbook = Workbook()
    workbook.remove(workbook.active)
    workbook.properties.creator = "MARVIS"
    workbook.properties.lastModifiedBy = "MARVIS"
    workbook.properties.title = str(
        _safe_text_projection(_field_display_value(bundle["title"]))
    )
    workbook.properties.description = STRATEGY_REPORT_OUTPUT_SCHEMA_VERSION
    workbook.properties.created = _FIXED_WORKBOOK_DATETIME
    workbook.properties.modified = _FIXED_WORKBOOK_DATETIME

    sections = {section["key"]: section for section in bundle["sections"]}
    tables_by_sheet: dict[str, list[tuple[str, Mapping[str, Any]]]] = {}
    for section in bundle["sections"]:
        for table in section["tables"]:
            tables_by_sheet.setdefault(table["sheet_key"], []).append(
                (section["key"], table)
            )
    appendix_keys = sorted(
        key for key in tables_by_sheet if key not in REPORT_CORE_SHEET_KEYS
    )
    try:
        for sheet_key in REPORT_CORE_SHEET_KEYS:
            sheet = workbook.create_sheet(sheet_key)
            if sheet_key == "00_summary":
                _write_summary_sheet(
                    sheet,
                    bundle=bundle,
                    tables=tables_by_sheet.get(sheet_key, ()),
                )
            elif sheet_key == "11_evidence":
                _write_evidence_sheet(sheet, bundle=bundle)
            else:
                _write_business_sheet(
                    sheet,
                    sheet_key=sheet_key,
                    sections=sections,
                    tables=tables_by_sheet.get(sheet_key, ()),
                )
        for sheet_key in appendix_keys:
            sheet = workbook.create_sheet(sheet_key)
            _write_appendix_sheet(
                sheet,
                sheet_key=sheet_key,
                tables=tables_by_sheet[sheet_key],
                sections=sections,
            )
        raw = BytesIO()
        workbook.save(raw)
    finally:
        workbook.close()
    return _canonicalize_xlsx_bytes(raw.getvalue())


def _write_summary_sheet(
    sheet: Any,
    *,
    bundle: Mapping[str, Any],
    tables: Sequence[tuple[str, Mapping[str, Any]]],
) -> None:
    _prepare_sheet(sheet)
    row = _write_sheet_title(
        sheet,
        title=STRATEGY_REPORT_SHEET_TITLES["00_summary"],
        subtitle="受治理证据投影；空白不等于 0，效果阶段不自动升级",
    )
    row = _write_key_values(
        sheet,
        start_row=row,
        title="报告身份与状态",
        rows=(
            ("报告标题", _field_display_value(bundle["title"])),
            ("报告 ID", bundle["report_id"]),
            ("修订", bundle["report_revision"]),
            ("前一修订", bundle["previous_report_id"]),
            ("状态", _REPORT_STATUS_LABELS[bundle["status"]]),
            ("策略 ID", bundle["strategy_id"]),
            ("策略版本", bundle["strategy_version"]),
            ("策略类型", bundle["strategy_type"]),
            ("效果阶段", _stage_labels(bundle["effect_stages"])),
            ("生成时间", bundle["generated_at"]),
            ("数据分级", bundle["data_classification"]),
            ("Producer", bundle["producer_version"]),
            ("内容 SHA-256", bundle["content_sha256"]),
        ),
    )
    row = _write_section_heading(sheet, row, "七步完成状态")
    _write_header_row(sheet, row, ("步骤", "模块", "状态", "红旗数", "证据数"))
    row += 1
    for index, section in enumerate(bundle["sections"], start=1):
        _write_row(
            sheet,
            row,
            (
                index,
                section["title"],
                _AVAILABILITY_LABELS[section["availability"]],
                len(section["red_flags"]),
                len(section["source_refs"]),
            ),
        )
        row += 1
    for section_key in (
        "current_project",
        "candidate_combinations",
        "impact_assessment",
        "final_document",
    ):
        row = _write_summary_fields(
            sheet,
            start_row=row,
            section=next(
                item
                for item in bundle["sections"]
                if item["key"] == section_key
            ),
        )
    row += 1
    row = _write_section_heading(sheet, row, "完整度")
    _write_header_row(sheet, row, ("类别", "状态", "数量"))
    row += 1
    completeness = bundle["completeness_summary"]
    for category, counts in (
        ("字段", completeness["field_counts"]),
        ("缺失信息", completeness["missing_information_counts"]),
        ("阻塞", completeness["blocking_counts"]),
    ):
        for key, count in counts.items():
            _write_row(sheet, row, (category, key, count))
            row += 1
    row += 1
    row = _write_red_flags(
        sheet,
        start_row=row,
        sections=bundle["sections"],
    )
    row = _write_routed_tables(
        sheet,
        start_row=row,
        tables=tables,
    )
    _guard_sheet_bounds(sheet)
    _set_widths(sheet, {1: 24, 2: 70, 3: 24, 4: 18, 5: 18})


def _write_business_sheet(
    sheet: Any,
    *,
    sheet_key: str,
    sections: Mapping[str, Mapping[str, Any]],
    tables: Sequence[tuple[str, Mapping[str, Any]]],
) -> None:
    _prepare_sheet(sheet)
    title = STRATEGY_REPORT_SHEET_TITLES[sheet_key]
    row = _write_sheet_title(
        sheet,
        title=title,
        subtitle=f"Sheet key: {sheet_key}",
    )
    relevant = [sections[key] for key in _SHEET_SECTION_KEYS[sheet_key]]
    row = _write_section_states(sheet, start_row=row, sections=relevant)
    for section in relevant:
        if _SECTION_PRIMARY_SHEETS[section["key"]] == sheet_key:
            row = _write_summary_fields(sheet, start_row=row, section=section)
            row = _write_red_flags(sheet, start_row=row, sections=(section,))
    row = _write_routed_tables(sheet, start_row=row, tables=tables)
    if not tables and all(
        _SECTION_PRIMARY_SHEETS[section["key"]] != sheet_key
        for section in relevant
    ):
        _write_row(sheet, row, ("当前没有路由到本 Sheet 的结构化表。",))
    _guard_sheet_bounds(sheet)
    _set_widths(sheet, {1: 24, 2: 40, 3: 18, 4: 42, 5: 54})


def _write_appendix_sheet(
    sheet: Any,
    *,
    sheet_key: str,
    tables: Sequence[tuple[str, Mapping[str, Any]]],
    sections: Mapping[str, Mapping[str, Any]],
) -> None:
    _prepare_sheet(sheet)
    row = _write_sheet_title(
        sheet,
        title=f"分析附件：{sheet_key}",
        subtitle="附件表仍引用同一受治理证据，不重新计算指标",
    )
    section_keys = tuple(dict.fromkeys(section_key for section_key, _ in tables))
    row = _write_section_states(
        sheet,
        start_row=row,
        sections=[sections[key] for key in section_keys],
    )
    _write_routed_tables(sheet, start_row=row, tables=tables)
    _guard_sheet_bounds(sheet)
    _set_widths(sheet, {1: 24, 2: 40, 3: 18, 4: 42, 5: 54})


def _write_evidence_sheet(sheet: Any, *, bundle: Mapping[str, Any]) -> None:
    _prepare_sheet(sheet)
    row = _write_sheet_title(
        sheet,
        title=STRATEGY_REPORT_SHEET_TITLES["11_evidence"],
        subtitle="完整 hash、缺失信息状态和字段级 lineage",
    )
    row = _write_section_heading(sheet, row, "缺失信息")
    _write_header_row(
        sheet,
        row,
        (
            "Missing ID",
            "字段",
            "阻塞级别",
            "状态",
            "原因",
            "问题",
            "已询问",
            "询问时间",
            "回答时间",
            "依赖 SHA-256",
            "回答来源",
        ),
    )
    row += 1
    for item in bundle["missing_information"]:
        _write_row(
            sheet,
            row,
            (
                item["missing_information_id"],
                item["field_path"],
                item["blocking"],
                item["status"],
                item["reason"],
                item["question"],
                item["asked_count"],
                item["asked_at"],
                item["answered_at"],
                item["dependency_hash"],
                _compact_refs(
                    []
                    if item["answer_source_ref"] is None
                    else [item["answer_source_ref"]]
                ),
            ),
        )
        row += 1
    if not bundle["missing_information"]:
        _write_row(sheet, row, ("", "", "", "", "", "", 0))
        row += 1

    row += 1
    row = _write_section_heading(sheet, row, "来源索引")
    _write_header_row(
        sheet,
        row,
        ("位置", "类型", "引用 ID", "内容 SHA-256"),
    )
    row += 1
    for location, ref in _source_locations(bundle):
        _write_row(
            sheet,
            row,
            (location, ref["kind"], ref["ref_id"], ref["content_hash"]),
        )
        row += 1

    row += 1
    row = _write_section_heading(sheet, row, "完整度")
    _write_header_row(sheet, row, ("类别", "状态", "数量"))
    row += 1
    completeness = bundle["completeness_summary"]
    for category, counts in (
        ("字段", completeness["field_counts"]),
        ("Section", completeness["section_counts"]),
        ("缺失信息", completeness["missing_information_counts"]),
        ("阻塞", completeness["blocking_counts"]),
    ):
        for key, count in counts.items():
            _write_row(sheet, row, (category, key, count))
            row += 1

    _guard_sheet_bounds(sheet)
    _set_widths(
        sheet,
        {
            1: 52,
            2: 32,
            3: 22,
            4: 20,
            5: 70,
            6: 70,
            7: 14,
            8: 28,
            9: 28,
            10: 68,
            11: 60,
        },
    )


def _write_section_states(
    sheet: Any,
    *,
    start_row: int,
    sections: Sequence[Mapping[str, Any]],
) -> int:
    row = _write_section_heading(sheet, start_row, "模块状态")
    _write_header_row(
        sheet,
        row,
        ("模块", "状态", "效果阶段", "红旗数", "来源"),
    )
    row += 1
    for section in sections:
        stages = sorted(
            {item["effect_stage"] for item in section["stage_evidence"]},
            key=(
                "estimated",
                "backtested",
                "oot_validated",
                "post_launch_observed",
            ).index,
        )
        _write_row(
            sheet,
            row,
            (
                section["title"],
                _AVAILABILITY_LABELS[section["availability"]],
                _stage_labels(stages),
                len(section["red_flags"]),
                _compact_refs(section["source_refs"]),
            ),
        )
        row += 1
    return row + 1


def _write_summary_fields(
    sheet: Any,
    *,
    start_row: int,
    section: Mapping[str, Any],
) -> int:
    if not section["summary_fields"]:
        return start_row
    row = _write_section_heading(sheet, start_row, f"{section['title']}摘要")
    _write_header_row(
        sheet,
        row,
        (
            "字段 ID",
            "字段",
            "值",
            "状态",
            "来源类型",
            "阻塞级别",
            "统计时点",
            "说明",
            "来源引用",
        ),
    )
    row += 1
    for item in section["summary_fields"]:
        field = item["field"]
        _write_row(
            sheet,
            row,
            (
                item["field_id"],
                item["label"],
                _field_display_value(field),
                _AVAILABILITY_LABELS[field["availability"]],
                field["origin"],
                field["blocking"],
                field["as_of"],
                field["note"],
                _compact_refs(field["source_refs"]),
            ),
        )
        _style_availability_cell(sheet.cell(row=row, column=4), field["availability"])
        row += 1
    return row + 1


def _write_red_flags(
    sheet: Any,
    *,
    start_row: int,
    sections: Sequence[Mapping[str, Any]],
) -> int:
    rows = [
        (section["title"], flag)
        for section in sections
        for flag in section["red_flags"]
    ]
    if not rows:
        return start_row
    row = _write_section_heading(sheet, start_row, "红旗与限制")
    _write_header_row(sheet, row, ("模块", "级别", "编码", "说明", "来源"))
    row += 1
    for section_title, flag in rows:
        _write_row(
            sheet,
            row,
            (
                section_title,
                flag["level"],
                flag["code"],
                flag["message"],
                _compact_refs(flag["source_refs"]),
            ),
        )
        _style_flag_cell(sheet.cell(row=row, column=2), flag["level"])
        row += 1
    return row + 1


def _write_routed_tables(
    sheet: Any,
    *,
    start_row: int,
    tables: Sequence[tuple[str, Mapping[str, Any]]],
) -> int:
    row = start_row
    for section_key, table in tables:
        row = _write_report_table(
            sheet,
            start_row=row,
            section_key=section_key,
            table=table,
        )
    return row


def _write_report_table(
    sheet: Any,
    *,
    start_row: int,
    section_key: str,
    table: Mapping[str, Any],
) -> int:
    expanded_columns = len(table["columns"]) * 4
    if expanded_columns > _MAX_XLSX_COLUMNS:
        raise StrategyReportOutputError(
            f"table {table['table_id']} exceeds Excel's column limit"
        )
    row = _write_section_heading(
        sheet,
        start_row,
        table["title"],
    )
    _write_row(
        sheet,
        row,
        (
            "Table ID",
            table["table_id"],
            "模块",
            section_key,
            "效果阶段",
            ""
            if table["effect_stage"] is None
            else _EFFECT_STAGE_LABELS[table["effect_stage"]],
            "粒度",
            table["granularity"],
            "内容类型",
            table["content_class"],
            "来源",
            _compact_refs(table["source_refs"]),
        ),
    )
    row += 1
    headers: list[str] = []
    for column in table["columns"]:
        label = column["label"]
        if column["unit"] is not None:
            label = f"{label} [{column['unit']}]"
        headers.extend((label, f"{label}状态", f"{label}说明", f"{label}来源"))
    _write_header_row(sheet, row, headers)
    row += 1
    for item in table["rows"]:
        values: list[object] = []
        value_columns: list[tuple[int, Mapping[str, Any], Mapping[str, Any]]] = []
        for index, column in enumerate(table["columns"]):
            field = item["cells"][column["key"]]
            values.extend(
                (
                    _field_display_value(field),
                    _AVAILABILITY_LABELS[field["availability"]],
                    field["note"],
                    _compact_refs(field["source_refs"]),
                )
            )
            value_columns.append((index * 4 + 1, field, column))
        _write_row(sheet, row, values)
        for column_index, field, column in value_columns:
            _apply_number_format(
                sheet.cell(row=row, column=column_index),
                field=field,
                column=column,
            )
            _style_availability_cell(
                sheet.cell(row=row, column=column_index + 1),
                field["availability"],
            )
        row += 1
    if not table["rows"]:
        _write_row(sheet, row, tuple("" for _ in headers))
        row += 1
    return row + 2


def _write_sheet_title(sheet: Any, *, title: str, subtitle: str) -> int:
    sheet.cell(row=1, column=1, value=_xlsx_cell(title))
    sheet.cell(row=1, column=1).font = Font(
        name="Arial",
        size=18,
        bold=True,
        color=_WHITE,
    )
    sheet.cell(row=1, column=1).fill = PatternFill("solid", fgColor=_NAVY)
    sheet.cell(row=1, column=1).alignment = Alignment(vertical="center")
    sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=9)
    sheet.row_dimensions[1].height = 30
    sheet.cell(row=2, column=1, value=_xlsx_cell(subtitle))
    sheet.cell(row=2, column=1).font = Font(
        name="Arial",
        size=10,
        italic=True,
        color="536878",
    )
    sheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=9)
    return 4


def _write_key_values(
    sheet: Any,
    *,
    start_row: int,
    title: str,
    rows: Iterable[tuple[object, object]],
) -> int:
    row = _write_section_heading(sheet, start_row, title)
    _write_header_row(sheet, row, ("属性", "值"))
    row += 1
    for item in rows:
        _write_row(sheet, row, item)
        row += 1
    return row + 1


def _write_section_heading(sheet: Any, row: int, title: object) -> int:
    sheet.cell(row=row, column=1, value=_xlsx_cell(title))
    sheet.cell(row=row, column=1).font = Font(
        name="Arial",
        size=12,
        bold=True,
        color=_WHITE,
    )
    sheet.cell(row=row, column=1).fill = PatternFill("solid", fgColor=_BLUE)
    sheet.cell(row=row, column=1).alignment = Alignment(vertical="center")
    sheet.merge_cells(start_row=row, start_column=1, end_row=row, end_column=9)
    return row + 1


def _write_header_row(sheet: Any, row: int, values: Sequence[object]) -> None:
    _write_row(sheet, row, values)
    for cell in sheet[row][: len(values)]:
        cell.font = Font(name="Arial", size=10, bold=True, color=_NAVY)
        cell.fill = PatternFill("solid", fgColor=_PALE_BLUE)
        cell.alignment = Alignment(
            horizontal="center",
            vertical="center",
            wrap_text=True,
        )


def _write_row(sheet: Any, row: int, values: Sequence[object]) -> None:
    if row > _MAX_XLSX_ROWS:
        raise StrategyReportOutputError("report exceeds Excel's row limit")
    if len(values) > _MAX_XLSX_COLUMNS:
        raise StrategyReportOutputError("report exceeds Excel's column limit")
    for column, value in enumerate(values, start=1):
        cell = sheet.cell(row=row, column=column, value=_xlsx_cell(value))
        cell.font = Font(name="Arial", size=10, color="1F2933")
        cell.alignment = Alignment(vertical="top", wrap_text=True)
        cell.border = _BORDER


def _prepare_sheet(sheet: Any) -> None:
    sheet.sheet_view.showGridLines = False
    sheet.freeze_panes = "A4"
    sheet.auto_filter.ref = None
    sheet.sheet_properties.pageSetUpPr.fitToPage = True
    sheet.page_setup.fitToWidth = 1
    sheet.page_setup.fitToHeight = 0
    sheet.sheet_properties.outlinePr.summaryBelow = True


def _set_widths(sheet: Any, widths: Mapping[int, float]) -> None:
    for column, width in widths.items():
        sheet.column_dimensions[get_column_letter(column)].width = width
    for column in range(1, sheet.max_column + 1):
        letter = get_column_letter(column)
        if sheet.column_dimensions[letter].width == 13.0:
            sheet.column_dimensions[letter].width = 20


def _guard_sheet_bounds(sheet: Any) -> None:
    if sheet.max_row > _MAX_XLSX_ROWS:
        raise StrategyReportOutputError(
            f"sheet {sheet.title} exceeds Excel's row limit"
        )
    if sheet.max_column > _MAX_XLSX_COLUMNS:
        raise StrategyReportOutputError(
            f"sheet {sheet.title} exceeds Excel's column limit"
        )


def _field_display_value(field: Mapping[str, Any]) -> object:
    availability = field["availability"]
    if availability in {"unavailable", "not_matured"}:
        return ""
    if availability == "not_applicable":
        return "本次不涉及"
    return _presentation_value(field["value"])


def _table_field_display_value(
    field: Mapping[str, Any],
    *,
    column: Mapping[str, Any],
) -> object:
    value = _field_display_value(field)
    if field["availability"] != "present":
        return value
    raw = field["value"]
    if isinstance(raw, bool) or not isinstance(raw, (int, float)):
        return value
    precision = column["precision"]
    if precision is None:
        return value
    if column["unit"] == "%":
        return f"{raw * 100:.{precision}f}%"
    return f"{raw:.{precision}f}"


def _presentation_value(value: object) -> object:
    if value is None or isinstance(value, (str, bool, int, float)):
        return value
    return json.dumps(
        value,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        allow_nan=False,
    )


def _source_locations(
    bundle: Mapping[str, Any],
) -> Iterable[tuple[str, Mapping[str, str]]]:
    for ref in bundle["title"]["source_refs"]:
        yield "report.title", ref
    for field_name in (
        "dataset_refs",
        "strategy_artifact_refs",
        "tool_run_refs",
    ):
        for ref in bundle[field_name]:
            yield field_name, ref
    for section in bundle["sections"]:
        prefix = f"sections.{section['key']}"
        for ref in section["source_refs"]:
            yield prefix, ref
        for item in section["summary_fields"]:
            for ref in item["field"]["source_refs"]:
                yield f"{prefix}.summary_fields.{item['field_id']}", ref
        for table in section["tables"]:
            table_prefix = f"{prefix}.tables.{table['table_id']}"
            for ref in table["source_refs"]:
                yield table_prefix, ref
            for row in table["rows"]:
                for key, field in row["cells"].items():
                    for ref in field["source_refs"]:
                        yield f"{table_prefix}.rows.{row['row_id']}.{key}", ref
        for index, item in enumerate(section["stage_evidence"]):
            binding = item["binding"]
            for key, ref in binding.items():
                if key in {"kind", "effective_period"}:
                    continue
                yield f"{prefix}.stage_evidence[{index}].{key}", ref
        for flag in section["red_flags"]:
            for ref in flag["source_refs"]:
                yield f"{prefix}.red_flags.{flag['code']}", ref
    for item in bundle["missing_information"]:
        if item["answer_source_ref"] is not None:
            yield (
                f"missing_information.{item['missing_information_id']}.answer",
                item["answer_source_ref"],
            )


def _compact_refs(refs: Sequence[Mapping[str, str]]) -> str:
    return "; ".join(
        f"{item['kind']}:{item['ref_id']}@{item['content_hash']}"
        for item in refs
    )


def _docx_compact_refs(refs: Sequence[Mapping[str, str]]) -> str:
    shown = min(len(refs), _DOCX_MAX_REFS_PER_CELL)
    text = "; ".join(
        f"{item['kind']}:{item['ref_id']}@{item['content_hash']}"
        for item in islice(refs, shown)
    )
    if shown < len(refs):
        suffix = "…[引用已截断；完整内容见 JSON/XLSX]"
        return f"{text}; {suffix}" if text else suffix
    return text


def _stage_labels(stages: Sequence[str]) -> str:
    return " / ".join(_EFFECT_STAGE_LABELS[item] for item in stages)


def _apply_number_format(
    cell: Any,
    *,
    field: Mapping[str, Any],
    column: Mapping[str, Any],
) -> None:
    if field["availability"] != "present":
        return
    value = field["value"]
    if isinstance(value, bool) or not isinstance(value, (int, float)):
        return
    precision = column["precision"]
    if precision is None:
        return
    decimals = "" if precision == 0 else "." + ("0" * precision)
    if column["unit"] == "%":
        cell.number_format = f"0{decimals}%"
    else:
        cell.number_format = f"#,##0{decimals}"


def _style_availability_cell(cell: Any, availability: str) -> None:
    fill = {
        "present": _PALE_GREEN,
        "unavailable": _PALE_AMBER,
        "not_applicable": _PALE_BLUE,
        "not_matured": _PALE_AMBER,
    }[availability]
    cell.fill = PatternFill("solid", fgColor=fill)


def _style_flag_cell(cell: Any, level: str) -> None:
    fill = {
        "info": _PALE_BLUE,
        "amber": _PALE_AMBER,
        "red": _PALE_RED,
    }[level]
    cell.fill = PatternFill("solid", fgColor=fill)


def _xlsx_cell(value: object) -> object:
    value = _presentation_value(value)
    if value is None or isinstance(value, (bool, float)):
        return value
    if isinstance(value, int):
        # Excel has only 15 significant decimal digits.  Preserve longer
        # identifiers and counts as text instead of silently rounding them.
        return str(value) if len(str(abs(value))) > 15 else value
    if not isinstance(value, str):
        raise StrategyReportOutputError(
            f"report cell contains unsupported {type(value).__name__}"
        )
    text = _safe_text_projection(value)
    if _looks_like_formula(text):
        text = "'" + text
    if len(text) > _MAX_XLSX_CELL_CHARACTERS:
        raise StrategyReportOutputError(
            "report cell exceeds Excel's 32767 character limit"
        )
    return text


def _looks_like_formula(value: str) -> bool:
    for character in value:
        if character.isspace() or unicodedata.category(character).startswith("C"):
            continue
        return character in _FORMULA_PREFIXES
    return False


def _markdown_cell(value: object) -> str:
    text = _safe_text_projection(
        str(_presentation_value(value) if value is not None else "")
    )
    escaped = (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
    escaped = _MARKDOWN_INLINE_META.sub(r"\\\1", escaped)
    return (
        escaped
        .replace("|", "\\|")
        .replace("\r\n", " / ")
        .replace("\n", " / ")
        .replace("\r", " / ")
    )


def _safe_text_projection(value: object) -> str:
    return _ILLEGAL_PRESENTATION_CONTROL.sub(
        lambda match: f"\\u{ord(match.group(0)):04x}",
        str(value),
    )


def _canonicalize_xlsx_bytes(raw: bytes) -> bytes:
    source = BytesIO(raw)
    destination = BytesIO()
    with ZipFile(source, "r") as input_archive:
        members = sorted(input_archive.infolist(), key=lambda item: item.filename)
        with ZipFile(
            destination,
            "w",
            compression=ZIP_DEFLATED,
            compresslevel=9,
            allowZip64=True,
        ) as output_archive:
            for source_member in members:
                member = ZipInfo(
                    source_member.filename,
                    date_time=_FIXED_ZIP_DATETIME,
                )
                member.compress_type = ZIP_DEFLATED
                member.create_system = 0
                member.external_attr = 0
                member.internal_attr = 0
                member.comment = b""
                member.extra = b""
                payload = input_archive.read(source_member.filename)
                if source_member.filename == "docProps/core.xml":
                    payload, replacements = _CORE_MODIFIED_TIMESTAMP.subn(
                        rb"\g<1>2000-01-01T00:00:00Z\g<2>",
                        payload,
                    )
                    if replacements != 1:
                        raise StrategyReportOutputError(
                            "generated workbook has an invalid modified "
                            "timestamp contract"
                        )
                output_archive.writestr(
                    member,
                    payload,
                    compress_type=ZIP_DEFLATED,
                    compresslevel=9,
                )
    return destination.getvalue()


__all__ = [
    "STRATEGY_REPORT_OUTPUT_SCHEMA_VERSION",
    "STRATEGY_REPORT_SHEET_TITLES",
    "StrategyReportOutputError",
    "render_strategy_report_bundle",
    "render_strategy_report_bundle_docx",
    "render_strategy_report_bundle_json",
    "render_strategy_report_bundle_markdown",
    "render_strategy_report_bundle_xlsx",
]
