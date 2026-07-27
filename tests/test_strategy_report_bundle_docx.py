from __future__ import annotations

from copy import deepcopy
from hashlib import sha256
from io import BytesIO
from xml.etree import ElementTree
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
import pytest

from marvis.output.strategy_report_bundle import (
    StrategyReportOutputError,
    _assert_safe_docx_package,
    render_strategy_report_bundle_docx,
)
from marvis.packs.strategy.report_bundle import (
    build_named_report_field,
    build_strategy_report_table,
)
from tests.test_strategy_report_bundle import (
    _bundle,
    _present,
    _sections,
    _source,
)


_RELATIONSHIP_NAMESPACE = (
    "http://schemas.openxmlformats.org/package/2006/relationships"
)


def _document(raw: bytes):
    return Document(BytesIO(raw))


def _table_rows(document):
    return [
        [cell.text for cell in row.cells]
        for table in document.tables
        for row in table.rows
    ]


def _replace_document_xml(raw: bytes, payload: bytes) -> bytes:
    rebuilt = BytesIO()
    with ZipFile(BytesIO(raw), "r") as source:
        with ZipFile(rebuilt, "w", compression=ZIP_DEFLATED) as destination:
            for member in source.infolist():
                destination.writestr(
                    member,
                    payload
                    if member.filename == "word/document.xml"
                    else source.read(member.filename),
                )
    return rebuilt.getvalue()


def test_docx_projection_is_bit_for_bit_deterministic_and_parseable():
    bundle = _bundle()

    first = render_strategy_report_bundle_docx(bundle)
    second = render_strategy_report_bundle_docx(bundle)

    assert first == second
    assert sha256(first).hexdigest() == sha256(second).hexdigest()
    document = _document(first)
    assert document.core_properties.author == "MARVIS"
    assert document.core_properties.identifier == bundle["report_id"]
    assert document.core_properties.created.year == 2000
    assert document.core_properties.modified.year == 2000


def test_docx_uses_business_brief_geometry_styles_and_memo_identity():
    bundle = _bundle()
    document = _document(render_strategy_report_bundle_docx(bundle))
    section = document.sections[0]

    assert section.page_width == Inches(8.5)
    assert section.page_height == Inches(11)
    assert section.top_margin == Inches(1)
    assert section.right_margin == Inches(1)
    assert section.bottom_margin == Inches(1)
    assert section.left_margin == Inches(1)
    assert section.header_distance.twips == 708
    assert section.footer_distance.twips == 708
    assert "MARVIS | 策略评审" in section.header.paragraphs[0].text
    assert bundle["report_id"] in section.header.paragraphs[0].text
    assert str(bundle["report_revision"]) in section.header.paragraphs[0].text

    styles = document.styles
    normal = styles["Normal"]
    assert normal.font.name == "Calibri"
    assert normal.font.size == Pt(11)
    assert normal.paragraph_format.space_after == Pt(6)
    assert normal.paragraph_format.line_spacing == 1.1
    normal_fonts = normal.element.rPr.rFonts
    assert normal_fonts.get(qn("w:ascii")) == "Calibri"
    assert normal_fonts.get(qn("w:hAnsi")) == "Calibri"
    assert normal_fonts.get(qn("w:eastAsia")) == "Microsoft YaHei"

    expected_headings = {
        "Heading 1": (16, "2E74B5", 16, 8),
        "Heading 2": (13, "2E74B5", 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for style_name, (size, color, before, after) in expected_headings.items():
        style = styles[style_name]
        assert style.font.name == "Calibri"
        assert style.font.size == Pt(size)
        assert str(style.font.color.rgb) == color
        assert style.paragraph_format.space_before == Pt(before)
        assert style.paragraph_format.space_after == Pt(after)

    heading_text = [
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.style.name == "Heading 1"
    ]
    assert "1. 当前项目状况" in heading_text
    assert "6. 策略影响测算" in heading_text
    assert "证据、缺失信息与完整度" in heading_text
    assert any(
        paragraph.text == "风险策略迭代评审"
        and paragraph.style.name == "Title"
        for paragraph in document.paragraphs
    )


def test_docx_tables_have_fixed_geometry_header_fill_and_no_fixed_rows():
    document = _document(render_strategy_report_bundle_docx(_bundle()))

    assert document.tables
    for table in document.tables:
        properties = table._tbl.tblPr
        width = properties.find(qn("w:tblW"))
        indent = properties.find(qn("w:tblInd"))
        layout = properties.find(qn("w:tblLayout"))
        margins = properties.find(qn("w:tblCellMar"))
        assert width is not None
        assert width.get(qn("w:w")) == "9360"
        assert width.get(qn("w:type")) == "dxa"
        assert indent is not None
        assert indent.get(qn("w:w")) == "120"
        assert indent.get(qn("w:type")) == "dxa"
        assert layout is not None
        assert layout.get(qn("w:type")) == "fixed"
        assert margins is not None
        assert {
            name: margins.find(qn(f"w:{name}")).get(qn("w:w"))
            for name in ("top", "bottom", "start", "end")
        } == {
            "top": "80",
            "bottom": "80",
            "start": "120",
            "end": "120",
        }

        grid_widths = [
            int(column.get(qn("w:w")))
            for column in table._tbl.tblGrid.findall(qn("w:gridCol"))
        ]
        assert sum(grid_widths) == 9360
        assert table.rows[0]._tr.trPr.find(qn("w:tblHeader")) is not None
        assert all(
            cell._tc.tcPr.find(qn("w:shd")).get(qn("w:fill")) == "F2F4F7"
            for cell in table.rows[0].cells
        )
        assert all(
            row._tr.trPr is None
            or row._tr.trPr.find(qn("w:trHeight")) is None
            for row in table.rows
        )


def test_docx_preserves_zero_blank_maturity_sections_and_evidence():
    document = _document(render_strategy_report_bundle_docx(_bundle()))
    rows = _table_rows(document)

    assert any(
        row[0] == "通过率"
        and row[1] == "0"
        and row[2] == "已有可信值"
        for row in rows
    )
    assert any(
        row[0] == "业务背景"
        and row[1] == ""
        and row[2] == "暂缺"
        and row[3] == "用户明确表示暂时没有"
        for row in rows
    )
    assert any(
        row[1] == "通过率"
        and row[2] == "0.0000%"
        and row[3] == "已有可信值"
        for row in rows
    )
    assert any(
        row[1] == "风险率"
        and row[2] == ""
        and row[3] == "尚未成熟"
        and "MOB3" in row[4]
        for row in rows
    )
    assert any(
        row[0] == "report.title"
        and row[1] == "report_context"
        and len(row[3]) == 64
        for row in rows
    )


def test_docx_plain_text_cannot_create_fields_external_relationships_or_media():
    sections = _sections()
    sections[0]["summary_fields"].append(
        build_named_report_field(
            field_id="external_text",
            label="外部说明",
            field=_present(
                'https://invalid.example { HYPERLINK "https://invalid.example" }',
                source=_source("report_context", "external-text"),
                origin="user",
            ),
        )
    )
    raw = render_strategy_report_bundle_docx(
        _bundle(
            title="外部链接 https://invalid.example",
            sections=deepcopy(sections),
        )
    )

    with ZipFile(BytesIO(raw), "r") as archive:
        assert archive.comment == b""
        infos = archive.infolist()
        assert [item.filename for item in infos] == sorted(archive.namelist())
        assert all(item.date_time == (1980, 1, 1, 0, 0, 0) for item in infos)
        assert all(item.compress_type == ZIP_DEFLATED for item in infos)
        assert all(item.create_system == 0 for item in infos)
        assert all(item.external_attr == 0o600 << 16 for item in infos)
        assert all(item.extra == b"" and item.comment == b"" for item in infos)
        assert not any(
            name.lower().startswith("word/media/")
            or name.lower().endswith("vbaproject.bin")
            for name in archive.namelist()
        )
        content_types = archive.read("[Content_Types].xml").lower()
        assert b"macroenabled" not in content_types
        assert b"vbaproject" not in content_types

        for name in archive.namelist():
            payload = archive.read(name)
            if name.endswith(".rels"):
                relationships = ElementTree.fromstring(payload)
                assert all(
                    relationship.attrib.get("TargetMode", "").lower()
                    != "external"
                    for relationship in relationships.findall(
                        f"{{{_RELATIONSHIP_NAMESPACE}}}Relationship"
                    )
                )
            if name.endswith(".xml"):
                assert b"<w:fldChar" not in payload
                assert b"<w:fldSimple" not in payload
                assert b"<w:instrText" not in payload

        core = archive.read("docProps/core.xml")
        assert b"2000-01-01T00:00:00Z" in core

    document = _document(raw)
    assert any(
        "https://invalid.example" in paragraph.text
        for paragraph in document.paragraphs
    )
    assert all(
        paragraph.alignment in {None, WD_ALIGN_PARAGRAPH.LEFT}
        for paragraph in document.sections[0].header.paragraphs
    )


def test_docx_large_bundle_is_deterministically_bounded_and_discloses_truncation():
    sections = _sections()
    sections[0]["summary_fields"][0]["field"]["source_refs"] = [
        _source("report_context", f"summary-ref-{index}")
        for index in range(40)
    ]
    sections[0]["summary_fields"][1]["field"]["note"] = "x" * 3_000

    source_table = sections[5]["tables"][0]
    tables = []
    for table_index in range(4):
        table = deepcopy(source_table)
        table["table_id"] = f"monthly_impact_{table_index}"
        template = source_table["rows"][0]
        table["rows"] = []
        for row_index in range(100):
            row = deepcopy(template)
            row["row_id"] = f"table-{table_index}-row-{row_index:03d}"
            row["cells"]["month"]["value"] = (
                f"table-{table_index}-row-{row_index:03d}"
            )
            table["rows"].append(row)
        tables.append(table)
    sections[5]["tables"] = tables
    bundle = _bundle(sections=sections)

    first = render_strategy_report_bundle_docx(bundle)
    second = render_strategy_report_bundle_docx(bundle)

    assert first == second
    assert len(first) < 1_000_000
    document = _document(first)
    fact_tables = [
        table
        for table in document.tables
        if table.rows[0].cells[0].text == "行 ID"
    ]
    assert len(fact_tables) == 4
    assert all(len(table.rows) - 1 <= 120 for table in fact_tables)
    assert sum(len(table.rows) - 1 for table in fact_tables) == 360
    assert sum(len(table.rows) for table in document.tables) < 800

    paragraphs = "\n".join(item.text for item in document.paragraphs)
    assert "已截断" in paragraphs
    assert "完整内容见同一报告修订的 JSON/XLSX 输出" in paragraphs
    assert "表 monthly_impact_0 事实行" in paragraphs
    assert "来源索引仅展示前 300 项" in paragraphs
    rows = _table_rows(document)
    assert any("引用已截断" in cell for row in rows for cell in row)
    assert any("内容已截断" in cell for row in rows for cell in row)


def test_docx_voting_rows_share_the_global_fact_budget():
    sections = _sections()
    voting_ref = _source(
        "voting_candidate_search",
        "voting-search-evidence",
        "d" * 64,
    )
    column_labels = {
        "search_id": "搜索ID",
        "combo_id": "组合ID",
        "member_ids": "成员规则ID",
        "n": "命中阈值 n",
        "eligible": "约束是否通过",
        "objective_metric": "目标指标",
        "objective_direction": "目标方向",
        "objective_value": "目标值",
        "constraint_failures": "约束未通过明细",
        "metrics": "完整指标",
    }
    voting_rows = []
    for index in range(20):
        combo_id = f"combo-{index:02d}"
        voting_rows.append(
            {
                "row_id": combo_id,
                "cells": {
                    "search_id": _present(
                        "voting-search-" + ("1" * 32),
                        source=voting_ref,
                    ),
                    "combo_id": _present(combo_id, source=voting_ref),
                    "member_ids": _present(
                        ["rule-a", "rule-b"],
                        source=voting_ref,
                    ),
                    "n": _present(1, source=voting_ref),
                    "eligible": _present(True, source=voting_ref),
                    "objective_metric": _present(
                        "bad_capture_rate",
                        source=voting_ref,
                    ),
                    "objective_direction": _present(
                        "maximize",
                        source=voting_ref,
                    ),
                    "objective_value": _present(
                        0.5 + index / 100,
                        source=voting_ref,
                    ),
                    "constraint_failures": _present([], source=voting_ref),
                    "metrics": _present(
                        {"hit_share": 0.1 + index / 100},
                        source=voting_ref,
                    ),
                },
            }
        )
    sections[4]["tables"] = [
        build_strategy_report_table(
            table_id="voting_candidate_search_combinations",
            title=(
                "Voting候选组合搜索结果"
                "（开发回测，仅供选择，未构建/未入池）"
            ),
            sheet_key="appendix_voting_search",
            granularity="aggregate",
            content_class="metric_summary",
            effect_stage="backtested",
            columns=[
                {
                    "key": key,
                    "label": label,
                    "unit": None,
                    "precision": None,
                }
                for key, label in column_labels.items()
            ],
            rows=voting_rows,
            source_refs=[voting_ref],
        )
    ]
    sections[4]["stage_evidence"] = [
        {
            "effect_stage": "backtested",
            "population": "risk",
            "partition": "development",
            "binding": {
                "kind": "development_backtest",
                "dataset_ref": _source(
                    "dataset",
                    "dataset-1",
                    "a" * 64,
                ),
                "frozen_artifact_ref": _source(
                    "strategy_candidate_pool",
                    "pool-1",
                    "b" * 64,
                ),
                "result_ref": voting_ref,
            },
        }
    ]
    sections[4]["source_refs"].extend(
        [
            _source(
                "dataset",
                "dataset-1",
                "a" * 64,
            ),
            voting_ref,
        ]
    )
    source_table = sections[5]["tables"][0]
    fact_tables = []
    for table_index in range(4):
        table = deepcopy(source_table)
        table["table_id"] = f"monthly_impact_{table_index}"
        template = source_table["rows"][0]
        table["rows"] = []
        for row_index in range(100):
            row = deepcopy(template)
            row["row_id"] = f"table-{table_index}-row-{row_index:03d}"
            table["rows"].append(row)
        fact_tables.append(table)
    sections[5]["tables"] = fact_tables

    document = _document(
        render_strategy_report_bundle_docx(_bundle(sections=sections))
    )
    voting_tables = [
        table
        for table in document.tables
        if table.rows[0].cells[0].text == "搜索ID"
    ]
    generic_fact_tables = [
        table
        for table in document.tables
        if table.rows[0].cells[0].text == "行 ID"
    ]

    assert len(voting_tables) == 1
    assert len(voting_tables[0].rows) - 1 == 20
    assert (
        sum(len(table.rows) - 1 for table in generic_fact_tables)
        + len(voting_tables[0].rows)
        - 1
        == 360
    )


def test_docx_field_code_rejection_is_namespace_aware_not_prefix_dependent():
    raw = render_strategy_report_bundle_docx(_bundle())
    with ZipFile(BytesIO(raw), "r") as archive:
        document_xml = archive.read("word/document.xml")
    aliased = document_xml.replace(
        b"xmlns:w=",
        b"xmlns:guard=",
    ).replace(
        b"w:",
        b"guard:",
    )
    field = (
        b'<guard:fldSimple guard:instr="HYPERLINK">'
        b"<guard:r><guard:t>unsafe</guard:t></guard:r>"
        b"</guard:fldSimple>"
    )
    aliased = aliased.replace(
        b"</guard:body>",
        field + b"</guard:body>",
        1,
    )
    assert b"<w:fldSimple" not in aliased
    assert b"<guard:fldSimple" in aliased

    malicious = _replace_document_xml(raw, aliased)
    with pytest.raises(
        StrategyReportOutputError,
        match="forbidden field code",
    ):
        _assert_safe_docx_package(malicious)
