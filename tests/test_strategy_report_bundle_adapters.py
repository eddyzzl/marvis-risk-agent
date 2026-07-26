from __future__ import annotations

from dataclasses import replace
import hashlib
from io import BytesIO
import json
from pathlib import Path
from types import SimpleNamespace
from zipfile import ZipFile

from docx import Document
import numpy as np
import pandas as pd
import pytest

from marvis.artifacts.model_score_vector import write_model_score_vector
from marvis.output.strategy_report_bundle import render_strategy_report_bundle
from marvis.packs.modeling.evidence import (
    canonical_modeling_training_evidence_json,
)
from marvis.packs.modeling.evidence_tools import (
    ModelingTrainingEvidenceArtifactBinding,
    build_training_evidence_ref,
)
from marvis.packs.modeling.score_evidence import (
    build_model_score_evidence_envelope,
    build_single_model_score_evidence,
    canonical_model_score_evidence_json,
)
from marvis.packs.modeling.score_evidence_tools import (
    ModelScoreEvidenceArtifactBinding,
)
from marvis.packs.strategy.candidate_fragment import (
    build_verified_candidate_fragment,
)
from marvis.packs.strategy.errors import StrategyError
from marvis.packs.strategy.candidate_stability import (
    CANDIDATE_STABILITY_PRODUCER_VERSION,
    build_candidate_stability_artifact,
    canonical_candidate_stability_artifact_json,
)
from marvis.packs.strategy.candidate_stability_tools import (
    ARTIFACT_SCHEMA_VERSION as CANDIDATE_STABILITY_ARTIFACT_SCHEMA_VERSION,
    StrategyCandidateStabilityArtifactBinding,
)
from marvis.packs.strategy.impact_cube import (
    build_strategy_impact_cube,
    canonical_strategy_impact_cube_json,
)
from marvis.packs.strategy.impact_cube_tools import (
    IMPACT_CUBE_ARTIFACT_SCHEMA_VERSION,
    IMPACT_CUBE_REQUIREMENTS_ARTIFACT_SCHEMA_VERSION,
    build_impact_cube_producer_run,
)
from marvis.packs.strategy.model_evidence import (
    build_model_comparison_evidence,
    build_model_comparison_metric,
    build_model_selection,
    build_strategy_model_evidence_bundle,
    canonical_strategy_model_evidence_bundle_json,
)
from marvis.packs.strategy.model_evidence_tools import (
    StrategyModelEvidenceV2ArtifactBinding,
)
from marvis.packs.strategy.pool import (
    add_verified_candidate_fragment,
    canonical_strategy_pool_json,
    compile_strategy_pool,
)
from marvis.packs.strategy.pool_impact import (
    build_strategy_pool_impact_assessment,
    canonical_strategy_pool_impact_json,
)
from marvis.packs.strategy.pool_impact_tools import (
    StrategyPoolImpactArtifactBinding,
)
from marvis.packs.strategy.pool_validation import (
    STRATEGY_POOL_VALIDATION_PRODUCER_VERSION,
    build_strategy_pool_validation_evidence,
    canonical_strategy_pool_validation_json,
)
from marvis.packs.strategy.pool_validation_tools import (
    POOL_VALIDATION_ARTIFACT_SCHEMA_VERSION,
    POOL_VALIDATION_REQUIREMENTS_ARTIFACT_SCHEMA_VERSION,
    StrategyPoolValidationArtifactBinding,
)
from marvis.packs.strategy.pool_requirement_resolver import (
    normalize_pool_requirements,
    validate_pool_requirement_bindings_provenance,
)
from marvis.packs.strategy import pool_tools as strategy_pool_tools
from marvis.packs.strategy.pool_tools import (
    StrategyCandidatePoolArtifactBinding,
    StrategyPoolDevelopmentDatasetBinding,
    StrategyPoolDevelopmentExecutionBinding,
    project_scorecard_report_evidence,
)
from marvis.packs.strategy.project_context import (
    build_context_field,
    build_current_project_snapshot,
    build_missing_information_record,
    build_report_field,
    build_source_ref,
    build_strategy_project_context_revision,
    build_strategy_project_context_state,
    canonical_strategy_project_context_revision_json,
)
from marvis.packs.strategy.project_context_tools import (
    StrategyProjectContextArtifactBinding,
)
from marvis.packs.strategy.report_bundle import (
    StrategyReportBundleError,
    build_strategy_report_bundle,
)
from marvis.packs.strategy.report_bundle_adapters import (
    StrategyImpactCubeArtifactBinding,
    build_strategy_report_bundle_source_inputs,
    validate_candidate_stability_report_compatibility,
)
import marvis.packs.strategy.report_bundle_adapters as report_adapters
from marvis.packs.strategy.sample_design_v2 import (
    build_strategy_sample_design_v2,
    build_strategy_sample_design_v2_bundle,
    canonical_strategy_sample_design_v2_bundle_json,
)
from marvis.packs.strategy.sample_design_v2_tools import (
    StrategySampleDesignV2ArtifactBinding,
)
from marvis.packs.strategy.sample_design_binding import (
    StrategySampleDesignExecutionBinding,
    StrategySampleDesignRef,
)
from marvis.packs.strategy.scorecard_candidate import (
    build_scorecard_cutoff_selection,
    canonical_scorecard_band_asset_json,
    canonical_scorecard_cutoff_selection_json,
    scorecard_cutoff_selection_to_verified_candidate_fragment,
)
from marvis.packs.strategy.scorecard_candidate_tools import (
    ScorecardBandAssetArtifactBinding,
    ScorecardCutoffSelectionArtifactBinding,
)
from marvis.packs.strategy.voting_candidate import (
    build_voting_candidate_asset,
)
from marvis.packs.strategy.voting_candidate_fragment import (
    VOTING_CANDIDATE_ARTIFACT_KIND,
    VOTING_CANDIDATE_ORIGIN_TOOL,
    voting_candidate_to_verified_fragment,
)
from marvis.packs.strategy import (
    voting_candidate_tools as strategy_voting_tools,
)
from marvis.packs.strategy.voting_candidate_tools import (
    VerifiedVotingCandidateArtifact,
    build_voting_candidate_artifact_document,
    canonical_voting_candidate_artifact_json,
)
from marvis.packs.strategy.voting_candidate_search import (
    VOTING_CANDIDATE_SEARCH_PRODUCER_VERSION,
    VOTING_CANDIDATE_SEARCH_REQUEST_SCHEMA_VERSION,
    canonical_voting_candidate_search_result_json,
    search_voting_candidate_combinations,
)
from marvis.packs.strategy.voting_candidate_search_tools import (
    VOTING_CANDIDATE_SEARCH_ARTIFACT_SCHEMA_VERSION,
    VotingCandidateSearchArtifactBinding,
)
from tests.test_modeling_evidence_contract import (
    _artifact as _model_artifact,
    _evidence as _training_evidence,
    _experiment as _model_experiment,
)
from tests.test_strategy_model_evidence import (
    _comparison as _model_comparison,
    _model as _single_model_evidence,
    _univariate as _univariate_evidence,
)
from tests.test_strategy_sample_design_v2 import (
    _components as _sample_components,
    _design_kwargs as _sample_design_kwargs,
    _diagnostic_statistics as _sample_diagnostic_statistics,
    _decoded_membership,
    _metric_observations as _sample_metric_observations,
    _policy as _sample_policy,
    _source_ref as _sample_source_ref,
)
from tests.test_strategy_scorecard_candidate import (
    _build as _build_scorecard_band_asset,
)
from tests.test_strategy_voting_candidate import (
    SAMPLE_DESIGN_REF as _PLAIN_VOTING_SAMPLE_REF,
    _pool as _plain_voting_parent_pool,
)


def _hash(label: str) -> str:
    return hashlib.sha256(label.encode("utf-8")).hexdigest()


def _file_hash(canonical: str) -> str:
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


def _source(
    label: str,
    *,
    kind: str = "tool_output",
) -> dict[str, str]:
    return build_source_ref(
        kind=kind,
        ref_id=label,
        content_hash=_hash(label),
    )


def _present(value, source: dict[str, str]) -> dict:
    return build_report_field(
        value=value,
        availability="present",
        origin="tool_output",
        source_refs=[source],
    )


def _unavailable(note: str = "No governed evidence.") -> dict:
    return build_report_field(
        value=None,
        availability="unavailable",
        origin="repository",
        source_refs=[],
        note=note,
    )


def _project_binding(
    tmp_path: Path,
    *,
    task_id: str,
    dataset_ref: dict[str, str],
) -> StrategyProjectContextArtifactBinding:
    source = _source("current-project-metrics", kind="metric_observation")
    snapshot = build_current_project_snapshot(
        task_id=task_id,
        as_of="2026-07-23",
        scope=_present("存量经营复借策略", source),
        dataset_refs=[dataset_ref],
        workspace_ref=None,
        champion_strategy_ref=None,
        status_fields={
            "volume": _present(8, source),
            "approval": _present(0, source),
            "risk": _unavailable("风险表现暂缺。"),
            "economics": _unavailable("收益口径暂缺。"),
        },
        metric_definition_refs=[],
        metric_observation_refs=[source],
        monthly_observation_refs=[],
        segment_observation_refs=[],
        maturity_summary=_present("confirmed_matured", source),
        user_context_fields=[
            build_context_field(
                field_path="sensitive.customer_id",
                field=_present("PII-CUSTOMER-0001", source),
            )
        ],
        red_flags=[],
        tool_run_refs=[_source("project-tool-run", kind="tool_run")],
    )
    answer = _source("missing-answer", kind="agent_message")
    missing = build_missing_information_record(
        task_id=task_id,
        field_path="historical_strategy_reviews",
        reason="No historical review is currently available.",
        blocking="report_optional",
        question="请提供历史版本策略评审材料。",
        status="unavailable",
        asked_count=1,
        asked_at="2026-07-23T08:00:00+00:00",
        answered_at="2026-07-23T08:05:00+00:00",
        answer_source_ref=answer,
        dependency_hash=_hash("history-missing-dependency"),
    )
    state = build_strategy_project_context_state(
        task_id=task_id,
        as_of="2026-07-23",
        current_project_snapshot=snapshot,
        historical_strategy_reviews=[],
        missing_information_records=[missing],
        red_flags=[],
    )
    revision = build_strategy_project_context_revision(
        state=state,
        revision=1,
        parent_revision_id=None,
        parent_state_hash=None,
        operation_kind="test_report_projection",
    )
    canonical = canonical_strategy_project_context_revision_json(revision)
    return StrategyProjectContextArtifactBinding(
        task_id=task_id,
        artifact_id=_hash("project-context-artifact"),
        artifact_path=tmp_path / "project-context.json",
        artifact_content_hash=_file_hash(canonical),
        provenance={},
        revision=revision,
        tasks_root=tmp_path,
        datasets_root=tmp_path,
        db_path=tmp_path / "marvis.sqlite",
    )


def _sample_binding(
    tmp_path: Path,
    *,
    maturity_status: str = "confirmed_matured",
    native_source: bool = False,
) -> StrategySampleDesignV2ArtifactBinding:
    membership = _decoded_membership()
    approval, risk, target, historical = _sample_components(
        membership,
        maturity_status=maturity_status,
    )
    design_kwargs = _sample_design_kwargs(maturity_status=maturity_status)
    if native_source:
        design_kwargs["legacy_development_ref"] = None
        design_kwargs["source_mode"] = "native_active_dataset"
    else:
        design_kwargs["legacy_development_ref"] = {
            "artifact_id": _hash("legacy-artifact-id"),
            "artifact_content_hash": _hash("legacy-artifact"),
            "sample_design_id": (
                "strategy-sample-design-" + _hash("legacy-design-id")[:24]
            ),
            "sample_design_content_hash": _hash("legacy-design"),
            "partition": "development",
        }
    source_refs = [
        _sample_source_ref("design-a"),
        _sample_source_ref("design-b"),
    ]
    design = build_strategy_sample_design_v2(
        task_id="task-v2",
        membership_header=membership["header"],
        relationship="nested_same_cohort",
        target_selector=target,
        approval_population=approval,
        risk_population=risk,
        historical_score=historical,
        policy=_sample_policy(),
        source_refs=source_refs,
        **design_kwargs,
    )
    bundle = build_strategy_sample_design_v2_bundle(
        task_id="task-v2",
        membership_header=membership["header"],
        membership_masks=membership["masks"],
        relationship="nested_same_cohort",
        target_selector=target,
        approval_population=approval,
        risk_population=risk,
        historical_score=historical,
        policy=_sample_policy(),
        diagnostic_statistics=_sample_diagnostic_statistics(membership),
        metric_observations=_sample_metric_observations(
            membership,
            design,
            maturity_status=maturity_status,
        ),
        source_refs=source_refs,
        **design_kwargs,
    )
    canonical = canonical_strategy_sample_design_v2_bundle_json(bundle)
    return StrategySampleDesignV2ArtifactBinding(
        task_id="task-v2",
        membership_artifact_id=_hash("membership-artifact-id"),
        membership_path=tmp_path / "membership.bin",
        membership_artifact_content_hash=_hash("membership-artifact-bytes"),
        bundle_artifact_id=_hash("bundle-artifact-id"),
        bundle_path=tmp_path / "sample-bundle.json",
        bundle_artifact_content_hash=_file_hash(canonical),
        provenance={},
        membership_provenance={},
        membership=_decoded_membership(),
        bundle=bundle,
        source_binding=object(),
    )


def _action(action_type: str) -> dict:
    values = {
        "approval": "approve",
        "reject": "reject",
        "review": "review",
    }
    return {
        "type": action_type,
        "value": values[action_type],
        "reason_code": None if action_type == "approval" else action_type.upper(),
        "stop": True,
    }


def _pool_and_impact_bindings(
    tmp_path: Path,
    sample: StrategySampleDesignV2ArtifactBinding,
    *,
    candidate_count: int = 1,
) -> tuple[
    StrategyCandidatePoolArtifactBinding,
    StrategyPoolImpactArtifactBinding,
]:
    design = sample.bundle["sample_design"]
    dataset = design["identity"]["dataset_ref"]
    workspace = design["identity"]["workspace_ref"]
    evidence_identity = {
        "dataset_id": dataset["dataset_id"],
        "dataset_content_hash": dataset["content_hash"],
        "workspace_revision": workspace["revision"],
        "workspace_generation": workspace["generation"],
        "semantic_mapping_hash": workspace["semantic_mapping_hash"],
        "sample_context_hash": _hash("sample-context"),
    }
    pool = None
    for ordinal in range(1, candidate_count + 1):
        fragment = build_verified_candidate_fragment(
            artifact={
                "artifact_id": f"candidate-artifact-{ordinal}",
                "artifact_kind": "test_candidate_json",
                "artifact_schema_version": "test.candidate-artifact.v1",
                "artifact_content_hash": _hash(
                    f"candidate-artifact-{ordinal}"
                ),
                "origin_tool": "strategy.test_candidate",
            },
            asset={
                "schema_version": "test.candidate.v1",
                "asset_id": f"candidate-asset-{ordinal}",
                "asset_hash": _hash(f"candidate-asset-{ordinal}"),
                "asset_type": "univariate_refinement",
            },
            fragment_type="strategy_rule",
            rule_id=f"rule-risk-{ordinal}",
            condition={
                "op": "compare",
                "field": "customer_id",
                "operator": "==",
                "value": f"PII-CUSTOMER-{ordinal:04d}",
                "missing": "no_match",
            },
            requirements=[],
            effect_id=f"candidate-effect-{ordinal}",
            evidence_id=f"candidate-evidence-{ordinal}",
            evidence_hash=_hash(f"candidate-evidence-{ordinal}"),
            evidence_identity=evidence_identity,
        )
        pool = add_verified_candidate_fragment(
            pool,
            task_id="task-v2",
            strategy_type="approval",
            default_action=_action("approval"),
            verified_candidate_fragment=fragment,
            action=_action("reject"),
        )
    assert pool is not None
    compiled = compile_strategy_pool(pool)
    pool_canonical = canonical_strategy_pool_json(pool)
    pool_binding = StrategyCandidatePoolArtifactBinding(
        task_id="task-v2",
        strategy_type="approval",
        pool=pool,
        compiled_design=compiled,
        artifact_id=_hash("candidate-pool-artifact"),
        artifact_path=tmp_path / "candidate-pool.json",
        artifact_content_hash=_file_hash(pool_canonical),
        artifact_origin_tool="strategy.add_candidate_to_pool",
        artifact_provenance={},
        artifact_provenance_json="{}",
        lineages=(),
        tasks_root=tmp_path,
        datasets_root=tmp_path,
        db_path=tmp_path / "marvis.sqlite",
    )
    frame = pd.DataFrame(
        {
            "customer_id": [
                "PII-CUSTOMER-0001",
                "PII-CUSTOMER-0002",
                "PII-CUSTOMER-0003",
                "PII-CUSTOMER-0004",
                "PII-CUSTOMER-0005",
                "PII-CUSTOMER-0006",
                "PII-CUSTOMER-0007",
                "PII-CUSTOMER-0008",
            ],
            "target": [1, 0, 1, 0, 1, 0, 1, 0],
            "apply_month": [
                "202601",
                "202601",
                "202601",
                "202601",
                "202602",
                "202602",
                "202602",
                "202602",
            ],
            "loan_amount": [100.0] * 8,
            "overdue_amount": [10.0, 0.0, 5.0, 0.0, 8.0, 0.0, 4.0, 0.0],
        }
    )
    sample_ref = design["compatibility"]["legacy_development_ref"]
    assessment = build_strategy_pool_impact_assessment(
        pool=pool,
        frame=frame,
        sample_binding={"task_id": "task-v2", **evidence_identity},
        sample_design_ref=sample_ref,
        target_col="target",
        target_bad_value=1,
        month_col="apply_month",
        loan_amount_col="loan_amount",
        overdue_amount_col="overdue_amount",
    )
    impact_binding = StrategyPoolImpactArtifactBinding(
        task_id="task-v2",
        artifact_id=_hash("pool-impact-artifact"),
        artifact_path=tmp_path / "pool-impact.json",
        artifact_content_hash=_file_hash(
            canonical_strategy_pool_impact_json(assessment)
        ),
        artifact_provenance={},
        artifact_provenance_json="{}",
        assessment=assessment,
        request={},
        pool=pool_binding,
        dataset=object(),
        sample_design=SimpleNamespace(to_ref_dict=lambda: sample_ref),
        baseline=None,
        stage="development_backtest",
        validation_status="unvalidated",
        tasks_root=tmp_path,
        db_path=tmp_path / "marvis.sqlite",
    )
    return pool_binding, impact_binding


def _scorecard_report_pool_binding(
    tmp_path: Path,
    sample: StrategySampleDesignV2ArtifactBinding,
    *,
    selection_ordinal: int = 0,
    selection_reason: str | None = "风险上限方案",
) -> tuple[
    StrategyCandidatePoolArtifactBinding,
    ScorecardCutoffSelectionArtifactBinding,
]:
    design = sample.bundle["sample_design"]
    dataset = design["identity"]["dataset_ref"]
    workspace = design["identity"]["workspace_ref"]
    identity = {
        "task_id": sample.task_id,
        "dataset_id": dataset["dataset_id"],
        "dataset_content_hash": dataset["content_hash"],
        "workspace_revision": workspace["revision"],
        "workspace_generation": workspace["generation"],
        "semantic_mapping_hash": workspace["semantic_mapping_hash"],
        "sample_context_hash": _hash("scorecard-report-sample-context"),
    }
    sample_ref = {
        "membership_artifact_id": sample.membership_artifact_id,
        "expected_membership_artifact_content_hash": (
            sample.membership_artifact_content_hash
        ),
        "bundle_artifact_id": sample.bundle_artifact_id,
        "expected_bundle_artifact_content_hash": (
            sample.bundle_artifact_content_hash
        ),
        "expected_bundle_id": sample.bundle["bundle_id"],
        "expected_sample_design_id": design["sample_design_id"],
        "expected_sample_design_content_hash": design["content_hash"],
    }
    asset = _build_scorecard_band_asset(
        identity=identity,
        sample_design_ref=sample_ref,
    )
    asset_canonical = canonical_scorecard_band_asset_json(asset).encode("utf-8")
    band_binding = ScorecardBandAssetArtifactBinding(
        task_id=sample.task_id,
        artifact_id=_hash("scorecard-report-band-artifact"),
        path=tmp_path / "scorecard-band.json",
        content_hash=hashlib.sha256(asset_canonical).hexdigest(),
        provenance={},
        canonical_bytes=asset_canonical,
        asset=asset,
        score_evidence=None,  # type: ignore[arg-type]
        sample_design=sample,
    )
    cutoff = asset["cutoffs"][selection_ordinal]
    selection = build_scorecard_cutoff_selection(
        asset,
        source_artifact_binding=band_binding.to_domain_binding(),
        cutoff_id=cutoff["cutoff_id"],
        selection_reason=selection_reason,
    )
    selection_canonical = canonical_scorecard_cutoff_selection_json(
        selection
    ).encode("utf-8")
    selection_binding = ScorecardCutoffSelectionArtifactBinding(
        task_id=sample.task_id,
        artifact_id=_hash(
            f"scorecard-report-selection-{selection_ordinal}-artifact"
        ),
        path=tmp_path / f"scorecard-selection-{selection_ordinal}.json",
        content_hash=hashlib.sha256(selection_canonical).hexdigest(),
        provenance={},
        canonical_bytes=selection_canonical,
        selection=selection,
        source_asset_binding=band_binding,
    )
    fragment = scorecard_cutoff_selection_to_verified_candidate_fragment(
        selection,
        asset,
        selection_artifact_binding=selection_binding.to_domain_binding(),
        source_artifact_binding=band_binding.to_domain_binding(),
    )
    pool = add_verified_candidate_fragment(
        None,
        task_id=sample.task_id,
        strategy_type="approval",
        default_action=_action("approval"),
        verified_candidate_fragment=fragment,
        action=_action("reject"),
    )
    compiled = compile_strategy_pool(pool)
    lineage = strategy_pool_tools._ScorecardCandidateLineage(
        selection=selection_binding,
        asset=band_binding,
        dataset=SimpleNamespace(),
        verified_fragment=fragment,
        source_binding=pool["entries"][0]["source"],
    )
    canonical = canonical_strategy_pool_json(pool)
    return (
        StrategyCandidatePoolArtifactBinding(
            task_id=sample.task_id,
            strategy_type="approval",
            pool=pool,
            compiled_design=compiled,
            artifact_id=_hash("scorecard-report-pool-artifact"),
            artifact_path=tmp_path / "scorecard-report-pool.json",
            artifact_content_hash=_file_hash(canonical),
            artifact_origin_tool="strategy.add_candidate_to_pool",
            artifact_provenance={},
            artifact_provenance_json="{}",
            lineages=(lineage,),
            tasks_root=tmp_path,
            datasets_root=tmp_path,
            db_path=tmp_path / "marvis.sqlite",
        ),
        selection_binding,
    )


def _scorecard_voting_report_pool_binding(
    tmp_path: Path,
    sample: StrategySampleDesignV2ArtifactBinding,
) -> StrategyCandidatePoolArtifactBinding:
    direct, first_selection = _scorecard_report_pool_binding(
        tmp_path,
        sample,
        selection_ordinal=0,
        selection_reason="低风险方案",
    )
    first_lineage = direct.lineages[0]
    band_binding = first_selection.source_asset_binding
    asset = band_binding.asset
    second_selection_value = build_scorecard_cutoff_selection(
        asset,
        source_artifact_binding=band_binding.to_domain_binding(),
        cutoff_id=asset["cutoffs"][1]["cutoff_id"],
        selection_reason="高风险方案",
    )
    second_canonical = canonical_scorecard_cutoff_selection_json(
        second_selection_value
    ).encode("utf-8")
    second_selection = ScorecardCutoffSelectionArtifactBinding(
        task_id=sample.task_id,
        artifact_id=_hash("scorecard-report-selection-1-artifact"),
        path=tmp_path / "scorecard-selection-1.json",
        content_hash=hashlib.sha256(second_canonical).hexdigest(),
        provenance={},
        canonical_bytes=second_canonical,
        selection=second_selection_value,
        source_asset_binding=band_binding,
    )
    second_fragment = scorecard_cutoff_selection_to_verified_candidate_fragment(
        second_selection_value,
        asset,
        selection_artifact_binding=second_selection.to_domain_binding(),
        source_artifact_binding=band_binding.to_domain_binding(),
    )
    parent_pool = add_verified_candidate_fragment(
        direct.pool,
        task_id=sample.task_id,
        strategy_type="approval",
        default_action=_action("approval"),
        verified_candidate_fragment=second_fragment,
        action=_action("reject"),
    )
    second_lineage = strategy_pool_tools._ScorecardCandidateLineage(
        selection=second_selection,
        asset=band_binding,
        dataset=SimpleNamespace(),
        verified_fragment=second_fragment,
        source_binding=parent_pool["entries"][1]["source"],
    )

    target = np.asarray([0, 0, 1, 1, 1, 0], dtype=np.int64)
    hit_count = np.asarray([0, 1, 2, 2, 1, 0], dtype=np.int64)
    voting_mask = hit_count >= 1
    effect = strategy_voting_tools._effect_from_mask(
        voting_mask,
        target=target,
        population_count=len(target),
    )
    voting_asset = build_voting_candidate_asset(
        parent_pool,
        selected_entry_ids=[
            entry["entry_id"] for entry in parent_pool["entries"]
        ],
        n=1,
        target_col="target",
        sample_design_ref=sample.bundle["sample_design"]["compatibility"][
            "legacy_development_ref"
        ],
        effect=effect,
    )
    document = build_voting_candidate_artifact_document(
        voting_asset,
        target_col="target",
        drop_nan_labels=False,
        nan_labels_dropped=0,
        population_count=len(target),
        labeled_count=len(target),
        hit_distribution=strategy_voting_tools._hit_distribution(
            hit_count,
            target=target,
            k=2,
        ),
        metric_observations=strategy_voting_tools._metric_observations(
            voting_mask,
            hit_count=hit_count,
            target=target,
            amount_values={
                "loan_amount": None,
                "overdue_amount": None,
            },
            k=2,
        ),
    )
    voting_canonical = canonical_voting_candidate_artifact_json(
        document
    ).encode("utf-8")
    parent_pool_artifact_id = _hash(
        "scorecard-report-voting-parent-pool-artifact"
    )
    parent_pool_artifact_hash = _file_hash(
        canonical_strategy_pool_json(parent_pool)
    )
    candidate = VerifiedVotingCandidateArtifact(
        artifact_id=_hash("scorecard-report-voting-artifact"),
        task_id=sample.task_id,
        kind=VOTING_CANDIDATE_ARTIFACT_KIND,
        path=tmp_path / "scorecard-voting.json",
        content_hash=hashlib.sha256(voting_canonical).hexdigest(),
        origin_tool=VOTING_CANDIDATE_ORIGIN_TOOL,
        provenance={
            "schema_version": document["schema_version"],
            "pool_artifact_id": parent_pool_artifact_id,
            "pool_artifact_content_hash": parent_pool_artifact_hash,
        },
        canonical_bytes=voting_canonical,
        document=document,
        asset=voting_asset,
    )
    voting_fragment = voting_candidate_to_verified_fragment(
        voting_asset,
        artifact_binding=candidate.artifact_binding(),
    )
    current_pool = add_verified_candidate_fragment(
        parent_pool,
        task_id=sample.task_id,
        strategy_type="approval",
        default_action=_action("approval"),
        verified_candidate_fragment=voting_fragment,
        action=_action("reject"),
        placement_mode="replace_selected_members",
        selected_entry_ids=[
            entry["entry_id"] for entry in parent_pool["entries"]
        ],
    )
    lineage = strategy_pool_tools._VotingCandidateLineage(
        candidate=candidate,
        parent_pool=parent_pool,
        parent_pool_artifact=SimpleNamespace(
            artifact_id=parent_pool_artifact_id,
            task_id=sample.task_id,
            kind=strategy_pool_tools.POOL_ARTIFACT_KIND,
            path=tmp_path / "scorecard-report-voting-parent-pool.json",
            content_hash=parent_pool_artifact_hash,
            origin_tool="strategy.add_candidate_to_pool",
            provenance={},
            provenance_json="{}",
        ),
        parent_lineages=(first_lineage, second_lineage),
        verified_fragment=voting_fragment,
        source_binding=current_pool["entries"][0]["source"],
    )
    canonical = canonical_strategy_pool_json(current_pool)
    return StrategyCandidatePoolArtifactBinding(
        task_id=sample.task_id,
        strategy_type="approval",
        pool=current_pool,
        compiled_design=compile_strategy_pool(current_pool),
        artifact_id=_hash("scorecard-report-voting-pool-artifact"),
        artifact_path=tmp_path / "scorecard-report-voting-pool.json",
        artifact_content_hash=_file_hash(canonical),
        artifact_origin_tool="strategy.add_candidate_to_pool",
        artifact_provenance={},
        artifact_provenance_json="{}",
        lineages=(lineage,),
        tasks_root=tmp_path,
        datasets_root=tmp_path,
        db_path=tmp_path / "marvis.sqlite",
    )


def _non_scorecard_voting_report_pool_binding(
    tmp_path: Path,
) -> StrategyCandidatePoolArtifactBinding:
    parent_pool = _plain_voting_parent_pool()
    selected_entry_ids = [
        entry["entry_id"] for entry in parent_pool["entries"]
    ]
    target = np.asarray([0, 0, 1, 1, 1, 0], dtype=np.int64)
    hit_count = np.asarray([0, 1, 2, 3, 2, 1], dtype=np.int64)
    voting_mask = hit_count >= 2
    effect = strategy_voting_tools._effect_from_mask(
        voting_mask,
        target=target,
        population_count=len(target),
    )
    voting_asset = build_voting_candidate_asset(
        parent_pool,
        selected_entry_ids=selected_entry_ids,
        n=2,
        target_col="bad",
        sample_design_ref=_PLAIN_VOTING_SAMPLE_REF,
        effect=effect,
    )
    document = build_voting_candidate_artifact_document(
        voting_asset,
        target_col="bad",
        drop_nan_labels=False,
        nan_labels_dropped=0,
        population_count=len(target),
        labeled_count=len(target),
        hit_distribution=strategy_voting_tools._hit_distribution(
            hit_count,
            target=target,
            k=3,
        ),
        metric_observations=strategy_voting_tools._metric_observations(
            voting_mask,
            hit_count=hit_count,
            target=target,
            amount_values={
                "loan_amount": None,
                "overdue_amount": None,
            },
            k=3,
        ),
    )
    voting_canonical = canonical_voting_candidate_artifact_json(
        document
    ).encode("utf-8")
    parent_pool_artifact_id = _hash(
        "non-scorecard-voting-parent-pool-artifact"
    )
    parent_pool_artifact_hash = _file_hash(
        canonical_strategy_pool_json(parent_pool)
    )
    candidate = VerifiedVotingCandidateArtifact(
        artifact_id=_hash("non-scorecard-voting-artifact"),
        task_id=parent_pool["task_id"],
        kind=VOTING_CANDIDATE_ARTIFACT_KIND,
        path=tmp_path / "non-scorecard-voting.json",
        content_hash=hashlib.sha256(voting_canonical).hexdigest(),
        origin_tool=VOTING_CANDIDATE_ORIGIN_TOOL,
        provenance={
            "schema_version": document["schema_version"],
            "pool_artifact_id": parent_pool_artifact_id,
            "pool_artifact_content_hash": parent_pool_artifact_hash,
        },
        canonical_bytes=voting_canonical,
        document=document,
        asset=voting_asset,
    )
    voting_fragment = voting_candidate_to_verified_fragment(
        voting_asset,
        artifact_binding=candidate.artifact_binding(),
    )
    current_pool = add_verified_candidate_fragment(
        parent_pool,
        task_id=parent_pool["task_id"],
        strategy_type=parent_pool["strategy_type"],
        default_action=_action("approval"),
        verified_candidate_fragment=voting_fragment,
        action=_action("reject"),
        placement_mode="replace_selected_members",
        selected_entry_ids=selected_entry_ids,
    )
    parent_lineages = tuple(
        strategy_pool_tools._UnivariateCandidateLineage(
            asset_record=None,
            asset={},
            parent_record=None,
            evidence={},
            dataset=SimpleNamespace(),
            verified_fragment={},
            source_binding=entry["source"],
        )
        for entry in parent_pool["entries"]
    )
    lineage = strategy_pool_tools._VotingCandidateLineage(
        candidate=candidate,
        parent_pool=parent_pool,
        parent_pool_artifact=SimpleNamespace(
            artifact_id=parent_pool_artifact_id,
            task_id=parent_pool["task_id"],
            kind=strategy_pool_tools.POOL_ARTIFACT_KIND,
            path=tmp_path / "non-scorecard-voting-parent-pool.json",
            content_hash=parent_pool_artifact_hash,
            origin_tool="strategy.add_candidate_to_pool",
            provenance={},
            provenance_json="{}",
        ),
        parent_lineages=parent_lineages,
        verified_fragment=voting_fragment,
        source_binding=current_pool["entries"][0]["source"],
    )
    canonical = canonical_strategy_pool_json(current_pool)
    return StrategyCandidatePoolArtifactBinding(
        task_id=parent_pool["task_id"],
        strategy_type=parent_pool["strategy_type"],
        pool=current_pool,
        compiled_design=compile_strategy_pool(current_pool),
        artifact_id=_hash("non-scorecard-voting-pool-artifact"),
        artifact_path=tmp_path / "non-scorecard-voting-pool.json",
        artifact_content_hash=_file_hash(canonical),
        artifact_origin_tool="strategy.add_candidate_to_pool",
        artifact_provenance={},
        artifact_provenance_json="{}",
        lineages=(lineage,),
        tasks_root=tmp_path,
        datasets_root=tmp_path,
        db_path=tmp_path / "marvis.sqlite",
    )


def _scorecard_reused_selection_voting_report_pool_binding(
    tmp_path: Path,
    sample: StrategySampleDesignV2ArtifactBinding,
) -> StrategyCandidatePoolArtifactBinding:
    one_of_two = _scorecard_voting_report_pool_binding(tmp_path, sample)
    one_of_two_lineage = one_of_two.lineages[0]
    assert isinstance(
        one_of_two_lineage,
        strategy_pool_tools._VotingCandidateLineage,
    )
    parent_pool = one_of_two_lineage.parent_pool
    target = np.asarray([0, 0, 1, 1, 1, 0], dtype=np.int64)
    hit_count = np.asarray([0, 1, 2, 2, 1, 0], dtype=np.int64)
    voting_mask = hit_count >= 2
    effect = strategy_voting_tools._effect_from_mask(
        voting_mask,
        target=target,
        population_count=len(target),
    )
    voting_asset = build_voting_candidate_asset(
        parent_pool,
        selected_entry_ids=[
            entry["entry_id"] for entry in parent_pool["entries"]
        ],
        n=2,
        target_col="target",
        sample_design_ref=sample.bundle["sample_design"]["compatibility"][
            "legacy_development_ref"
        ],
        effect=effect,
    )
    document = build_voting_candidate_artifact_document(
        voting_asset,
        target_col="target",
        drop_nan_labels=False,
        nan_labels_dropped=0,
        population_count=len(target),
        labeled_count=len(target),
        hit_distribution=strategy_voting_tools._hit_distribution(
            hit_count,
            target=target,
            k=2,
        ),
        metric_observations=strategy_voting_tools._metric_observations(
            voting_mask,
            hit_count=hit_count,
            target=target,
            amount_values={
                "loan_amount": None,
                "overdue_amount": None,
            },
            k=2,
        ),
    )
    canonical = canonical_voting_candidate_artifact_json(document).encode(
        "utf-8"
    )
    candidate = VerifiedVotingCandidateArtifact(
        artifact_id=_hash("scorecard-report-voting-n2-artifact"),
        task_id=sample.task_id,
        kind=VOTING_CANDIDATE_ARTIFACT_KIND,
        path=tmp_path / "scorecard-voting-n2.json",
        content_hash=hashlib.sha256(canonical).hexdigest(),
        origin_tool=VOTING_CANDIDATE_ORIGIN_TOOL,
        provenance={
            "schema_version": document["schema_version"],
            "pool_artifact_id": (
                one_of_two_lineage.parent_pool_artifact.artifact_id
            ),
            "pool_artifact_content_hash": (
                one_of_two_lineage.parent_pool_artifact.content_hash
            ),
        },
        canonical_bytes=canonical,
        document=document,
        asset=voting_asset,
    )
    two_of_two_fragment = voting_candidate_to_verified_fragment(
        voting_asset,
        artifact_binding=candidate.artifact_binding(),
    )
    selected_entry_ids = [
        entry["entry_id"] for entry in parent_pool["entries"]
    ]
    current_pool = add_verified_candidate_fragment(
        parent_pool,
        task_id=sample.task_id,
        strategy_type="approval",
        default_action=_action("approval"),
        verified_candidate_fragment=two_of_two_fragment,
        action=_action("reject"),
        placement_mode="before_selected_members",
        selected_entry_ids=selected_entry_ids,
    )
    current_pool = add_verified_candidate_fragment(
        current_pool,
        task_id=sample.task_id,
        strategy_type="approval",
        default_action=_action("approval"),
        verified_candidate_fragment=one_of_two_lineage.verified_fragment,
        action=_action("reject"),
        placement_mode="replace_selected_members",
        selected_entry_ids=selected_entry_ids,
    )
    two_of_two_lineage = strategy_pool_tools._VotingCandidateLineage(
        candidate=candidate,
        parent_pool=parent_pool,
        parent_pool_artifact=one_of_two_lineage.parent_pool_artifact,
        parent_lineages=one_of_two_lineage.parent_lineages,
        verified_fragment=two_of_two_fragment,
        source_binding=current_pool["entries"][0]["source"],
    )
    one_of_two_lineage = replace(
        one_of_two_lineage,
        source_binding=current_pool["entries"][1]["source"],
    )
    pool_canonical = canonical_strategy_pool_json(current_pool)
    return StrategyCandidatePoolArtifactBinding(
        task_id=sample.task_id,
        strategy_type="approval",
        pool=current_pool,
        compiled_design=compile_strategy_pool(current_pool),
        artifact_id=_hash("scorecard-report-reused-selection-pool-artifact"),
        artifact_path=tmp_path / "scorecard-reused-selection-pool.json",
        artifact_content_hash=_file_hash(pool_canonical),
        artifact_origin_tool="strategy.add_candidate_to_pool",
        artifact_provenance={},
        artifact_provenance_json="{}",
        lineages=(two_of_two_lineage, one_of_two_lineage),
        tasks_root=tmp_path,
        datasets_root=tmp_path,
        db_path=tmp_path / "marvis.sqlite",
    )


def _candidate_stability_binding(
    tmp_path: Path,
    sample: StrategySampleDesignV2ArtifactBinding,
    pool: StrategyCandidatePoolArtifactBinding,
    *,
    source_kind: str = "pool_entry",
    low_sample_month: bool = False,
) -> StrategyCandidateStabilityArtifactBinding:
    entry = pool.pool["entries"][0]
    entry_source = entry["source"]
    month_counts = (20, 40) if low_sample_month else (30, 30)
    months = [
        month
        for month, count in zip(
            ("202601", "202602"),
            month_counts,
            strict=True,
        )
        for _ in range(count)
    ]
    frame = pd.DataFrame(
        {
            "apply_month": months,
            "target": [index % 2 for index in range(len(months))],
        }
    )
    identity = {
        "task_id": sample.task_id,
        **entry_source["evidence_identity"],
    }
    if source_kind == "pool_entry":
        basis = "pool_entry_incremental_first_match"
        source_ref = {
            "source_kind": "pool_entry",
            "artifact_id": pool.artifact_id,
            "artifact_content_hash": pool.artifact_content_hash,
            "pool_id": pool.pool["pool_id"],
            "revision": pool.pool["revision"],
            "revision_id": pool.pool["revision_id"],
            "snapshot_hash": pool.pool["snapshot_hash"],
            "entry_id": entry["entry_id"],
            "rule_id": entry["rule_id"],
        }
    elif source_kind == "univariate_asset":
        basis = "asset_rule_hit"
        source_ref = {
            "source_kind": "univariate_asset",
            "artifact_id": entry_source["artifact_id"],
            "artifact_content_hash": entry_source["artifact_content_hash"],
            "asset_id": entry_source["asset_id"],
            "asset_hash": entry_source["asset_hash"],
            "rule_id": entry["rule_id"],
        }
    else:  # pragma: no cover - fixture contract
        raise AssertionError(f"unsupported source kind: {source_kind}")
    stability = build_candidate_stability_artifact(
        frame=frame,
        month_col="apply_month",
        target_col="target",
        hit_mask=np.asarray(
            [index % 3 == 0 for index in range(len(frame))],
            dtype=bool,
        ),
        basis=basis,
        identity=identity,
        source_ref=source_ref,
        sample_design_ref=sample.bundle["sample_design"]["compatibility"][
            "legacy_development_ref"
        ],
    )
    canonical = canonical_candidate_stability_artifact_json(stability)
    provenance = {
        "schema_version": CANDIDATE_STABILITY_ARTIFACT_SCHEMA_VERSION,
        "producer_version": CANDIDATE_STABILITY_PRODUCER_VERSION,
        "task_id": sample.task_id,
        "stability_id": stability["stability_id"],
        "stability_content_hash": stability["content_hash"],
        "basis": stability["basis"],
        "source_kind": source_ref["source_kind"],
        "source_artifact_id": source_ref["artifact_id"],
        "source_artifact_content_hash": source_ref["artifact_content_hash"],
        "source_id": (
            source_ref["asset_id"]
            if source_kind == "univariate_asset"
            else source_ref["pool_id"]
        ),
        "source_hash": (
            source_ref["asset_hash"]
            if source_kind == "univariate_asset"
            else source_ref["snapshot_hash"]
        ),
        "rule_id": source_ref["rule_id"],
        "entry_id": source_ref.get("entry_id"),
        "pool_id": source_ref.get("pool_id"),
        "pool_revision": source_ref.get("revision"),
        "pool_revision_id": source_ref.get("revision_id"),
        "dataset_id": identity["dataset_id"],
        "dataset_content_hash": identity["dataset_content_hash"],
        "workspace_revision": identity["workspace_revision"],
        "workspace_generation": identity["workspace_generation"],
        "semantic_mapping_hash": identity["semantic_mapping_hash"],
        "target_col": "target",
        "month_col": "apply_month",
        "sample_design_ref": stability["sample_design_ref"],
        "sample_context_hash": identity["sample_context_hash"],
        "sample_partition": "development",
    }
    provenance_json = json.dumps(
        provenance,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        allow_nan=False,
    )
    return StrategyCandidateStabilityArtifactBinding(
        task_id=sample.task_id,
        artifact_id=_hash(
            "candidate-stability-artifact"
            f"-{source_kind}-{low_sample_month}"
        ),
        artifact_path=tmp_path / f"candidate-stability-{source_kind}.json",
        artifact_content_hash=_file_hash(canonical),
        artifact_provenance=provenance,
        artifact_provenance_json=provenance_json,
        stability=stability,
        tasks_root=tmp_path,
        db_path=tmp_path / "marvis.sqlite",
    )


def _bindings(
    tmp_path: Path,
    *,
    maturity_status: str = "confirmed_matured",
    candidate_count: int = 1,
) -> tuple[
    StrategyProjectContextArtifactBinding,
    StrategySampleDesignV2ArtifactBinding,
    StrategyCandidatePoolArtifactBinding,
    StrategyPoolImpactArtifactBinding,
]:
    sample = _sample_binding(tmp_path, maturity_status=maturity_status)
    dataset = sample.bundle["sample_design"]["identity"]["dataset_ref"]
    project = _project_binding(
        tmp_path,
        task_id=sample.task_id,
        dataset_ref=build_source_ref(
            kind="dataset",
            ref_id=dataset["dataset_id"],
            content_hash=dataset["content_hash"],
        ),
    )
    pool, impact = _pool_and_impact_bindings(
        tmp_path,
        sample,
        candidate_count=candidate_count,
    )
    return project, sample, pool, impact


def _voting_search_binding(
    tmp_path: Path,
    sample: StrategySampleDesignV2ArtifactBinding,
    pool: StrategyCandidatePoolArtifactBinding,
) -> VotingCandidateSearchArtifactBinding:
    design = sample.bundle["sample_design"]
    dataset = design["identity"]["dataset_ref"]
    workspace = design["identity"]["workspace_ref"]
    semantics = design["sample_semantics"]
    target = design["target_selector"]
    legacy_ref = design["compatibility"]["legacy_development_ref"]
    risk_population = next(
        item for item in sample.bundle["populations"] if item["role"] == "risk"
    )
    development_count = next(
        item["row_count"]
        for item in risk_population["partitions"]
        if item["name"] == "development"
    )
    candidate_ids = sorted(
        entry["rule_id"]
        for entry in pool.pool["entries"]
        if entry["enabled"]
        and entry["source"]["asset_type"] != "voting_candidate"
    )
    assert len(candidate_ids) == 7
    request = {
        "schema_version": VOTING_CANDIDATE_SEARCH_REQUEST_SCHEMA_VERSION,
        "candidate_ids": candidate_ids,
        "hit_matrix": [
            [
                bool((candidate_index + row_index) % 2)
                for row_index in range(development_count)
            ]
            for candidate_index in range(len(candidate_ids))
        ],
        "target": [row_index % 2 for row_index in range(development_count)],
        "weights": None,
        "amounts": [
            float((row_index + 1) * 100)
            for row_index in range(development_count)
        ],
        "member_count": 3,
        "n": 2,
        "objective": {
            "metric": "bad_capture_rate",
            "direction": "maximize",
        },
        "constraints": [
            {
                "metric": "hit_share",
                "operator": "lte",
                "value": 0.0,
            }
        ],
        "include": [],
        "exclude": [],
        "max_combinations": 20,
    }
    result = search_voting_candidate_combinations(request)
    assert result["evaluated"] == 20
    assert all(not item["eligible"] for item in result["combinations"])

    reference = StrategySampleDesignRef.from_value(legacy_ref)
    legacy_sample = StrategySampleDesignExecutionBinding(
        reference=reference,
        artifact=SimpleNamespace(),
        task_id=sample.task_id,
        dataset_id=dataset["dataset_id"],
        dataset_content_hash=dataset["content_hash"],
        workspace_revision=workspace["revision"],
        workspace_generation=workspace["generation"],
        semantic_mapping_hash=workspace["semantic_mapping_hash"],
        target_col=target["column"],
        target_bad_value=target["bad_value"],
        drop_nan_labels=target["drop_missing"],
        split_column=semantics["split_definition"]["column"],
        development_values=tuple(
            semantics["split_definition"]["development_values"]
        ),
        development_population_count=development_count,
        active_population_count=risk_population["total_count"],
        month_col=semantics["field_bindings"]["month_field"],
        weight_col=semantics["field_bindings"]["weight_field"],
        loan_amount_col=semantics["field_bindings"]["loan_amount_field"],
        overdue_amount_col=semantics["field_bindings"][
            "overdue_amount_field"
        ],
    )
    development_dataset = StrategyPoolDevelopmentDatasetBinding(
        task_id=sample.task_id,
        dataset_id=dataset["dataset_id"],
        source_path="datasets/dataset-v2.parquet",
        path=tmp_path / "dataset-v2.parquet",
        content_hash=dataset["content_hash"],
        registry_metadata_hash=_hash("dataset-v2-registry-metadata"),
        columns=(
            "customer_id",
            "target",
            "apply_month",
            "loan_amount",
            "overdue_amount",
        ),
        row_count=8,
    )
    evidence_identity = {
        "dataset_id": dataset["dataset_id"],
        "dataset_content_hash": dataset["content_hash"],
        "workspace_revision": workspace["revision"],
        "workspace_generation": workspace["generation"],
        "semantic_mapping_hash": workspace["semantic_mapping_hash"],
        "sample_context_hash": _hash("sample-context"),
    }
    development = StrategyPoolDevelopmentExecutionBinding(
        task_id=sample.task_id,
        pool=pool,
        dataset=development_dataset,
        sample_design=legacy_sample,
        sample_design_v2=sample,
        evidence_identity=evidence_identity,
        target_col=target["column"],
        month_col=semantics["field_bindings"]["month_field"],
    )
    provenance = {
        "schema_version": VOTING_CANDIDATE_SEARCH_ARTIFACT_SCHEMA_VERSION,
        "producer_version": VOTING_CANDIDATE_SEARCH_PRODUCER_VERSION,
        "task_id": sample.task_id,
        "search_id": result["search_id"],
        "search_content_hash": result["content_hash"],
        "request_hash": result["request_hash"],
        "pool_ref": {
            "artifact_id": pool.artifact_id,
            "artifact_content_hash": pool.artifact_content_hash,
            "pool_id": pool.pool["pool_id"],
            "strategy_type": pool.pool["strategy_type"],
            "revision": pool.pool["revision"],
            "revision_id": pool.pool["revision_id"],
            "snapshot_hash": pool.pool["snapshot_hash"],
        },
        "dataset_binding": {
            "task_id": sample.task_id,
            "dataset_id": development_dataset.dataset_id,
            "dataset_source_path": development_dataset.source_path,
            "dataset_content_hash": development_dataset.content_hash,
            "dataset_registry_metadata_hash": (
                development_dataset.registry_metadata_hash
            ),
            "workspace_revision": legacy_sample.workspace_revision,
            "workspace_generation": legacy_sample.workspace_generation,
            "semantic_mapping_hash": legacy_sample.semantic_mapping_hash,
        },
        "sample_design_ref": legacy_sample.to_ref_dict(),
        "sample_context_hash": evidence_identity["sample_context_hash"],
        "target_binding": {
            "column": legacy_sample.target_col,
            "raw_bad_value": legacy_sample.target_bad_value,
            "normalized_bad_value": 1,
            "drop_nan_labels": legacy_sample.drop_nan_labels,
            "nan_labels_dropped": 0,
            "labeled_count": development_count,
            "sample_partition": "development",
        },
        "observation_bindings": {
            "weight_col": legacy_sample.weight_col,
            "amount_col": legacy_sample.loan_amount_col,
        },
        "requirement_bindings": None,
        "excluded_unsupported_rule_ids": [],
        "lifecycle": {
            "mutated_pool": False,
            "selected": False,
            "admitted": False,
            "applied": False,
            "adopted": False,
            "deployed": False,
        },
    }
    provenance_json = json.dumps(
        provenance,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        allow_nan=False,
    )
    canonical = canonical_voting_candidate_search_result_json(result)
    return VotingCandidateSearchArtifactBinding(
        task_id=sample.task_id,
        artifact_id=_hash("voting-search-artifact"),
        artifact_path=tmp_path / "voting-search.json",
        artifact_content_hash=_file_hash(canonical),
        artifact_provenance=provenance,
        artifact_provenance_json=provenance_json,
        result=result,
        pool_development=development,
        resolved_requirements=None,
        tasks_root=tmp_path,
        db_path=tmp_path / "marvis.sqlite",
    )


def _zip_xml_text(raw: bytes) -> str:
    with ZipFile(BytesIO(raw), "r") as archive:
        return "".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.endswith(".xml")
        )


def _docx_visible_text(raw: bytes) -> str:
    document = Document(BytesIO(raw))
    return "\n".join(
        [
            *(paragraph.text for paragraph in document.paragraphs),
            *(
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ),
        ]
    )


def _pool_binding_for_strategy_type(
    tmp_path: Path,
    sample: StrategySampleDesignV2ArtifactBinding,
    strategy_type: str,
    *,
    score_requirement: dict | None = None,
) -> StrategyCandidatePoolArtifactBinding:
    design = sample.bundle["sample_design"]
    dataset = design["identity"]["dataset_ref"]
    workspace = design["identity"]["workspace_ref"]
    evidence_identity = {
        "dataset_id": dataset["dataset_id"],
        "dataset_content_hash": dataset["content_hash"],
        "workspace_revision": workspace["revision"],
        "workspace_generation": workspace["generation"],
        "semantic_mapping_hash": workspace["semantic_mapping_hash"],
        "sample_context_hash": _hash("sample-context"),
    }
    fragment = build_verified_candidate_fragment(
        artifact={
            "artifact_id": f"{strategy_type}-candidate-artifact",
            "artifact_kind": "test_candidate_json",
            "artifact_schema_version": "test.candidate-artifact.v1",
            "artifact_content_hash": _hash(
                f"{strategy_type}-candidate-artifact"
            ),
            "origin_tool": "strategy.test_candidate",
        },
        asset={
            "schema_version": "test.candidate.v1",
            "asset_id": f"{strategy_type}-candidate-asset",
            "asset_hash": _hash(f"{strategy_type}-candidate-asset"),
            "asset_type": "univariate_refinement",
        },
        fragment_type="strategy_rule",
        rule_id=f"rule-{strategy_type}",
        condition=(
            {
                "op": "compare",
                "field": "customer_id",
                "operator": "==",
                "value": "PII-CUSTOMER-0001",
                "missing": "no_match",
            }
            if score_requirement is None
            else {
                "op": "compare",
                "field": score_requirement["virtual_field"],
                "operator": ">=",
                "value": 0.5,
                "missing": "no_match",
            }
        ),
        requirements=([] if score_requirement is None else [score_requirement]),
        effect_id=f"{strategy_type}-effect",
        evidence_id=f"{strategy_type}-evidence",
        evidence_hash=_hash(f"{strategy_type}-evidence"),
        evidence_identity=evidence_identity,
    )
    actions = {
        "approval": (_action("approval"), _action("reject")),
        "reject": (_action("approval"), _action("reject")),
        "limit": (
            {
                "type": "limit",
                "value": 1_000.0,
                "reason_code": None,
                "stop": True,
            },
            {
                "type": "limit",
                "value": 2_000.0,
                "reason_code": None,
                "stop": True,
            },
        ),
        "pricing": (
            {
                "type": "pricing",
                "value": 0.10,
                "reason_code": None,
                "stop": True,
            },
            {
                "type": "pricing",
                "value": 0.20,
                "reason_code": None,
                "stop": True,
            },
        ),
        "segmentation": (
            {
                "type": "segment",
                "value": "A",
                "reason_code": None,
                "stop": True,
            },
            {
                "type": "segment",
                "value": "B",
                "reason_code": None,
                "stop": True,
            },
        ),
    }
    default_action, action = actions[strategy_type]
    pool = add_verified_candidate_fragment(
        None,
        task_id=sample.task_id,
        strategy_type=strategy_type,
        default_action=default_action,
        verified_candidate_fragment=fragment,
        action=action,
    )
    compiled = compile_strategy_pool(pool)
    canonical = canonical_strategy_pool_json(pool)
    return StrategyCandidatePoolArtifactBinding(
        task_id=sample.task_id,
        strategy_type=strategy_type,
        pool=pool,
        compiled_design=compiled,
        artifact_id=_hash(f"{strategy_type}-pool-artifact"),
        artifact_path=tmp_path / f"{strategy_type}-pool.json",
        artifact_content_hash=_file_hash(canonical),
        artifact_origin_tool="strategy.add_candidate_to_pool",
        artifact_provenance={},
        artifact_provenance_json="{}",
        lineages=(),
        tasks_root=tmp_path,
        datasets_root=tmp_path,
        db_path=tmp_path / "marvis.sqlite",
    )


def _sample_with_dataset_source(
    sample: StrategySampleDesignV2ArtifactBinding,
    tmp_path: Path,
) -> StrategySampleDesignV2ArtifactBinding:
    identity = sample.bundle["sample_design"]["identity"]
    dataset = identity["dataset_ref"]
    workspace = identity["workspace_ref"]
    return replace(
        sample,
        source_binding=SimpleNamespace(
            task_id=sample.task_id,
            dataset_id=dataset["dataset_id"],
            dataset_content_hash=dataset["content_hash"],
            dataset_source_path=str(tmp_path / "dataset.parquet"),
            dataset_registry_metadata_hash=_hash("dataset-registry-metadata"),
            workspace_revision=workspace["revision"],
            workspace_generation=workspace["generation"],
            semantic_mapping_hash=workspace["semantic_mapping_hash"],
        ),
    )


def _pool_validation_binding(
    tmp_path: Path,
    sample: StrategySampleDesignV2ArtifactBinding,
    pool: StrategyCandidatePoolArtifactBinding,
    *,
    partition: str,
) -> StrategyPoolValidationArtifactBinding:
    design = sample.bundle["sample_design"]
    header = sample.membership["header"]
    target = design["target_selector"]
    fields = design["sample_semantics"]["field_bindings"]
    count = header["counts"]["risk"][partition]
    frame = pd.DataFrame(
        {
            "customer_id": [
                f"PII-CUSTOMER-{index + 1:04d}" for index in range(count)
            ],
            target["column"]: [
                target["bad_value"] if index % 2 == 0 else target["good_value"]
                for index in range(count)
            ],
            fields["month_field"]: ["202601"] * count,
            fields["loan_amount_field"]: [100.0] * count,
            fields["overdue_amount_field"]: [10.0] * count,
        }
    )
    normalized_requirements = list(
        normalize_pool_requirements(pool.compiled_design["requirements"])
    )
    for outer in normalized_requirements:
        frame[outer["requirement"]["virtual_field"]] = [0.8] * count
    source = sample.source_binding
    dataset_binding = {
        "task_id": source.task_id,
        "dataset_id": source.dataset_id,
        "dataset_content_hash": source.dataset_content_hash,
        "dataset_source_path": source.dataset_source_path,
        "dataset_registry_metadata_hash": (
            source.dataset_registry_metadata_hash
        ),
        "workspace_revision": source.workspace_revision,
        "workspace_generation": source.workspace_generation,
        "semantic_mapping_hash": source.semantic_mapping_hash,
    }
    sample_ref = {
        "membership_artifact_id": sample.membership_artifact_id,
        "membership_artifact_content_hash": (
            sample.membership_artifact_content_hash
        ),
        "membership_id": header["membership_id"],
        "membership_content_hash": header["content_hash"],
        "bundle_artifact_id": sample.bundle_artifact_id,
        "bundle_artifact_content_hash": (
            sample.bundle_artifact_content_hash
        ),
        "bundle_id": sample.bundle["bundle_id"],
        "bundle_content_hash": sample.bundle["content_hash"],
        "sample_design_id": design["sample_design_id"],
        "sample_design_content_hash": design["content_hash"],
        "partition_key": f"risk/{partition}",
        "partition_count": count,
        "analysis_universe_row_count": header["row_count"],
    }
    evidence = build_strategy_pool_validation_evidence(
        pool=pool.pool,
        frame=frame,
        pool_artifact_ref={
            "artifact_id": pool.artifact_id,
            "artifact_content_hash": pool.artifact_content_hash,
        },
        sample_design_v2_ref=sample_ref,
        dataset_binding=dataset_binding,
        legacy_development_ref=design["compatibility"][
            "legacy_development_ref"
        ],
        partition=partition,
        population="risk",
        comparison_mode="absolute",
        target_col=target["column"],
        target_bad_value=target["bad_value"],
        month_col=fields["month_field"],
        loan_amount_col=fields["loan_amount_field"],
        overdue_amount_col=fields["overdue_amount_field"],
        development_rows_excluded=True,
    )
    requirement_bindings = (
        None
        if not normalized_requirements
        else validate_pool_requirement_bindings_provenance(
            {
                "requirements_hash": hashlib.sha256(
                    json.dumps(
                        normalized_requirements,
                        ensure_ascii=False,
                        sort_keys=True,
                        separators=(",", ":"),
                    ).encode("utf-8")
                ).hexdigest(),
                "requirements": normalized_requirements,
                "virtual_fields": [
                    outer["requirement"]["virtual_field"]
                    for outer in normalized_requirements
                ],
            }
        )
    )
    provenance = {
        "schema_version": (
            POOL_VALIDATION_ARTIFACT_SCHEMA_VERSION
            if requirement_bindings is None
            else POOL_VALIDATION_REQUIREMENTS_ARTIFACT_SCHEMA_VERSION
        ),
        "producer_version": STRATEGY_POOL_VALIDATION_PRODUCER_VERSION,
        "task_id": sample.task_id,
        "evidence_id": evidence["evidence_id"],
        "evidence_content_hash": evidence["content_hash"],
        "pool_ref": {
            "artifact_id": pool.artifact_id,
            "expected_artifact_content_hash": pool.artifact_content_hash,
            "expected_pool_id": pool.pool["pool_id"],
            "expected_revision": pool.pool["revision"],
            "expected_revision_id": pool.pool["revision_id"],
            "expected_snapshot_hash": pool.pool["snapshot_hash"],
            "pool_id": pool.pool["pool_id"],
            "revision_id": pool.pool["revision_id"],
        },
        "sample_design_ref": {
            "membership_artifact_id": sample.membership_artifact_id,
            "expected_membership_artifact_content_hash": (
                sample.membership_artifact_content_hash
            ),
            "bundle_artifact_id": sample.bundle_artifact_id,
            "expected_bundle_artifact_content_hash": (
                sample.bundle_artifact_content_hash
            ),
            "expected_bundle_id": sample.bundle["bundle_id"],
            "expected_sample_design_id": design["sample_design_id"],
            "expected_sample_design_content_hash": design["content_hash"],
        },
        "dataset_binding": dataset_binding,
        "target_binding": evidence["source_bindings"]["target"],
        "field_bindings": {
            **evidence["source_bindings"]["fields"],
            **(
                {}
                if requirement_bindings is None
                else {"requirements": requirement_bindings}
            ),
        },
        "partition": partition,
        "population": "risk",
        "comparison_mode": "absolute",
        "lifecycle_stage": partition,
        "validation_status": "independent_evidence",
    }
    provenance_json = json.dumps(
        provenance,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    )
    return StrategyPoolValidationArtifactBinding(
        task_id=sample.task_id,
        artifact_id=_hash(f"pool-validation-{partition}-artifact"),
        artifact_path=tmp_path / f"pool-validation-{partition}.json",
        artifact_content_hash=_file_hash(
            canonical_strategy_pool_validation_json(evidence)
        ),
        artifact_provenance=provenance,
        artifact_provenance_json=provenance_json,
        evidence=evidence,
        tasks_root=tmp_path,
        db_path=tmp_path / "marvis.sqlite",
    )


def _impact_cube_binding(
    tmp_path: Path,
    sample: StrategySampleDesignV2ArtifactBinding,
    pool: StrategyCandidatePoolArtifactBinding,
    *,
    economics: bool = True,
) -> StrategyImpactCubeArtifactBinding:
    design = sample.bundle["sample_design"]
    header = sample.membership["header"]
    counts = header["counts"]
    partitions = ("development", "validation", "oot")
    frame = pd.DataFrame(
        {
            "customer_id": [
                "PII-CUSTOMER-0001",
                "PII-CUSTOMER-0002",
                "PII-CUSTOMER-0003",
            ],
            "target": [1, 0, 1],
            "apply_month": ["202601", "202601", "202602"],
            "channel": ["web", "web", "web"],
            "segment": ["repeat", "repeat", "repeat"],
            "loan_amount": [100.0, 200.0, None],
            "overdue_amount": [10.0, 0.0, 20.0],
            "ead": [1_000.0, 2_000.0, 1_500.0],
            "pd": [0.20, 0.01, 0.30],
            "utilization": [0.50, 0.40, 0.60],
        }
    )
    for outer in pool.compiled_design["requirements"]:
        virtual_field = outer["requirement"]["virtual_field"]
        frame[virtual_field] = [0.8, 0.2, 0.9]
    approval_frames = {
        partition: frame.iloc[: counts["approval"][partition]].reset_index(
            drop=True
        )
        for partition in partitions
    }
    risk_frames = {
        partition: frame.iloc[: counts["risk"][partition]].reset_index(
            drop=True
        )
        for partition in partitions
    }
    sample_ref = {
        "membership_artifact_id": sample.membership_artifact_id,
        "membership_artifact_content_hash": (
            sample.membership_artifact_content_hash
        ),
        "membership_id": header["membership_id"],
        "membership_content_hash": header["content_hash"],
        "bundle_artifact_id": sample.bundle_artifact_id,
        "bundle_artifact_content_hash": sample.bundle_artifact_content_hash,
        "bundle_id": sample.bundle["bundle_id"],
        "bundle_content_hash": sample.bundle["content_hash"],
        "sample_design_id": design["sample_design_id"],
        "sample_design_content_hash": design["content_hash"],
        "analysis_universe_row_count": counts["analysis_universe"],
        "partition_counts": {
            partition: counts["risk"][partition]
            for partition in partitions
        },
        "population_partition_counts": {
            role: {
                partition: counts[role][partition]
                for partition in partitions
            }
            for role in ("approval", "risk")
        },
    }
    dataset = design["identity"]["dataset_ref"]
    workspace = design["identity"]["workspace_ref"]
    dataset_binding = {
        "task_id": sample.task_id,
        "dataset_id": dataset["dataset_id"],
        "dataset_content_hash": dataset["content_hash"],
        "dataset_source_path": f"{sample.task_id}/sample.parquet",
        "dataset_registry_metadata_hash": _hash("dataset-registry"),
        "workspace_revision": workspace["revision"],
        "workspace_generation": workspace["generation"],
        "semantic_mapping_hash": workspace["semantic_mapping_hash"],
    }
    economics_by_type = {
        "approval": {
            "ead": {"kind": "column", "column": "ead"},
            "pd": {"kind": "column", "column": "pd"},
            "annual_rate": {"kind": "scalar", "value": 0.20},
            "funding_rate": {"kind": "scalar", "value": 0.05},
            "lgd": {"kind": "scalar", "value": 0.50},
            "operating_cost_per_loan": {
                "kind": "scalar",
                "value": 10.0,
            },
            "term_months": {"kind": "scalar", "value": 12},
        },
        "reject": {
            "ead": {"kind": "column", "column": "ead"},
            "pd": {"kind": "column", "column": "pd"},
            "annual_rate": {"kind": "scalar", "value": 0.20},
            "funding_rate": {"kind": "scalar", "value": 0.05},
            "lgd": {"kind": "scalar", "value": 0.50},
            "operating_cost_per_loan": {
                "kind": "scalar",
                "value": 10.0,
            },
            "term_months": {"kind": "scalar", "value": 12},
        },
        "limit": {
            "pd": {"kind": "column", "column": "pd"},
            "lgd": {"kind": "scalar", "value": 0.50},
            "utilization": {
                "kind": "column",
                "column": "utilization",
            },
        },
        "pricing": {
            "ead": {"kind": "column", "column": "ead"},
            "pd": {"kind": "column", "column": "pd"},
            "lgd": {"kind": "scalar", "value": 0.50},
            "funding_rate": {"kind": "scalar", "value": 0.05},
            "term_months": {"kind": "scalar", "value": 12},
            "operating_cost_per_loan": {
                "kind": "scalar",
                "value": 10.0,
            },
        },
        "segmentation": None,
    }
    economics_inputs = (
        economics_by_type[pool.strategy_type] if economics else None
    )
    cube = build_strategy_impact_cube(
        pool=pool.pool,
        approval_partition_frames=approval_frames,
        partition_frames=risk_frames,
        pool_artifact_ref={
            "artifact_id": pool.artifact_id,
            "artifact_content_hash": pool.artifact_content_hash,
        },
        sample_design_v2_ref=sample_ref,
        dataset_binding=dataset_binding,
        legacy_development_ref=design["compatibility"][
            "legacy_development_ref"
        ],
        target_col="target",
        target_bad_value=1,
        month_col="apply_month",
        group_col="channel",
        segment_col="segment",
        current_strategy_spec=None,
        current_strategy_ref=None,
        economics_bindings=economics_inputs,
        loan_amount_col="loan_amount",
        overdue_amount_col="overdue_amount",
    )
    canonical = canonical_strategy_impact_cube_json(cube)
    artifact_hash = _file_hash(canonical)
    artifact_id = _hash("impact-cube-artifact")
    request_sample_ref = {
        "membership_artifact_id": sample.membership_artifact_id,
        "expected_membership_artifact_content_hash": (
            sample.membership_artifact_content_hash
        ),
        "bundle_artifact_id": sample.bundle_artifact_id,
        "expected_bundle_artifact_content_hash": (
            sample.bundle_artifact_content_hash
        ),
        "expected_bundle_id": sample.bundle["bundle_id"],
        "expected_sample_design_id": design["sample_design_id"],
        "expected_sample_design_content_hash": design["content_hash"],
    }
    pool_ref = {
        "artifact_id": pool.artifact_id,
        "expected_artifact_content_hash": pool.artifact_content_hash,
        "expected_pool_id": pool.pool["pool_id"],
        "expected_revision": pool.pool["revision"],
        "expected_revision_id": pool.pool["revision_id"],
        "expected_snapshot_hash": pool.pool["snapshot_hash"],
    }
    dimensions = {
        "month_col": "apply_month",
        "group_col": "channel",
        "segment_col": "segment",
    }
    producer_run = build_impact_cube_producer_run(
        task_id=sample.task_id,
        request={
            "strategy_type": pool.strategy_type,
            "pool_ref": pool_ref,
            "sample_design_ref": request_sample_ref,
            "partitions": list(partitions),
            "population": "risk",
            "dimension_bindings": dimensions,
            "current_strategy_ref": None,
            "economics_inputs": economics_inputs,
        },
        cube_id=cube["cube_id"],
        cube_content_hash=cube["content_hash"],
        artifact_id=artifact_id,
        artifact_filename=f"{cube['cube_id']}.json",
        artifact_content_hash=artifact_hash,
    )
    requirements = pool.compiled_design["requirements"]
    provenance = {
        "schema_version": (
            IMPACT_CUBE_REQUIREMENTS_ARTIFACT_SCHEMA_VERSION
            if requirements
            else IMPACT_CUBE_ARTIFACT_SCHEMA_VERSION
        ),
        "producer_version": cube["producer_version"],
        "task_id": sample.task_id,
        "cube_id": cube["cube_id"],
        "cube_content_hash": cube["content_hash"],
        "pool_ref": pool_ref,
        "sample_design_ref": request_sample_ref,
        "dataset_binding": dataset_binding,
        "target_binding": {
            "column": "target",
            "good_value": 0,
            "bad_value": 1,
            "missing_policy": (
                "retain_population_exclude_risk_denominator"
            ),
        },
        "dimension_bindings": dimensions,
        "current_strategy_ref": None,
        "economics_inputs": economics_inputs,
        "partitions": list(partitions),
        "populations": ["approval", "risk"],
        "lifecycle": dict(cube["lifecycle"]),
        "producer_run": producer_run,
    }
    if requirements:
        provenance["requirement_bindings"] = {
            "requirements_hash": hashlib.sha256(
                json.dumps(
                    requirements,
                    ensure_ascii=False,
                    sort_keys=True,
                    separators=(",", ":"),
                ).encode("utf-8")
            ).hexdigest(),
            "requirements": requirements,
            "virtual_fields": list(
                dict.fromkeys(
                    item["requirement"]["virtual_field"]
                    for item in requirements
                )
            ),
        }
    return StrategyImpactCubeArtifactBinding(
        task_id=sample.task_id,
        artifact_id=artifact_id,
        artifact_path=(
            tmp_path
            / sample.task_id
            / "strategy_impact_cubes"
            / f"{cube['cube_id']}.json"
        ),
        artifact_content_hash=artifact_hash,
        artifact_provenance=provenance,
        artifact_provenance_json=json.dumps(
            provenance,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
            allow_nan=False,
        ),
        cube=cube,
        tasks_root=tmp_path,
        db_path=tmp_path / "marvis.sqlite",
    )


def _training_binding(
    tmp_path: Path,
    sample: StrategySampleDesignV2ArtifactBinding,
) -> ModelingTrainingEvidenceArtifactBinding:
    experiment = _model_experiment()
    artifact = _model_artifact()
    evidence = _training_evidence(
        bundle=sample.bundle,
        experiment=experiment,
        artifact=artifact,
    )
    canonical = canonical_modeling_training_evidence_json(
        evidence,
        sample_design_bundle=sample.bundle,
    )
    model_binary = evidence["model_artifact"]["model_binary_ref"]
    return ModelingTrainingEvidenceArtifactBinding(
        task_id=sample.task_id,
        sample=sample,
        experiment=experiment,
        model_artifact=artifact,
        model_binary_record={
            "id": model_binary["artifact_id"],
            "content_hash": model_binary["content_hash"],
        },
        evidence_record={
            "id": _hash("training-evidence-record"),
            "content_hash": _file_hash(canonical),
        },
        evidence=evidence,
        model_binary_path=tmp_path / "model.bin",
        evidence_path=tmp_path / "training-evidence.json",
    )


def _model_binding(
    tmp_path: Path,
    sample: StrategySampleDesignV2ArtifactBinding,
    *,
    selection_status: str = "selected",
    comparison_metric_status: str = "present",
) -> StrategyModelEvidenceV2ArtifactBinding:
    univariate = _univariate_evidence(sample.bundle)
    first_model = _single_model_evidence(
        sample.bundle,
        "first",
        validation_ks=0.38,
    )
    second_model = _single_model_evidence(
        sample.bundle,
        "second",
        auc=0.71,
        validation_ks=0.35,
    )
    comparison = _model_comparison(sample.bundle, first_model, second_model)
    if comparison_metric_status != "present":
        evaluation = comparison["evaluation_sample_ref"]
        source_metric = comparison["metrics"][0]
        unavailable_metric = build_model_comparison_metric(
            sample_design_bundle=sample.bundle,
            population=evaluation["population"],
            partition=evaluation["partition"],
            metric_key=source_metric["metric_key"],
            status=comparison_metric_status,
            unit=source_metric["unit"],
            source_ref=source_metric["source_ref"],
            model_values=None,
            delta=None,
            period=source_metric["period"],
            reason="验证样本尚未形成可用比较指标。",
        )
        comparison = build_model_comparison_evidence(
            sample_design_bundle=sample.bundle,
            population=evaluation["population"],
            partition=evaluation["partition"],
            comparison_ref=comparison["comparison_ref"],
            model_evidence_refs=comparison["model_evidence_refs"],
            metrics=[unavailable_metric],
            selection=build_model_selection(
                status="no_selection",
                reason="比较指标不可用，未选择模型。",
            ),
        )
    if selection_status == "no_selection":
        evaluation = comparison["evaluation_sample_ref"]
        comparison = build_model_comparison_evidence(
            sample_design_bundle=sample.bundle,
            population=evaluation["population"],
            partition=evaluation["partition"],
            comparison_ref=comparison["comparison_ref"],
            model_evidence_refs=comparison["model_evidence_refs"],
            metrics=comparison["metrics"],
            selection=build_model_selection(
                status="no_selection",
                reason="业务尚未确认最终模型。",
            ),
        )
    bundle = build_strategy_model_evidence_bundle(
        sample_design_bundle=sample.bundle,
        univariate_evidence=[univariate],
        model_evidence=[first_model, second_model],
        comparison_evidence=[comparison],
    )
    canonical = canonical_strategy_model_evidence_bundle_json(
        bundle,
        sample_design_bundle=sample.bundle,
    )
    return StrategyModelEvidenceV2ArtifactBinding(
        task_id=sample.task_id,
        artifact_id=_hash("model-evidence-artifact"),
        path=tmp_path / "model-evidence.json",
        artifact_content_hash=_file_hash(canonical),
        provenance={},
        bundle=bundle,
        sample_design_binding=sample,
        sources=(),
        warnings=(),
        tasks_root=tmp_path,
        db_path=tmp_path / "marvis.sqlite",
    )


def _score_binding(
    tmp_path: Path,
    training: ModelingTrainingEvidenceArtifactBinding,
) -> ModelScoreEvidenceArtifactBinding:
    frame = pd.DataFrame(
        {
            "target": [0, 1, 0, 1, 0, 1, 1, 0],
            "apply_month": [
                "202601",
                "202601",
                "202601",
                "202602",
                "202603",
                "202602",
                "202603",
                "202604",
            ],
            "income": [10.0, 20.0, 30.0, 40.0, 50.0, 60.0, 70.0, 80.0],
            "age": [21, 22, 23, 24, 25, 26, 27, 28],
        }
    )
    scores = np.asarray([0.1, 0.9, 0.2, 0.8, 0.3, 0.7, 0.6, 0.4])
    vector = write_model_score_vector(tmp_path / "scores.parquet", scores)
    training_ref = build_training_evidence_ref(training)
    model_binary = training.evidence["model_artifact"]["model_binary_ref"]
    model_ref = {
        "kind": model_binary["kind"],
        "ref_id": model_binary["artifact_id"],
        "content_hash": model_binary["content_hash"],
    }
    vector_record_id = _hash("score-vector-record")
    score_ref = {
        "kind": "model_score_vector_parquet",
        "ref_id": vector_record_id,
        "content_hash": vector.content_hash,
    }
    single = build_single_model_score_evidence(
        sample_design_bundle=training.sample.bundle,
        membership_masks=_decoded_membership()["masks"],
        frame=frame,
        scores=scores,
        training_evidence_ref={
            "kind": "modeling_training_evidence_json",
            "ref_id": training_ref["evidence_artifact_id"],
            "content_hash": training_ref[
                "expected_evidence_artifact_content_hash"
            ],
        },
        model_ref=model_ref,
        score_ref=score_ref,
        features=training.evidence["training_contract"]["features"],
    )
    envelope = build_model_score_evidence_envelope(
        task_id=training.task_id,
        training_evidence_ref=training_ref,
        training_evidence=training.evidence,
        sample_design_bundle=training.sample.bundle,
        model_ref=model_ref,
        score_ref=score_ref,
        score_vector=vector,
        single_model_evidence=single,
    )
    canonical = canonical_model_score_evidence_json(
        envelope,
        sample_design_bundle=training.sample.bundle,
        training_evidence=training.evidence,
        expected_training_evidence_ref=training_ref,
        score_vector=vector,
    )
    return ModelScoreEvidenceArtifactBinding(
        task_id=training.task_id,
        training=training,
        vector_record={
            "id": vector_record_id,
            "content_hash": vector.content_hash,
        },
        evidence_record={
            "id": _hash("score-evidence-record"),
            "content_hash": _file_hash(canonical),
        },
        vector=vector,
        envelope=envelope,
        vector_path=vector.path,
        evidence_path=tmp_path / "score-evidence.json",
    )


def _project(
    bindings: tuple[
        StrategyProjectContextArtifactBinding,
        StrategySampleDesignV2ArtifactBinding,
        StrategyCandidatePoolArtifactBinding,
        StrategyPoolImpactArtifactBinding,
    ],
    **optional,
) -> dict:
    project, sample, pool, impact = bindings
    return build_strategy_report_bundle_source_inputs(
        project_context=project,
        sample_design=sample,
        candidate_pool=pool,
        pool_impact=impact,
        **optional,
    )


def test_adapter_is_deterministic_bundle_ready_and_uses_exact_source_identities(
    tmp_path: Path,
) -> None:
    bindings = _bindings(tmp_path)
    project, sample, pool, impact = bindings

    first = _project(bindings)
    second = _project(bindings)

    assert first == second
    assert [item["key"] for item in first["sections"]] == [
        "current_project",
        "historical_versions",
        "sample_design",
        "univariate_and_models",
        "candidate_combinations",
        "impact_assessment",
        "final_document",
    ]
    assert first["strategy_artifact_refs"] == [
        {
            "kind": "pool_impact",
            "ref_id": impact.artifact_id,
            "content_hash": impact.artifact_content_hash,
        },
        {
            "kind": "strategy_candidate_pool",
            "ref_id": pool.artifact_id,
            "content_hash": pool.artifact_content_hash,
        },
    ]
    bundle = build_strategy_report_bundle(
        task_id=project.task_id,
        report_revision=1,
        strategy_id=None,
        strategy_version=None,
        strategy_type=None,
        title=_present("V2策略开发评审", first["strategy_artifact_refs"][1]),
        status="partial",
        generated_at="2026-07-23T16:00:00+08:00",
        **first,
    )
    assert bundle["effect_stages"] == ["backtested"]
    assert bundle["missing_information"] == project.revision["state"][
        "missing_information_records"
    ]
    sample_refs = first["sections"][2]["source_refs"]
    assert sample_refs == [
        {
            "kind": "sample_design",
            "ref_id": sample.bundle_artifact_id,
            "content_hash": sample.bundle_artifact_content_hash,
        }
    ]


def test_full_report_native_sample_fails_closed_at_legacy_evidence_boundary(
    tmp_path: Path,
) -> None:
    project, _legacy_sample, pool, impact = _bindings(tmp_path)
    native_sample = _sample_binding(tmp_path, native_source=True)

    with pytest.raises(StrategyError) as raised:
        build_strategy_report_bundle_source_inputs(
            project_context=project,
            sample_design=native_sample,
            candidate_pool=pool,
            pool_impact=impact,
        )

    assert (
        getattr(raised.value, "code", None)
        == "strategy_sample_design_v2_native_source_unsupported"
    )
    assert getattr(raised.value, "consumer", None) == "strategy_report_bundle"


def test_adapter_requires_typed_authenticated_pool_validation_bindings(
    tmp_path: Path,
) -> None:
    bindings = _bindings(tmp_path)

    with pytest.raises(
        StrategyReportBundleError,
        match="authenticated StrategyPoolValidationArtifactBinding",
    ):
        _project(
            bindings,
            pool_validations=(
                SimpleNamespace(task_id=bindings[0].task_id),
            ),
        )


def test_adapter_projects_independent_pool_validation_as_oot_stage_and_conclusion(
    tmp_path: Path,
) -> None:
    project, original_sample, pool, impact = _bindings(tmp_path)
    sample = _sample_with_dataset_source(original_sample, tmp_path)
    validations = tuple(
        _pool_validation_binding(
            tmp_path,
            sample,
            pool,
            partition=partition,
        )
        for partition in ("validation", "oot")
    )

    result = build_strategy_report_bundle_source_inputs(
        project_context=project,
        sample_design=sample,
        candidate_pool=pool,
        pool_validations=validations,
        pool_impact=impact,
    )

    impact_section = result["sections"][5]
    validation_stages = [
        item
        for item in impact_section["stage_evidence"]
        if item["binding"]["result_ref"]["kind"] == "strategy_validation"
    ]
    assert [
        (item["effect_stage"], item["population"], item["partition"])
        for item in validation_stages
    ] == [
        ("oot_validated", "risk", "validation"),
        ("oot_validated", "risk", "oot"),
    ]
    assert all(
        table["effect_stage"] == "oot_validated"
        for table in impact_section["tables"]
        if table["table_id"].startswith(
            "strategy_pool_independent_validation"
        )
    )
    final_fields = {
        item["field_id"]: item["field"]["value"]
        for item in result["sections"][6]["summary_fields"]
    }
    assert final_fields["evidence_stages"] == [
        "backtested",
        "oot_validated",
    ]
    assert final_fields["validation_statuses"] == ["independent_evidence"]
    assert (
        final_fields["validation_conclusion"]
        == "independent_replay_evidence_only"
    )
    bundle = build_strategy_report_bundle(
        task_id=project.task_id,
        report_revision=1,
        strategy_id=None,
        strategy_version=None,
        strategy_type="approval",
        title=_present(
            "独立策略验证",
            result["strategy_artifact_refs"][0],
        ),
        status="partial",
        generated_at="2026-07-26T12:00:00+08:00",
        **result,
    )
    assert bundle["effect_stages"] == ["backtested", "oot_validated"]
    assert bundle["strategy_id"] is None


def test_adapter_validation_only_uses_umbrella_stage_without_claiming_oot_partition(
    tmp_path: Path,
) -> None:
    project, original_sample, pool, impact = _bindings(tmp_path)
    sample = _sample_with_dataset_source(original_sample, tmp_path)
    validation = _pool_validation_binding(
        tmp_path,
        sample,
        pool,
        partition="validation",
    )

    result = build_strategy_report_bundle_source_inputs(
        project_context=project,
        sample_design=sample,
        candidate_pool=pool,
        pool_validations=(validation,),
        pool_impact=impact,
    )

    impact_section = result["sections"][5]
    validation_stages = [
        item
        for item in impact_section["stage_evidence"]
        if item["binding"]["result_ref"]["kind"] == "strategy_validation"
    ]
    assert [
        (item["effect_stage"], item["population"], item["partition"])
        for item in validation_stages
    ] == [("oot_validated", "risk", "validation")]
    assert '"partition":"oot"' not in json.dumps(
        validation_stages,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
    )
    validation_tables = [
        table
        for table in impact_section["tables"]
        if table["table_id"].startswith(
            "strategy_pool_independent_validation"
        )
    ]
    assert {
        row["cells"]["partition"]["value"]
        for table in validation_tables
        for row in table["rows"]
    } == {"validation"}
    final_fields = {
        item["field_id"]: item["field"]["value"]
        for item in result["sections"][6]["summary_fields"]
    }
    assert final_fields["evidence_stages"] == [
        "backtested",
        "oot_validated",
    ]
    assert final_fields["independent_replay_partitions"] == ["validation"]
    validation_refs = [
        ref
        for ref in result["strategy_artifact_refs"]
        if ref["kind"] == "strategy_validation"
    ]
    assert len(validation_refs) == 1
    assert validation_refs[0]["ref_id"] == validation.artifact_id


def test_pool_validation_amount_projection_keeps_unavailable_values_blank() -> None:
    source_ref = _source("pool-validation", kind="strategy_validation")
    amounts = {
        "loan_amount": {
            "status": "unavailable",
            "column": None,
            "coverage_count": None,
            "coverage_rate": None,
            "sum": None,
        }
    }

    field = report_adapters._pool_validation_amount_field(
        amounts,
        amount_key="loan_amount",
        field="sum",
        source_ref=source_ref,
    )

    assert field["availability"] == "unavailable"
    assert field["value"] is None
    assert field["source_refs"] == []


def test_adapter_binds_requirement_validation_to_normalized_current_pool(
    tmp_path: Path,
) -> None:
    original_sample = _sample_binding(tmp_path)
    sample = _sample_with_dataset_source(original_sample, tmp_path)
    dataset = sample.bundle["sample_design"]["identity"]["dataset_ref"]
    project = _project_binding(
        tmp_path,
        task_id=sample.task_id,
        dataset_ref=build_source_ref(
            kind="dataset",
            ref_id=dataset["dataset_id"],
            content_hash=dataset["content_hash"],
        ),
    )
    vector_id = "1" * 64
    requirement = {
        "type": "model_score_vector.v1",
        "virtual_field": "__marvis_model_pd_" + vector_id[:16],
        "score_product": "raw_native_uncalibrated_bad_probability",
        "score_evidence_artifact_id": "2" * 64,
        "score_evidence_artifact_content_hash": "3" * 64,
        "score_vector_artifact_id": vector_id,
        "score_vector_artifact_content_hash": "4" * 64,
    }
    pool = _pool_binding_for_strategy_type(
        tmp_path,
        sample,
        "approval",
        score_requirement=requirement,
    )
    impact_cube = _impact_cube_binding(tmp_path, sample, pool)
    validation = _pool_validation_binding(
        tmp_path,
        sample,
        pool,
        partition="validation",
    )

    result = build_strategy_report_bundle_source_inputs(
        project_context=project,
        sample_design=sample,
        candidate_pool=pool,
        pool_validations=(validation,),
        impact_cube=impact_cube,
    )
    assert any(
        ref["kind"] == "strategy_validation"
        for ref in result["strategy_artifact_refs"]
    )

    compiled_outer = pool.compiled_design["requirements"][0]
    voting_style = {
        "rule_id": compiled_outer["rule_id"],
        "fragment_id": compiled_outer["fragment_id"],
        "requirement": {
            "entry_id": "voting-entry-1",
            "rule_id": compiled_outer["rule_id"],
            "fragment_id": compiled_outer["fragment_id"],
            "requirement": compiled_outer["requirement"],
        },
    }
    report_adapters._require_pool_validation_provenance_requirements(
        validation,
        compiled_design={
            **pool.compiled_design,
            "requirements": [voting_style],
        },
    )

    forged_provenance = json.loads(validation.artifact_provenance_json)
    forged_requirements = forged_provenance["field_bindings"][
        "requirements"
    ]["requirements"]
    forged_requirements[0]["requirement"][
        "score_evidence_artifact_id"
    ] = "f" * 64
    forged_requirements[0]["requirement"][
        "score_vector_artifact_content_hash"
    ] = "e" * 64
    forged_provenance["field_bindings"]["requirements"][
        "requirements_hash"
    ] = hashlib.sha256(
        json.dumps(
            forged_requirements,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        ).encode("utf-8")
    ).hexdigest()
    forged = replace(
        validation,
        artifact_provenance=forged_provenance,
        artifact_provenance_json=json.dumps(
            forged_provenance,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
        ),
    )
    with pytest.raises(
        StrategyReportBundleError,
        match="requirements differ from the Candidate Pool",
    ):
        build_strategy_report_bundle_source_inputs(
            project_context=project,
            sample_design=sample,
            candidate_pool=pool,
            pool_validations=(forged,),
            impact_cube=impact_cube,
        )


def test_adapter_projects_voting_search_only_as_twenty_row_development_evidence(
    tmp_path: Path,
) -> None:
    bindings = _bindings(tmp_path, candidate_count=7)
    project, sample, pool, _impact = bindings
    search = _voting_search_binding(tmp_path, sample, pool)

    without_search = _project(bindings)
    projected = _project(bindings, voting_candidate_search=search)

    assert len(projected["sections"]) == 7
    candidate_without = without_search["sections"][4]
    candidate_with = projected["sections"][4]
    base_tables = {
        table["table_id"]: table for table in candidate_without["tables"]
    }
    projected_tables = {
        table["table_id"]: table for table in candidate_with["tables"]
    }
    assert (
        projected_tables["candidate_pool_entries"]
        == base_tables["candidate_pool_entries"]
    )
    assert (
        projected_tables["compiled_candidate_design"]
        == base_tables["compiled_candidate_design"]
    )
    assert projected["sections"][5:] == without_search["sections"][5:]
    assert all(
        ref["kind"] != "voting_candidate_search"
        for ref in without_search["strategy_artifact_refs"]
    )

    table = projected_tables["voting_candidate_search_combinations"]
    assert table["title"] == (
        "Voting候选组合搜索结果（开发回测，仅供选择，未构建/未入池）"
    )
    assert table["sheet_key"] == "appendix_voting_search"
    expected_combinations = search.result["combinations"][:20]
    assert [
        row["cells"]["combo_id"]["value"] for row in table["rows"]
    ] == [item["combo_id"] for item in expected_combinations]
    assert [
        row["cells"]["eligible"]["value"] for row in table["rows"]
    ] == [item["eligible"] for item in expected_combinations]
    assert "selected" not in {
        column["key"] for column in table["columns"]
    }
    first = table["rows"][0]["cells"]
    assert {
        "search_id",
        "combo_id",
        "member_ids",
        "n",
        "eligible",
        "objective_metric",
        "objective_direction",
        "objective_value",
        "constraint_failures",
        "metrics",
    } <= set(first)

    summary = {
        item["field_id"]: item["field"]["value"]
        for item in candidate_with["summary_fields"]
    }
    assert summary["voting_search_search_space"] == search.result["search_space"]
    assert summary["voting_search_evaluated"] == search.result["evaluated"]
    assert summary["voting_search_truncated"] is True
    assert summary["voting_search_eligible"] == search.result["eligible"]
    assert summary["voting_search_displayed"] == 20
    assert any(
        item["binding"]["result_ref"]["kind"]
        == "voting_candidate_search"
        for item in candidate_with["stage_evidence"]
    )

    bundle = build_strategy_report_bundle(
        task_id=project.task_id,
        report_revision=1,
        strategy_id=None,
        strategy_version=None,
        strategy_type="approval",
        title=_present(
            "Voting搜索开发回测",
            projected["strategy_artifact_refs"][0],
        ),
        status="partial",
        generated_at="2026-07-25T12:00:00+08:00",
        **projected,
    )
    rendered = render_strategy_report_bundle(bundle)
    markdown = rendered["markdown"].decode("utf-8")
    docx_text = _docx_visible_text(rendered["docx"])
    xlsx_text = _zip_xml_text(rendered["xlsx"])
    json_text = rendered["json"].decode("utf-8")
    for combination in expected_combinations:
        combo_id = combination["combo_id"]
        assert combo_id in json_text
        assert combo_id in markdown
        assert combo_id in xlsx_text
        assert combo_id in docx_text
    assert table["title"] in markdown
    assert table["title"] in docx_text
    for forbidden in ("winner", "champion", "selected", "冠军", "最佳"):
        assert forbidden not in markdown.lower()
        assert forbidden not in docx_text.lower()


def test_adapter_rejects_voting_search_sample_context_drift(
    tmp_path: Path,
) -> None:
    bindings = _bindings(tmp_path, candidate_count=7)
    _project_binding_value, sample, pool, _impact = bindings
    search = _voting_search_binding(tmp_path, sample, pool)
    development = replace(
        search.pool_development,
        evidence_identity={
            **search.pool_development.evidence_identity,
            "sample_context_hash": _hash("different-sample-context"),
        },
    )

    with pytest.raises(
        StrategyReportBundleError,
        match="Voting search.*sample context",
    ):
        _project(
            bindings,
            voting_candidate_search=replace(
                search,
                pool_development=development,
            ),
        )


def test_adapter_prefers_impact_cube_and_projects_all_populations_partitions(
    tmp_path: Path,
) -> None:
    project, sample, pool, legacy_impact = _bindings(tmp_path)
    impact_cube = _impact_cube_binding(tmp_path, sample, pool)

    result = build_strategy_report_bundle_source_inputs(
        project_context=project,
        sample_design=sample,
        candidate_pool=pool,
        pool_impact=legacy_impact,
        impact_cube=impact_cube,
    )

    assert result["strategy_artifact_refs"] == [
        {
            "kind": "strategy_candidate_pool",
            "ref_id": pool.artifact_id,
            "content_hash": pool.artifact_content_hash,
        },
        {
            "kind": "strategy_impact",
            "ref_id": impact_cube.artifact_id,
            "content_hash": impact_cube.artifact_content_hash,
        },
    ]
    producer_run = impact_cube.artifact_provenance["producer_run"]
    assert result["tool_run_refs"].count(
        {
            "kind": "tool_run",
            "ref_id": producer_run["run_id"],
            "content_hash": producer_run["content_hash"],
        }
    ) == 1
    impact = result["sections"][5]
    table_ids = {table["table_id"] for table in impact["tables"]}
    assert {
        "impact_cube_partitions",
        "impact_cube_slices",
        "impact_cube_waterfall",
        "impact_cube_transitions",
        "impact_cube_economics",
    } <= table_ids
    assert {
        (item["effect_stage"], item["population"], item["partition"])
        for item in impact["stage_evidence"]
    } == {
        ("backtested", "approval", "development"),
        ("backtested", "risk", "development"),
        ("oot_validated", "approval", "validation"),
        ("oot_validated", "risk", "validation"),
        ("oot_validated", "approval", "oot"),
        ("oot_validated", "risk", "oot"),
    }
    slices = next(
        table
        for table in impact["tables"]
        if table["table_id"] == "impact_cube_slices"
    )
    assert {
        row["cells"]["population"]["value"]
        for row in slices["rows"]
    } == {"approval", "risk"}
    assert {
        row["cells"]["partition"]["value"]
        for row in slices["rows"]
    } == {"development", "validation", "oot"}
    assert {
        row["cells"]["family"]["value"]
        for row in slices["rows"]
    } >= {
        "overall",
        "month",
        "group",
        "segment",
        "group_month",
        "segment_month",
        "new_action",
    }
    serialized = json.dumps(result, ensure_ascii=False, sort_keys=True)
    assert legacy_impact.artifact_id not in serialized
    assert "PII-CUSTOMER" not in serialized


def test_adapter_accepts_exact_score_requirements_and_rejects_ref_drift(
    tmp_path: Path,
) -> None:
    project, sample, _plain_pool, legacy_impact = _bindings(tmp_path)
    vector_id = "1" * 64
    requirement = {
        "type": "model_score_vector.v1",
        "virtual_field": "__marvis_model_pd_" + vector_id[:16],
        "score_product": "raw_native_uncalibrated_bad_probability",
        "score_evidence_artifact_id": "2" * 64,
        "score_evidence_artifact_content_hash": "3" * 64,
        "score_vector_artifact_id": vector_id,
        "score_vector_artifact_content_hash": "4" * 64,
    }
    pool = _pool_binding_for_strategy_type(
        tmp_path,
        sample,
        "approval",
        score_requirement=requirement,
    )
    impact_cube = _impact_cube_binding(tmp_path, sample, pool)

    result = build_strategy_report_bundle_source_inputs(
        project_context=project,
        sample_design=sample,
        candidate_pool=pool,
        pool_impact=legacy_impact,
        impact_cube=impact_cube,
    )

    assert any(
        item["kind"] == "strategy_impact"
        for item in result["strategy_artifact_refs"]
    )

    forged_provenance = json.loads(
        json.dumps(impact_cube.artifact_provenance)
    )
    forged_requirements = forged_provenance["requirement_bindings"][
        "requirements"
    ]
    forged_requirements[0]["requirement"][
        "score_evidence_artifact_id"
    ] = "f" * 64
    forged_provenance["requirement_bindings"]["requirements_hash"] = (
        hashlib.sha256(
            json.dumps(
                forged_requirements,
                ensure_ascii=False,
                sort_keys=True,
                separators=(",", ":"),
            ).encode("utf-8")
        ).hexdigest()
    )
    forged = replace(
        impact_cube,
        artifact_provenance=forged_provenance,
        artifact_provenance_json=json.dumps(
            forged_provenance,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
            allow_nan=False,
        ),
    )
    with pytest.raises(
        StrategyReportBundleError,
        match="requirement",
    ):
        build_strategy_report_bundle_source_inputs(
            project_context=project,
            sample_design=sample,
            candidate_pool=pool,
            pool_impact=legacy_impact,
            impact_cube=forged,
        )


def test_adapter_suppresses_oot_claims_but_keeps_partition_results_when_blocked(
    tmp_path: Path,
) -> None:
    sample = _sample_binding(tmp_path, maturity_status="not_matured")
    dataset = sample.bundle["sample_design"]["identity"]["dataset_ref"]
    project = _project_binding(
        tmp_path,
        task_id=sample.task_id,
        dataset_ref=build_source_ref(
            kind="dataset",
            ref_id=dataset["dataset_id"],
            content_hash=dataset["content_hash"],
        ),
    )
    pool = _pool_binding_for_strategy_type(
        tmp_path,
        sample,
        "approval",
    )
    impact_cube = _impact_cube_binding(tmp_path, sample, pool)

    result = build_strategy_report_bundle_source_inputs(
        project_context=project,
        sample_design=sample,
        candidate_pool=pool,
        impact_cube=impact_cube,
    )

    impact = result["sections"][5]
    assert {
        (item["population"], item["partition"])
        for item in impact["stage_evidence"]
    } == {
        ("approval", "development"),
        ("risk", "development"),
    }
    partitions = next(
        table
        for table in impact["tables"]
        if table["table_id"] == "impact_cube_partitions"
    )
    blocked_rows = [
        row
        for row in partitions["rows"]
        if row["cells"]["partition"]["value"] in {"validation", "oot"}
    ]
    assert blocked_rows
    assert all(
        row["cells"][field]["availability"] == "unavailable"
        and row["cells"][field]["value"] is None
        and row["cells"][field]["note"]
        == "claim_suppressed_by_validation_blocker"
        for row in blocked_rows
        for field in ("effect_stage", "validation_status")
    )
    slices = next(
        table
        for table in impact["tables"]
        if table["table_id"] == "impact_cube_slices"
    )
    assert {
        row["cells"]["partition"]["value"] for row in slices["rows"]
    } == {"development", "validation", "oot"}
    assert {
        item["code"] for item in impact["red_flags"]
    } >= {"oot_claim_suppressed_by_validation_blocker"}
    final_fields = {
        item["field_id"]: item["field"]["value"]
        for item in result["sections"][6]["summary_fields"]
    }
    assert final_fields["evidence_stages"] == ["backtested"]
    assert final_fields["validation_statuses"] == ["unvalidated"]
    bundle = build_strategy_report_bundle(
        task_id=sample.task_id,
        report_revision=1,
        strategy_id=None,
        strategy_version=None,
        strategy_type="approval",
        title=_present(
            "Blocked ImpactCube",
            result["strategy_artifact_refs"][0],
        ),
        status="partial",
        generated_at="2026-07-24T16:00:00+08:00",
        **result,
    )
    assert bundle["effect_stages"] == ["backtested"]
    assert bundle["completeness_summary"]["has_validation_blocker"] is True


@pytest.mark.parametrize(
    ("strategy_type", "metric_key"),
    [
        ("approval", "approve_count"),
        ("reject", "bad_capture_rate"),
        ("limit", "total_limit"),
        ("pricing", "mean_rate"),
        ("segmentation", "segment_count"),
    ],
)
def test_impact_cube_projection_preserves_all_five_strategy_metric_contracts(
    tmp_path: Path,
    strategy_type: str,
    metric_key: str,
) -> None:
    sample = _sample_binding(tmp_path)
    dataset = sample.bundle["sample_design"]["identity"]["dataset_ref"]
    project = _project_binding(
        tmp_path,
        task_id=sample.task_id,
        dataset_ref=build_source_ref(
            kind="dataset",
            ref_id=dataset["dataset_id"],
            content_hash=dataset["content_hash"],
        ),
    )
    pool = _pool_binding_for_strategy_type(
        tmp_path,
        sample,
        strategy_type,
    )
    impact_cube = _impact_cube_binding(tmp_path, sample, pool)

    result = build_strategy_report_bundle_source_inputs(
        project_context=project,
        sample_design=sample,
        candidate_pool=pool,
        impact_cube=impact_cube,
    )

    table = next(
        item
        for item in result["sections"][5]["tables"]
        if item["table_id"] == "impact_cube_slices"
    )
    row = next(
        item
        for item in table["rows"]
        if item["cells"]["population"]["value"] == "risk"
        and item["cells"]["partition"]["value"] == "development"
        and item["cells"]["family"]["value"] == "overall"
    )
    assert row["cells"]["new_metrics"]["availability"] == "present"
    assert metric_key in row["cells"]["new_metrics"]["value"]
    bundle = build_strategy_report_bundle(
        task_id=sample.task_id,
        report_revision=1,
        strategy_id=None,
        strategy_version=None,
        strategy_type=strategy_type,
        title=_present(
            f"{strategy_type} ImpactCube",
            result["strategy_artifact_refs"][0],
        ),
        status="partial",
        generated_at="2026-07-24T16:00:00+08:00",
        **result,
    )
    assert bundle["strategy_type"] == strategy_type
    assert bundle["effect_stages"] == ["backtested", "oot_validated"]


def test_impact_cube_missing_economics_stays_blank_never_synthetic_zero(
    tmp_path: Path,
) -> None:
    project, sample, pool, _legacy_impact = _bindings(tmp_path)
    impact_cube = _impact_cube_binding(
        tmp_path,
        sample,
        pool,
        economics=False,
    )

    result = build_strategy_report_bundle_source_inputs(
        project_context=project,
        sample_design=sample,
        candidate_pool=pool,
        pool_impact=None,
        impact_cube=impact_cube,
    )

    impact = result["sections"][5]
    economics = next(
        table
        for table in impact["tables"]
        if table["table_id"] == "impact_cube_economics"
    )
    assert economics["rows"]
    for row in economics["rows"]:
        for field in ("current", "new", "delta"):
            assert row["cells"][field]["value"] is None
            assert row["cells"][field]["availability"] in {
                "unavailable",
                "not_applicable",
            }


def test_adapter_rejects_impact_cube_bound_to_another_pool_artifact(
    tmp_path: Path,
) -> None:
    project, sample, pool, _legacy_impact = _bindings(tmp_path)
    impact_cube = _impact_cube_binding(tmp_path, sample, pool)
    selected_pool = replace(pool, artifact_id="a" * 64)

    with pytest.raises(
        StrategyReportBundleError,
        match="another candidate-pool artifact",
    ):
        build_strategy_report_bundle_source_inputs(
            project_context=project,
            sample_design=sample,
            candidate_pool=selected_pool,
            impact_cube=impact_cube,
        )


def test_adapter_rejects_impact_cube_producer_run_self_hash_tamper(
    tmp_path: Path,
) -> None:
    project, sample, pool, _legacy_impact = _bindings(tmp_path)
    impact_cube = _impact_cube_binding(tmp_path, sample, pool)
    forged_provenance = json.loads(
        impact_cube.artifact_provenance_json
    )
    forged_provenance["producer_run"]["content_hash"] = "0" * 64
    forged = replace(
        impact_cube,
        artifact_provenance=forged_provenance,
        artifact_provenance_json=json.dumps(
            forged_provenance,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
            allow_nan=False,
        ),
    )

    with pytest.raises(
        StrategyReportBundleError,
        match="producer_run|self hash",
    ):
        build_strategy_report_bundle_source_inputs(
            project_context=project,
            sample_design=sample,
            candidate_pool=pool,
            impact_cube=forged,
        )


def test_adapter_rejects_impact_cube_with_forged_development_lineage(
    tmp_path: Path,
) -> None:
    project, sample, pool, _legacy_impact = _bindings(tmp_path)
    impact_cube = _impact_cube_binding(tmp_path, sample, pool)
    forged_cube = json.loads(
        canonical_strategy_impact_cube_json(impact_cube.cube)
    )
    forged_cube["source_bindings"]["development_lineage"][
        "sample_binding"
    ]["sample_context_hash"] = _hash("forged-sample-context")
    body = {
        key: value
        for key, value in forged_cube.items()
        if key not in {"cube_id", "content_hash"}
    }
    canonical_body = json.dumps(
        body,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        allow_nan=False,
    )
    forged_cube["cube_id"] = (
        "strategy-impact-cube-"
        + hashlib.sha256(canonical_body.encode("utf-8")).hexdigest()[:24]
    )
    without_hash = {
        key: value
        for key, value in forged_cube.items()
        if key != "content_hash"
    }
    forged_cube["content_hash"] = hashlib.sha256(
        json.dumps(
            without_hash,
            ensure_ascii=False,
            sort_keys=True,
            separators=(",", ":"),
            allow_nan=False,
        ).encode("utf-8")
    ).hexdigest()
    forged_canonical = canonical_strategy_impact_cube_json(forged_cube)
    forged_artifact_hash = _file_hash(forged_canonical)
    forged_producer_run = build_impact_cube_producer_run(
        task_id=impact_cube.task_id,
        request=impact_cube.artifact_provenance["producer_run"]["request"],
        cube_id=forged_cube["cube_id"],
        cube_content_hash=forged_cube["content_hash"],
        artifact_id=impact_cube.artifact_id,
        artifact_filename=impact_cube.artifact_path.name,
        artifact_content_hash=forged_artifact_hash,
    )
    forged_provenance = {
        **impact_cube.artifact_provenance,
        "cube_id": forged_cube["cube_id"],
        "cube_content_hash": forged_cube["content_hash"],
        "producer_run": forged_producer_run,
    }
    forged_provenance_json = json.dumps(
        forged_provenance,
        ensure_ascii=False,
        sort_keys=True,
        separators=(",", ":"),
        allow_nan=False,
    )
    forged_binding = replace(
        impact_cube,
        artifact_content_hash=forged_artifact_hash,
        artifact_provenance=forged_provenance,
        artifact_provenance_json=forged_provenance_json,
        cube=forged_cube,
    )

    with pytest.raises(
        StrategyReportBundleError,
        match="development lineage differs",
    ):
        build_strategy_report_bundle_source_inputs(
            project_context=project,
            sample_design=sample,
            candidate_pool=pool,
            impact_cube=forged_binding,
        )


def test_adapter_never_leaks_raw_rows_pii_or_overclaims_lifecycle(
    tmp_path: Path,
) -> None:
    result = _project(_bindings(tmp_path))
    serialized = json.dumps(result, ensure_ascii=False, sort_keys=True)
    impact = result["sections"][5]
    candidates = result["sections"][4]
    final = result["sections"][6]

    assert "PII-CUSTOMER" not in serialized
    assert "customer_id" not in serialized
    assert "row_ordinal" not in serialized
    assert "oot_validated" not in serialized
    assert "post_launch_observed" not in serialized
    assert impact["stage_evidence"] == [
        {
            "effect_stage": "backtested",
            "population": "risk",
            "partition": "development",
            "binding": {
                "kind": "development_backtest",
                "dataset_ref": impact["stage_evidence"][0]["binding"][
                    "dataset_ref"
                ],
                "frozen_artifact_ref": impact["stage_evidence"][0]["binding"][
                    "frozen_artifact_ref"
                ],
                "result_ref": impact["stage_evidence"][0]["binding"][
                    "result_ref"
                ],
            },
        }
    ]
    assert {
        item["field_id"]: item["field"]["value"]
        for item in candidates["summary_fields"]
    }["adoption_status"] == "not_adopted"
    final_fields = {
        item["field_id"]: item["field"]["value"]
        for item in final["summary_fields"]
    }
    assert final_fields["adoption_status"] == "not_adopted"
    assert final_fields["deployment_status"] == "not_deployed"
    assert final_fields["creates_strategy"] is False


def test_adapter_preserves_zero_unavailable_and_not_matured_without_filling(
    tmp_path: Path,
) -> None:
    result = _project(
        _bindings(tmp_path, maturity_status="not_matured")
    )
    current = {
        item["field_id"]: item["field"]
        for item in result["sections"][0]["summary_fields"]
    }
    sample = {
        item["field_id"]: item["field"]
        for item in result["sections"][2]["summary_fields"]
    }
    sample_rows = result["sections"][2]["tables"][0]["rows"]

    assert current["current_approval"] == {
        "value": 0,
        "availability": "present",
        "origin": "tool_output",
        "source_refs": current["current_approval"]["source_refs"],
        "as_of": None,
        "blocking": "none",
        "note": None,
    }
    assert current["current_risk"]["value"] is None
    assert current["current_risk"]["availability"] == "unavailable"
    assert sample["risk_maturity"]["value"] is None
    assert sample["risk_maturity"]["availability"] == "not_matured"
    assert sample["risk_maturity"]["blocking"] == "validation"
    assert any(
        row["cells"]["value"]["value"] is None
        and row["cells"]["value"]["availability"] == "not_matured"
        for row in sample_rows
    )


def test_optional_model_training_and_score_evidence_project_only_aggregates(
    tmp_path: Path,
) -> None:
    bindings = _bindings(tmp_path)
    sample = bindings[1]
    model = _model_binding(tmp_path, sample)
    training = _training_binding(tmp_path, sample)
    score = _score_binding(tmp_path, training)

    result = _project(
        bindings,
        model_evidence=model,
        training_evidence=training,
        score_evidence=score,
    )
    section = result["sections"][3]
    table_ids = {item["table_id"] for item in section["tables"]}
    serialized = json.dumps(section, ensure_ascii=False, sort_keys=True)

    assert section["availability"] == "present"
    assert {
        "univariate_model_observations",
        "model_comparison_metrics",
        "model_selection_results",
        "training_metrics",
        "training_feature_importance",
        "governed_score_observations",
    } <= table_ids
    selection_table = next(
        item
        for item in section["tables"]
        if item["table_id"] == "model_selection_results"
    )
    assert selection_table["rows"][0]["cells"]["status"]["value"] == "selected"
    assert (
        selection_table["rows"][0]["cells"]["selected_model_evidence_id"][
            "value"
        ]
        is not None
    )
    fields = {
        item["field_id"]: item["field"]
        for item in section["summary_fields"]
    }
    assert fields["training_selection_status"]["availability"] == "unavailable"
    assert fields["score_selection_status"]["availability"] == "unavailable"
    assert fields["training_selection_status"]["value"] is None
    assert fields["score_selection_status"]["value"] is None
    assert score.evidence_record["id"] in serialized
    assert score.vector_record["id"] not in serialized
    assert "row_ordinal" not in serialized
    assert "score_min" not in serialized
    assert "score_max" not in serialized
    assert "PII-CUSTOMER" not in serialized


def test_model_no_selection_is_projected_without_inventing_a_selected_model(
    tmp_path: Path,
) -> None:
    bindings = _bindings(tmp_path)
    model = _model_binding(
        tmp_path,
        bindings[1],
        selection_status="no_selection",
    )

    result = _project(bindings, model_evidence=model)
    section = result["sections"][3]
    selection_table = next(
        item
        for item in section["tables"]
        if item["table_id"] == "model_selection_results"
    )
    row = selection_table["rows"][0]["cells"]

    assert row["status"]["value"] == "no_selection"
    assert row["selected_model_evidence_id"]["value"] is None
    assert row["selected_model_evidence_id"]["availability"] == "not_applicable"
    assert row["reason"]["value"] == "业务尚未确认最终模型。"
    comparison_table = next(
        item
        for item in section["tables"]
        if item["table_id"] == "model_comparison_metrics"
    )
    assert all(
        item["cells"]["selected"]["value"] is False
        for item in comparison_table["rows"]
    )


def test_non_present_model_comparison_projects_one_empty_row_per_model(
    tmp_path: Path,
) -> None:
    bindings = _bindings(tmp_path)
    model = _model_binding(
        tmp_path,
        bindings[1],
        comparison_metric_status="unavailable",
    )

    result = _project(bindings, model_evidence=model)
    table = next(
        item
        for item in result["sections"][3]["tables"]
        if item["table_id"] == "model_comparison_metrics"
    )

    assert len(table["rows"]) == 2
    assert all(
        row["cells"]["value"]["value"] is None
        and row["cells"]["value"]["availability"] == "unavailable"
        and row["cells"]["delta"]["value"] is None
        and row["cells"]["delta"]["availability"] == "unavailable"
        for row in table["rows"]
    )


@pytest.mark.parametrize(
    ("mutation", "message"),
    [
        ("untyped_project", "authenticated"),
        ("sample_hash", "canonical evidence"),
        ("compiled_design", "compiled design"),
        ("impact_stage", "development"),
        ("cross_task", "another task"),
    ],
)
def test_adapter_rejects_untyped_or_forged_bindings(
    tmp_path: Path,
    mutation: str,
    message: str,
) -> None:
    project, sample, pool, impact = _bindings(tmp_path)
    if mutation == "untyped_project":
        project = {"task_id": project.task_id}  # type: ignore[assignment]
    elif mutation == "sample_hash":
        sample = replace(
            sample,
            bundle_artifact_content_hash="f" * 64,
        )
    elif mutation == "compiled_design":
        pool = replace(
            pool,
            compiled_design={
                **pool.compiled_design,
                "requirements": [{"forged": True}],
            },
        )
    elif mutation == "impact_stage":
        impact = replace(impact, stage="independent_oot")
    else:
        impact = replace(impact, task_id="foreign-task")

    with pytest.raises(StrategyReportBundleError, match=message):
        build_strategy_report_bundle_source_inputs(
            project_context=project,  # type: ignore[arg-type]
            sample_design=sample,
            candidate_pool=pool,
            pool_impact=impact,
        )


@pytest.mark.parametrize("source_kind", ["pool_entry", "univariate_asset"])
def test_candidate_stability_compatibility_accepts_exact_current_pool_sources(
    tmp_path: Path,
    source_kind: str,
) -> None:
    _project, sample, pool, _impact = _bindings(tmp_path)
    binding = _candidate_stability_binding(
        tmp_path,
        sample,
        pool,
        source_kind=source_kind,
    )

    stability = validate_candidate_stability_report_compatibility(
        candidate_stability=binding,
        sample_design=sample,
        candidate_pool=pool,
    )

    assert stability == binding.stability
    assert stability["source_ref"]["source_kind"] == source_kind
    assert stability["lifecycle"] == {
        "candidate_stage": "development",
        "observation_stage": "backtested",
        "validation_status": "unvalidated",
        "not_created_strategy": True,
        "not_adopted": True,
        "not_deployed": True,
    }


def test_candidate_stability_compatibility_rejects_another_current_pool_artifact(
    tmp_path: Path,
) -> None:
    _project, sample, pool, _impact = _bindings(tmp_path)
    binding = _candidate_stability_binding(tmp_path, sample, pool)
    another_pool_artifact = replace(pool, artifact_id="a" * 64)

    with pytest.raises(
        StrategyReportBundleError,
        match="exact current Candidate Pool entry",
    ):
        validate_candidate_stability_report_compatibility(
            candidate_stability=binding,
            sample_design=sample,
            candidate_pool=another_pool_artifact,
        )


def test_adapter_projects_candidate_stability_summary_table_and_sources(
    tmp_path: Path,
) -> None:
    project, sample, pool, impact = _bindings(tmp_path)
    binding = _candidate_stability_binding(tmp_path, sample, pool)

    result = build_strategy_report_bundle_source_inputs(
        project_context=project,
        sample_design=sample,
        candidate_pool=pool,
        candidate_stability=binding,
        pool_impact=impact,
    )

    section = result["sections"][4]
    fields = {
        item["field_id"]: item["field"]["value"]
        for item in section["summary_fields"]
    }
    assert fields["candidate_stability_population_count"] == 60
    assert fields["candidate_stability_month_count"] == 2
    assert fields["candidate_stability_max_psi_month"] in {
        "202601",
        "202602",
    }
    table = next(
        item
        for item in section["tables"]
        if item["table_id"] == "candidate_monthly_stability"
    )
    assert table["title"] == "候选逐月稳定性（开发回测，未独立验证）"
    assert table["sheet_key"] == "appendix_candidate_stability"
    assert table["content_class"] == "monthly_summary"
    assert table["effect_stage"] == "backtested"
    assert [row["row_id"] for row in table["rows"]] == [
        "candidate-stability-baseline",
        "candidate-stability-month-202601",
        "candidate-stability-month-202602",
    ]
    assert [
        row["cells"]["sample_count"]["value"]
        for row in table["rows"]
    ] == [60, 30, 30]
    assert set(table["rows"][0]["cells"]) == {
        "period",
        "sample_count",
        "hit_count",
        "not_hit_count",
        "hit_share",
        "not_hit_share",
        "labeled_count",
        "label_coverage",
        "hit_labeled_count",
        "hit_bad_count",
        "hit_bad_rate",
        "psi_vs_development",
    }
    stability_ref = {
        "kind": "backtest",
        "ref_id": binding.artifact_id,
        "content_hash": binding.artifact_content_hash,
    }
    assert table["source_refs"] == [stability_ref]
    assert stability_ref in result["strategy_artifact_refs"]
    assert section["stage_evidence"] == [
        {
            "effect_stage": "backtested",
            "population": "risk",
            "partition": "development",
            "binding": {
                "kind": "development_backtest",
                "dataset_ref": result["dataset_refs"][0],
                "frozen_artifact_ref": {
                    "kind": "strategy_candidate_pool",
                    "ref_id": pool.artifact_id,
                    "content_hash": pool.artifact_content_hash,
                },
                "result_ref": stability_ref,
            },
        }
    ]
    bundle = build_strategy_report_bundle(
        task_id=project.task_id,
        report_revision=1,
        strategy_id=None,
        strategy_version=None,
        strategy_type="approval",
        title=_present("候选稳定性报告", stability_ref),
        status="partial",
        generated_at="2026-07-25T12:00:00+08:00",
        **result,
    )
    assert bundle["schema_version"] == "strategy.report-bundle.v2"
    assert "backtested" in bundle["effect_stages"]


def test_adapter_binds_standalone_stability_to_exact_candidate_asset(
    tmp_path: Path,
) -> None:
    project, sample, pool, impact = _bindings(tmp_path)
    binding = _candidate_stability_binding(
        tmp_path,
        sample,
        pool,
        source_kind="univariate_asset",
    )

    result = build_strategy_report_bundle_source_inputs(
        project_context=project,
        sample_design=sample,
        candidate_pool=pool,
        candidate_stability=binding,
        pool_impact=impact,
    )

    source = binding.stability["source_ref"]
    asset_ref = {
        "kind": "strategy_candidate_asset",
        "ref_id": source["artifact_id"],
        "content_hash": source["artifact_content_hash"],
    }
    stability_ref = {
        "kind": "backtest",
        "ref_id": binding.artifact_id,
        "content_hash": binding.artifact_content_hash,
    }
    candidate_section = result["sections"][4]
    stage = candidate_section["stage_evidence"][0]
    assert stage["binding"]["frozen_artifact_ref"] == asset_ref
    assert asset_ref in candidate_section["source_refs"]
    assert asset_ref in result["strategy_artifact_refs"]
    assert stage["binding"]["frozen_artifact_ref"] != {
        "kind": "strategy_candidate_pool",
        "ref_id": pool.artifact_id,
        "content_hash": pool.artifact_content_hash,
    }

    bundle = build_strategy_report_bundle(
        task_id=project.task_id,
        report_revision=1,
        strategy_id=None,
        strategy_version=None,
        strategy_type="approval",
        title=_present("独立候选稳定性报告", stability_ref),
        status="partial",
        generated_at="2026-07-25T12:00:00+08:00",
        **result,
    )
    assert bundle["effect_stages"] == ["backtested"]


def test_adapter_projects_low_sample_stability_flags_as_amber(
    tmp_path: Path,
) -> None:
    project, sample, pool, impact = _bindings(tmp_path)
    binding = _candidate_stability_binding(
        tmp_path,
        sample,
        pool,
        low_sample_month=True,
    )

    result = build_strategy_report_bundle_source_inputs(
        project_context=project,
        sample_design=sample,
        candidate_pool=pool,
        candidate_stability=binding,
        pool_impact=impact,
    )

    red_flags = result["sections"][4]["red_flags"]
    assert len(red_flags) == 1
    assert red_flags[0]["code"] == (
        "candidate_stability_insufficient_month_rows"
    )
    assert red_flags[0]["level"] == "amber"
    assert "202601" in red_flags[0]["message"]
    assert "低样本" in red_flags[0]["message"]


def test_adapter_projects_complete_scorecard_tables_from_pool_lineage(
    tmp_path: Path,
) -> None:
    project, sample, _pool, _legacy_impact = _bindings(tmp_path)
    pool, selection = _scorecard_report_pool_binding(tmp_path, sample)
    impact_cube = _impact_cube_binding(tmp_path, sample, pool)

    result = build_strategy_report_bundle_source_inputs(
        project_context=project,
        sample_design=sample,
        candidate_pool=pool,
        impact_cube=impact_cube,
    )

    section = result["sections"][4]
    tables = {table["table_id"]: table for table in section["tables"]}
    assert {
        "scorecard_model_summary",
        "scorecard_points",
        "scorecard_bands",
        "scorecard_cutoff_evaluations",
    } <= set(tables)
    assert tables["scorecard_points"]["sheet_key"] == "appendix_scorecard"
    scorecard_tables = [
        tables[table_id]
        for table_id in (
            "scorecard_model_summary",
            "scorecard_points",
            "scorecard_bands",
            "scorecard_cutoff_evaluations",
        )
    ]
    assert {
        table["effect_stage"] for table in scorecard_tables
    } == {"backtested"}
    backtest_stages = [
        stage
        for stage in section["stage_evidence"]
        if stage["effect_stage"] == "backtested"
    ]
    assert len(backtest_stages) == 1
    assert backtest_stages[0]["binding"]["result_ref"] == {
        "kind": "backtest",
        "ref_id": selection.source_asset_binding.artifact_id,
        "content_hash": selection.source_asset_binding.content_hash,
    }
    assert backtest_stages[0]["binding"]["frozen_artifact_ref"] == {
        "kind": "strategy_candidate_asset",
        "ref_id": selection.source_asset_binding.artifact_id,
        "content_hash": selection.source_asset_binding.content_hash,
    }
    assert {
        (
            ref["kind"],
            ref["ref_id"],
            ref["content_hash"],
        )
        for ref in result["strategy_artifact_refs"]
    } >= {
        (
            "strategy_scorecard_band_asset",
            selection.source_asset_binding.artifact_id,
            selection.source_asset_binding.content_hash,
        ),
        (
            "strategy_candidate_asset",
            selection.source_asset_binding.artifact_id,
            selection.source_asset_binding.content_hash,
        ),
        (
            "backtest",
            selection.source_asset_binding.artifact_id,
            selection.source_asset_binding.content_hash,
        ),
    }
    summary = tables["scorecard_model_summary"]["rows"][0]["cells"]
    assert summary["score_direction"]["value"] == "higher_is_riskier"
    assert summary["points_direction"]["value"] == "higher_is_better"
    assert summary["auc"]["value"] == 1.0
    assert summary["ks"]["value"] == 1.0
    pool_ref = {
        "kind": "strategy_candidate_pool",
        "ref_id": pool.artifact_id,
        "content_hash": pool.artifact_content_hash,
    }
    assert pool_ref in summary["model_index"]["source_refs"]
    assert {
        ref["kind"] for ref in summary["model_index"]["source_refs"]
    } >= {
        "strategy_candidate_pool",
        "strategy_scorecard_band_asset",
        "strategy_scorecard_cutoff_selection",
    }
    for table in scorecard_tables[1:]:
        assert pool_ref in table["rows"][0]["cells"]["model_index"][
            "source_refs"
        ]

    asset = selection.source_asset_binding.asset
    assert len(tables["scorecard_points"]["rows"]) == len(
        asset["score_contract"]["scorecard_table"]
    )
    assert len(tables["scorecard_bands"]["rows"]) == len(asset["bands"])
    cutoff_rows = tables["scorecard_cutoff_evaluations"]["rows"]
    assert len(cutoff_rows) == len(asset["cutoffs"])
    selected = next(
        row for row in cutoff_rows if row["cells"]["selected"]["value"] is True
    )
    unselected = next(
        row
        for row in cutoff_rows
        if row["cells"]["selected"]["value"] is False
    )
    assert {
        ref["kind"]
        for ref in unselected["cells"]["selected"]["source_refs"]
    } >= {
        "strategy_candidate_pool",
        "strategy_scorecard_band_asset",
        "strategy_scorecard_cutoff_selection",
    }
    assert selected["cells"]["cutoff_id"]["value"] == selection.selection[
        "cutoff_id"
    ]
    assert selected["cells"]["selection_reason"]["value"] == ["风险上限方案"]
    assert selected["cells"]["lower_risk_count"]["value"] == 2
    assert selected["cells"]["lower_risk_bad_count"]["value"] == 0
    assert selected["cells"]["lower_risk_bad_rate"]["value"] == 0.0
    assert selected["cells"]["higher_risk_count"]["value"] == 4
    assert selected["cells"]["higher_risk_bad_count"]["value"] == 2
    assert selected["cells"]["higher_risk_bad_rate"]["value"] == pytest.approx(
        2.0 / 3.0
    )

    bundle = build_strategy_report_bundle(
        task_id=project.task_id,
        report_revision=1,
        strategy_id=None,
        strategy_version=None,
        strategy_type="approval",
        title=_present("评分卡策略报告", result["strategy_artifact_refs"][0]),
        status="partial",
        generated_at="2026-07-25T12:00:00+08:00",
        **result,
    )
    assert bundle["schema_version"] == "strategy.report-bundle.v2"
    assert "backtested" in bundle["effect_stages"]
    rendered = render_strategy_report_bundle(bundle)
    assert set(rendered) == {"json", "markdown", "xlsx", "docx"}
    assert b"scorecard_model_summary" in rendered["json"]
    assert "评分卡模型汇总" in rendered["markdown"].decode("utf-8")
    assert rendered["xlsx"].startswith(b"PK")
    assert rendered["docx"].startswith(b"PK")


def test_scorecard_report_projection_recurses_voting_dedupes_and_redacts(
    tmp_path: Path,
) -> None:
    _project, sample, _pool, _impact = _bindings(tmp_path)
    pool = _scorecard_voting_report_pool_binding(tmp_path, sample)

    projection = project_scorecard_report_evidence(pool)

    assert len(projection["models"]) == 1
    assert len(projection["usages"]) == 2
    assert [
        [
            node["scope"]
            for node in usage["usage_paths"][0]["path"]
        ]
        for usage in projection["usages"]
    ] == [
        ["current_pool_entry", "voting_parent_entry"],
        ["current_pool_entry", "voting_parent_entry"],
    ]
    assert [
        usage["selection_reason"] for usage in projection["usages"]
    ] == ["低风险方案", "高风险方案"]
    assert len(projection["artifact_refs"]) == 5
    assert {
        ref["kind"] for ref in projection["artifact_refs"]
    } == {
        "strategy_candidate_pool",
        "strategy_scorecard_band_asset",
        "strategy_scorecard_cutoff_selection",
        "strategy_voting_candidate",
    }
    for usage in projection["usages"]:
        assert {
            ref["kind"] for ref in usage["usage_artifact_refs"]
        } == {
            "strategy_candidate_pool",
            "strategy_scorecard_band_asset",
            "strategy_scorecard_cutoff_selection",
            "strategy_voting_candidate",
        }
        assert (
            usage["usage_paths"][0]["artifact_refs"]
            == usage["usage_artifact_refs"]
        )
    serialized = json.dumps(
        projection,
        ensure_ascii=False,
        sort_keys=True,
    )
    for forbidden in (
        "asset_hash",
        "selection_hash",
        "raw_pd_content_hash",
        "canonical_bytes",
        "score_vector",
        "dataset_content_hash",
    ):
        assert forbidden not in serialized


def test_non_scorecard_voting_pool_keeps_scorecard_projection_empty(
    tmp_path: Path,
) -> None:
    pool = _non_scorecard_voting_report_pool_binding(tmp_path)

    projection = project_scorecard_report_evidence(pool)

    assert projection == {
        "schema_version": "strategy.scorecard-report-projection.v1",
        "models": [],
        "usages": [],
        "artifact_refs": [],
    }


def test_scorecard_projection_keeps_each_reused_usage_path_audit_chain(
    tmp_path: Path,
) -> None:
    _project, sample, _pool, _impact = _bindings(tmp_path)
    pool = _scorecard_reused_selection_voting_report_pool_binding(
        tmp_path,
        sample,
    )

    projection = project_scorecard_report_evidence(pool)

    assert len(projection["models"]) == 1
    assert len(projection["usages"]) == 2
    for usage in projection["usages"]:
        assert len(usage["usage_paths"]) == 2
        path_voting_refs = [
            [
                ref["ref_id"]
                for ref in path["artifact_refs"]
                if ref["kind"] == "strategy_voting_candidate"
            ]
            for path in usage["usage_paths"]
        ]
        assert all(len(refs) == 1 for refs in path_voting_refs)
        assert len({refs[0] for refs in path_voting_refs}) == 2
        assert {
            ref["ref_id"]
            for ref in usage["usage_artifact_refs"]
            if ref["kind"] == "strategy_voting_candidate"
        } == {refs[0] for refs in path_voting_refs}


def test_scorecard_report_projection_rejects_tampered_selection_and_voting_order(
    tmp_path: Path,
) -> None:
    _project, sample, _pool, _impact = _bindings(tmp_path)
    direct, _selection = _scorecard_report_pool_binding(tmp_path, sample)
    direct_lineage = direct.lineages[0]
    tampered_selection = replace(
        direct_lineage.selection,
        selection={
            **direct_lineage.selection.selection,
            "selection_reason": "伪造理由",
        },
    )
    tampered_direct = replace(
        direct,
        lineages=(
            replace(direct_lineage, selection=tampered_selection),
        ),
    )

    with pytest.raises(StrategyError):
        project_scorecard_report_evidence(tampered_direct)

    voting = _scorecard_voting_report_pool_binding(tmp_path, sample)
    voting_lineage = voting.lineages[0]
    tampered_voting = replace(
        voting,
        lineages=(
            replace(
                voting_lineage,
                parent_lineages=tuple(
                    reversed(voting_lineage.parent_lineages)
                ),
            ),
        ),
    )
    with pytest.raises(StrategyError, match="parent entry"):
        project_scorecard_report_evidence(tampered_voting)

    parent_artifact = voting_lineage.parent_pool_artifact
    forged_parent_artifact = SimpleNamespace(**vars(parent_artifact))
    forged_parent_artifact.content_hash = "f" * 64
    tampered_parent_ref = replace(
        voting,
        lineages=(
            replace(
                voting_lineage,
                parent_pool_artifact=forged_parent_artifact,
            ),
        ),
    )
    with pytest.raises(StrategyError, match="parent Pool artifact"):
        project_scorecard_report_evidence(tampered_parent_ref)


def test_scorecard_report_projection_rejects_relevant_lineage_type_drift(
    tmp_path: Path,
) -> None:
    _project, sample, _pool, _impact = _bindings(tmp_path)
    direct, _selection = _scorecard_report_pool_binding(tmp_path, sample)
    source = direct.pool["entries"][0]["source"]
    wrong_lineage = strategy_pool_tools._UnivariateCandidateLineage(
        asset_record=None,
        asset={},
        parent_record=None,
        evidence={},
        dataset=SimpleNamespace(),
        verified_fragment={},
        source_binding=source,
    )

    with pytest.raises(StrategyError, match="lineage type"):
        project_scorecard_report_evidence(
            replace(direct, lineages=(wrong_lineage,))
        )

    voting = _scorecard_voting_report_pool_binding(tmp_path, sample)
    voting_source = voting.pool["entries"][0]["source"]
    with pytest.raises(StrategyError, match="lineage type"):
        project_scorecard_report_evidence(
            replace(
                voting,
                lineages=(
                    replace(wrong_lineage, source_binding=voting_source),
                ),
            )
        )


def test_adapter_preserves_scorecard_voting_usage_paths_and_all_cutoffs(
    tmp_path: Path,
) -> None:
    project, sample, _pool, _impact = _bindings(tmp_path)
    pool = _scorecard_voting_report_pool_binding(tmp_path, sample)
    impact_cube = _impact_cube_binding(tmp_path, sample, pool)

    result = build_strategy_report_bundle_source_inputs(
        project_context=project,
        sample_design=sample,
        candidate_pool=pool,
        impact_cube=impact_cube,
    )

    tables = {
        table["table_id"]: table
        for table in result["sections"][4]["tables"]
    }
    summary = tables["scorecard_model_summary"]["rows"][0]["cells"]
    assert summary["usage_count"]["value"] == 2
    assert {
        ref["kind"] for ref in summary["usage_paths"]["source_refs"]
    } >= {
        "strategy_candidate_pool",
        "strategy_scorecard_band_asset",
        "strategy_scorecard_cutoff_selection",
        "strategy_voting_candidate",
    }
    assert all(
        "Pool[1]" in path["label"] and "Voting[" in path["label"]
        for path in summary["usage_paths"]["value"]
    )
    assert all(
        {
            ref["kind"] for ref in path["artifact_refs"]
        } == {
            "strategy_candidate_pool",
            "strategy_scorecard_band_asset",
            "strategy_scorecard_cutoff_selection",
            "strategy_voting_candidate",
        }
        for path in summary["usage_paths"]["value"]
    )
    cutoff_rows = tables["scorecard_cutoff_evaluations"]["rows"]
    assert len(cutoff_rows) == 2
    assert [
        row["cells"]["selected"]["value"] for row in cutoff_rows
    ] == [True, True]
    assert [
        row["cells"]["selection_reason"]["value"] for row in cutoff_rows
    ] == [["低风险方案"], ["高风险方案"]]
    assert {
        ref["kind"] for ref in result["strategy_artifact_refs"]
    } >= {
        "strategy_candidate_pool",
        "strategy_scorecard_band_asset",
        "strategy_scorecard_cutoff_selection",
        "strategy_voting_candidate",
    }


def test_adapter_keeps_each_reused_scorecard_path_audit_chain(
    tmp_path: Path,
) -> None:
    project, sample, _pool, _impact = _bindings(tmp_path)
    pool = _scorecard_reused_selection_voting_report_pool_binding(
        tmp_path,
        sample,
    )
    impact_cube = _impact_cube_binding(tmp_path, sample, pool)

    result = build_strategy_report_bundle_source_inputs(
        project_context=project,
        sample_design=sample,
        candidate_pool=pool,
        impact_cube=impact_cube,
    )

    summary = next(
        table
        for table in result["sections"][4]["tables"]
        if table["table_id"] == "scorecard_model_summary"
    )["rows"][0]["cells"]
    paths = summary["usage_paths"]["value"]
    assert len(paths) == 4
    voting_ref_ids = [
        [
            ref["ref_id"]
            for ref in path["artifact_refs"]
            if ref["kind"] == "strategy_voting_candidate"
        ]
        for path in paths
    ]
    assert all(len(refs) == 1 for refs in voting_ref_ids)
    assert len({refs[0] for refs in voting_ref_ids}) == 2
    assert all(
        {
            (ref["kind"], ref["ref_id"], ref["content_hash"])
            for ref in path["artifact_refs"]
        }
        <= {
            (ref["kind"], ref["ref_id"], ref["content_hash"])
            for ref in summary["usage_paths"]["source_refs"]
        }
        for path in paths
    )


def test_scorecard_report_limits_fail_closed_without_trimming(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    _project, sample, _pool, _impact = _bindings(tmp_path)
    direct, _selection = _scorecard_report_pool_binding(tmp_path, sample)
    monkeypatch.setattr(
        strategy_pool_tools,
        "MAX_SCORECARD_REPORT_DETAIL_ROWS",
        1,
    )
    with pytest.raises(StrategyError, match="detail rows exceed budget"):
        project_scorecard_report_evidence(direct)

    voting = _scorecard_voting_report_pool_binding(tmp_path, sample)
    monkeypatch.setattr(
        strategy_pool_tools,
        "MAX_SCORECARD_REPORT_DETAIL_ROWS",
        512,
    )
    monkeypatch.setattr(
        strategy_pool_tools,
        "MAX_SCORECARD_REPORT_USAGES",
        1,
    )
    with pytest.raises(StrategyError, match="usage paths exceed budget"):
        project_scorecard_report_evidence(voting)

    monkeypatch.setattr(
        strategy_pool_tools,
        "MAX_SCORECARD_REPORT_TABLE_REFS",
        1,
        raising=False,
    )
    with pytest.raises(
        StrategyError,
        match="report reference footprint exceeds budget",
    ):
        project_scorecard_report_evidence(direct)


def test_scorecard_report_adapter_fails_before_global_report_budget(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    project, sample, _pool, _impact = _bindings(tmp_path)
    pool, _selection = _scorecard_report_pool_binding(tmp_path, sample)
    impact_cube = _impact_cube_binding(tmp_path, sample, pool)
    monkeypatch.setattr(
        "marvis.packs.strategy.report_bundle_adapters."
        "_MAX_SCORECARD_REPORT_TABLE_REFS",
        1,
    )

    with pytest.raises(
        StrategyReportBundleError,
        match="scorecard report table references exceed reserved budget",
    ):
        build_strategy_report_bundle_source_inputs(
            project_context=project,
            sample_design=sample,
            candidate_pool=pool,
            impact_cube=impact_cube,
        )
