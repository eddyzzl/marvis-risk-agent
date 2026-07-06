from base64 import b64decode
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage

from marvis.template_reports import (
    TemplateTable,
    TemplateReportPayload,
    _column_widths_twips,
    find_placeholders,
    render_template_report,
)


def test_column_widths_never_negative_for_many_or_narrow_columns():
    # Many columns whose 360-twip floors already exceed the table width must not
    # produce a negative last-column width (invalid OOXML).
    for column_count in (40, 30, 12, 3, 1):
        table = TemplateTable(rows=[["x"] * column_count])
        widths = _column_widths_twips(table, column_count, table_width_twips=9000)
        assert len(widths) == column_count
        assert min(widths) >= 360, (column_count, widths)


PNG_BYTES = b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _save_template(path: Path, paragraphs: list[str]) -> Path:
    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    document.save(path)
    return path


def _document_xml(path: Path) -> ET.Element:
    with ZipFile(path) as archive:
        return ET.fromstring(archive.read("word/document.xml"))


def _first_table(path: Path) -> ET.Element:
    table = _document_xml(path).find(f".//{W_NS}tbl")
    assert table is not None
    return table


def _w_attr(element: ET.Element, name: str) -> str:
    return element.attrib[f"{W_NS}{name}"]


def _document_content_width(document: Document):
    section = document.sections[0]
    return section.page_width - section.left_margin - section.right_margin


def test_render_template_report_replaces_inline_and_repeated_text(tmp_path):
    template = _save_template(
        tmp_path / "template.docx",
        [
            "模型：{{TEXT:model_name}}",
            "{{TEXT:model_name}}验证结论：{{TEXT:final_validation_conclusion}}",
        ],
    )
    output = tmp_path / "out.docx"

    result = render_template_report(
        TemplateReportPayload(
            template_path=template,
            output_path=output,
            text_values={
                "TEXT:model_name": "贷前评分卡",
                "TEXT:final_validation_conclusion": "通过验证",
            },
            image_values={},
        )
    )

    generated = Document(result.output_path)
    text = "\n".join(paragraph.text for paragraph in generated.paragraphs)

    assert result.unresolved_placeholders == []
    assert "模型：贷前评分卡" in text
    assert "贷前评分卡验证结论：通过验证" in text
    assert "{{" not in text


def test_render_template_report_preserves_text_placeholder_run_formatting(tmp_path):
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("模型：")
    placeholder = paragraph.add_run("{{TEXT:model_name}}")
    placeholder.bold = True
    placeholder.font.name = "Arial"
    placeholder.font.size = Pt(16)
    placeholder.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    paragraph.add_run(" 验证报告")
    template = tmp_path / "template.docx"
    document.save(template)
    output = tmp_path / "out.docx"

    render_template_report(
        TemplateReportPayload(
            template_path=template,
            output_path=output,
            text_values={"TEXT:model_name": "贷前评分卡"},
            image_values={},
        )
    )

    generated = Document(output)
    replacement = next(
        run for run in generated.paragraphs[0].runs
        if run.text == "贷前评分卡"
    )

    assert generated.paragraphs[0].text == "模型：贷前评分卡 验证报告"
    assert replacement.bold is True
    assert replacement.font.name == "Arial"
    assert replacement.font.size == Pt(16)
    assert str(replacement.font.color.rgb) == "C00000"


def test_render_template_report_preserves_color_when_text_placeholder_is_split_across_runs(tmp_path):
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("训练说明：")
    paragraph.add_run("{{")
    placeholder_body = paragraph.add_run("TEXT:model_training_description")
    placeholder_body.bold = True
    placeholder_body.font.size = Pt(14)
    placeholder_body.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    paragraph.add_run("}}")
    paragraph.add_run("。")
    template = tmp_path / "template.docx"
    document.save(template)
    output = tmp_path / "out.docx"

    render_template_report(
        TemplateReportPayload(
            template_path=template,
            output_path=output,
            text_values={"TEXT:model_training_description": "固定算法介绍"},
            image_values={},
        )
    )

    generated = Document(output)
    replacement = next(
        run for run in generated.paragraphs[0].runs
        if run.text == "固定算法介绍"
    )

    assert generated.paragraphs[0].text == "训练说明：固定算法介绍。"
    assert replacement.bold is True
    assert replacement.font.size == Pt(14)
    assert str(replacement.font.color.rgb) == "1F4E79"


def test_render_template_report_inserts_image_at_image_placeholder(tmp_path):
    template = _save_template(tmp_path / "template.docx", ["图：{{IMAGE:ks_table}}"])
    image_path = tmp_path / "image.png"
    image_path.write_bytes(PNG_BYTES)
    output = tmp_path / "out.docx"

    result = render_template_report(
        TemplateReportPayload(
            template_path=template,
            output_path=output,
            text_values={},
            image_values={"IMAGE:ks_table": image_path},
        )
    )

    generated = Document(result.output_path)

    assert result.unresolved_placeholders == []
    assert len(generated.inline_shapes) == 1
    assert generated.inline_shapes[0].width == _document_content_width(generated)
    assert "{{IMAGE:ks_table}}" not in "\n".join(
        paragraph.text for paragraph in generated.paragraphs
    )


def test_render_template_report_keeps_generated_image_natural_width(tmp_path):
    template = _save_template(tmp_path / "template.docx", ["图：{{IMAGE:small_table}}"])
    image_path = tmp_path / "small_table.png"
    PILImage.new("RGB", (540, 180), "white").save(image_path, dpi=(180, 180))
    output = tmp_path / "out.docx"

    render_template_report(
        TemplateReportPayload(
            template_path=template,
            output_path=output,
            text_values={},
            image_values={"IMAGE:small_table": image_path},
        )
    )

    generated = Document(output)

    assert generated.inline_shapes[0].width == Inches(3.0)


def test_render_template_report_caps_wide_generated_image_to_page_width(tmp_path):
    template = _save_template(tmp_path / "template.docx", ["图：{{IMAGE:wide_table}}"])
    image_path = tmp_path / "wide_table.png"
    PILImage.new("RGB", (1800, 180), "white").save(image_path, dpi=(180, 180))
    output = tmp_path / "out.docx"

    render_template_report(
        TemplateReportPayload(
            template_path=template,
            output_path=output,
            text_values={},
            image_values={"IMAGE:wide_table": image_path},
        )
    )

    generated = Document(output)

    assert generated.inline_shapes[0].width == _document_content_width(generated)


def test_render_template_report_centers_wide_images_within_template_body_width(tmp_path):
    document = Document()
    paragraph = document.add_paragraph("{{IMAGE:wide_table}}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    template = tmp_path / "template.docx"
    document.save(template)
    image_path = tmp_path / "wide_table.png"
    PILImage.new("RGB", (1920, 240), "white").save(image_path, dpi=(180, 180))
    output = tmp_path / "out.docx"

    render_template_report(
        TemplateReportPayload(
            template_path=template,
            output_path=output,
            text_values={},
            image_values={"IMAGE:wide_table": image_path},
        )
    )

    generated = Document(output)

    assert generated.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert generated.inline_shapes[0].width == _document_content_width(generated)


def test_render_template_report_expands_multiple_images_at_one_placeholder(tmp_path):
    template = _save_template(
        tmp_path / "template.docx",
        ["图：{{IMAGE:pressure_score_shift}}"],
    )
    first_image = tmp_path / "first.png"
    second_image = tmp_path / "second.png"
    first_image.write_bytes(PNG_BYTES)
    second_image.write_bytes(PNG_BYTES)
    output = tmp_path / "out.docx"

    result = render_template_report(
        TemplateReportPayload(
            template_path=template,
            output_path=output,
            text_values={},
            image_values={"IMAGE:pressure_score_shift": [first_image, second_image]},
        )
    )

    generated = Document(result.output_path)

    assert result.unresolved_placeholders == []
    assert len(generated.inline_shapes) == 2
    assert "{{IMAGE:pressure_score_shift}}" not in "\n".join(
        paragraph.text for paragraph in generated.paragraphs
    )


def test_render_template_report_inserts_native_table_at_image_placeholder(tmp_path):
    template = _save_template(tmp_path / "template.docx", ["表：{{IMAGE:psi_table}}"])
    output = tmp_path / "out.docx"

    result = render_template_report(
        TemplateReportPayload(
            template_path=template,
            output_path=output,
            text_values={},
            image_values={},
            table_values={
                "IMAGE:psi_table": TemplateTable(
                    rows=[
                        ["分箱", "训练集样本量", "训练集PSI"],
                        ["1", "100", "0.0123"],
                    ]
                )
            },
        )
    )

    generated = Document(result.output_path)
    table_text = [
        [cell.text for cell in row.cells]
        for row in generated.tables[0].rows
    ]
    paragraph_text = "\n".join(paragraph.text for paragraph in generated.paragraphs)

    assert result.unresolved_placeholders == []
    assert table_text == [
        ["分箱", "训练集样本量", "训练集PSI"],
        ["1", "100", "0.0123"],
    ]
    assert "{{IMAGE:psi_table}}" not in paragraph_text
    assert "待插入图片" not in paragraph_text


def test_render_template_report_keeps_same_paragraph_table_order(tmp_path):
    template = _save_template(
        tmp_path / "template.docx",
        ["表：{{IMAGE:first_table}}{{IMAGE:second_table}}"],
    )
    output = tmp_path / "out.docx"

    render_template_report(
        TemplateReportPayload(
            template_path=template,
            output_path=output,
            text_values={},
            image_values={},
            table_values={
                "IMAGE:first_table": TemplateTable(rows=[["first"]]),
                "IMAGE:second_table": TemplateTable(rows=[["second"]]),
            },
        )
    )

    generated = Document(output)

    assert [table.cell(0, 0).text for table in generated.tables] == [
        "first",
        "second",
    ]


def test_render_template_report_formats_native_table_like_report_artifact(tmp_path):
    template = _save_template(tmp_path / "template.docx", ["表：{{IMAGE:psi_table}}"])
    output = tmp_path / "out.docx"

    render_template_report(
        TemplateReportPayload(
            template_path=template,
            output_path=output,
            text_values={},
            image_values={},
            table_values={
                "IMAGE:psi_table": TemplateTable(
                    rows=[
                        ["月份", "样本量", "逾期率"],
                        ["2026/04", "100", "5.00%"],
                    ]
                )
            },
        )
    )

    table = _first_table(output)
    table_properties = table.find(f"{W_NS}tblPr")
    assert table_properties is not None
    table_width = table_properties.find(f"{W_NS}tblW")
    table_layout = table_properties.find(f"{W_NS}tblLayout")
    table_borders = table_properties.find(f"{W_NS}tblBorders")
    assert table_width is not None
    assert table_layout is not None
    assert table_borders is not None
    assert _w_attr(table_width, "type") == "dxa"
    assert int(_w_attr(table_width, "w")) >= 8500
    assert _w_attr(table_layout, "type") == "fixed"
    assert len(table.findall(f"{W_NS}tblGrid/{W_NS}gridCol")) == 3

    header_cell_properties = table.find(f"{W_NS}tr/{W_NS}tc/{W_NS}tcPr")
    assert header_cell_properties is not None
    header_fill = header_cell_properties.find(f"{W_NS}shd")
    header_width = header_cell_properties.find(f"{W_NS}tcW")
    header_vertical_alignment = header_cell_properties.find(f"{W_NS}vAlign")
    assert header_fill is not None
    assert header_width is not None
    assert header_vertical_alignment is not None
    assert _w_attr(header_fill, "fill") == "C00000"
    assert _w_attr(header_vertical_alignment, "val") == "center"

    header_paragraph_properties = table.find(f"{W_NS}tr/{W_NS}tc/{W_NS}p/{W_NS}pPr")
    header_run_properties = table.find(f"{W_NS}tr/{W_NS}tc/{W_NS}p/{W_NS}r/{W_NS}rPr")
    assert header_paragraph_properties is not None
    assert header_run_properties is not None
    header_alignment = header_paragraph_properties.find(f"{W_NS}jc")
    header_spacing = header_paragraph_properties.find(f"{W_NS}spacing")
    header_fonts = header_run_properties.find(f"{W_NS}rFonts")
    header_size = header_run_properties.find(f"{W_NS}sz")
    header_color = header_run_properties.find(f"{W_NS}color")
    assert header_alignment is not None
    assert header_spacing is not None
    assert header_fonts is not None
    assert header_size is not None
    assert header_color is not None
    assert _w_attr(header_alignment, "val") == "center"
    assert _w_attr(header_spacing, "before") == "0"
    assert _w_attr(header_spacing, "after") == "0"
    assert _w_attr(header_fonts, "eastAsia") == "微软雅黑"
    assert _w_attr(header_size, "val") == "16"
    assert _w_attr(header_color, "val") == "FFFFFF"


def test_render_template_report_inserts_table_when_style_is_missing(tmp_path):
    template = _save_template(tmp_path / "template.docx", ["表：{{IMAGE:psi_table}}"])
    output = tmp_path / "out.docx"

    result = render_template_report(
        TemplateReportPayload(
            template_path=template,
            output_path=output,
            text_values={},
            image_values={},
            table_values={
                "IMAGE:psi_table": TemplateTable(
                    rows=[["分箱"], ["1"]],
                    style="Missing Table Style",
                )
            },
        )
    )

    generated = Document(result.output_path)

    assert len(generated.tables) == 1
    assert generated.tables[0].cell(0, 0).text == "分箱"


def test_render_template_report_reports_missing_text_placeholder(tmp_path):
    template = _save_template(tmp_path / "template.docx", ["KS={{TEXT:oot_ks}}"])
    output = tmp_path / "out.docx"

    result = render_template_report(
        TemplateReportPayload(
            template_path=template,
            output_path=output,
            text_values={},
            image_values={},
        )
    )

    assert result.unresolved_placeholders == ["{{TEXT:oot_ks}}"]


def test_render_template_report_reports_missing_image_and_table_placeholders(tmp_path):
    template = _save_template(
        tmp_path / "template.docx",
        [
            "图={{IMAGE:missing_chart}}",
            "表={{TABLE:missing_table}}",
        ],
    )
    output = tmp_path / "out.docx"

    result = render_template_report(
        TemplateReportPayload(
            template_path=template,
            output_path=output,
            text_values={},
            image_values={},
            table_values={},
        )
    )

    generated = Document(result.output_path)
    text = "\n".join(paragraph.text for paragraph in generated.paragraphs)

    assert result.unresolved_placeholders == [
        "{{IMAGE:missing_chart}}",
        "{{TABLE:missing_table}}",
    ]
    assert "待插入图片：IMAGE:missing_chart" in text
    assert "待插入表格：TABLE:missing_table" in text


def test_find_placeholders_reads_paragraphs_and_tables(tmp_path):
    document = Document()
    document.add_paragraph("正文 {{TEXT:model_name}}")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "表格 {{TEXT:oot_ks}}"
    template = tmp_path / "template.docx"
    document.save(template)

    placeholders = find_placeholders(template)

    assert placeholders == ["{{TEXT:model_name}}", "{{TEXT:oot_ks}}"]


def test_render_template_report_replaces_placeholders_in_later_table_cells(tmp_path):
    document = Document()
    table = document.add_table(rows=2, cols=5)
    for column in range(5):
        table.cell(0, column).text = f"标签 {column}"
        table.cell(1, column).text = f"{{{{TEXT:revision_{column}}}}}"
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    document.save(template)

    result = render_template_report(
        TemplateReportPayload(
            template_path=template,
            output_path=output,
            text_values={
                f"TEXT:revision_{column}": f"值 {column}"
                for column in range(5)
            },
            image_values={},
        )
    )

    generated = Document(output)
    table_text = "\n".join(
        paragraph.text
        for table in generated.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    placeholders = find_placeholders(output)
    assert result.unresolved_placeholders == []
    assert placeholders == []
    for column in range(5):
        assert f"值 {column}" in table_text
        assert f"{{{{TEXT:revision_{column}}}}}" not in table_text


def test_render_template_report_replaces_placeholders_inside_nested_tables(tmp_path):
    document = Document()
    outer = document.add_table(rows=1, cols=1)
    nested = outer.cell(0, 0).add_table(rows=1, cols=1)
    nested.cell(0, 0).text = "嵌套 {{TEXT:model_name}}"
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    document.save(template)

    result = render_template_report(
        TemplateReportPayload(
            template_path=template,
            output_path=output,
            text_values={"TEXT:model_name": "贷前评分卡"},
            image_values={},
        )
    )

    generated = Document(output)
    assert result.unresolved_placeholders == []
    assert "贷前评分卡" in "\n".join(
        paragraph.text
        for table in generated.tables
        for row in table.rows
        for cell in row.cells
        for nested_table in cell.tables
        for nested_row in nested_table.rows
        for nested_cell in nested_row.cells
        for paragraph in nested_cell.paragraphs
    )


def test_render_template_report_clears_non_run_paragraph_children(tmp_path):
    document = Document()
    paragraph = document.add_paragraph("模型：{{TEXT:model_name}}")
    bookmark = OxmlElement("w:bookmarkStart")
    bookmark.set(qn("w:id"), "0")
    bookmark.set(qn("w:name"), "bookmark_to_remove")
    paragraph._p.append(bookmark)
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    document.save(template)

    result = render_template_report(
        TemplateReportPayload(
            template_path=template,
            output_path=output,
            text_values={"TEXT:model_name": "贷前评分卡"},
            image_values={},
        )
    )

    document_xml = _document_xml(output)
    assert result.unresolved_placeholders == []
    assert document_xml.find(f".//{W_NS}bookmarkStart") is None
