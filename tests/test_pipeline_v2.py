import json
import math
import shutil
import sys
from dataclasses import replace
from pathlib import Path
from types import SimpleNamespace

import nbformat
import pandas as pd
import pytest
from docx import Document

from marvis.agent_memory.store import AgentMemoryStore
from marvis.db import TaskRepository, init_db
from marvis.domain import FileArtifact, FileRole, TaskCreate, TaskStatus
from marvis.notebook_contract import RuntimeContract
from marvis.notebook_cancellation import request_notebook_cancellation
from marvis.notebooks import close_live_notebook_session, register_live_notebook_session
from marvis import pipeline as pipeline_module
from marvis.pipeline import (
    LEGACY_LIVE_NOTEBOOK_ENV_VAR,
    NOTEBOOK_STAGE_FAILURE_PREFIX,
    REPORT_STAGE_FAILURE_PREFIX,
    REPRODUCIBILITY_RESULT_JSON,
    SCAN_STAGE_FAILURE_PREFIX,
    PipelineError,
    PipelineSettings,
    _build_metrics_cell_source,
    _clear_generated_artifacts,
    _feature_columns,
    _load_sample,
    _required_path,
    _scan_artifacts,
    _stage_failure_message,
    _write_metrics_results_in_session,
    _write_reproducibility_result_in_session,
    run_metrics_stage,
    run_notebook_stage,
    run_pipeline,
    run_report_stage,
    run_staged_pipeline,
)


FIXTURES = Path(__file__).resolve().parent / "fixtures"


def _allow_legacy_live_notebook(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setenv(LEGACY_LIVE_NOTEBOOK_ENV_VAR, "1")


def test_legacy_live_notebook_execution_requires_process_env(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
):
    settings = PipelineSettings(
        workspace=tmp_path,
        db_path=tmp_path / "marvis.sqlite",
        report_template_path=tmp_path / "template.docx",
        notebook_isolated_execution=False,
        allow_legacy_live_notebook_execution=True,
    )

    monkeypatch.delenv(LEGACY_LIVE_NOTEBOOK_ENV_VAR, raising=False)
    assert not pipeline_module.legacy_live_notebook_execution_allowed(settings)

    monkeypatch.setenv(LEGACY_LIVE_NOTEBOOK_ENV_VAR, "yes")
    assert pipeline_module.legacy_live_notebook_execution_allowed(settings)


def test_scan_artifacts_tags_limit_breach_as_scan_failure(monkeypatch):
    # A scan-limit ValueError (max_files / max_depth) raised mid-pipeline must be
    # wrapped as a PipelineError carrying the scan-stage prefix, so downstream
    # stage handlers do not mislabel it as a notebook failure.
    def _raise_limit(_path):
        raise ValueError("source_dir has too many files: max_files=2000")

    monkeypatch.setattr(pipeline_module, "scan_source_dir", _raise_limit)
    task = SimpleNamespace(source_dir="/tmp/whatever")

    try:
        _scan_artifacts(task)
    except PipelineError as exc:
        assert str(exc).startswith(SCAN_STAGE_FAILURE_PREFIX)
        assert "too many files" in str(exc)
    else:  # pragma: no cover - defensive
        raise AssertionError("expected PipelineError")


def test_memory_failure_kind_classifies_scan_limit_failures():
    kind = pipeline_module._memory_failure_kind
    assert kind(f"{SCAN_STAGE_FAILURE_PREFIX}source_dir has too many files: max_files=2000",
                default="notebook") == "scan"
    assert kind("source dir invalid: too deep", default="notebook") == "scan"
    # unrelated failures keep their existing classification
    assert kind("PMML mismatch", default="notebook") == "pmml"
    assert kind("live notebook kernel is not available", default="notebook") == "notebook"


def test_stage_failure_message_preserves_scan_attribution():
    scan_message = f"{SCAN_STAGE_FAILURE_PREFIX}source_dir is too deep"
    # Bubbling through the notebook stage handler keeps the scan prefix intact...
    assert (
        _stage_failure_message(NOTEBOOK_STAGE_FAILURE_PREFIX, scan_message) == scan_message
    )
    # ...while a genuine notebook failure still receives the notebook prefix.
    assert _stage_failure_message(NOTEBOOK_STAGE_FAILURE_PREFIX, "kernel died") == (
        f"{NOTEBOOK_STAGE_FAILURE_PREFIX}kernel died"
    )


def test_stage_artifact_cleanup_invalidates_downstream_outputs(tmp_path: Path):
    task_dir = tmp_path / "tasks" / "t1"
    execution_dir = task_dir / "execution"
    outputs_dir = task_dir / "outputs"
    images_dir = task_dir / "images"
    for directory in (execution_dir, outputs_dir, images_dir):
        directory.mkdir(parents=True)
    (execution_dir / "scan_result.json").write_text("{}", encoding="utf-8")
    (execution_dir / "runtime_contract.json").write_text("{}", encoding="utf-8")
    (outputs_dir / "reproducibility_result.json").write_text("{}", encoding="utf-8")
    (outputs_dir / "validation_results.json").write_text("{}", encoding="utf-8")
    (outputs_dir / "validation.xlsx").write_bytes(b"xlsx")
    (outputs_dir / "validation_report.docx").write_bytes(b"docx")
    (images_dir / "chart.png").write_bytes(b"png")

    _clear_generated_artifacts(task_dir, stage="notebook")

    assert (execution_dir / "scan_result.json").exists()
    assert not (execution_dir / "runtime_contract.json").exists()
    assert not (outputs_dir / "reproducibility_result.json").exists()
    assert not (outputs_dir / "validation_results.json").exists()
    assert not (outputs_dir / "validation.xlsx").exists()
    assert not (outputs_dir / "validation_report.docx").exists()
    assert not images_dir.exists()


def _sigmoid(value: float) -> float:
    return 1 / (1 + math.exp(-value))


def _write_contract_notebook(path: Path) -> None:
    notebook = nbformat.v4.new_notebook(
        cells=[
            nbformat.v4.new_code_cell(
                "\n".join(
                    [
                        "import pandas as pd",
                        "RMC_SAMPLE_DF = pd.read_csv('sample.csv')",
                        "RMC_TARGET_COL = 'y'",
                        "RMC_ALGORITHM = 'lgb'",
                        "RMC_SPLIT_COL = 'split'",
                        "RMC_TIME_COL = 'apply_month'",
                        "def RMC_SCORE_FN(df):",
                        "    return df['pred']",
                        "RMC_FEATURE_IMPORTANCE = pd.DataFrame({",
                        "    'feature': ['x1', 'x2'],",
                        "    'importance': [0.8, 0.2],",
                        "})",
                        "RMC_MODEL_PARAMS = {'learning_rate': 0.05, 'max_depth': 5}",
                    ]
                )
            )
        ]
    )
    nbformat.write(notebook, path)


def _write_live_sample_contract_notebook(path: Path) -> None:
    rows_source = repr(
        [
            {
                "x1": (index + 1) / 41,
                "x2": 0.0,
                "pred": _sigmoid((index + 1) / 41),
                "y": int(index >= 20),
                "split": split,
                "apply_month": "202503" if split == "train" else "202505",
            }
            for split in ("train", "test", "oot")
            for index in range(40)
        ]
    )
    notebook = nbformat.v4.new_notebook(
        cells=[
            nbformat.v4.new_code_cell(
                "\n".join(
                    [
                        "from pathlib import Path",
                        "import pandas as pd",
                        "counter = Path('notebook_run_count.txt')",
                        "counter.write_text(str(int(counter.read_text()) + 1 if counter.exists() else 1))",
                        f"RMC_SAMPLE_DF = pd.DataFrame({rows_source})",
                        "RMC_TARGET_COL = 'y'",
                        "RMC_ALGORITHM = 'lgb'",
                        "RMC_SPLIT_COL = 'split'",
                        "RMC_TIME_COL = 'apply_month'",
                        "def RMC_SCORE_FN(df):",
                        "    return df['pred']",
                        "RMC_FEATURE_IMPORTANCE = pd.DataFrame({",
                        "    'feature': ['x1', 'x2'],",
                        "    'importance': [0.8, 0.2],",
                        "})",
                        "RMC_MODEL_PARAMS = {'learning_rate': 0.05, 'max_depth': 5}",
                    ]
                )
            )
        ]
    )
    nbformat.write(notebook, path)


def _build_project(tmp_path: Path) -> Path:
    project = tmp_path / "project"
    project.mkdir()
    rows = []
    for split in ("train", "test", "oot"):
        for index in range(40):
            x1 = (index + 1) / 41
            x2 = 0.0
            rows.append(
                {
                    "x1": x1,
                    "x2": x2,
                    "pred": _sigmoid(x1),
                    "y": int(index >= 20),
                    "split": split,
                    "apply_month": "202503" if split == "train" else "202505",
                }
            )
    pd.DataFrame(rows).to_csv(project / "sample.csv", index=False)
    shutil.copy(FIXTURES / "min_lr.pmml", project / "fr_final.pmml")
    pd.DataFrame(
        {
            "特征名": ["x1", "x2"],
            "类别": ["征信", "基础信息"],
        }
    ).to_excel(project / "data_dictionary.xlsx", index=False)
    _write_contract_notebook(project / "dev.ipynb")
    return project


def _write_template(path: Path) -> Path:
    document = Document()
    document.add_paragraph("{{TEXT:report_title}}")
    document.add_paragraph("OOT KS：{{TEXT:oot_ks}}")
    document.add_paragraph("{{TEXT:reproducibility_summary}}")
    document.add_paragraph("{{IMAGE:overall_model_effect}}")
    document.add_paragraph("{{IMAGE:pressure_ks_table}}")
    document.save(path)
    return path


class _PredColumnScorer:
    def score(self, df: pd.DataFrame) -> list[float]:
        return df["pred"].astype(float).tolist()


def test_metrics_cell_reuses_notebook_scorer_and_saved_reproducibility(tmp_path: Path):
    repo = TaskRepository(tmp_path / "marvis.sqlite")
    init_db(repo.db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(tmp_path),
        )
    )
    source = _build_metrics_cell_source(
        package_root=tmp_path,
        task=task,
        settings=PipelineSettings(
            workspace=tmp_path,
            db_path=repo.db_path,
            report_template_path=tmp_path / "template.docx",
        ),
        dictionary_path=tmp_path / "dictionary.csv",
        input_pmml_path=tmp_path / "model.pmml",
        contract=RuntimeContract(
            target_col="y",
            split_col="split",
            time_col="apply_month",
            pmml_output_field="probability_1",
            score_decimal_places=6,
            code_model_scores_path=tmp_path / "code_model_scores.csv",
            feature_importance_path=None,
            model_params_path=None,
            algorithm="lgb",
        ),
        model_meta_path=tmp_path / "model_meta.json",
        reproducibility_json_path=tmp_path / "reproducibility_result.json",
        results_json_path=tmp_path / "validation_results.json",
        excel_path=tmp_path / "validation.xlsx",
    )

    assert "class _RmcNotebookScorer" in source
    assert "RMC_SCORE_FN" in source
    assert "reproducibility_json_path" in source
    assert "load_pmml_scorer" not in source
    assert "def _rmc_raise_if_metrics_cancelled()" in source
    assert "cancellation_check=_rmc_raise_if_metrics_cancelled" in source
    assert "dataframe.index.equals(_rmc_sample.index)" in source
    assert "import pickle" not in source
    assert "results_pickle_path" not in source
    assert f"_rmc_package_root = {tmp_path.as_posix()!r}" in source


def test_v1_validation_appended_policy_matches_generated_cell_kinds(tmp_path: Path):
    repo = TaskRepository(tmp_path / "marvis.sqlite")
    init_db(repo.db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(tmp_path),
        )
    )
    settings = PipelineSettings(
        workspace=tmp_path,
        db_path=repo.db_path,
        report_template_path=tmp_path / "template.docx",
    )
    contract = RuntimeContract(
        target_col="y",
        split_col="split",
        time_col="apply_month",
        pmml_output_field="probability_1",
        score_decimal_places=6,
        code_model_scores_path=tmp_path / "code_model_scores.csv",
        feature_importance_path=None,
        model_params_path=None,
        algorithm="lgb",
    )

    reproducibility_sources = dict(
        pipeline_module._build_reproducibility_cell_sources(
            package_root=tmp_path,
            task=task,
            settings=settings,
            input_pmml_path=tmp_path / "model.pmml",
            contract_meta_path=tmp_path / "runtime_contract.json",
            output_path=tmp_path / "reproducibility_result.json",
        )
    )
    reproducibility_kinds = set(reproducibility_sources)
    metrics_kinds = {
        kind
        for kind, _source in pipeline_module._build_metrics_cell_sources(
            package_root=tmp_path,
            task=task,
            settings=settings,
            dictionary_path=tmp_path / "dictionary.csv",
            input_pmml_path=tmp_path / "model.pmml",
            contract=contract,
            model_meta_path=tmp_path / "model_meta.json",
            reproducibility_json_path=tmp_path / "reproducibility_result.json",
            results_json_path=tmp_path / "validation_results.json",
            excel_path=tmp_path / "validation.xlsx",
        )
    }
    generated_kinds = reproducibility_kinds | metrics_kinds

    assert generated_kinds == set(pipeline_module.V1_VALIDATION_APPENDED_CELL_KINDS)
    assert (
        set(
            pipeline_module.V1_VALIDATION_APPENDED_EXECUTION_POLICY.allowed_marvis_kinds
        )
        == generated_kinds
    )
    assert "_rmc_consistency_status(" in reproducibility_sources["repro-compare"]
    assert "PASS if _rmc_mismatch_count == 0" not in reproducibility_sources["repro-compare"]


def test_metrics_cell_handles_null_split_and_time_columns_in_history(tmp_path: Path):
    repo = TaskRepository(tmp_path / "marvis.sqlite")
    init_db(repo.db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(tmp_path),
        )
    )
    task = replace(task, split_col=None, time_col=None)

    source = _build_metrics_cell_source(
        package_root=tmp_path,
        task=task,
        settings=PipelineSettings(
            workspace=tmp_path,
            db_path=repo.db_path,
            report_template_path=tmp_path / "template.docx",
        ),
        dictionary_path=tmp_path / "dictionary.csv",
        input_pmml_path=tmp_path / "model.pmml",
        contract=RuntimeContract(
            target_col="y",
            split_col=None,
            time_col=None,
            pmml_output_field="probability_1",
            score_decimal_places=6,
            code_model_scores_path=tmp_path / "code_model_scores.csv",
            feature_importance_path=None,
            model_params_path=None,
            algorithm="lgb",
        ),
        model_meta_path=tmp_path / "model_meta.json",
        reproducibility_json_path=tmp_path / "reproducibility_result.json",
        results_json_path=tmp_path / "validation_results.json",
        excel_path=tmp_path / "validation.xlsx",
    )

    assert '"split_col": ""' in source
    assert '"time_col": ""' in source
    assert "if column and column not in _rmc_sample.columns" in source


def test_metrics_cell_uses_runtime_contract_algorithm_not_create_task_placeholder(tmp_path: Path):
    repo = TaskRepository(tmp_path / "marvis.sqlite")
    init_db(repo.db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(tmp_path),
            algorithm="",
        )
    )

    source = _build_metrics_cell_source(
        package_root=tmp_path,
        task=task,
        settings=PipelineSettings(
            workspace=tmp_path,
            db_path=repo.db_path,
            report_template_path=tmp_path / "template.docx",
        ),
        dictionary_path=tmp_path / "dictionary.csv",
        input_pmml_path=tmp_path / "model.pmml",
        contract=RuntimeContract(
            target_col="y",
            split_col="split",
            time_col="apply_month",
            pmml_output_field="probability_1",
            score_decimal_places=6,
            code_model_scores_path=tmp_path / "code_model_scores.csv",
            feature_importance_path=None,
            model_params_path=None,
            algorithm="CatBoostClassifier",
        ),
        model_meta_path=tmp_path / "model_meta.json",
        reproducibility_json_path=tmp_path / "reproducibility_result.json",
        results_json_path=tmp_path / "validation_results.json",
        excel_path=tmp_path / "validation.xlsx",
    )

    assert '"algorithm": "catboost"' in source


def test_reproducibility_stage_shows_pmml_scoring_and_compare_progress(tmp_path: Path):
    repo = TaskRepository(tmp_path / "marvis.sqlite")
    init_db(repo.db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(tmp_path),
        )
    )
    contract_path = tmp_path / "runtime_contract.json"
    contract_path.write_text(
        json.dumps(
            {
                "target_col": "y",
                "split_col": "split",
                "time_col": "apply_month",
                "pmml_output_field": "probability_1",
                "score_decimal_places": 6,
                "code_model_scores_path": str(tmp_path / "code_scores.csv"),
            }
        ),
        encoding="utf-8",
    )
    output_path = tmp_path / "reproducibility_result.json"
    calls = []

    class FakeSession:
        def append_code_cell(self, source, **kwargs):
            calls.append(("append", kwargs["metadata"]["marvis"]))
            return len(calls) - 1

        def execute_existing_code_cell(self, cell_index, **kwargs):
            calls.append(("execute", cell_index))
            if cell_index == 1:
                output_path.write_text("{}", encoding="utf-8")
            return SimpleNamespace(succeeded=True, cancelled=False)

    _write_reproducibility_result_in_session(
        session=FakeSession(),
        task=task,
        settings=PipelineSettings(
            workspace=tmp_path,
            db_path=repo.db_path,
            report_template_path=tmp_path / "template.docx",
        ),
        input_pmml_path=tmp_path / "model.pmml",
        contract_meta_path=contract_path,
        output_path=output_path,
    )

    assert calls == [
        ("append", "repro-pmml"),
        ("append", "repro-compare"),
        ("execute", 0),
        ("execute", 1),
    ]


def test_metrics_stage_shows_named_internal_progress_steps(tmp_path: Path):
    repo = TaskRepository(tmp_path / "marvis.sqlite")
    init_db(repo.db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(tmp_path),
        )
    )
    outputs_dir = tmp_path / "outputs"
    outputs_dir.mkdir()
    calls = []

    class FakeSession:
        def append_code_cell(self, source, **kwargs):
            calls.append(("append", kwargs["metadata"]["marvis"]))
            return len(calls) - 1

        def execute_existing_code_cell(self, cell_index, **kwargs):
            calls.append(("execute", cell_index))
            if cell_index == 7:
                (outputs_dir / "validation_results.json").write_text("{}", encoding="utf-8")
                (outputs_dir / "validation.xlsx").write_bytes(b"xlsx")
            return SimpleNamespace(succeeded=True, cancelled=False)

    _write_metrics_results_in_session(
        session=FakeSession(),
        task=task,
        settings=PipelineSettings(
            workspace=tmp_path,
            db_path=repo.db_path,
            report_template_path=tmp_path / "template.docx",
        ),
        dictionary_path=tmp_path / "dictionary.csv",
        input_pmml_path=tmp_path / "model.pmml",
        contract=RuntimeContract(
            target_col="y",
            split_col="split",
            time_col="apply_month",
            pmml_output_field="probability_1",
            score_decimal_places=6,
            code_model_scores_path=tmp_path / "code_model_scores.csv",
            feature_importance_path=None,
            model_params_path=None,
            algorithm="lgb",
        ),
        model_meta_path=tmp_path / "model_meta.json",
        outputs_dir=outputs_dir,
    )

    assert calls == [
        ("append", "metrics-prepare"),
        ("append", "metrics-score"),
        ("append", "metrics-basic"),
        ("append", "metrics-ks"),
        ("append", "metrics-psi"),
        ("append", "metrics-binning"),
        ("append", "metrics-stress"),
        ("append", "metrics-output"),
        ("execute", 0),
        ("execute", 1),
        ("execute", 2),
        ("execute", 3),
        ("execute", 4),
        ("execute", 5),
        ("execute", 6),
        ("execute", 7),
    ]


@pytest.mark.slow
def test_pipeline_end_to_end(tmp_path: Path):
    project = _build_project(tmp_path)
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    template = _write_template(tmp_path / "template.docx")
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(project),
            algorithm="lgb",
            target_col="y",
            score_col="pred",
            split_col="split",
            time_col="apply_month",
            report_values={"TEXT:report_title": "人工确认标题"},
        )
    )

    run_pipeline(
        task_id=task.id,
        settings=PipelineSettings(
            workspace=workspace,
            db_path=db_path,
            report_template_path=template,
            feature_columns=["x1", "x2"],
            random_sample_size=12,
        ),
    )

    final = repo.get_task(task.id)
    task_dir = workspace / "tasks" / task.id
    assert final.status == TaskStatus.SUCCEEDED
    assert (task_dir / "execution" / "code_model_scores.csv").exists()
    assert (task_dir / "execution" / "runtime_contract.json").exists()
    assert (task_dir / "execution" / "notebook_steps.json").exists()
    assert (task_dir / "execution" / "model_meta.json").exists()
    assert (task_dir / "outputs" / "validation.xlsx").exists()
    report_path = task_dir / "outputs" / "validation_report.docx"
    assert report_path.exists()
    assert Document(report_path).paragraphs[0].text == "人工确认标题"
    result_json = json.loads(
        (task_dir / "outputs" / "validation_results.json").read_text(encoding="utf-8")
    )
    assert result_json["reproducibility"]["summary"]["status"] == "pass"


def test_notebook_stage_writes_reproducibility_evidence_before_metrics(
    tmp_path: Path,
    monkeypatch,
):
    _allow_legacy_live_notebook(monkeypatch)
    project = _build_project(tmp_path)
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(project),
            algorithm="lgb",
            target_col="y",
            score_col="pred",
            split_col="split",
            time_col="apply_month",
        )
    )
    repo.update_status(task.id, TaskStatus.SCANNED, message="source scanned")
    repo.update_status(task.id, TaskStatus.RUNNING, message="notebook queued")

    def fake_notebook_step_v3(
        *,
        repo,
        task,
        contract_meta_path,
        code_scores_path,
        **_kwargs,
    ):
        assert _kwargs["mark_executed"] is False
        sample = pd.read_csv(project / "sample.csv")
        pd.DataFrame(
            {
                "row_index": sample.index,
                "code_model_score": sample["pred"],
            }
        ).to_csv(
            code_scores_path,
            index=False,
        )
        contract_meta_path.write_text(
            json.dumps(
                {
                    "target_col": "y",
                    "split_col": "split",
                    "time_col": "apply_month",
                    "pmml_output_field": "probability_1",
                    "score_decimal_places": 6,
                    "code_model_scores_path": str(code_scores_path),
                    "feature_importance_path": None,
                    "model_params_path": None,
                    "algorithm": "lgb",
                }
            ),
            encoding="utf-8",
        )
        return SimpleNamespace(closed=False, close=lambda: None)

    monkeypatch.setattr("marvis.pipeline._notebook_step_v3", fake_notebook_step_v3)

    def fake_write_reproducibility_result_in_session(*, output_path, settings, **_kwargs):
        assert repo.get_task(task.id).status is TaskStatus.RUNNING
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(
            json.dumps(
                {
                    "sample_size": settings.random_sample_size,
                    "seed": settings.random_seed,
                    "rows": [
                        {
                            "row_index": 0,
                            "score_code_model": 0.1,
                            "score_submitted_pmml": 0.1,
                            "abs_diff": 0.0,
                            "matched": True,
                        }
                    ],
                    "summary": {
                        "match_count": 1,
                        "mismatch_count": 0,
                        "max_abs_diff": 0.0,
                        "status": "pass",
                    },
                }
            ),
            encoding="utf-8",
        )

    monkeypatch.setattr(
        "marvis.pipeline._write_reproducibility_result_in_session",
        fake_write_reproducibility_result_in_session,
    )

    run_notebook_stage(
        task_id=task.id,
        settings=PipelineSettings(
            workspace=workspace,
            db_path=db_path,
            report_template_path=tmp_path / "template.docx",
            random_sample_size=2,
            notebook_isolated_execution=False,
            allow_legacy_live_notebook_execution=True,
        ),
        stage_claimed=True,
    )

    evidence_path = workspace / "tasks" / task.id / "outputs" / "reproducibility_result.json"
    evidence = json.loads(evidence_path.read_text(encoding="utf-8"))
    assert evidence["summary"]["status"] == "pass"
    assert evidence["sample_size"] == 2
    assert evidence["rows"][0]["score_code_model"] == evidence["rows"][0]["score_submitted_pmml"]
    assert repo.get_task(task.id).status is TaskStatus.EXECUTED


def test_notebook_stage_can_write_reproducibility_in_isolated_worker(
    tmp_path: Path,
    monkeypatch,
):
    project = tmp_path / "project"
    project.mkdir()
    notebook_path = project / "dev.ipynb"
    sample_path = project / "sample.csv"
    pmml_path = project / "model.pmml"
    nbformat.write(nbformat.v4.new_notebook(cells=[]), notebook_path)
    sample_path.write_text("x1,pred,y,split,apply_month\n1,0.1,0,train,202501\n")
    pmml_path.write_text("<PMML/>", encoding="utf-8")

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(project),
            notebook_path=str(notebook_path),
            sample_path=str(sample_path),
            pmml_path=str(pmml_path),
            algorithm="lgb",
        )
    )
    repo.update_status(task.id, TaskStatus.SCANNED, message="source scanned")
    repo.update_status(task.id, TaskStatus.RUNNING, message="notebook queued")

    def fake_notebook_step_v3(
        *,
        contract_meta_path,
        code_scores_path,
        extra_code_cells,
        **kwargs,
    ):
        assert kwargs["keep_alive"] is False
        assert kwargs["isolated"] is True
        assert kwargs["mark_executed"] is False
        assert [kind for kind, _source in extra_code_cells] == [
            "repro-pmml",
            "repro-compare",
        ]
        contract_meta_path.write_text(
            json.dumps(
                {
                    "target_col": "y",
                    "split_col": "split",
                    "time_col": "apply_month",
                    "pmml_output_field": "probability_1",
                    "score_decimal_places": 6,
                    "code_model_scores_path": str(code_scores_path),
                    "feature_importance_path": None,
                    "model_params_path": None,
                    "algorithm": "lgb",
                }
            ),
            encoding="utf-8",
        )
        code_scores_path.write_text("row_index,code_model_score\n0,0.1\n")
        output_path = (
            workspace
            / "tasks"
            / task.id
            / "outputs"
            / REPRODUCIBILITY_RESULT_JSON
        )
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(
            json.dumps({"summary": {"status": "pass"}, "rows": []}),
            encoding="utf-8",
        )

    monkeypatch.setattr("marvis.pipeline._notebook_step_v3", fake_notebook_step_v3)

    run_notebook_stage(
        task_id=task.id,
        settings=PipelineSettings(
            workspace=workspace,
            db_path=db_path,
            report_template_path=tmp_path / "template.docx",
        ),
        stage_claimed=True,
    )

    assert repo.get_task(task.id).status is TaskStatus.EXECUTED


def test_metrics_stage_reruns_isolated_notebook_even_with_stale_live_session(
    tmp_path: Path,
    monkeypatch,
):
    project = tmp_path / "project"
    project.mkdir()
    notebook_path = project / "dev.ipynb"
    sample_path = project / "sample.csv"
    pmml_path = project / "model.pmml"
    dictionary_path = project / "dictionary.csv"
    nbformat.write(nbformat.v4.new_notebook(cells=[]), notebook_path)
    sample_path.write_text("x1,pred,y,split,apply_month\n1,0.1,0,train,202501\n")
    pmml_path.write_text("<PMML/>", encoding="utf-8")
    dictionary_path.write_text("特征名,类别\nx1,基础信息\n", encoding="utf-8")

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(project),
            notebook_path=str(notebook_path),
            sample_path=str(sample_path),
            pmml_path=str(pmml_path),
            dictionary_path=str(dictionary_path),
            algorithm="lgb",
        )
    )
    repo.update_status(task.id, TaskStatus.SCANNED, message="source scanned")
    repo.update_status(task.id, TaskStatus.RUNNING, message="notebook queued")
    repo.update_status(task.id, TaskStatus.EXECUTED, message="notebook executed")
    repo.update_status(task.id, TaskStatus.COMPUTING_METRICS, message="metrics queued")
    task_dir = workspace / "tasks" / task.id
    execution_dir = task_dir / "execution"
    outputs_dir = task_dir / "outputs"
    execution_dir.mkdir(parents=True)
    outputs_dir.mkdir(parents=True)
    code_scores_path = execution_dir / "code_model_scores.csv"
    code_scores_path.write_text("row_index,code_model_score\n0,0.1\n")
    (execution_dir / "runtime_contract.json").write_text(
        json.dumps(
            {
                "target_col": "y",
                "split_col": "split",
                "time_col": "apply_month",
                "pmml_output_field": "probability_1",
                "score_decimal_places": 6,
                "code_model_scores_path": str(code_scores_path),
                "feature_importance_path": None,
                "model_params_path": None,
                "algorithm": "lgb",
            }
        ),
        encoding="utf-8",
    )
    (outputs_dir / REPRODUCIBILITY_RESULT_JSON).write_text(
        json.dumps({"summary": {"status": "pass"}, "rows": []}),
        encoding="utf-8",
    )
    live_closed = {"value": False}
    register_live_notebook_session(
        task.id,
        SimpleNamespace(
            closed=False,
            close=lambda: live_closed.__setitem__("value", True),
        ),
    )

    def fake_notebook_step_v3(*, extra_code_cells, **kwargs):
        assert kwargs["keep_alive"] is False
        assert kwargs["isolated"] is True
        assert kwargs["stage_claimed"] is True
        assert kwargs["notebook_steps_path"] == execution_dir / "metrics_steps.json"
        assert [kind for kind, _source in extra_code_cells] == [
            "metrics-prepare",
            "metrics-score",
            "metrics-basic",
            "metrics-ks",
            "metrics-psi",
            "metrics-binning",
            "metrics-stress",
            "metrics-output",
        ]
        metrics_work_dir = outputs_dir / ".metrics-stage-work"
        (metrics_work_dir / "validation_results.json").write_text("{}", encoding="utf-8")
        (metrics_work_dir / "validation.xlsx").write_bytes(b"xlsx")

    monkeypatch.setattr("marvis.pipeline._notebook_step_v3", fake_notebook_step_v3)

    run_metrics_stage(
        task_id=task.id,
        settings=PipelineSettings(
            workspace=workspace,
            db_path=db_path,
            report_template_path=tmp_path / "template.docx",
        ),
        stage_claimed=True,
    )

    assert live_closed["value"] is True
    assert repo.get_task(task.id).status is TaskStatus.WRITING_ARTIFACTS


def test_metrics_stage_marks_sample_column_failure_as_metrics_failure(
    tmp_path: Path,
    monkeypatch,
):
    _allow_legacy_live_notebook(monkeypatch)
    project = tmp_path / "project"
    project.mkdir()
    sample_path = project / "sample.csv"
    pmml_path = project / "model.pmml"
    dictionary_path = project / "dictionary.csv"
    sample_path.write_text(
        "y,split,apply_month,pred\n0,train,202501,0.1\n",
        encoding="utf-8",
    )
    pmml_path.write_text("<PMML />", encoding="utf-8")
    dictionary_path.write_text("特征名,类别\nx1,数值\n", encoding="utf-8")

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(project),
            sample_path=str(sample_path),
            pmml_path=str(pmml_path),
            dictionary_path=str(dictionary_path),
        )
    )
    repo.update_status(task.id, TaskStatus.SCANNED, message="source scanned")
    repo.update_status(task.id, TaskStatus.RUNNING, message="notebook queued")
    repo.update_status(task.id, TaskStatus.EXECUTED, message="notebook executed")
    repo.update_status(task.id, TaskStatus.COMPUTING_METRICS, message="metrics queued")
    execution_dir = workspace / "tasks" / task.id / "execution"
    execution_dir.mkdir(parents=True)
    code_scores_path = execution_dir / "code_model_scores.csv"
    code_scores_path.write_text("row_index,code_model_score\n0,0.1\n", encoding="utf-8")
    (execution_dir / "runtime_contract.json").write_text(
        json.dumps(
            {
                "target_col": "y",
                "split_col": "new_flag",
                "time_col": "apply_month",
                "pmml_output_field": "probability_1",
                "score_decimal_places": 6,
                "code_model_scores_path": str(code_scores_path),
                "feature_importance_path": None,
                "model_params_path": None,
                "algorithm": "lgb",
            }
        ),
        encoding="utf-8",
    )
    fake_session = SimpleNamespace(
        closed=False,
        execute_code_cell=lambda source, **_kwargs: SimpleNamespace(
            succeeded=False,
            failed_cell_index=14,
            error_name="ValueError",
            error_value="sample column check failed: split_col='new_flag'",
        ),
        close=lambda: None,
    )
    register_live_notebook_session(task.id, fake_session)

    try:
        run_metrics_stage(
            task_id=task.id,
            settings=PipelineSettings(
                workspace=workspace,
                db_path=db_path,
                report_template_path=tmp_path / "template.docx",
                notebook_isolated_execution=False,
                allow_legacy_live_notebook_execution=True,
            ),
            stage_claimed=True,
        )
    except PipelineError as exc:
        assert "sample column check failed" in str(exc)
    else:
        raise AssertionError("metrics stage should fail when split column is missing")

    final = repo.get_task(task.id)
    assert final.status == TaskStatus.FAILED
    assert final.status_message.startswith("模型效果&稳定性验证失败：")
    assert "split_col='new_flag'" in final.status_message
    store = AgentMemoryStore(db_path)
    pitfalls = store.list_entries(memory_type="validation_pitfall")
    assert len(pitfalls) == 1
    assert pitfalls[0].payload["failure_kind"] == "field"
    task_memories = store.list_entries(memory_type="task_experience")
    assert len(task_memories) == 1
    assert task_memories[0].payload["status"] == "failed"
    close_live_notebook_session(task.id)


def test_metrics_stage_success_captures_model_experience_memory(
    tmp_path: Path,
    monkeypatch,
):
    _allow_legacy_live_notebook(monkeypatch)
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    project = tmp_path / "project"
    project.mkdir()
    sample_path = project / "sample.csv"
    pmml_path = project / "model.pmml"
    dictionary_path = project / "dictionary.csv"
    sample_path.write_text("y,score,split,apply_month\n0,0.1,train,202601\n", encoding="utf-8")
    pmml_path.write_text("<PMML/>", encoding="utf-8")
    dictionary_path.write_text("特征名,类别\nx1,征信\n", encoding="utf-8")
    task = repo.create_task(
        TaskCreate(
            model_name="A卡模型",
            model_version="V2026",
            validator="qa",
            source_dir=str(project),
            sample_path=str(sample_path),
            pmml_path=str(pmml_path),
            dictionary_path=str(dictionary_path),
            algorithm="lgb",
        )
    )
    repo.update_status(task.id, TaskStatus.SCANNED, message="source scanned")
    repo.update_status(task.id, TaskStatus.RUNNING, message="notebook queued")
    repo.update_status(task.id, TaskStatus.EXECUTED, message="notebook executed")
    repo.update_status(task.id, TaskStatus.COMPUTING_METRICS, message="metrics queued")
    execution_dir = workspace / "tasks" / task.id / "execution"
    outputs_dir = workspace / "tasks" / task.id / "outputs"
    images_dir = workspace / "tasks" / task.id / "images"
    execution_dir.mkdir(parents=True)
    outputs_dir.mkdir(parents=True)
    images_dir.mkdir(parents=True)
    (outputs_dir / "validation_report.docx").write_bytes(b"old-report")
    (images_dir / "old.png").write_bytes(b"old-image")
    code_scores_path = execution_dir / "code_model_scores.csv"
    code_scores_path.write_text("row_index,code_model_score\n0,0.1\n", encoding="utf-8")
    (execution_dir / "runtime_contract.json").write_text(
        json.dumps(
            {
                "target_col": "y",
                "split_col": "split",
                "time_col": "apply_month",
                "pmml_output_field": "probability_1",
                "score_decimal_places": 6,
                "code_model_scores_path": str(code_scores_path),
                "feature_importance_path": None,
                "model_params_path": None,
                "algorithm": "lgb",
            }
        ),
        encoding="utf-8",
    )

    def fake_write_metrics_results_in_session(*, outputs_dir, task, **_kwargs):
        (outputs_dir / "validation_results.json").write_text(
            json.dumps(
                {
                    "model_name": task.model_name,
                    "model_version": task.model_version,
                    "scope": "贷前A卡",
                    "channel": "自营",
                    "month": "202601",
                    "effectiveness": {
                        "overall": [
                            {
                                "split": "oot",
                                "ks": 0.30,
                                "auc": 0.72,
                                "psi_vs_train": 0.08,
                            }
                        ],
                        "monthly_ks": [{"month": "202601", "ks": 0.30}],
                    },
                    "basic_info": {
                        "feature_importance": [
                            {"rank": 1, "feature": "x1", "category": "征信"}
                        ]
                    },
                }
            ),
            encoding="utf-8",
        )
        (outputs_dir / "validation.xlsx").write_bytes(b"xlsx")

    monkeypatch.setattr(
        "marvis.pipeline._write_metrics_results_in_session",
        fake_write_metrics_results_in_session,
    )
    register_live_notebook_session(
        task.id,
        SimpleNamespace(closed=False, close=lambda: None),
    )

    run_metrics_stage(
        task_id=task.id,
        settings=PipelineSettings(
            workspace=workspace,
            db_path=db_path,
            report_template_path=tmp_path / "template.docx",
            notebook_isolated_execution=False,
            allow_legacy_live_notebook_execution=True,
        ),
        stage_claimed=True,
    )

    memories = AgentMemoryStore(db_path).list_entries(memory_type="model_experience")
    assert len(memories) == 1
    assert memories[0].payload["ks"] == 0.30
    assert memories[0].payload["auc"] == 0.72
    assert memories[0].payload["psi"] == 0.08
    assert memories[0].payload["month"] == "202601"
    assert memories[0].payload["channel"] == "自营"
    assert memories[0].payload["scope"] == "贷前A卡"
    assert memories[0].payload["important_feature_sources"] == ["征信"]
    assert not (outputs_dir / "validation_report.docx").exists()
    assert not images_dir.exists()
    assert not (outputs_dir / ".metrics-stage-work").exists()
    assert not (outputs_dir / ".staging").exists()


def test_metrics_stage_status_failure_rolls_back_outputs_report_and_images(
    tmp_path: Path,
    monkeypatch,
):
    _allow_legacy_live_notebook(monkeypatch)
    project = tmp_path / "project"
    project.mkdir()
    sample_path = project / "sample.csv"
    pmml_path = project / "model.pmml"
    dictionary_path = project / "dictionary.csv"
    sample_path.write_text("y,score,split,apply_month\n0,0.1,train,202601\n", encoding="utf-8")
    pmml_path.write_text("<PMML/>", encoding="utf-8")
    dictionary_path.write_text("特征名,类别\nx1,征信\n", encoding="utf-8")

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡模型",
            model_version="V2026",
            validator="qa",
            source_dir=str(project),
            sample_path=str(sample_path),
            pmml_path=str(pmml_path),
            dictionary_path=str(dictionary_path),
            algorithm="lgb",
        )
    )
    repo.update_status(task.id, TaskStatus.SCANNED, message="source scanned")
    repo.update_status(task.id, TaskStatus.RUNNING, message="notebook queued")
    repo.update_status(task.id, TaskStatus.EXECUTED, message="notebook executed")
    repo.update_status(task.id, TaskStatus.COMPUTING_METRICS, message="metrics queued")
    task_dir = workspace / "tasks" / task.id
    execution_dir = task_dir / "execution"
    outputs_dir = task_dir / "outputs"
    images_dir = task_dir / "images"
    execution_dir.mkdir(parents=True)
    outputs_dir.mkdir(parents=True)
    images_dir.mkdir(parents=True)
    old_results = {"old": True}
    (outputs_dir / "validation_results.json").write_text(json.dumps(old_results), encoding="utf-8")
    (outputs_dir / "validation.xlsx").write_bytes(b"old-xlsx")
    (outputs_dir / "validation_report.docx").write_bytes(b"old-report")
    (images_dir / "old.png").write_bytes(b"old-image")
    code_scores_path = execution_dir / "code_model_scores.csv"
    code_scores_path.write_text("row_index,code_model_score\n0,0.1\n", encoding="utf-8")
    (execution_dir / "runtime_contract.json").write_text(
        json.dumps(
            {
                "target_col": "y",
                "split_col": "split",
                "time_col": "apply_month",
                "pmml_output_field": "probability_1",
                "score_decimal_places": 6,
                "code_model_scores_path": str(code_scores_path),
                "feature_importance_path": None,
                "model_params_path": None,
                "algorithm": "lgb",
            }
        ),
        encoding="utf-8",
    )
    old_model_meta = {"algorithm": "old", "feature_importance": []}
    (execution_dir / "model_meta.json").write_text(
        json.dumps(old_model_meta),
        encoding="utf-8",
    )
    (outputs_dir / REPRODUCIBILITY_RESULT_JSON).write_text(
        json.dumps({"summary": {"status": "pass"}, "rows": []}),
        encoding="utf-8",
    )

    def fake_write_metrics_results_in_session(*, outputs_dir, **_kwargs):
        (outputs_dir / "validation_results.json").write_text(
            json.dumps({"new": True}),
            encoding="utf-8",
        )
        (outputs_dir / "validation.xlsx").write_bytes(b"new-xlsx")

    monkeypatch.setattr(
        "marvis.pipeline._write_metrics_results_in_session",
        fake_write_metrics_results_in_session,
    )
    register_live_notebook_session(
        task.id,
        SimpleNamespace(closed=False, close=lambda: None),
    )
    original_update = TaskRepository.update_status_on_connection

    def failing_status_update(self, conn, task_id, status, *args, **kwargs):
        original_update(self, conn, task_id, status, *args, **kwargs)
        if status is TaskStatus.WRITING_ARTIFACTS:
            raise RuntimeError("simulated metrics status failure")

    monkeypatch.setattr(TaskRepository, "update_status_on_connection", failing_status_update)

    with pytest.raises(RuntimeError, match="simulated metrics status failure"):
        run_metrics_stage(
            task_id=task.id,
            settings=PipelineSettings(
                workspace=workspace,
                db_path=db_path,
                report_template_path=tmp_path / "template.docx",
                notebook_isolated_execution=False,
                allow_legacy_live_notebook_execution=True,
            ),
            stage_claimed=True,
        )

    assert json.loads((outputs_dir / "validation_results.json").read_text(encoding="utf-8")) == old_results
    assert (outputs_dir / "validation.xlsx").read_bytes() == b"old-xlsx"
    assert (outputs_dir / "validation_report.docx").read_bytes() == b"old-report"
    assert (images_dir / "old.png").read_bytes() == b"old-image"
    assert json.loads((execution_dir / "model_meta.json").read_text(encoding="utf-8")) == old_model_meta
    assert not (execution_dir / ".staging").exists()
    assert not (outputs_dir / ".metrics-stage-work").exists()
    assert not (outputs_dir / ".staging").exists()
    assert not (task_dir / ".staging").exists()


def test_metrics_stage_cancel_returns_to_executed_status(
    tmp_path: Path,
    monkeypatch,
):
    _allow_legacy_live_notebook(monkeypatch)
    project = tmp_path / "project"
    project.mkdir()
    sample_path = project / "sample.csv"
    pmml_path = project / "model.pmml"
    dictionary_path = project / "dictionary.csv"
    sample_path.write_text(
        "y,split,apply_month,pred\n0,oot,202501,0.1\n",
        encoding="utf-8",
    )
    pmml_path.write_text("<PMML />", encoding="utf-8")
    dictionary_path.write_text("特征名,类别\nx1,数值\n", encoding="utf-8")

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(project),
            sample_path=str(sample_path),
            pmml_path=str(pmml_path),
            dictionary_path=str(dictionary_path),
        )
    )
    repo.update_status(task.id, TaskStatus.SCANNED, message="source scanned")
    repo.update_status(task.id, TaskStatus.RUNNING, message="notebook queued")
    repo.update_status(task.id, TaskStatus.EXECUTED, message="notebook executed")
    repo.update_status(task.id, TaskStatus.COMPUTING_METRICS, message="metrics queued")
    execution_dir = workspace / "tasks" / task.id / "execution"
    execution_dir.mkdir(parents=True)
    code_scores_path = execution_dir / "code_model_scores.csv"
    code_scores_path.write_text("row_index,code_model_score\n0,0.1\n", encoding="utf-8")
    (execution_dir / "runtime_contract.json").write_text(
        json.dumps(
            {
                "target_col": "y",
                "split_col": "split",
                "time_col": "apply_month",
                "pmml_output_field": "probability_1",
                "score_decimal_places": 6,
                "code_model_scores_path": str(code_scores_path),
                "feature_importance_path": None,
                "model_params_path": None,
                "algorithm": "lgb",
            }
        ),
        encoding="utf-8",
    )
    fake_session = SimpleNamespace(
        closed=False,
        cancellation_token=None,
        client=SimpleNamespace(),
        execute_code_cell=lambda source, **_kwargs: SimpleNamespace(
            succeeded=False,
            failed_cell_index=14,
            error_name="NotebookCancelled",
            error_value="notebook execution cancelled",
            cancelled=True,
        ),
        close=lambda: None,
    )
    register_live_notebook_session(task.id, fake_session)

    run_metrics_stage(
        task_id=task.id,
        settings=PipelineSettings(
            workspace=workspace,
            db_path=db_path,
            report_template_path=tmp_path / "template.docx",
            notebook_isolated_execution=False,
            allow_legacy_live_notebook_execution=True,
        ),
        stage_claimed=True,
    )

    final = repo.get_task(task.id)
    assert final.status == TaskStatus.EXECUTED
    assert final.status_message == "metrics cancelled"
    close_live_notebook_session(task.id)


def test_report_stage_cancel_returns_to_review_required_status(tmp_path: Path):
    project = tmp_path / "project"
    project.mkdir()
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(project),
        )
    )
    repo.update_status(task.id, TaskStatus.SCANNED, message="source scanned")
    repo.update_status(task.id, TaskStatus.RUNNING, message="notebook queued")
    repo.update_status(task.id, TaskStatus.EXECUTED, message="notebook executed")
    repo.update_status(task.id, TaskStatus.COMPUTING_METRICS, message="metrics queued")
    repo.update_status(
        task.id,
        TaskStatus.WRITING_ARTIFACTS,
        message="metrics and excel generated",
    )
    request_notebook_cancellation(task.id)

    run_report_stage(
        task_id=task.id,
        settings=PipelineSettings(
            workspace=workspace,
            db_path=db_path,
            report_template_path=tmp_path / "template.docx",
        ),
    )

    final = repo.get_task(task.id)
    assert final.status == TaskStatus.REVIEW_REQUIRED
    assert final.status_message == "report cancelled"


def test_report_stage_cancel_during_word_write_does_not_promote_partial_docx(
    tmp_path: Path,
    monkeypatch,
):
    project = tmp_path / "project"
    project.mkdir()
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(project),
        )
    )
    repo.update_status(task.id, TaskStatus.SCANNED, message="source scanned")
    repo.update_status(task.id, TaskStatus.RUNNING, message="notebook queued")
    repo.update_status(task.id, TaskStatus.EXECUTED, message="notebook executed")
    repo.update_status(task.id, TaskStatus.COMPUTING_METRICS, message="metrics queued")
    repo.update_status(
        task.id,
        TaskStatus.WRITING_ARTIFACTS,
        message="metrics and excel generated",
    )
    outputs_dir = workspace / "tasks" / task.id / "outputs"
    outputs_dir.mkdir(parents=True)
    report_path = outputs_dir / "validation_report.docx"
    report_path.write_bytes(b"previous-docx")

    monkeypatch.setattr(
        "marvis.pipeline._load_validation_results",
        lambda _outputs_dir: SimpleNamespace(),
    )

    def fake_word_writer(*_args, output_path, **_kwargs):
        Path(output_path).write_bytes(b"partial-docx")
        request_notebook_cancellation(task.id)
        return SimpleNamespace(unresolved_placeholders=[])

    monkeypatch.setattr("marvis.pipeline.write_validation_word", fake_word_writer)

    run_report_stage(
        task_id=task.id,
        settings=PipelineSettings(
            workspace=workspace,
            db_path=db_path,
            report_template_path=tmp_path / "template.docx",
        ),
    )

    final = repo.get_task(task.id)
    assert final.status == TaskStatus.REVIEW_REQUIRED
    assert final.status_message == "report cancelled"
    assert report_path.read_bytes() == b"previous-docx"
    assert not (outputs_dir / ".validation_report.docx.tmp").exists()


def test_report_stage_status_failure_rolls_back_promoted_docx_and_images(
    tmp_path: Path,
    monkeypatch,
):
    project = tmp_path / "project"
    project.mkdir()
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(project),
        )
    )
    repo.update_status(task.id, TaskStatus.SCANNED, message="source scanned")
    repo.update_status(task.id, TaskStatus.RUNNING, message="notebook queued")
    repo.update_status(task.id, TaskStatus.EXECUTED, message="notebook executed")
    repo.update_status(task.id, TaskStatus.COMPUTING_METRICS, message="metrics queued")
    repo.update_status(
        task.id,
        TaskStatus.WRITING_ARTIFACTS,
        message="metrics and excel generated",
    )
    task_dir = workspace / "tasks" / task.id
    outputs_dir = task_dir / "outputs"
    images_dir = task_dir / "images"
    outputs_dir.mkdir(parents=True)
    images_dir.mkdir(parents=True)
    report_path = outputs_dir / "validation_report.docx"
    report_path.write_bytes(b"previous-docx")
    (images_dir / "old.png").write_bytes(b"old-image")

    monkeypatch.setattr(
        "marvis.pipeline._load_validation_results",
        lambda _outputs_dir: SimpleNamespace(
            reproducibility=SimpleNamespace(
                summary=SimpleNamespace(status=pipeline_module.ConsistencyStatus.PASS)
            )
        ),
    )

    def fake_word_writer(*_args, output_path, image_output_dir, **_kwargs):
        Path(output_path).write_bytes(b"new-docx")
        Path(image_output_dir).mkdir(parents=True, exist_ok=True)
        (Path(image_output_dir) / "new.png").write_bytes(b"new-image")
        return SimpleNamespace(unresolved_placeholders=[])

    monkeypatch.setattr("marvis.pipeline.write_validation_word", fake_word_writer)
    original_update = TaskRepository.update_status_on_connection

    def failing_status_update(self, conn, *args, **kwargs):
        original_update(self, conn, *args, **kwargs)
        raise RuntimeError("simulated status commit failure")

    monkeypatch.setattr(TaskRepository, "update_status_on_connection", failing_status_update)

    with pytest.raises(RuntimeError, match="simulated status commit failure"):
        run_report_stage(
            task_id=task.id,
            settings=PipelineSettings(
                workspace=workspace,
                db_path=db_path,
                report_template_path=tmp_path / "template.docx",
            ),
        )

    final = repo.get_task(task.id)
    assert final.status == TaskStatus.WRITING_ARTIFACTS
    assert final.status_message == "metrics and excel generated"
    assert report_path.read_bytes() == b"previous-docx"
    assert (images_dir / "old.png").read_bytes() == b"old-image"
    assert not (images_dir / "new.png").exists()
    assert not (outputs_dir / ".validation_report.docx.tmp").exists()
    assert not (outputs_dir / ".staging").exists()
    assert not (task_dir / ".staging").exists()


def test_staged_metrics_use_live_notebook_sample_without_rerunning_notebook(
    tmp_path: Path,
    monkeypatch,
):
    _allow_legacy_live_notebook(monkeypatch)
    project = tmp_path / "project"
    project.mkdir()
    (project / "sample.csv").write_text("placeholder\n1\n", encoding="utf-8")
    shutil.copy(FIXTURES / "min_lr.pmml", project / "fr_final.pmml")
    pd.DataFrame(
        {
            "特征名": ["x1", "x2"],
            "类别": ["征信", "基础信息"],
        }
    ).to_csv(project / "data_dictionary.csv", index=False)
    _write_live_sample_contract_notebook(project / "dev.ipynb")

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(project),
            algorithm="lgb",
            target_col="y",
            score_col="pred",
            split_col="split",
            time_col="apply_month",
        )
    )
    settings = PipelineSettings(
        workspace=workspace,
        db_path=db_path,
        report_template_path=tmp_path / "template.docx",
        random_sample_size=12,
        notebook_isolated_execution=False,
        allow_legacy_live_notebook_execution=True,
    )

    run_notebook_stage(task_id=task.id, settings=settings)
    run_metrics_stage(task_id=task.id, settings=settings)

    assert repo.get_task(task.id).status == TaskStatus.WRITING_ARTIFACTS
    assert (project / "notebook_run_count.txt").read_text(encoding="utf-8") == "1"
    task_dir = workspace / "tasks" / task.id
    assert (task_dir / "outputs" / "validation_results.json").exists()
    result_json = json.loads(
        (task_dir / "outputs" / "validation_results.json").read_text(encoding="utf-8")
    )
    assert result_json["reproducibility"]["summary"]["status"] == "pass"


def test_completed_task_cannot_rerun_metrics_after_live_notebook_session_closed(
    tmp_path: Path,
    monkeypatch,
):
    _allow_legacy_live_notebook(monkeypatch)
    project = tmp_path / "project"
    project.mkdir()
    (project / "sample.csv").write_text("placeholder\n1\n", encoding="utf-8")
    shutil.copy(FIXTURES / "min_lr.pmml", project / "fr_final.pmml")
    pd.DataFrame(
        {
            "特征名": ["x1", "x2"],
            "类别": ["征信", "基础信息"],
        }
    ).to_csv(project / "data_dictionary.csv", index=False)
    _write_live_sample_contract_notebook(project / "dev.ipynb")

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(project),
            algorithm="lgb",
            target_col="y",
            score_col="pred",
            split_col="split",
            time_col="apply_month",
        )
    )
    settings = PipelineSettings(
        workspace=workspace,
        db_path=db_path,
        report_template_path=tmp_path / "template.docx",
        random_sample_size=12,
        notebook_isolated_execution=False,
        allow_legacy_live_notebook_execution=True,
    )

    run_notebook_stage(task_id=task.id, settings=settings)
    run_metrics_stage(task_id=task.id, settings=settings)
    repo.update_status(task.id, TaskStatus.SUCCEEDED, message="pipeline succeeded")
    task_dir = workspace / "tasks" / task.id
    report_path = task_dir / "outputs" / "validation_report.docx"
    report_path.write_bytes(b"generated report")
    repo.update_status(task.id, TaskStatus.COMPUTING_METRICS, message="metrics queued")
    try:
        run_metrics_stage(task_id=task.id, settings=settings, stage_claimed=True)
    except PipelineError as exc:
        assert "live notebook kernel is not available" in str(exc)
    else:
        raise AssertionError("metrics should fail when the live notebook session is gone")

    final = repo.get_task(task.id)
    assert final.status == TaskStatus.FAILED
    assert final.status_message.startswith("模型效果&稳定性验证失败：")
    assert (project / "notebook_run_count.txt").read_text(encoding="utf-8") == "1"
    assert (task_dir / "outputs" / "validation_results.json").exists()
    assert report_path.read_bytes() == b"generated report"


def test_full_pipeline_marks_word_failures_as_report_stage_failures(
    tmp_path: Path,
    monkeypatch,
):
    _allow_legacy_live_notebook(monkeypatch)
    project = tmp_path / "project"
    project.mkdir()
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(project),
            algorithm="lgb",
        )
    )
    artifact_paths = {
        FileRole.NOTEBOOK: project / "model.ipynb",
        FileRole.SAMPLE: project / "sample.csv",
        FileRole.MODEL_PMML: project / "model.pmml",
        FileRole.DATA_DICTIONARY: project / "dictionary.csv",
    }
    artifacts = [
        FileArtifact(role, path, 1, None)
        for role, path in artifact_paths.items()
    ]

    def fake_scan_step(repo_arg, task_arg):
        repo_arg.update_status(
            task_arg.id,
            TaskStatus.SCANNED,
            "source scanned",
            expected={TaskStatus.CREATED, TaskStatus.SCANNED, TaskStatus.FAILED},
        )
        return artifacts

    class FakeSession:
        def close(self):
            pass

    def fake_notebook_step_v3(*, repo, task, **_kwargs):
        repo.update_status(task.id, TaskStatus.RUNNING, "notebook running")
        repo.update_status(task.id, TaskStatus.EXECUTED, "notebook executed")
        return FakeSession()

    def fake_metrics_writer(*, outputs_dir, **_kwargs):
        outputs_dir.mkdir(parents=True, exist_ok=True)
        (outputs_dir / "validation_results.json").write_text("{}", encoding="utf-8")
        (outputs_dir / "validation.xlsx").write_bytes(b"xlsx")

    monkeypatch.setattr("marvis.pipeline._scan_step", fake_scan_step)
    monkeypatch.setattr("marvis.pipeline._notebook_step_v3", fake_notebook_step_v3)
    monkeypatch.setattr("marvis.pipeline._write_reproducibility_result_in_session", lambda **_kwargs: None)
    monkeypatch.setattr("marvis.pipeline._write_metrics_results_in_session", fake_metrics_writer)
    monkeypatch.setattr(
        "marvis.pipeline.load_runtime_contract",
        lambda _path: RuntimeContract(
            target_col="y",
            split_col="split",
            time_col="apply_month",
            pmml_output_field="probability_1",
            score_decimal_places=6,
            code_model_scores_path=tmp_path / "scores.csv",
            feature_importance_path=None,
            model_params_path=None,
            algorithm="lgb",
        ),
    )
    monkeypatch.setattr(
        "marvis.pipeline._load_validation_results",
        lambda _outputs_dir: SimpleNamespace(
            reproducibility=SimpleNamespace(
                summary=SimpleNamespace(status=None)
            )
        ),
    )
    monkeypatch.setattr(
        "marvis.pipeline.write_validation_word",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(RuntimeError("docx failed")),
    )

    try:
        run_pipeline(
            task_id=task.id,
            settings=PipelineSettings(
                workspace=workspace,
                db_path=db_path,
                report_template_path=tmp_path / "template.docx",
                notebook_isolated_execution=False,
                allow_legacy_live_notebook_execution=True,
            ),
        )
    except RuntimeError as exc:
        assert "docx failed" in str(exc)
    else:
        raise AssertionError("word generation failure should bubble to the job runner")

    final = repo.get_task(task.id)
    assert final.status == TaskStatus.FAILED
    assert final.status_message == f"{REPORT_STAGE_FAILURE_PREFIX}RuntimeError: docx failed"
    assert (workspace / "tasks" / task.id / "outputs" / "validation.xlsx").exists()


def test_legacy_run_pipeline_metrics_status_failure_does_not_promote_outputs(
    tmp_path: Path,
    monkeypatch,
):
    _allow_legacy_live_notebook(monkeypatch)
    project = tmp_path / "project"
    project.mkdir()
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(project),
            algorithm="lgb",
        )
    )
    artifact_paths = {
        FileRole.NOTEBOOK: project / "model.ipynb",
        FileRole.SAMPLE: project / "sample.csv",
        FileRole.MODEL_PMML: project / "model.pmml",
        FileRole.DATA_DICTIONARY: project / "dictionary.csv",
    }
    artifacts = [
        FileArtifact(role, path, 1, None)
        for role, path in artifact_paths.items()
    ]

    def fake_scan_step(repo_arg, task_arg):
        repo_arg.update_status(
            task_arg.id,
            TaskStatus.SCANNED,
            "source scanned",
            expected={TaskStatus.CREATED, TaskStatus.SCANNED, TaskStatus.FAILED},
        )
        return artifacts

    class FakeSession:
        def close(self):
            pass

    def fake_notebook_step_v3(*, repo, task, **_kwargs):
        repo.update_status(task.id, TaskStatus.RUNNING, "notebook running")
        repo.update_status(task.id, TaskStatus.EXECUTED, "notebook executed")
        return FakeSession()

    def fake_metrics_writer(*, outputs_dir, **_kwargs):
        outputs_dir.mkdir(parents=True, exist_ok=True)
        (outputs_dir / "validation_results.json").write_text(
            json.dumps({"new": True}),
            encoding="utf-8",
        )
        (outputs_dir / "validation.xlsx").write_bytes(b"new-xlsx")

    monkeypatch.setattr("marvis.pipeline._scan_step", fake_scan_step)
    monkeypatch.setattr("marvis.pipeline._notebook_step_v3", fake_notebook_step_v3)
    monkeypatch.setattr("marvis.pipeline._write_reproducibility_result_in_session", lambda **_kwargs: None)
    monkeypatch.setattr("marvis.pipeline._write_metrics_results_in_session", fake_metrics_writer)
    monkeypatch.setattr(
        "marvis.pipeline.load_runtime_contract",
        lambda _path: RuntimeContract(
            target_col="y",
            split_col="split",
            time_col="apply_month",
            pmml_output_field="probability_1",
            score_decimal_places=6,
            code_model_scores_path=tmp_path / "scores.csv",
            feature_importance_path=None,
            model_params_path=None,
            algorithm="lgb",
        ),
    )
    original_update = TaskRepository.update_status_on_connection

    def failing_status_update(self, conn, task_id, status, *args, **kwargs):
        original_update(self, conn, task_id, status, *args, **kwargs)
        if status is TaskStatus.WRITING_ARTIFACTS:
            raise RuntimeError("simulated legacy metrics status failure")

    monkeypatch.setattr(TaskRepository, "update_status_on_connection", failing_status_update)

    with pytest.raises(RuntimeError, match="simulated legacy metrics status failure"):
        run_pipeline(
            task_id=task.id,
            settings=PipelineSettings(
                workspace=workspace,
                db_path=db_path,
                report_template_path=tmp_path / "template.docx",
                notebook_isolated_execution=False,
                allow_legacy_live_notebook_execution=True,
            ),
        )

    task_dir = workspace / "tasks" / task.id
    outputs_dir = task_dir / "outputs"
    execution_dir = task_dir / "execution"
    assert not (outputs_dir / "validation_results.json").exists()
    assert not (outputs_dir / "validation.xlsx").exists()
    assert not (outputs_dir / ".metrics-stage-work").exists()
    assert not (outputs_dir / ".staging").exists()
    assert not (execution_dir / "model_meta.json").exists()
    assert not (execution_dir / ".staging").exists()


def test_pipeline_marks_missing_required_input_failed(tmp_path: Path):
    project = tmp_path / "project"
    project.mkdir()
    pd.DataFrame({"x1": [0.1], "pred": [0.1], "y": [0], "split": ["train"], "apply_month": ["202503"]}).to_csv(
        project / "sample.csv",
        index=False,
    )
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(project),
        )
    )

    try:
        run_pipeline(
            task_id=task.id,
            settings=PipelineSettings(
                workspace=workspace,
                db_path=db_path,
                report_template_path=tmp_path / "missing.docx",
                feature_columns=["x1"],
            ),
        )
    except Exception as exc:
        assert "missing required input" in str(exc)
    else:
        raise AssertionError("pipeline should fail when required input is missing")

    final = repo.get_task(task.id)
    assert final.status == TaskStatus.FAILED
    assert "missing required input" in final.status_message


def test_pipeline_rejects_ambiguous_notebooks_without_explicit_path(tmp_path: Path):
    project = tmp_path / "project"
    project.mkdir()
    first = project / "first.ipynb"
    second = project / "second.ipynb"
    first.write_text("{}", encoding="utf-8")
    second.write_text("{}", encoding="utf-8")
    task = SimpleNamespace(source_dir=str(project), notebook_path=None)
    artifacts = [
        FileArtifact(FileRole.NOTEBOOK, first, 2, None),
        FileArtifact(FileRole.NOTEBOOK, second, 2, None),
    ]

    try:
        _required_path(task, artifacts, FileRole.NOTEBOOK, "notebook", "notebook_path")
    except PipelineError as exc:
        message = str(exc)
    else:
        raise AssertionError("ambiguous notebook candidates should fail")

    assert "notebook role ambiguous" in message
    assert str(first) in message
    assert str(second) in message
    assert "configure notebook_path" in message


def test_pipeline_uses_explicit_notebook_path_when_multiple_exist(tmp_path: Path):
    project = tmp_path / "project"
    project.mkdir()
    first = project / "first.ipynb"
    chosen = project / "chosen.ipynb"
    first.write_text("{}", encoding="utf-8")
    chosen.write_text("{}", encoding="utf-8")
    task = SimpleNamespace(source_dir=str(project), notebook_path="chosen.ipynb")
    artifacts = [
        FileArtifact(FileRole.NOTEBOOK, first, 2, None),
        FileArtifact(FileRole.NOTEBOOK, chosen, 2, None),
    ]

    result = _required_path(
        task, artifacts, FileRole.NOTEBOOK, "notebook", "notebook_path"
    )

    assert result == chosen.resolve()


def test_pipeline_rejects_completed_task_without_marking_failed(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(tmp_path / "project"),
        )
    )
    repo.update_status(task.id, TaskStatus.SCANNED, message="source scanned")
    repo.update_status(task.id, TaskStatus.RUNNING, message="notebook running")
    repo.update_status(task.id, TaskStatus.EXECUTED, message="notebook executed")
    repo.update_status(task.id, TaskStatus.SUCCEEDED, message="pipeline succeeded")

    try:
        run_pipeline(
            task_id=task.id,
            settings=PipelineSettings(
                workspace=workspace,
                db_path=db_path,
                report_template_path=tmp_path / "template.docx",
                feature_columns=["x1"],
            ),
        )
    except PipelineError as exc:
        assert "already succeeded" in str(exc)
    else:
        raise AssertionError("completed tasks should not rerun")

    final = repo.get_task(task.id)
    assert final.status == TaskStatus.SUCCEEDED
    assert final.status_message == "pipeline succeeded"


def test_pipeline_feature_columns_fall_back_to_task_when_settings_empty(
    tmp_path: Path,
):
    settings = PipelineSettings(
        workspace=tmp_path / "workspace",
        db_path=tmp_path / "db.sqlite",
        report_template_path=tmp_path / "template.docx",
        feature_columns=[],
    )
    task = SimpleNamespace(feature_columns=("x1", "x2"))

    assert _feature_columns(settings, task) == ["x1", "x2"]


def test_load_sample_falls_back_to_selected_python_for_arrow_files(
    tmp_path: Path,
    monkeypatch,
):
    sample_path = tmp_path / "sample.feather"
    sample_path.write_bytes(b"not used by fake fallback")
    fallback_python = tmp_path / "selected-python"
    fallback_python.write_text(
        "\n".join(
            [
                f"#!{sys.executable}",
                "import os",
                "import sys",
                "assert 'OPENAI_API_KEY' not in os.environ",
                "assert 'ANTHROPIC_API_KEY' not in os.environ",
                "assert os.environ.get('PYTHONUNBUFFERED') == '1'",
                "import pandas as pd",
                "pd.DataFrame({",
                "    'x1': [9007199254740993, 9007199254740995],",
                "    'y': [0, 1],",
                "    'missing': [float('nan'), 1.5],",
                "    'as_of': pd.to_datetime(['2026-01-01', '2026-01-02']),",
                "}).to_pickle(sys.argv[4])",
            ]
        ),
        encoding="utf-8",
    )
    fallback_python.chmod(0o755)
    monkeypatch.setenv("OPENAI_API_KEY", "sk-secret")
    monkeypatch.setenv("ANTHROPIC_API_KEY", "anthropic-secret")

    def fake_read_feather(path):
        raise ImportError("pyarrow is unavailable in platform env")

    monkeypatch.setattr("marvis.pipeline.pd.read_feather", fake_read_feather)

    sample = _load_sample(sample_path, fallback_python=fallback_python)

    assert sample["x1"].tolist() == [9007199254740993, 9007199254740995]
    assert sample["y"].tolist() == [0, 1]
    assert pd.isna(sample["missing"].iloc[0])
    assert str(sample["as_of"].dtype).startswith("datetime64")


def test_load_sample_supports_excel_files(tmp_path: Path):
    sample_path = tmp_path / "model_sample.xlsx"
    pd.DataFrame({"x1": [1, 2], "y": [0, 1]}).to_excel(sample_path, index=False)

    sample = _load_sample(sample_path)

    assert sample.to_dict(orient="list") == {"x1": [1, 2], "y": [0, 1]}


@pytest.mark.slow
def test_staged_pipeline_isolated_mode_executes_notebook_once(tmp_path: Path):
    # PERF-3 regression: the default isolated-mode staged pipeline must run
    # the user notebook exactly once. Before the fix, the metrics stage
    # force-closed the (nonexistent, in isolated mode) live session and
    # replayed the entire notebook from scratch, doubling wall-clock time.
    project = tmp_path / "project"
    project.mkdir()
    (project / "sample.csv").write_text("placeholder\n1\n", encoding="utf-8")
    shutil.copy(FIXTURES / "min_lr.pmml", project / "fr_final.pmml")
    pd.DataFrame(
        {
            "特征名": ["x1", "x2"],
            "类别": ["征信", "基础信息"],
        }
    ).to_csv(project / "data_dictionary.csv", index=False)
    _write_live_sample_contract_notebook(project / "dev.ipynb")

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(project),
            algorithm="lgb",
            target_col="y",
            score_col="pred",
            split_col="split",
            time_col="apply_month",
        )
    )
    settings = PipelineSettings(
        workspace=workspace,
        db_path=db_path,
        report_template_path=_write_template(tmp_path / "template.docx"),
        random_sample_size=12,
    )
    assert settings.notebook_isolated_execution is True

    run_staged_pipeline(task_id=task.id, settings=settings)

    final = repo.get_task(task.id)
    assert final.status in {
        TaskStatus.WRITING_ARTIFACTS,
        TaskStatus.REVIEW_REQUIRED,
        TaskStatus.SUCCEEDED,
    }
    assert (project / "notebook_run_count.txt").read_text(encoding="utf-8") == "1"
    task_dir = workspace / "tasks" / task.id
    assert (task_dir / "execution" / "model_meta.json").exists()
    assert (task_dir / "outputs" / "validation_results.json").exists()
    result_json = json.loads(
        (task_dir / "outputs" / "validation_results.json").read_text(encoding="utf-8")
    )
    assert result_json["reproducibility"]["summary"]["status"] == "pass"


def test_staged_pipeline_isolated_mode_merged_run_calls_notebook_step_once(
    tmp_path: Path,
    monkeypatch,
):
    # Same guarantee as above, verified at the mock level: _notebook_step_v3
    # (the isolated subprocess launcher) must be invoked exactly once for
    # the default consecutive notebook+metrics path, with a combined
    # extra_code_cells list covering both reproducibility and metrics cells.
    project = tmp_path / "project"
    project.mkdir()
    notebook_path = project / "dev.ipynb"
    sample_path = project / "sample.csv"
    pmml_path = project / "fr_final.pmml"
    dictionary_path = project / "data_dictionary.csv"
    nbformat.write(nbformat.v4.new_notebook(cells=[]), notebook_path)
    sample_path.write_text("x1,pred,y,split,apply_month\n1,0.1,0,train,202501\n")
    pmml_path.write_text("<PMML/>", encoding="utf-8")
    dictionary_path.write_text("特征名,类别\nx1,基础信息\n", encoding="utf-8")

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    db_path = workspace / "marvis.sqlite"
    init_db(db_path)
    repo = TaskRepository(db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="A卡",
            model_version="v1",
            validator="qa",
            source_dir=str(project),
            algorithm="lgb",
        )
    )

    call_count = {"value": 0}

    def fake_notebook_step_v3(
        *,
        repo,
        task,
        contract_meta_path,
        code_scores_path,
        execution_dir,
        extra_code_cells,
        stage_claimed,
        **kwargs,
    ):
        call_count["value"] += 1
        assert kwargs["keep_alive"] is False
        assert kwargs["isolated"] is True
        kinds = [kind for kind, _source in extra_code_cells]
        assert kinds[:2] == ["repro-pmml", "repro-compare"]
        assert "metrics-prepare" in kinds
        assert "metrics-output" in kinds
        if not stage_claimed:
            repo.update_status(
                task.id,
                TaskStatus.RUNNING,
                message="notebook queued",
                expected={
                    TaskStatus.SCANNED,
                    TaskStatus.RUNNING,
                    TaskStatus.EXECUTED,
                    TaskStatus.FAILED,
                },
            )
        contract_meta_path.write_text(
            json.dumps(
                {
                    "target_col": "y",
                    "split_col": "split",
                    "time_col": "apply_month",
                    "pmml_output_field": "probability_1",
                    "score_decimal_places": 6,
                    "code_model_scores_path": str(code_scores_path),
                    "feature_importance_path": None,
                    "model_params_path": None,
                    "algorithm": "lgb",
                }
            ),
            encoding="utf-8",
        )
        code_scores_path.write_text("row_index,code_model_score\n0,0.1\n")
        (execution_dir / "model_meta.json").write_text(
            json.dumps({"algorithm": "lgb", "feature_importance": [], "hyperparameters": {}}),
            encoding="utf-8",
        )
        outputs_dir = workspace / "tasks" / task.id / "outputs"
        outputs_dir.mkdir(parents=True, exist_ok=True)
        (outputs_dir / REPRODUCIBILITY_RESULT_JSON).write_text(
            json.dumps({"summary": {"status": "pass"}, "rows": []}),
            encoding="utf-8",
        )
        metrics_work_dir = outputs_dir / ".metrics-stage-work"
        metrics_work_dir.mkdir(parents=True, exist_ok=True)
        (metrics_work_dir / "validation_results.json").write_text("{}", encoding="utf-8")
        (metrics_work_dir / "validation.xlsx").write_bytes(b"xlsx")

    monkeypatch.setattr("marvis.pipeline._notebook_step_v3", fake_notebook_step_v3)

    settings = PipelineSettings(
        workspace=workspace,
        db_path=db_path,
        report_template_path=tmp_path / "template.docx",
    )
    run_notebook_stage(
        task_id=task.id,
        settings=settings,
        also_prepare_metrics=True,
    )
    assert repo.get_task(task.id).status == TaskStatus.EXECUTED

    run_metrics_stage(task_id=task.id, settings=settings)

    assert call_count["value"] == 1
    assert repo.get_task(task.id).status == TaskStatus.WRITING_ARTIFACTS
