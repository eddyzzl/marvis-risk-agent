import json
from pathlib import Path

import pandas as pd
from docx import Document
from openpyxl import load_workbook

from marvis.output.excel import write_validation_excel
from marvis.output.word import write_validation_word
from marvis.validation.checks import finite_score_series, validate_binary_target
from marvis.validation.config import ValidationConfig
from marvis.validation.effectiveness import run_effectiveness
from marvis.validation.feature_categories import resolve_feature_categories
from marvis.validation.reproducibility import run_reproducibility
from marvis.validation.results import ValidationResults
from marvis.validation.sample_stats import run_basic_info
from marvis.validation.stress_test import run_stress_test


class _IdentityScorer:
    def __init__(self, feature: str = "x1") -> None:
        self.feature = feature

    def score(self, df: pd.DataFrame) -> list[float]:
        return df[self.feature].astype(float).tolist()


def _make_template(path: Path) -> Path:
    document = Document()
    document.add_paragraph("{{TEXT:report_title}}")
    document.add_paragraph("OOT KS：{{TEXT:oot_ks}}")
    document.add_paragraph("{{IMAGE:overall_model_effect}}")
    document.add_paragraph("{{IMAGE:pressure_ks_table}}")
    document.save(path)
    return path


def _compose_results(
    *,
    sample: pd.DataFrame,
    dictionary: pd.DataFrame,
    model_meta_path: Path,
    scorer: _IdentityScorer,
    config: ValidationConfig,
) -> ValidationResults:
    # Mirrors the live composition that pipeline_cellgen injects into the
    # validation notebook (the former validation.engine wrapper was test-only).
    validate_binary_target(sample, config.target_col)
    code_scores = sample[config.score_col].astype(float)

    scored = sample.copy()
    scored[config.score_col] = finite_score_series(
        scorer.score(scored.copy()),
        index=scored.index,
        label="submitted PMML scorer",
    )

    reproducibility = run_reproducibility(
        sample=sample,
        config=config,
        code_scores=code_scores,
        submitted_pmml_scorer=scorer,
    )
    basic_info = run_basic_info(
        sample=scored,
        config=config,
        model_meta_path=model_meta_path,
    )
    effectiveness = run_effectiveness(sample=scored, config=config)

    oot_sample = sample[sample[config.split_col] == config.split_values["oot"]]
    category_resolution = resolve_feature_categories(
        model_features=[
            (row.feature, row.category) for row in basic_info.feature_importance
        ],
        dictionary=dictionary,
        feature_col=config.data_dict_feature_col,
        category_col=config.data_dict_category_col,
    )
    stress_test = run_stress_test(
        oot_sample=oot_sample,
        config=config,
        feature_categories=category_resolution.per_category,
        input_scorer=scorer,
        unclassified_features=category_resolution.unclassified_features,
        category_source_counts=category_resolution.source_counts,
    )

    return ValidationResults(
        model_name="A卡",
        model_version="v1",
        algorithm="lgb",
        target_type="binary",
        schema_version="marvis.validation_results.v1",
        reproducibility=reproducibility,
        basic_info=basic_info,
        effectiveness=effectiveness,
        stress_test=stress_test,
    )


def test_engine_output_round_trip(tmp_path: Path):
    rows = []
    for split in ("train", "test", "oot"):
        for i in range(20):
            x1 = (i + 1) / 21
            rows.append({
                "x1": x1, "sample_score": x1, "y": int(i >= 10),
                "split": split, "apply_month": "202503" if split == "train" else "202505",
            })
    sample = pd.DataFrame(rows)
    dictionary = pd.DataFrame({"特征名": ["x1"], "类别": ["征信"]})
    meta_path = tmp_path / "model_meta.json"
    meta_path.write_text(json.dumps({
        "feature_importance": [{"feature": "x1", "importance": 1.0}],
        "hyperparameters": {"max_depth": 4, "learning_rate": 0.1},
    }), encoding="utf-8")

    results = _compose_results(
        sample=sample,
        dictionary=dictionary,
        model_meta_path=meta_path,
        scorer=_IdentityScorer(),
        config=ValidationConfig(
            target_col="y", score_col="sample_score", split_col="split",
            time_col="apply_month", feature_columns=["x1"],
            bin_count=5, random_sample_size=10,
        ),
    )

    excel_path = tmp_path / "out.xlsx"
    write_validation_excel(results, excel_path)
    wb = load_workbook(excel_path)
    assert "验证总览" in wb.sheetnames
    assert "压力测试_分箱_征信" in wb.sheetnames

    template = _make_template(tmp_path / "template.docx")
    word_path = tmp_path / "report.docx"
    word_result = write_validation_word(
        results, template_path=template,
        output_path=word_path, image_output_dir=tmp_path / "images",
    )
    assert word_result.unresolved_placeholders == []
    assert word_path.exists()


def test_transformed_notebook_feature_keeps_stress_category_through_excel(
    tmp_path: Path,
):
    feature = "BH_A044_C0580"
    rows = []
    for split in ("train", "test", "oot"):
        for index in range(20):
            value = (index + 1) / 21
            rows.append(
                {
                    feature: value,
                    "sample_score": value,
                    "y": int(index >= 10),
                    "split": split,
                    "apply_month": "202503" if split == "train" else "202505",
                }
            )
    sample = pd.DataFrame(rows)
    dictionary = pd.DataFrame({"特征名": ["BH_A044"], "类别": ["睿智"]})
    meta_path = tmp_path / "model_meta.json"
    meta_path.write_text(
        json.dumps(
            {
                "feature_importance": [
                    {"feature": feature, "category": "睿智", "importance": 1.0}
                ],
                "hyperparameters": {},
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    results = _compose_results(
        sample=sample,
        dictionary=dictionary,
        model_meta_path=meta_path,
        scorer=_IdentityScorer(feature),
        config=ValidationConfig(
            target_col="y",
            score_col="sample_score",
            split_col="split",
            time_col="apply_month",
            feature_columns=[feature],
            bin_count=5,
            random_sample_size=10,
        ),
    )

    assert results.stress_test.status == "completed"
    assert results.stress_test.unclassified_features == []
    assert results.stress_test.per_category[0].category == "睿智"
    assert results.stress_test.per_category[0].dropped_features == [feature]

    excel_path = tmp_path / "transformed.xlsx"
    write_validation_excel(results, excel_path)
    workbook = load_workbook(excel_path, data_only=True)
    assert "压力测试_分箱_睿智" in workbook.sheetnames
    assert workbook["压力测试_汇总"]["A3"].value == "睿智"
