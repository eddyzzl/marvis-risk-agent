from pathlib import Path

from docx import Document

from marvis.output.word import write_validation_word
from tests.output.test_excel import _make_pmml_results, _make_results


def _make_template(path: Path) -> Path:
    document = Document()
    document.add_paragraph("模型：{{TEXT:model_name}} {{TEXT:model_version}}")
    document.add_paragraph("样本：{{TEXT:sample_period}}")
    document.add_paragraph("OOT KS：{{TEXT:oot_ks}}")
    document.add_paragraph("{{IMAGE:overall_model_effect}}")
    document.add_paragraph("{{IMAGE:pressure_ks_table}}")
    document.save(path)
    return path


def _make_manual_template(path: Path) -> Path:
    document = Document()
    document.add_paragraph("标题：{{TEXT:report_title}}")
    document.add_paragraph("概述：{{TEXT:model_overview}}")
    document.add_paragraph("日期：{{TEXT:draft_date}}")
    document.add_paragraph("KS：{{TEXT:oot_ks}}")
    document.save(path)
    return path


def _make_pmml_template(path: Path) -> Path:
    document = Document()
    document.add_paragraph("原模板：{{TEXT:reproducibility_summary}}")
    document.add_paragraph("新模板：{{TEXT:pmml_scoring_summary}}")
    document.add_paragraph("{{IMAGE:overall_model_effect}}")
    document.save(path)
    return path


def test_word_report_has_no_unresolved_text_placeholders(tmp_path: Path):
    template = _make_template(tmp_path / "template.docx")
    output = tmp_path / "report.docx"
    image_dir = tmp_path / "images"

    result = write_validation_word(
        _make_results(), template_path=template,
        output_path=output, image_output_dir=image_dir,
    )

    assert result.output_path == output
    assert result.unresolved_placeholders == []
    document = Document(output)
    text = "\n".join(p.text for p in document.paragraphs)
    assert "A卡" in text
    assert "0.2500" in text  # oot_ks formatted to four decimals


def test_word_report_inlines_pngs(tmp_path: Path):
    template = _make_template(tmp_path / "template.docx")
    output = tmp_path / "report.docx"
    image_dir = tmp_path / "images"
    write_validation_word(
        _make_results(), template_path=template,
        output_path=output, image_output_dir=image_dir,
    )
    document = Document(output)
    inline_shapes = document.inline_shapes
    assert len(inline_shapes) >= 2


def test_word_report_merges_manual_report_values_with_metric_text(tmp_path: Path):
    template = _make_manual_template(tmp_path / "template.docx")
    output = tmp_path / "report.docx"

    result = write_validation_word(
        _make_results(),
        template_path=template,
        output_path=output,
        image_output_dir=tmp_path / "images",
        report_values={
            "TEXT:report_title": "人工标题",
            "TEXT:model_overview": "人工模型概述",
        },
        manual_values={"draft_date": "2026-05-21"},
    )

    document = Document(result.output_path)
    text = "\n".join(p.text for p in document.paragraphs)

    assert result.unresolved_placeholders == []
    assert "人工标题" in text
    assert "人工模型概述" in text
    assert "2026-05-21" in text
    assert "0.2500" in text


def test_word_report_renders_v2_pmml_scoring_summary_for_old_and_new_templates(
    tmp_path: Path,
):
    output = tmp_path / "pmml-report.docx"

    result = write_validation_word(
        _make_pmml_results(),
        template_path=_make_pmml_template(tmp_path / "pmml-template.docx"),
        output_path=output,
        image_output_dir=tmp_path / "pmml-images",
    )

    assert result.unresolved_placeholders == []
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert text.count("PMML打分测试") == 2
    assert text.count("全量 3 行") == 2
    assert "三方分数对比" not in text
    assert len(document.inline_shapes) >= 1
