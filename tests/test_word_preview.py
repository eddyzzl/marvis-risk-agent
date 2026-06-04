from base64 import b64decode
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
import xml.etree.ElementTree as ET

from docx import Document

from riskmodel_checker.output.word_preview import docx_to_html_preview


PNG_BYTES = b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
)
CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
PACKAGE_REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
OFFICE_REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


def test_docx_preview_replaces_browser_unsupported_images_with_notice(tmp_path: Path):
    report_path = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("报告图片")
    document.add_paragraph().add_run().add_picture(BytesIO(PNG_BYTES))
    document.save(report_path)
    _append_body_image_with_unsupported_content_type(report_path)

    preview_html = docx_to_html_preview(report_path)

    assert 'class="doc-image"' in preview_html
    assert "data:image/png;base64" in preview_html
    assert "data:image/x-emf" not in preview_html
    assert "unsupported-image-note" in preview_html
    assert "预览暂不支持 EMF 图片" in preview_html
    assert "1 张图片仅在 Word 中保留" in preview_html


def _append_body_image_with_unsupported_content_type(docx_path: Path) -> None:
    unsupported_rid = "rIdUnsupportedImage"
    replacement_entries: dict[str, bytes] = {}
    with ZipFile(docx_path) as archive:
        for name in archive.namelist():
            replacement_entries[name] = archive.read(name)

    replacement_entries["[Content_Types].xml"] = _with_emf_content_type(
        replacement_entries["[Content_Types].xml"]
    )
    replacement_entries["word/_rels/document.xml.rels"] = _with_image_relationship(
        replacement_entries["word/_rels/document.xml.rels"],
        unsupported_rid,
    )
    replacement_entries["word/document.xml"] = _with_unsupported_image_paragraph(
        replacement_entries["word/document.xml"],
        unsupported_rid,
    )
    replacement_entries["word/media/unsupported.emf"] = b"not-a-browser-image"

    rebuilt_path = docx_path.with_suffix(".rebuilt.docx")
    with ZipFile(rebuilt_path, "w", ZIP_DEFLATED) as archive:
        for name, data in replacement_entries.items():
            archive.writestr(name, data)
    rebuilt_path.replace(docx_path)


def _with_emf_content_type(content_types_xml: bytes) -> bytes:
    ET.register_namespace("", CONTENT_TYPES_NS)
    root = ET.fromstring(content_types_xml)
    has_emf_default = any(
        child.tag == f"{{{CONTENT_TYPES_NS}}}Default"
        and child.attrib.get("Extension") == "emf"
        for child in root
    )
    if not has_emf_default:
        ET.SubElement(
            root,
            f"{{{CONTENT_TYPES_NS}}}Default",
            Extension="emf",
            ContentType="image/x-emf",
        )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _with_image_relationship(rels_xml: bytes, rid: str) -> bytes:
    ET.register_namespace("", PACKAGE_REL_NS)
    root = ET.fromstring(rels_xml)
    ET.SubElement(
        root,
        f"{{{PACKAGE_REL_NS}}}Relationship",
        Id=rid,
        Type=f"{OFFICE_REL_NS}/image",
        Target="media/unsupported.emf",
    )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _with_unsupported_image_paragraph(document_xml: bytes, rid: str) -> bytes:
    ET.register_namespace("w", W_NS)
    ET.register_namespace("a", A_NS)
    ET.register_namespace("r", OFFICE_REL_NS)
    root = ET.fromstring(document_xml)
    body = root.find(f"{{{W_NS}}}body")
    assert body is not None

    paragraph = ET.Element(f"{{{W_NS}}}p")
    run = ET.SubElement(paragraph, f"{{{W_NS}}}r")
    drawing = ET.SubElement(run, f"{{{W_NS}}}drawing")
    blip = ET.SubElement(drawing, f"{{{A_NS}}}blip")
    blip.set(f"{{{OFFICE_REL_NS}}}embed", rid)

    sect_pr = body.find(f"{{{W_NS}}}sectPr")
    if sect_pr is None:
        body.append(paragraph)
    else:
        body.insert(list(body).index(sect_pr), paragraph)
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)
