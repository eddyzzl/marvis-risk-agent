from __future__ import annotations

from concurrent.futures import ThreadPoolExecutor
from copy import deepcopy
import hashlib
from io import BytesIO
import json
from pathlib import Path
import sqlite3
from types import SimpleNamespace

from docx import Document
from openpyxl import load_workbook
import pytest

from marvis.data.workspace import data_semantic_mapping_hash
from marvis.db import TaskRepository
from marvis.output.strategy_report_bundle import render_strategy_report_bundle
from marvis.packs.strategy import tools as strategy_tools
from marvis.packs.strategy.candidate_stability_tools import (
    ARTIFACT_KIND as CANDIDATE_STABILITY_ARTIFACT_KIND,
    resolve_candidate_monthly_stability_inputs,
    run_measure_candidate_monthly_stability,
)
from marvis.packs.strategy.errors import StrategyError
from marvis.packs.strategy.impact_cube_tools import (
    IMPACT_CUBE_MEASUREMENT_AUDIT_KIND,
    run_measure_strategy_impact_cube,
)
from marvis.packs.strategy.pool_impact_tools import run_measure_pool_impact
from marvis.packs.strategy.pool_stability_tools import (
    POOL_STABILITY_MEASUREMENT_AUDIT_KIND,
    run_measure_strategy_pool_stability,
)
from marvis.packs.strategy.pool_validation_tools import (
    run_measure_strategy_pool_validation,
)
from marvis.packs.strategy.pool_tools import (
    load_current_strategy_candidate_pool_artifact,
)
from marvis.packs.strategy.project_context_tools import (
    load_current_strategy_project_context_artifact,
    run_materialize_project_context,
)
from marvis.packs.strategy.report_bundle_tools import (
    BUILD_STRATEGY_REPORT_BUNDLE_V2_AUDIT_KIND,
    BUILD_STRATEGY_REPORT_BUNDLE_V2_TOOL_SCHEMA_VERSION,
    run_build_strategy_report_bundle_v2,
    validate_build_strategy_report_bundle_v2_tool_output,
)
from marvis.packs.strategy.sample_design_v2_tools import (
    SAMPLE_DESIGN_V2_BUNDLE_ARTIFACT_KIND,
    SAMPLE_DESIGN_V2_MEMBERSHIP_ARTIFACT_KIND,
    run_materialize_sample_design_v2,
)
from marvis.plugins.errors import SchemaValidationError
from marvis.plugins.loader import load_manifest
from marvis.plugins.schema_validation import validate_against_schema
from marvis.repositories.strategy_reports import (
    STRATEGY_REPORT_ORIGIN_TOOL,
    STRATEGY_REPORT_OUTPUT_KINDS,
    StrategyReportRepository,
)
from marvis.repositories.task_artifacts import TaskArtifactRepository
import marvis.packs.strategy.report_bundle_tools as report_tools
from test_strategy_pool_impact_tools import _setup as _impact_setup
from test_strategy_impact_cube_tools import _setup as _cube_setup
from test_strategy_pool_validation_tools import (
    _setup as _pool_validation_setup,
)


def _eq(column: str, value: object) -> dict:
    return {
        "op": "eq",
        "left": {"column": column},
        "right": {"literal": value},
    }


def _setup(tmp_path: Path) -> dict:
    fixture = _impact_setup(tmp_path, partitioned=True)
    impact_output = run_measure_pool_impact(
        fixture["request"],
        fixture["ctx"],
        fixture["runtime"],
    )
    v2_output = run_materialize_sample_design_v2(
        {
            "legacy_sample_design_ref": fixture["sample_design_ref"],
            "relationship": "nested_same_cohort",
            "scope": "exploration_only",
            "approval_population": {"inclusion": None, "exclusion": None},
            "risk_population": {"inclusion": None, "exclusion": None},
            "partitioning": {
                "method": "predicate_ast",
                "selectors": {
                    "development": _eq("sample_split", "development"),
                    "validation": _eq("sample_split", "validation"),
                    "oot": _eq("sample_split", "oot"),
                },
            },
            "maturity": {
                "status": "unavailable",
                "performance_window_days": None,
                "cutoff_date": None,
                "reason": "No row-level application date field was supplied.",
            },
            "performance_window": {"status": "provided", "days": 90},
            "observation_window": {
                "status": "unavailable",
                "start": None,
                "end": None,
            },
            "field_bindings": {
                "entity_field": None,
                "time_field": None,
                "group_field": None,
                "month_field": "apply_month",
                "weight_field": None,
                "loan_amount_field": "loan_amount",
                "overdue_amount_field": "overdue_amount",
            },
            "historical_score": {
                "status": "unavailable",
                "column": None,
                "direction": None,
                "reason": "No historical score field was supplied.",
            },
            "policy": {
                "minimum_partition_count": 1,
                "minimum_bad_count": 1,
                "minimum_label_coverage": 0.5,
                "minimum_historical_score_coverage": 0.0,
                "maximum_group_coverage_gap": 1.0,
                "diagnostic_severities": {
                    "entity_overlap": "warn",
                    "temporal_oot": "warn",
                    "risk_outside_approval": "fail",
                    "maturity": "fail",
                    "label_coverage": "fail",
                    "historical_score_coverage": "warn",
                    "group_coverage_gap": "warn",
                    "sufficiency": "warn",
                },
            },
        },
        fixture["ctx"],
        fixture["runtime"],
    )
    message = TaskRepository(fixture["settings"].db_path).add_agent_message(
        fixture["task"].id,
        role="user",
        stage="chat",
        content="生成当前准入策略的受治理报告，历史补充材料暂时没有。",
    )
    run_materialize_project_context(
        {
            "expected_revision": 0,
            "expected_revision_id": None,
            "expected_state_hash": None,
            "user_message_ref": {
                "message_id": message["id"],
                "content_hash": hashlib.sha256(
                    message["content"].encode("utf-8")
                ).hexdigest(),
            },
            "as_of": "2026-07-23",
            "scope": "贷前准入策略",
            "business_context": {"project.goal": "准入策略迭代"},
            "explicit_unavailable": ["historical_strategy_reviews"],
            "external_report_filenames": [],
        },
        fixture["ctx"],
        fixture["runtime"],
    )
    records = TaskArtifactRepository(
        fixture["settings"].db_path
    ).list_for_task(fixture["task"].id)
    membership_record = next(
        item
        for item in records
        if item["kind"] == SAMPLE_DESIGN_V2_MEMBERSHIP_ARTIFACT_KIND
    )
    bundle_record = next(
        item
        for item in records
        if item["kind"] == SAMPLE_DESIGN_V2_BUNDLE_ARTIFACT_KIND
    )
    pool = load_current_strategy_candidate_pool_artifact(
        fixture["runtime"],
        task_id=fixture["task"].id,
        strategy_type="approval",
    )
    impact_artifact = impact_output["artifacts"][0]
    design = v2_output["bundle"]["sample_design"]
    project_context = load_current_strategy_project_context_artifact(
        fixture["runtime"],
        task_id=fixture["task"].id,
    )
    assert project_context is not None
    request = {
        "title": "准入策略迭代报告",
        "status": "partial",
        "report_revision": 1,
        "previous_report_id": None,
        "previous_report_content_hash": None,
        "generated_at": "2026-07-23T08:00:00+00:00",
        "project_context_ref": {
            "artifact_id": project_context.artifact_id,
            "expected_artifact_content_hash": (
                project_context.artifact_content_hash
            ),
            "expected_revision": project_context.revision["revision"],
            "expected_revision_id": project_context.revision["revision_id"],
            "expected_state_hash": project_context.revision["state_hash"],
        },
        "sample_design_ref": {
            "membership_artifact_id": membership_record["id"],
            "expected_membership_artifact_content_hash": membership_record[
                "content_hash"
            ],
            "bundle_artifact_id": bundle_record["id"],
            "expected_bundle_artifact_content_hash": bundle_record[
                "content_hash"
            ],
            "expected_bundle_id": v2_output["bundle"]["bundle_id"],
            "expected_sample_design_id": design["sample_design_id"],
            "expected_sample_design_content_hash": design["content_hash"],
        },
        "candidate_pool_ref": {
            "strategy_type": pool.strategy_type,
            "expected_pool_revision": pool.pool["revision"],
            "expected_pool_snapshot_hash": pool.pool["snapshot_hash"],
            "expected_artifact_id": pool.artifact_id,
            "expected_artifact_content_hash": pool.artifact_content_hash,
        },
        "pool_validation_refs": [],
        "pool_impact_ref": {
            "artifact_id": impact_artifact["artifact_id"],
            "expected_artifact_content_hash": impact_artifact["content_hash"],
            "expected_assessment_id": impact_output["assessment_id"],
            "expected_assessment_content_hash": impact_output["content_hash"],
        },
        "strategy_identity": None,
        "model_evidence_ref": None,
        "training_evidence_ref": None,
        "score_evidence_ref": None,
    }
    return {
        **fixture,
        "impact_output": impact_output,
        "v2_output": v2_output,
        "request": request,
    }


def _setup_impact_cube_report(tmp_path: Path) -> dict:
    fixture = _cube_setup(tmp_path)
    impact_output = run_measure_strategy_impact_cube(
        fixture["impact_request"],
        fixture["ctx"],
        fixture["runtime"],
    )
    message = TaskRepository(fixture["settings"].db_path).add_agent_message(
        fixture["task"].id,
        role="user",
        stage="chat",
        content="生成 ImpactCube V2 策略报告。",
    )
    run_materialize_project_context(
        {
            "expected_revision": 0,
            "expected_revision_id": None,
            "expected_state_hash": None,
            "user_message_ref": {
                "message_id": message["id"],
                "content_hash": hashlib.sha256(
                    message["content"].encode("utf-8")
                ).hexdigest(),
            },
            "as_of": "2026-07-24",
            "scope": "贷前准入策略",
            "business_context": {"project.goal": "准入策略迭代"},
            "explicit_unavailable": ["historical_strategy_reviews"],
            "external_report_filenames": [],
        },
        fixture["ctx"],
        fixture["runtime"],
    )
    project_context = load_current_strategy_project_context_artifact(
        fixture["runtime"],
        task_id=fixture["task"].id,
    )
    assert project_context is not None
    pool = load_current_strategy_candidate_pool_artifact(
        fixture["runtime"],
        task_id=fixture["task"].id,
        strategy_type="approval",
    )
    request = {
        "title": "ImpactCube V2 策略报告",
        "status": "partial",
        "report_revision": 1,
        "previous_report_id": None,
        "previous_report_content_hash": None,
        "generated_at": "2026-07-24T08:00:00+00:00",
        "project_context_ref": {
            "artifact_id": project_context.artifact_id,
            "expected_artifact_content_hash": (
                project_context.artifact_content_hash
            ),
            "expected_revision": project_context.revision["revision"],
            "expected_revision_id": project_context.revision["revision_id"],
            "expected_state_hash": project_context.revision["state_hash"],
        },
        "sample_design_ref": deepcopy(fixture["sample_ref"]),
        "candidate_pool_ref": {
            "strategy_type": pool.strategy_type,
            "expected_pool_revision": pool.pool["revision"],
            "expected_pool_snapshot_hash": pool.pool["snapshot_hash"],
            "expected_artifact_id": pool.artifact_id,
            "expected_artifact_content_hash": pool.artifact_content_hash,
        },
        "pool_validation_refs": [],
        # A structurally valid but nonexistent legacy ref proves the V2 source
        # wins without authenticating or projecting stale legacy evidence.
        "pool_impact_ref": {
            "artifact_id": "f" * 64,
            "expected_artifact_content_hash": "e" * 64,
            "expected_assessment_id": "unused-legacy-impact",
            "expected_assessment_content_hash": "d" * 64,
        },
        "impact_cube_ref": {
            "artifact_id": impact_output["artifact"]["artifact_id"],
            "expected_artifact_content_hash": impact_output["artifact"][
                "content_hash"
            ],
            "expected_cube_id": impact_output["cube_id"],
            "expected_cube_content_hash": impact_output["content_hash"],
        },
        "strategy_identity": None,
        "model_evidence_ref": None,
        "training_evidence_ref": None,
        "score_evidence_ref": None,
    }
    return {
        **fixture,
        "impact_output": impact_output,
        "request": request,
    }


def _setup_pool_validation_report(
    tmp_path: Path,
    *,
    partitions: tuple[str, ...] = ("validation", "oot"),
) -> dict:
    fixture = _pool_validation_setup(tmp_path)
    validation_outputs = {
        partition: run_measure_strategy_pool_validation(
            {**fixture["validation_request"], "partition": partition},
            fixture["ctx"],
            fixture["runtime"],
        )
        for partition in partitions
    }
    workspace = fixture["workspace"]
    impact_output = run_measure_pool_impact(
        {
            "strategy_type": "approval",
            "expected_pool_revision": fixture["pool"]["revision"],
            "expected_pool_snapshot_hash": fixture["pool"]["snapshot_hash"],
            "dataset_id": fixture["dataset"].id,
            "expected_dataset_content_hash": fixture["dataset"].content_hash,
            "workspace_revision": workspace.revision,
            "workspace_generation": workspace.analysis_generation,
            "semantic_mapping_hash": data_semantic_mapping_hash(
                workspace.semantic_mapping
            ),
            "target_col": "bad",
            "sample_design_ref": fixture["request"][
                "legacy_sample_design_ref"
            ],
            "month_col": "apply_month",
            "loan_amount_col": "loan_amount",
            "overdue_amount_col": "overdue_amount",
            "comparison_mode": "absolute",
            "drop_nan_labels": True,
        },
        fixture["ctx"],
        fixture["runtime"],
    )
    message = TaskRepository(fixture["settings"].db_path).add_agent_message(
        fixture["task"].id,
        role="user",
        stage="chat",
        content="生成包含独立验证与 OOT 重放证据的策略报告。",
    )
    run_materialize_project_context(
        {
            "expected_revision": 0,
            "expected_revision_id": None,
            "expected_state_hash": None,
            "user_message_ref": {
                "message_id": message["id"],
                "content_hash": hashlib.sha256(
                    message["content"].encode("utf-8")
                ).hexdigest(),
            },
            "as_of": "2026-07-25",
            "scope": "贷前准入策略",
            "business_context": {"project.goal": "准入策略独立验证"},
            "explicit_unavailable": ["historical_strategy_reviews"],
            "external_report_filenames": [],
        },
        fixture["ctx"],
        fixture["runtime"],
    )
    project_context = load_current_strategy_project_context_artifact(
        fixture["runtime"],
        task_id=fixture["task"].id,
    )
    assert project_context is not None
    pool = load_current_strategy_candidate_pool_artifact(
        fixture["runtime"],
        task_id=fixture["task"].id,
        strategy_type="approval",
    )
    impact_artifact = impact_output["artifacts"][0]
    request = {
        "title": "准入策略独立验证报告",
        "status": "partial",
        "report_revision": 1,
        "previous_report_id": None,
        "previous_report_content_hash": None,
        "generated_at": "2026-07-25T08:00:00+00:00",
        "project_context_ref": {
            "artifact_id": project_context.artifact_id,
            "expected_artifact_content_hash": (
                project_context.artifact_content_hash
            ),
            "expected_revision": project_context.revision["revision"],
            "expected_revision_id": project_context.revision["revision_id"],
            "expected_state_hash": project_context.revision["state_hash"],
        },
        "sample_design_ref": deepcopy(fixture["sample_ref"]),
        "candidate_pool_ref": {
            "strategy_type": pool.strategy_type,
            "expected_pool_revision": pool.pool["revision"],
            "expected_pool_snapshot_hash": pool.pool["snapshot_hash"],
            "expected_artifact_id": pool.artifact_id,
            "expected_artifact_content_hash": pool.artifact_content_hash,
        },
        "pool_validation_refs": [
            {
                "partition": partition,
                "artifact_id": output["artifact"]["artifact_id"],
                "expected_artifact_content_hash": output["artifact"][
                    "content_hash"
                ],
                "expected_evidence_id": output["evidence_id"],
                "expected_evidence_content_hash": output["content_hash"],
            }
            for partition, output in validation_outputs.items()
        ],
        "pool_impact_ref": {
            "artifact_id": impact_artifact["artifact_id"],
            "expected_artifact_content_hash": impact_artifact["content_hash"],
            "expected_assessment_id": impact_output["assessment_id"],
            "expected_assessment_content_hash": impact_output["content_hash"],
        },
        "strategy_identity": None,
        "model_evidence_ref": None,
        "training_evidence_ref": None,
        "score_evidence_ref": None,
    }
    return {
        **fixture,
        "validation_outputs": validation_outputs,
        "impact_output": impact_output,
        "request": request,
    }


def _attach_candidate_stability(fixture: dict) -> dict:
    pool = load_current_strategy_candidate_pool_artifact(
        fixture["runtime"],
        task_id=fixture["task"].id,
        strategy_type="approval",
    )
    entry = pool.pool["entries"][0]
    exact_inputs = resolve_candidate_monthly_stability_inputs(
        fixture["runtime"],
        task_id=fixture["task"].id,
        user_pointer={
            "source_kind": "pool_entry",
            "strategy_type": "approval",
            "entry_id": entry["entry_id"],
        },
    )
    output = run_measure_candidate_monthly_stability(
        exact_inputs,
        fixture["ctx"],
        fixture["runtime"],
    )
    artifact = output["artifacts"][0]
    fixture["request"]["candidate_stability_ref"] = {
        "artifact_id": artifact["artifact_id"],
        "expected_artifact_content_hash": artifact["content_hash"],
        "expected_stability_id": output["stability_id"],
        "expected_stability_content_hash": output["content_hash"],
    }
    return output


def _attach_pool_stability(fixture: dict) -> dict:
    output = run_measure_strategy_pool_stability(
        fixture["request"]["impact_cube_ref"],
        fixture["ctx"],
        fixture["runtime"],
    )
    artifact = output["artifact"]
    fixture["request"]["pool_stability_ref"] = {
        "artifact_id": artifact["artifact_id"],
        "expected_artifact_content_hash": artifact["content_hash"],
        "expected_stability_id": output["stability_id"],
        "expected_stability_content_hash": output["content_hash"],
    }
    return output


def _run(fixture: dict) -> dict:
    return run_build_strategy_report_bundle_v2(
        fixture["request"],
        fixture["ctx"],
        fixture["runtime"],
    )


def test_report_source_loading_blocks_native_before_candidate_or_pool_loaders(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    task_id = "task-native-report"
    project = SimpleNamespace(
        artifact_id="1" * 64,
        artifact_content_hash="2" * 64,
        revision={
            "revision": 1,
            "revision_id": "project-context-revision-" + "3" * 24,
            "state_hash": "4" * 64,
        },
    )
    native_sample = SimpleNamespace(
        bundle={
            "sample_design": {
                "compatibility": {
                    "source_mode": "native_active_dataset",
                    "development_partition": "risk/development",
                }
            }
        }
    )
    later_loader_calls: list[str] = []

    monkeypatch.setattr(
        report_tools,
        "load_current_strategy_project_context_artifact",
        lambda _runtime, *, task_id: project,
    )
    monkeypatch.setattr(
        report_tools,
        "load_any_strategy_sample_design_v2_artifacts",
        lambda _runtime, **_kwargs: native_sample,
    )

    def forbidden_candidate_loader(*_args, **_kwargs):
        later_loader_calls.append("candidate_pool")
        raise AssertionError("candidate/pool loaders must not run")

    monkeypatch.setattr(
        report_tools,
        "load_current_strategy_candidate_pool_artifact",
        forbidden_candidate_loader,
    )
    request = {
        "project_context_ref": {
            "artifact_id": project.artifact_id,
            "expected_artifact_content_hash": (
                project.artifact_content_hash
            ),
            "expected_revision": project.revision["revision"],
            "expected_revision_id": project.revision["revision_id"],
            "expected_state_hash": project.revision["state_hash"],
        },
        "sample_design_ref": {
            "membership_artifact_id": "5" * 64,
            "expected_membership_artifact_content_hash": "6" * 64,
            "bundle_artifact_id": "7" * 64,
            "expected_bundle_artifact_content_hash": "8" * 64,
            "expected_bundle_id": "strategy-sample-design-bundle-" + "9" * 24,
            "expected_sample_design_id": "strategy-sample-design-" + "a" * 24,
            "expected_sample_design_content_hash": "b" * 64,
        },
        "candidate_pool_ref": {},
    }

    with pytest.raises(StrategyError) as raised:
        report_tools._load_sources(
            SimpleNamespace(),
            task_id=task_id,
            request=request,
        )

    assert (
        getattr(raised.value, "code", None)
        == "strategy_sample_design_v2_native_source_unsupported"
    )
    assert getattr(raised.value, "consumer", None) == "strategy_report_bundle"
    assert later_loader_calls == []


def _report_rows(fixture: dict) -> list[dict]:
    return [
        item
        for item in fixture["runtime"].task_artifacts.list_for_task(
            fixture["task"].id
        )
        if item["kind"] in set(STRATEGY_REPORT_OUTPUT_KINDS.values())
    ]


def _audit_rows(fixture: dict) -> list[sqlite3.Row]:
    with fixture["runtime"].task_artifacts.transaction() as conn:
        return conn.execute(
            "SELECT * FROM audit WHERE kind = ? ORDER BY at, id",
            (BUILD_STRATEGY_REPORT_BUNDLE_V2_AUDIT_KIND,),
        ).fetchall()


def _measurement_audit_rows(fixture: dict) -> list[sqlite3.Row]:
    with fixture["runtime"].task_artifacts.transaction() as conn:
        return conn.execute(
            "SELECT * FROM audit WHERE kind = ? ORDER BY at, id",
            (IMPACT_CUBE_MEASUREMENT_AUDIT_KIND,),
        ).fetchall()


def test_build_report_bundle_publishes_four_exact_governed_outputs(
    tmp_path: Path,
) -> None:
    fixture = _setup(tmp_path)

    output = _run(fixture)

    assert output["schema_version"] == (
        BUILD_STRATEGY_REPORT_BUNDLE_V2_TOOL_SCHEMA_VERSION
    )
    assert validate_build_strategy_report_bundle_v2_tool_output(output) == output
    rendered = render_strategy_report_bundle(output["bundle"])
    rows = _report_rows(fixture)
    assert [artifact["format"] for artifact in output["artifacts"]] == [
        "json",
        "markdown",
        "xlsx",
        "docx",
    ]
    assert {row["kind"] for row in rows} == {
        "strategy_report_bundle_json",
        "strategy_report_markdown",
        "strategy_report_xlsx",
        "strategy_report_docx",
    }
    assert all(row["origin_tool"] == STRATEGY_REPORT_ORIGIN_TOOL for row in rows)
    for artifact in output["artifacts"]:
        row = next(item for item in rows if item["id"] == artifact["artifact_id"])
        assert Path(row["path"]).read_bytes() == rendered[artifact["format"]]
        assert row["provenance"]["report_id"] == output["report_id"]
        assert row["provenance"]["bundle_content_sha256"] == output["content_hash"]
        assert artifact["download_url"].endswith(
            f"?expected_content_hash={artifact['content_hash']}"
        )
    current = StrategyReportRepository(
        fixture["settings"].db_path
    ).get_current(task_id=fixture["task"].id, strategy_id=None)
    assert current is not None
    assert current["bundle"] == output["bundle"]
    audits = _audit_rows(fixture)
    assert len(audits) == 1
    assert audits[0]["target_ref"] == output["report_id"]
    audit_detail = json.loads(str(audits[0]["detail_json"]))
    assert set(audit_detail["output_artifacts"]) == {
        "json",
        "markdown",
        "xlsx",
        "docx",
    }
    assert audit_detail["output_artifacts"]["docx"]["kind"] == (
        "strategy_report_docx"
    )
    assert audit_detail["source_artifacts"]["pool_validations"] == {}
    impact = next(
        section
        for section in output["bundle"]["sections"]
        if section["key"] == "impact_assessment"
    )
    assert all(
        not table["table_id"].startswith(
            "strategy_pool_independent_validation"
        )
        for table in impact["tables"]
    )

    tool = next(
        item
        for item in load_manifest(
            Path(__file__).parents[1] / "marvis" / "packs" / "strategy",
            builtin=True,
        ).tools
        if item.name == "build_report_bundle_v2"
    )
    validate_against_schema(
        fixture["request"],
        tool.input_schema,
        label="report bundle input",
    )
    validate_against_schema(output, tool.output_schema, label="report bundle output")


def test_report_bundle_projects_platform_bound_pool_validation_refs_exactly(
    tmp_path: Path,
) -> None:
    fixture = _setup_pool_validation_report(tmp_path)

    output = _run(fixture)

    assert len(output["bundle"]["sections"]) == 7
    impact = next(
        section
        for section in output["bundle"]["sections"]
        if section["key"] == "impact_assessment"
    )
    summary = next(
        table
        for table in impact["tables"]
        if table["table_id"] == "strategy_pool_independent_validation_summary"
    )
    assert summary["sheet_key"] == "10_validation"
    assert [
        (
            row["cells"]["partition"]["value"],
            row["cells"]["population_count"]["value"],
            row["cells"]["labelled_count"]["value"],
            row["cells"]["validation_status"]["value"],
        )
        for row in summary["rows"]
    ] == [
        ("validation", 2, 2, "independent_evidence"),
        ("oot", 2, 1, "independent_evidence"),
    ]
    validation_summary_row = summary["rows"][0]["cells"]
    assert validation_summary_row["overall_bad_count"]["value"] == 1
    assert validation_summary_row["loan_amount_sum"]["value"] == 330.0
    assert validation_summary_row["loan_amount_coverage_count"]["value"] == 2
    assert validation_summary_row["overdue_amount_sum"]["value"] == 10.0
    assert validation_summary_row["paired_coverage_count"]["value"] == 2
    assert validation_summary_row["paired_loan_amount_sum"]["value"] == 330.0
    assert validation_summary_row["paired_overdue_amount_sum"]["value"] == 10.0
    assert validation_summary_row["paired_overdue_rate"]["value"] == pytest.approx(
        10.0 / 330.0
    )
    monthly = next(
        table
        for table in impact["tables"]
        if table["table_id"]
        == "strategy_pool_independent_validation_monthly"
    )
    validation_month = next(
        row["cells"]
        for row in monthly["rows"]
        if row["cells"]["partition"]["value"] == "validation"
    )
    assert validation_month["labelled_count"]["value"] == 2
    assert validation_month["label_coverage"]["value"] == 1.0
    assert validation_month["bad_count"]["value"] == 1
    assert validation_month["loan_amount_sum"]["value"] == 330.0
    assert validation_month["overdue_amount_sum"]["value"] == 10.0
    assert validation_month["paired_coverage_count"]["value"] == 2
    assert validation_month["paired_overdue_rate"]["value"] == pytest.approx(
        10.0 / 330.0
    )
    assert summary["effect_stage"] is None
    assert output["bundle"]["effect_stages"] == ["backtested"]
    validation_stages = [
        item
        for item in impact["stage_evidence"]
        if item["binding"]["result_ref"]["kind"] == "strategy_validation"
    ]
    assert validation_stages == []
    assert any(
        item["code"]
        == "pool_validation_claim_suppressed_by_validation_blocker"
        for item in impact["red_flags"]
    )
    final_document = next(
        section
        for section in output["bundle"]["sections"]
        if section["key"] == "final_document"
    )
    final_fields = {
        item["field_id"]: item["field"]["value"]
        for item in final_document["summary_fields"]
    }
    assert final_fields["evidence_stages"] == ["backtested"]
    assert final_fields["validation_statuses"] == ["independent_evidence"]
    assert (
        final_fields["validation_conclusion"]
        == (
            "independent_replay_evidence_available_"
            "claim_suppressed_by_validation_blocker"
        )
    )
    assert final_fields["adoption_status"] == "not_adopted"
    assert final_fields["deployment_status"] == "not_deployed"
    assert final_fields["creates_strategy"] is False
    validation_refs = [
        ref
        for ref in output["bundle"]["strategy_artifact_refs"]
        if ref["kind"] == "strategy_validation"
    ]
    assert {
        ref["ref_id"] for ref in validation_refs
    } == {
        value["artifact"]["artifact_id"]
        for value in fixture["validation_outputs"].values()
    }
    serialized = json.dumps(
        impact,
        ensure_ascii=False,
        sort_keys=True,
    ).lower()
    assert "psi" not in serialized
    assert "stability" not in serialized
    assert output["not_created_strategy"] is True
    assert output["not_adopted"] is True
    assert output["not_deployed"] is True
    audit = json.loads(str(_audit_rows(fixture)[0]["detail_json"]))
    assert set(audit["source_artifacts"]["pool_validations"]) == {
        "validation",
        "oot",
    }


def test_report_bundle_tool_schema_version_is_v6() -> None:
    assert (
        BUILD_STRATEGY_REPORT_BUNDLE_V2_TOOL_SCHEMA_VERSION
        == "strategy.build-report-bundle-v2-tool.v6"
    )


def test_legacy_report_plan_without_pool_stability_ref_replays_as_none(
    tmp_path: Path,
) -> None:
    fixture = _setup_impact_cube_report(tmp_path)
    legacy_request = deepcopy(fixture["request"])
    explicit_request = {**deepcopy(legacy_request), "pool_stability_ref": None}

    normalized_legacy = report_tools._validate_inputs(legacy_request)
    normalized_explicit = report_tools._validate_inputs(explicit_request)
    assert normalized_legacy == normalized_explicit
    assert normalized_legacy["pool_stability_ref"] is None
    assert report_tools._request_hash(
        normalized_legacy
    ) == report_tools._request_hash(normalized_explicit)

    fixture["request"] = legacy_request
    first = _run(fixture)
    fixture["request"] = explicit_request
    replay = _run(fixture)

    assert replay == first
    assert len(_report_rows(fixture)) == 4
    assert len(_audit_rows(fixture)) == 1


def test_legacy_report_plan_without_pool_validation_refs_replays_as_empty(
    tmp_path: Path,
) -> None:
    fixture = _setup(tmp_path)
    explicit_request = deepcopy(fixture["request"])
    legacy_request = deepcopy(explicit_request)
    legacy_request.pop("pool_validation_refs")

    normalized_legacy = report_tools._validate_inputs(legacy_request)
    normalized_explicit = report_tools._validate_inputs(explicit_request)
    assert normalized_legacy == normalized_explicit
    assert normalized_legacy["pool_validation_refs"] == []
    assert report_tools._request_hash(
        normalized_legacy
    ) == report_tools._request_hash(normalized_explicit)

    fixture["request"] = legacy_request
    first = _run(fixture)
    fixture["request"] = explicit_request
    replay = _run(fixture)

    assert replay == first
    assert len(_report_rows(fixture)) == 4
    assert len(_audit_rows(fixture)) == 1


def test_report_bundle_rejects_duplicate_pool_validation_partition_refs(
    tmp_path: Path,
) -> None:
    fixture = _setup_pool_validation_report(tmp_path)
    duplicate = deepcopy(fixture["request"]["pool_validation_refs"][0])
    duplicate["artifact_id"] = "a" * 64
    fixture["request"]["pool_validation_refs"] = [
        fixture["request"]["pool_validation_refs"][0],
        duplicate,
    ]

    with pytest.raises(
        StrategyError,
        match="pool_validation_refs duplicates validation",
    ):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []


def test_report_bundle_revalidates_exact_pool_validation_before_commit(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    fixture = _setup_pool_validation_report(tmp_path)
    artifact_id = fixture["validation_outputs"]["validation"]["artifact"][
        "artifact_id"
    ]
    record = fixture["runtime"].task_artifacts.get_for_task(
        fixture["task"].id,
        artifact_id,
    )
    assert record is not None
    source_path = Path(record["path"])
    original_render = report_tools.render_strategy_report_bundle

    def tamper_after_render(bundle):
        rendered = original_render(bundle)
        source_path.write_bytes(source_path.read_bytes() + b"\n")
        return rendered

    monkeypatch.setattr(
        report_tools,
        "render_strategy_report_bundle",
        tamper_after_render,
    )

    with pytest.raises(
        StrategyError,
        match="bytes changed|canonical bytes changed",
    ):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []


def test_report_bundle_exact_pool_validation_refs_are_idempotent_and_do_not_drift(
    tmp_path: Path,
) -> None:
    fixture = _setup_pool_validation_report(tmp_path)

    first = _run(fixture)
    replay = _run(fixture)

    assert replay == first
    assert len(_report_rows(fixture)) == 4
    assert len(_audit_rows(fixture)) == 1


def test_new_pool_validation_evidence_does_not_rebind_an_existing_report_plan(
    tmp_path: Path,
) -> None:
    fixture = _setup_pool_validation_report(
        tmp_path,
        partitions=("validation",),
    )
    planned_refs = deepcopy(fixture["request"]["pool_validation_refs"])
    run_measure_strategy_pool_validation(
        {**fixture["validation_request"], "partition": "oot"},
        fixture["ctx"],
        fixture["runtime"],
    )

    output = _run(fixture)

    assert fixture["request"]["pool_validation_refs"] == planned_refs
    impact = next(
        section
        for section in output["bundle"]["sections"]
        if section["key"] == "impact_assessment"
    )
    summary = next(
        table
        for table in impact["tables"]
        if table["table_id"] == "strategy_pool_independent_validation_summary"
    )
    assert [
        row["cells"]["partition"]["value"] for row in summary["rows"]
    ] == ["validation"]
    audit = json.loads(str(_audit_rows(fixture)[0]["detail_json"]))
    assert set(audit["source_artifacts"]["pool_validations"]) == {
        "validation"
    }


def test_report_bundle_projects_exact_candidate_stability_into_all_outputs(
    tmp_path: Path,
) -> None:
    fixture = _setup(tmp_path)
    stability = _attach_candidate_stability(fixture)

    output = _run(fixture)

    candidate_section = next(
        section
        for section in output["bundle"]["sections"]
        if section["key"] == "candidate_combinations"
    )
    table = next(
        item
        for item in candidate_section["tables"]
        if item["table_id"] == "candidate_monthly_stability"
    )
    assert table["sheet_key"] == "appendix_candidate_stability"
    assert table["content_class"] == "monthly_summary"
    assert table["effect_stage"] == "backtested"
    assert table["rows"][0]["row_id"] == "candidate-stability-baseline"
    assert len(table["rows"]) == stability["month_count"] + 1
    assert any(
        flag["level"] == "amber"
        and "低样本" in flag["message"]
        for flag in candidate_section["red_flags"]
    )
    rendered = render_strategy_report_bundle(output["bundle"])
    assert "候选逐月稳定性" in rendered["markdown"].decode("utf-8")
    assert set(rendered) == {"json", "markdown", "xlsx", "docx"}
    audit = json.loads(str(_audit_rows(fixture)[0]["detail_json"]))
    assert audit["source_artifacts"]["candidate_stability"] == {
        "artifact_id": fixture["request"]["candidate_stability_ref"][
            "artifact_id"
        ],
        "content_hash": fixture["request"]["candidate_stability_ref"][
            "expected_artifact_content_hash"
        ],
        "stability_id": stability["stability_id"],
        "stability_content_hash": stability["content_hash"],
    }


def test_report_bundle_projects_exact_pool_stability_into_all_outputs_and_audit(
    tmp_path: Path,
) -> None:
    fixture = _setup_impact_cube_report(tmp_path)
    stability = _attach_pool_stability(fixture)

    output = _run(fixture)

    impact = next(
        section
        for section in output["bundle"]["sections"]
        if section["key"] == "impact_assessment"
    )
    table = next(
        item
        for item in impact["tables"]
        if item["table_id"] == "strategy_pool_stability_summary"
    )
    assert table["sheet_key"] == "10_validation"
    assert table["effect_stage"] is None
    assert 0 < len(table["rows"]) <= 8
    assert all(
        stage["binding"]["result_ref"]["kind"] != "pool_stability"
        for stage in impact["stage_evidence"]
    )
    assert output["bundle"]["effect_stages"] == ["backtested"]
    assert {
        (stage["population"], stage["partition"])
        for stage in impact["stage_evidence"]
    } == {
        ("approval", "development"),
        ("risk", "development"),
    }
    assert "oot_claim_suppressed_by_validation_blocker" in {
        flag["code"] for flag in impact["red_flags"]
    }

    published = {}
    for artifact in output["artifacts"]:
        record = fixture["runtime"].task_artifacts.get_for_task(
            fixture["task"].id,
            artifact["artifact_id"],
        )
        assert record is not None
        published[artifact["format"]] = Path(record["path"]).read_bytes()
    title = "策略池跨分区分布稳定性摘要（非效果验证）"
    assert "strategy_pool_stability_summary" in published["json"].decode("utf-8")
    assert title in published["markdown"].decode("utf-8")
    workbook = load_workbook(BytesIO(published["xlsx"]), read_only=True)
    assert title in {
        cell.value
        for row in workbook["10_validation"].iter_rows()
        for cell in row
    }
    document = Document(BytesIO(published["docx"]))
    visible_docx = "\n".join(
        [
            *(paragraph.text for paragraph in document.paragraphs),
            *(
                cell.text
                for document_table in document.tables
                for row in document_table.rows
                for cell in row.cells
            ),
        ]
    )
    assert title in visible_docx

    audit = json.loads(str(_audit_rows(fixture)[0]["detail_json"]))
    assert audit["source_artifacts"]["pool_stability"] == {
        "artifact_id": fixture["request"]["pool_stability_ref"]["artifact_id"],
        "content_hash": fixture["request"]["pool_stability_ref"][
            "expected_artifact_content_hash"
        ],
        "stability_id": stability["stability_id"],
        "stability_content_hash": stability["content_hash"],
        "producer_run_ref": stability["producer_run_ref"],
    }


def test_report_bundle_rejects_unknown_pool_stability_ref_fields(
    tmp_path: Path,
) -> None:
    fixture = _setup_impact_cube_report(tmp_path)
    fixture["request"]["pool_stability_ref"] = {
        "artifact_id": "a" * 64,
        "expected_artifact_content_hash": "b" * 64,
        "expected_stability_id": "strategy-pool-stability-" + ("1" * 24),
        "expected_stability_content_hash": "c" * 64,
        "latest": True,
    }

    with pytest.raises(
        StrategyError,
        match="pool_stability_ref.*fields",
    ):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []


@pytest.mark.parametrize(
    "mutation",
    [
        "file_deleted",
        "file_tampered",
        "registry_deleted",
        "registry_tampered",
        "audit_deleted",
        "audit_tampered",
        "source_cube_tampered",
    ],
)
def test_report_bundle_pool_stability_source_mutations_fail_closed(
    tmp_path: Path,
    mutation: str,
) -> None:
    fixture = _setup_impact_cube_report(tmp_path)
    stability = _attach_pool_stability(fixture)
    task_id = fixture["task"].id
    artifact_id = stability["artifact"]["artifact_id"]
    record = fixture["runtime"].task_artifacts.get_for_task(
        task_id,
        artifact_id,
    )
    assert record is not None

    if mutation == "file_deleted":
        Path(record["path"]).unlink()
    elif mutation == "file_tampered":
        Path(record["path"]).write_bytes(b"{}")
    elif mutation == "source_cube_tampered":
        cube_record = fixture["runtime"].task_artifacts.get_for_task(
            task_id,
            fixture["impact_output"]["artifact"]["artifact_id"],
        )
        assert cube_record is not None
        Path(cube_record["path"]).write_bytes(b"{}")
    else:
        with fixture["runtime"].task_artifacts.transaction() as conn:
            if mutation == "registry_deleted":
                conn.execute(
                    "DELETE FROM task_artifacts WHERE id = ?",
                    (artifact_id,),
                )
            elif mutation == "registry_tampered":
                conn.execute(
                    "DROP TRIGGER trg_task_artifacts_immutable_update"
                )
                conn.execute(
                    "UPDATE task_artifacts SET content_hash = ? WHERE id = ?",
                    ("f" * 64, artifact_id),
                )
            elif mutation == "audit_deleted":
                conn.execute(
                    "DELETE FROM audit WHERE kind = ? AND target_ref = ?",
                    (
                        POOL_STABILITY_MEASUREMENT_AUDIT_KIND,
                        stability["producer_run_ref"]["ref_id"],
                    ),
                )
            elif mutation == "audit_tampered":
                conn.execute(
                    """
                    UPDATE audit
                       SET detail_json = ?
                     WHERE kind = ? AND target_ref = ?
                    """,
                    (
                        "{}",
                        POOL_STABILITY_MEASUREMENT_AUDIT_KIND,
                        stability["producer_run_ref"]["ref_id"],
                    ),
                )
            else:
                raise AssertionError(f"unsupported mutation {mutation}")
            conn.commit()

    with pytest.raises(StrategyError):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []


def test_report_bundle_exact_pool_stability_is_idempotent_and_never_rebinds_latest(
    tmp_path: Path,
) -> None:
    fixture = _setup_impact_cube_report(tmp_path)
    planned = _attach_pool_stability(fixture)
    planned_ref = deepcopy(fixture["request"]["pool_stability_ref"])

    newer_cube = run_measure_strategy_impact_cube(
        {
            **deepcopy(fixture["impact_request"]),
            "partitions": ["development", "validation", "oot"],
        },
        fixture["ctx"],
        fixture["runtime"],
    )
    newer_stability = run_measure_strategy_pool_stability(
        {
            "artifact_id": newer_cube["artifact"]["artifact_id"],
            "expected_artifact_content_hash": newer_cube["artifact"][
                "content_hash"
            ],
            "expected_cube_id": newer_cube["cube_id"],
            "expected_cube_content_hash": newer_cube["content_hash"],
        },
        fixture["ctx"],
        fixture["runtime"],
    )
    assert newer_stability["artifact"]["artifact_id"] != planned["artifact"][
        "artifact_id"
    ]
    assert fixture["request"]["pool_stability_ref"] == planned_ref

    first = _run(fixture)
    replay = _run(fixture)

    assert replay == first
    assert len(_report_rows(fixture)) == 4
    assert len(_audit_rows(fixture)) == 1
    audit = json.loads(str(_audit_rows(fixture)[0]["detail_json"]))
    assert audit["source_artifacts"]["pool_stability"][
        "artifact_id"
    ] == planned["artifact"]["artifact_id"]
    assert audit["source_artifacts"]["pool_stability"][
        "producer_run_ref"
    ] == planned["producer_run_ref"]


def test_report_bundle_pool_stability_rejects_legacy_pool_impact_fallback(
    tmp_path: Path,
) -> None:
    fixture = _setup_impact_cube_report(tmp_path)
    _attach_pool_stability(fixture)
    workspace = fixture["workspace"]
    legacy_sample_ref = fixture["v2_output"]["bundle"]["sample_design"][
        "compatibility"
    ]["legacy_development_ref"]
    legacy_impact = run_measure_pool_impact(
        {
            "strategy_type": "approval",
            "expected_pool_revision": fixture["pool"]["revision"],
            "expected_pool_snapshot_hash": fixture["pool"]["snapshot_hash"],
            "dataset_id": fixture["dataset"].id,
            "expected_dataset_content_hash": fixture["dataset"].content_hash,
            "workspace_revision": workspace.revision,
            "workspace_generation": workspace.analysis_generation,
            "semantic_mapping_hash": data_semantic_mapping_hash(
                workspace.semantic_mapping
            ),
            "target_col": "bad",
            "sample_design_ref": legacy_sample_ref,
            "month_col": "apply_month",
            "loan_amount_col": "loan_amount",
            "overdue_amount_col": "overdue_amount",
            "comparison_mode": "absolute",
            "drop_nan_labels": True,
        },
        fixture["ctx"],
        fixture["runtime"],
    )
    legacy_artifact = legacy_impact["artifacts"][0]
    fixture["request"].pop("impact_cube_ref")
    fixture["request"]["pool_impact_ref"] = {
        "artifact_id": legacy_artifact["artifact_id"],
        "expected_artifact_content_hash": legacy_artifact["content_hash"],
        "expected_assessment_id": legacy_impact["assessment_id"],
        "expected_assessment_content_hash": legacy_impact["content_hash"],
    }

    with pytest.raises(
        StrategyError,
        match="Pool stability requires the report's exact ImpactCube",
    ):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []


def test_report_bundle_passes_exact_voting_search_binding_and_audits_it(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    fixture = _setup(tmp_path)
    search_id = "voting-search-" + ("1" * 32)
    search_hash = "c" * 64
    voting_ref = {
        "artifact_id": "a" * 64,
        "expected_artifact_content_hash": "b" * 64,
        "expected_search_id": search_id,
        "expected_search_content_hash": search_hash,
    }
    fixture["request"]["voting_candidate_search_ref"] = voting_ref
    binding = SimpleNamespace(
        artifact_id=voting_ref["artifact_id"],
        artifact_content_hash=voting_ref[
            "expected_artifact_content_hash"
        ],
        result={
            "search_id": search_id,
            "content_hash": search_hash,
        },
    )
    observed = {}
    original_adapter = report_tools.build_strategy_report_bundle_source_inputs

    def load_search(runtime, **kwargs):
        observed["loader"] = (runtime, kwargs)
        return binding

    def capture_search_binding(**kwargs):
        observed["binding"] = kwargs["voting_candidate_search"]
        kwargs.pop("voting_candidate_search")
        return original_adapter(**kwargs)

    monkeypatch.setattr(
        report_tools,
        "load_voting_candidate_search_artifact",
        load_search,
        raising=False,
    )
    monkeypatch.setattr(
        report_tools,
        "require_voting_candidate_search_artifact_binding_on_connection",
        lambda conn, actual: None,
        raising=False,
    )
    monkeypatch.setattr(
        report_tools,
        "build_strategy_report_bundle_source_inputs",
        capture_search_binding,
    )

    output = _run(fixture)

    assert observed["binding"] is binding
    assert observed["loader"][1] == {
        "task_id": fixture["task"].id,
        **voting_ref,
    }
    audit = json.loads(str(_audit_rows(fixture)[0]["detail_json"]))
    assert audit["source_artifacts"]["voting_candidate_search"] == {
        "artifact_id": binding.artifact_id,
        "content_hash": binding.artifact_content_hash,
        "search_id": search_id,
        "search_content_hash": search_hash,
    }
    assert output["schema_version"] == (
        "strategy.build-report-bundle-v2-tool.v6"
    )


def test_report_bundle_rejects_unknown_voting_search_ref_fields(
    tmp_path: Path,
) -> None:
    fixture = _setup(tmp_path)
    fixture["request"]["voting_candidate_search_ref"] = {
        "artifact_id": "a" * 64,
        "expected_artifact_content_hash": "b" * 64,
        "expected_search_id": "voting-search-" + ("1" * 32),
        "expected_search_content_hash": "c" * 64,
        "schema_version": "forged.v1",
    }

    with pytest.raises(
        StrategyError,
        match="voting_candidate_search_ref.*fields",
    ):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []


def test_report_bundle_manifest_accepts_exactly_supported_impact_source_shapes(
    tmp_path: Path,
) -> None:
    tool = next(
        item
        for item in load_manifest(
            Path(__file__).parents[1] / "marvis" / "packs" / "strategy",
            builtin=True,
        ).tools
        if item.name == "build_report_bundle_v2"
    )
    legacy_only = deepcopy(_setup(tmp_path / "legacy")["request"])
    cube_and_legacy = deepcopy(
        _setup_impact_cube_report(tmp_path / "cube")["request"]
    )
    cube_only = deepcopy(cube_and_legacy)
    cube_only.pop("pool_impact_ref")
    neither = deepcopy(legacy_only)
    neither.pop("pool_impact_ref")

    validate_against_schema(
        legacy_only,
        tool.input_schema,
        label="legacy PoolImpact-only report input",
    )
    validate_against_schema(
        cube_only,
        tool.input_schema,
        label="ImpactCube-only report input",
    )
    validate_against_schema(
        cube_and_legacy,
        tool.input_schema,
        label="ImpactCube-preferred report input with legacy ref",
    )
    with pytest.raises(SchemaValidationError):
        validate_against_schema(
            neither,
            tool.input_schema,
            label="report input without impact evidence",
        )


def test_build_report_bundle_prefers_authenticated_impact_cube_v2(
    tmp_path: Path,
) -> None:
    fixture = _setup_impact_cube_report(tmp_path)

    output = _run(fixture)

    refs = output["bundle"]["strategy_artifact_refs"]
    assert {
        (item["kind"], item["ref_id"])
        for item in refs
    } == {
        (
            "strategy_candidate_pool",
            fixture["pool_artifact"]["artifact_id"],
        ),
        (
            "strategy_impact",
            fixture["impact_output"]["artifact"]["artifact_id"],
        ),
    }
    assert fixture["request"]["pool_impact_ref"]["artifact_id"] not in {
        item["ref_id"] for item in refs
    }
    assert output["bundle"]["effect_stages"] == ["backtested"]
    impact = output["bundle"]["sections"][5]
    assert {
        (item["population"], item["partition"])
        for item in impact["stage_evidence"]
    } == {
        ("approval", "development"),
        ("risk", "development"),
    }
    assert "oot_claim_suppressed_by_validation_blocker" in {
        item["code"] for item in impact["red_flags"]
    }
    final_fields = {
        item["field_id"]: item["field"]["value"]
        for item in output["bundle"]["sections"][6]["summary_fields"]
    }
    assert final_fields["evidence_stages"] == ["backtested"]
    assert final_fields["validation_statuses"] == ["unvalidated"]
    slice_table = next(
        table
        for table in impact["tables"]
        if table["table_id"] == "impact_cube_slices"
    )
    assert {
        row["cells"]["population"]["value"]
        for row in slice_table["rows"]
    } == {"approval", "risk"}
    assert {
        row["cells"]["partition"]["value"]
        for row in slice_table["rows"]
    } == {"development", "validation"}
    assert {
        row["cells"]["family"]["value"]
        for row in slice_table["rows"]
    } >= {"overall", "month", "group", "segment", "group_month"}
    economics = next(
        table
        for table in impact["tables"]
        if table["table_id"] == "impact_cube_economics"
    )
    assert all(
        row["cells"]["new"]["value"] is None
        for row in economics["rows"]
    )
    rendered = render_strategy_report_bundle(output["bundle"])
    assert fixture["impact_output"]["cube_id"].encode("utf-8") in rendered[
        "json"
    ]
    assert rendered["xlsx"].startswith(b"PK")
    assert rendered["docx"].startswith(b"PK")
    assert len(_report_rows(fixture)) == 4
    audit = _audit_rows(fixture)[0]
    detail = json.loads(str(audit["detail_json"]))
    assert detail["source_artifacts"]["impact_cube"] == {
        "artifact_id": fixture["impact_output"]["artifact"]["artifact_id"],
        "content_hash": fixture["impact_output"]["artifact"]["content_hash"],
        "producer_run_ref": fixture["impact_output"]["producer_run_ref"],
    }
    assert detail["source_artifacts"]["pool_impact"] is None


def test_build_report_bundle_rejects_impact_cube_provenance_drift(
    tmp_path: Path,
) -> None:
    fixture = _setup_impact_cube_report(tmp_path)
    artifact_id = fixture["impact_output"]["artifact"]["artifact_id"]
    with fixture["runtime"].task_artifacts.transaction() as conn:
        conn.execute("DROP TRIGGER trg_task_artifacts_immutable_update")
        row = conn.execute(
            "SELECT provenance_json FROM task_artifacts WHERE id = ?",
            (artifact_id,),
        ).fetchone()
        assert row is not None
        provenance = json.loads(str(row["provenance_json"]))
        provenance["cube_content_hash"] = "0" * 64
        conn.execute(
            "UPDATE task_artifacts SET provenance_json = ? WHERE id = ?",
            (
                json.dumps(
                    provenance,
                    ensure_ascii=False,
                    sort_keys=True,
                    separators=(",", ":"),
                    allow_nan=False,
                ),
                artifact_id,
            ),
        )
        conn.commit()

    with pytest.raises(StrategyError, match="provenance"):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []


def test_build_report_bundle_rejects_missing_or_duplicate_cube_run_audit(
    tmp_path: Path,
) -> None:
    missing = _setup_impact_cube_report(tmp_path / "missing")
    with missing["runtime"].task_artifacts.transaction() as conn:
        conn.execute(
            "DELETE FROM audit WHERE kind = ?",
            (IMPACT_CUBE_MEASUREMENT_AUDIT_KIND,),
        )
        conn.commit()

    with pytest.raises(StrategyError, match="measurement audit is missing"):
        _run(missing)
    assert _report_rows(missing) == []
    assert _audit_rows(missing) == []

    duplicate = _setup_impact_cube_report(tmp_path / "duplicate")
    original = _measurement_audit_rows(duplicate)[0]
    with duplicate["runtime"].task_artifacts.transaction() as conn:
        conn.execute(
            """
            INSERT INTO audit(
                id, kind, actor, target_ref, inputs_hash, outcome,
                detail_json, at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                "duplicate-impact-cube-measurement-audit",
                original["kind"],
                original["actor"],
                original["target_ref"],
                original["inputs_hash"],
                original["outcome"],
                original["detail_json"],
                original["at"],
            ),
        )
        conn.commit()

    with pytest.raises(
        StrategyError,
        match="measurement audit is duplicated",
    ):
        _run(duplicate)
    assert _report_rows(duplicate) == []
    assert _audit_rows(duplicate) == []


def test_build_report_bundle_rejects_tampered_cube_run_self_hash(
    tmp_path: Path,
) -> None:
    fixture = _setup_impact_cube_report(tmp_path)
    artifact_id = fixture["impact_output"]["artifact"]["artifact_id"]
    with fixture["runtime"].task_artifacts.transaction() as conn:
        conn.execute("DROP TRIGGER trg_task_artifacts_immutable_update")
        row = conn.execute(
            "SELECT provenance_json FROM task_artifacts WHERE id = ?",
            (artifact_id,),
        ).fetchone()
        assert row is not None
        provenance = json.loads(str(row["provenance_json"]))
        provenance["producer_run"]["content_hash"] = "0" * 64
        conn.execute(
            "UPDATE task_artifacts SET provenance_json = ? WHERE id = ?",
            (
                json.dumps(
                    provenance,
                    ensure_ascii=False,
                    sort_keys=True,
                    separators=(",", ":"),
                    allow_nan=False,
                ),
                artifact_id,
            ),
        )
        conn.commit()

    with pytest.raises(StrategyError, match="producer_run|self hash"):
        _run(fixture)
    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []


def test_build_report_bundle_requires_one_authenticated_impact_source(
    tmp_path: Path,
) -> None:
    fixture = _setup(tmp_path)
    fixture["request"]["pool_impact_ref"] = None
    fixture["request"]["impact_cube_ref"] = None

    with pytest.raises(
        StrategyError,
        match="requires impact_cube_ref or pool_impact_ref",
    ):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []


def test_first_report_revision_defaults_omitted_previous_head_to_null(
    tmp_path: Path,
) -> None:
    fixture = _setup(tmp_path)
    fixture["request"].pop("previous_report_id")
    fixture["request"].pop("previous_report_content_hash")

    output = _run(fixture)

    assert output["report_revision"] == 1
    assert output["bundle"]["previous_report_id"] is None
    assert validate_build_strategy_report_bundle_v2_tool_output(output) == output
    tool = next(
        item
        for item in load_manifest(
            Path(__file__).parents[1] / "marvis" / "packs" / "strategy",
            builtin=True,
        ).tools
        if item.name == "build_report_bundle_v2"
    )
    validate_against_schema(
        fixture["request"],
        tool.input_schema,
        label="report bundle input without previous head",
    )


def test_build_report_bundle_exact_retry_is_idempotent_and_validator_is_strict(
    tmp_path: Path,
) -> None:
    fixture = _setup(tmp_path)

    first = _run(fixture)
    replay = _run(fixture)

    assert replay == first
    assert len(_report_rows(fixture)) == 4
    assert len(_audit_rows(fixture)) == 1
    with fixture["runtime"].task_artifacts.transaction() as conn:
        assert conn.execute(
            "SELECT COUNT(*) FROM strategy_report_revisions"
        ).fetchone()[0] == 1

    extra = deepcopy(first)
    extra["caller_metric"] = 0.99
    with pytest.raises(StrategyError, match="output envelope"):
        validate_build_strategy_report_bundle_v2_tool_output(extra)
    forged = deepcopy(first)
    forged["artifacts"][0]["content_hash"] = "f" * 64
    with pytest.raises(StrategyError, match="artifact"):
        validate_build_strategy_report_bundle_v2_tool_output(forged)
    forged_url = deepcopy(first)
    forged_url["artifacts"][0]["download_url"] = forged_url["artifacts"][0][
        "download_url"
    ].replace(
        forged_url["artifacts"][0]["content_hash"],
        "0" * 64,
    )
    with pytest.raises(StrategyError, match="artifact"):
        validate_build_strategy_report_bundle_v2_tool_output(forged_url)
    missing_docx = deepcopy(first)
    missing_docx["artifacts"] = [
        item
        for item in missing_docx["artifacts"]
        if item["format"] != "docx"
    ]
    with pytest.raises(StrategyError, match="four canonical artifacts"):
        validate_build_strategy_report_bundle_v2_tool_output(missing_docx)
    forged_docx = deepcopy(first)
    docx = next(
        item for item in forged_docx["artifacts"] if item["format"] == "docx"
    )
    docx["content_hash"] = "f" * 64
    with pytest.raises(StrategyError, match="artifact"):
        validate_build_strategy_report_bundle_v2_tool_output(forged_docx)


def test_build_report_bundle_missing_docx_render_rolls_back_every_output(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    fixture = _setup(tmp_path)
    original_render = report_tools.render_strategy_report_bundle

    def omit_docx(bundle):
        rendered = original_render(bundle)
        rendered.pop("docx")
        return rendered

    monkeypatch.setattr(
        report_tools,
        "render_strategy_report_bundle",
        omit_docx,
    )

    with pytest.raises(StrategyError, match="invalid output set"):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []
    with fixture["runtime"].task_artifacts.transaction() as conn:
        assert conn.execute(
            "SELECT COUNT(*) FROM strategy_report_revisions"
        ).fetchone()[0] == 0


def test_build_report_bundle_rolls_back_files_rows_publication_and_audit(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    fixture = _setup(tmp_path)

    def fail_audit(*args, **kwargs):
        raise RuntimeError("audit unavailable")

    monkeypatch.setattr(
        fixture["runtime"].repo,
        "write_audit_on_connection",
        fail_audit,
    )
    with pytest.raises(RuntimeError, match="audit unavailable"):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []
    with fixture["runtime"].task_artifacts.transaction() as conn:
        assert conn.execute(
            "SELECT COUNT(*) FROM strategy_report_revisions"
        ).fetchone()[0] == 0
    report_root = (
        Path(fixture["settings"].tasks_dir)
        / fixture["task"].id
        / "strategy_reports"
    )
    assert not any(report_root.rglob("report.*")) if report_root.exists() else True


def test_build_report_bundle_identical_concurrent_writers_share_one_revision(
    tmp_path: Path,
) -> None:
    fixture = _setup(tmp_path)

    def execute() -> dict:
        runtime = strategy_tools._runtime(fixture["ctx"])
        return run_build_strategy_report_bundle_v2(
            fixture["request"],
            fixture["ctx"],
            runtime,
        )

    with ThreadPoolExecutor(max_workers=2) as executor:
        results = list(executor.map(lambda _: execute(), range(2)))

    assert results[0] == results[1]
    assert len(_report_rows(fixture)) == 4
    assert len(_audit_rows(fixture)) == 1
    with fixture["runtime"].task_artifacts.transaction() as conn:
        assert conn.execute(
            "SELECT COUNT(*) FROM strategy_report_revisions"
        ).fetchone()[0] == 1


def test_build_report_bundle_revalidates_source_tamper_before_commit(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    fixture = _setup(tmp_path)
    original_render = report_tools.render_strategy_report_bundle
    impact_path = Path(
        next(
            item
            for item in fixture["runtime"].task_artifacts.list_for_task(
                fixture["task"].id
            )
            if item["id"]
            == fixture["impact_output"]["artifacts"][0]["artifact_id"]
        )["path"]
    )

    def tampering_render(bundle):
        rendered = original_render(bundle)
        impact_path.write_bytes(b"{}")
        return rendered

    monkeypatch.setattr(
        report_tools,
        "render_strategy_report_bundle",
        tampering_render,
    )
    with pytest.raises(StrategyError, match="impact|artifact|hash|bytes"):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []


def test_build_report_bundle_revalidates_candidate_stability_before_commit(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    fixture = _setup(tmp_path)
    stability = _attach_candidate_stability(fixture)
    original_render = report_tools.render_strategy_report_bundle
    stability_path = Path(
        next(
            item
            for item in fixture["runtime"].task_artifacts.list_for_task(
                fixture["task"].id
            )
            if item["kind"] == CANDIDATE_STABILITY_ARTIFACT_KIND
            and item["id"] == stability["artifacts"][0]["artifact_id"]
        )["path"]
    )

    def tampering_render(bundle):
        rendered = original_render(bundle)
        stability_path.write_bytes(b"{}")
        return rendered

    monkeypatch.setattr(
        report_tools,
        "render_strategy_report_bundle",
        tampering_render,
    )
    with pytest.raises(
        StrategyError,
        match="stability|artifact|hash|bytes",
    ):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []


@pytest.mark.parametrize("source_kind", ["pool_stability", "impact_cube"])
def test_build_report_bundle_revalidates_pool_stability_and_exact_impact_cube_before_commit(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    source_kind: str,
) -> None:
    fixture = _setup_impact_cube_report(tmp_path)
    stability = _attach_pool_stability(fixture)
    artifact_id = (
        stability["artifact"]["artifact_id"]
        if source_kind == "pool_stability"
        else fixture["impact_output"]["artifact"]["artifact_id"]
    )
    record = fixture["runtime"].task_artifacts.get_for_task(
        fixture["task"].id,
        artifact_id,
    )
    assert record is not None
    source_path = Path(record["path"])
    original_render = report_tools.render_strategy_report_bundle

    def tampering_render(bundle):
        rendered = original_render(bundle)
        source_path.write_bytes(b"{}")
        return rendered

    monkeypatch.setattr(
        report_tools,
        "render_strategy_report_bundle",
        tampering_render,
    )
    with pytest.raises(
        StrategyError,
        match="stability|ImpactCube|artifact|hash|bytes",
    ):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []


def test_build_report_bundle_revalidates_voting_search_immediately_before_commit(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    fixture = _setup(tmp_path)
    voting_ref = {
        "artifact_id": "a" * 64,
        "expected_artifact_content_hash": "b" * 64,
        "expected_search_id": "voting-search-" + ("1" * 32),
        "expected_search_content_hash": "c" * 64,
    }
    fixture["request"]["voting_candidate_search_ref"] = voting_ref
    binding = SimpleNamespace(
        artifact_id=voting_ref["artifact_id"],
        artifact_content_hash=voting_ref[
            "expected_artifact_content_hash"
        ],
        result={
            "search_id": voting_ref["expected_search_id"],
            "content_hash": voting_ref["expected_search_content_hash"],
        },
    )
    original_adapter = report_tools.build_strategy_report_bundle_source_inputs
    guard_calls = 0

    def adapter_without_fake_binding(**kwargs):
        assert kwargs.pop("voting_candidate_search") is binding
        return original_adapter(**kwargs)

    def revalidate_search(conn, actual):
        nonlocal guard_calls
        assert actual is binding
        guard_calls += 1
        if guard_calls == 2:
            raise StrategyError("Voting search bytes changed")

    monkeypatch.setattr(
        report_tools,
        "load_voting_candidate_search_artifact",
        lambda runtime, **kwargs: binding,
        raising=False,
    )
    monkeypatch.setattr(
        report_tools,
        "require_voting_candidate_search_artifact_binding_on_connection",
        revalidate_search,
        raising=False,
    )
    monkeypatch.setattr(
        report_tools,
        "build_strategy_report_bundle_source_inputs",
        adapter_without_fake_binding,
    )

    with pytest.raises(
        StrategyError,
        match="Voting search|artifact|hash|bytes",
    ):
        _run(fixture)

    assert guard_calls == 2
    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []
    with fixture["runtime"].task_artifacts.transaction() as conn:
        assert conn.execute(
            "SELECT COUNT(*) FROM strategy_report_revisions"
        ).fetchone()[0] == 0
    report_root = (
        Path(fixture["settings"].tasks_dir)
        / fixture["task"].id
        / "strategy_reports"
    )
    assert not any(report_root.rglob("report.*")) if report_root.exists() else True


def test_build_report_bundle_revalidates_impact_cube_tamper_before_commit(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    fixture = _setup_impact_cube_report(tmp_path)
    original_render = report_tools.render_strategy_report_bundle
    impact_path = Path(
        next(
            item
            for item in fixture["runtime"].task_artifacts.list_for_task(
                fixture["task"].id
            )
            if item["id"]
            == fixture["impact_output"]["artifact"]["artifact_id"]
        )["path"]
    )

    def tampering_render(bundle):
        rendered = original_render(bundle)
        impact_path.write_bytes(b"{}")
        return rendered

    monkeypatch.setattr(
        report_tools,
        "render_strategy_report_bundle",
        tampering_render,
    )
    with pytest.raises(StrategyError, match="ImpactCube|artifact|hash|bytes"):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []


def test_build_report_bundle_revalidates_cube_run_audit_before_commit(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    fixture = _setup_impact_cube_report(tmp_path)
    original_render = report_tools.render_strategy_report_bundle

    def tampering_render(bundle):
        rendered = original_render(bundle)
        with fixture["runtime"].task_artifacts.transaction() as conn:
            conn.execute(
                """
                UPDATE audit
                   SET outcome = 'failed'
                 WHERE kind = ?
                """,
                (IMPACT_CUBE_MEASUREMENT_AUDIT_KIND,),
            )
            conn.commit()
        return rendered

    monkeypatch.setattr(
        report_tools,
        "render_strategy_report_bundle",
        tampering_render,
    )
    with pytest.raises(StrategyError, match="audit binding changed"):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []


def test_build_report_bundle_never_rebinds_a_planned_project_context(
    tmp_path: Path,
) -> None:
    fixture = _setup(tmp_path)
    previous = load_current_strategy_project_context_artifact(
        fixture["runtime"],
        task_id=fixture["task"].id,
    )
    assert previous is not None
    message = TaskRepository(fixture["settings"].db_path).add_agent_message(
        fixture["task"].id,
        role="user",
        stage="chat",
        content="补充项目目标：控制准入风险并保持通过率稳定。",
    )
    run_materialize_project_context(
        {
            "expected_revision": previous.revision["revision"],
            "expected_revision_id": previous.revision["revision_id"],
            "expected_state_hash": previous.revision["state_hash"],
            "user_message_ref": {
                "message_id": message["id"],
                "content_hash": hashlib.sha256(
                    message["content"].encode("utf-8")
                ).hexdigest(),
            },
            "as_of": "2026-07-23",
            "scope": "贷前准入策略",
            "business_context": {
                "project.goal": "控制准入风险并保持通过率稳定"
            },
            "explicit_unavailable": [],
            "external_report_filenames": [],
        },
        fixture["ctx"],
        fixture["runtime"],
    )

    with pytest.raises(
        StrategyError,
        match="project context.*exact planned revision",
    ):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []


def test_build_report_bundle_detects_docx_output_tamper_and_rolls_back(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    fixture = _setup(tmp_path)
    original_register = fixture["runtime"].task_artifacts.register_on_connection

    def tampering_register(*args, **kwargs):
        record = original_register(*args, **kwargs)
        if record["kind"] == "strategy_report_docx":
            Path(record["path"]).write_bytes(b"tampered")
        return record

    monkeypatch.setattr(
        fixture["runtime"].task_artifacts,
        "register_on_connection",
        tampering_register,
    )
    with pytest.raises(
        StrategyError,
        match="report output|artifact file is invalid|bytes|hash",
    ):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []


def test_build_report_bundle_fourth_registration_failure_leaves_no_partial_set(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    fixture = _setup(tmp_path)
    original_register = fixture["runtime"].task_artifacts.register_on_connection

    def fail_docx_registration(*args, **kwargs):
        if kwargs["kind"] == "strategy_report_docx":
            raise RuntimeError("DOCX registry unavailable")
        return original_register(*args, **kwargs)

    monkeypatch.setattr(
        fixture["runtime"].task_artifacts,
        "register_on_connection",
        fail_docx_registration,
    )

    with pytest.raises(RuntimeError, match="DOCX registry unavailable"):
        _run(fixture)

    assert _report_rows(fixture) == []
    assert _audit_rows(fixture) == []
    with fixture["runtime"].task_artifacts.transaction() as conn:
        assert conn.execute(
            "SELECT COUNT(*) FROM strategy_report_revisions"
        ).fetchone()[0] == 0


def test_build_report_bundle_never_follows_existing_output_symlink(
    tmp_path: Path,
) -> None:
    fixture = _setup(tmp_path)
    output = _run(fixture)
    json_artifact = next(
        item for item in output["artifacts"] if item["format"] == "json"
    )
    row = next(
        item
        for item in _report_rows(fixture)
        if item["id"] == json_artifact["artifact_id"]
    )
    path = Path(row["path"])
    target = tmp_path / "outside.json"
    target.write_bytes(path.read_bytes())
    path.unlink()
    path.symlink_to(target)

    with pytest.raises(
        StrategyError,
        match="must stay under|symlink|regular|unavailable|invalid",
    ):
        _run(fixture)

    assert len(_audit_rows(fixture)) == 1
