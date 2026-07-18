"""Shared pytest fixtures.

Pre-warm the pypmml JVM gateway once per PMML-runtime test session, in the clean
startup environment, before any test runs. Pure strategy/data/frontend selections
must not pay the JVM cold-start cost.

Why: the suite uses pytest-randomly (random test order). pypmml's JVM gateway
is a process-wide singleton that is cold-started lazily on the first PMML use.
If that first use happens *after* another test has mutated the process state
(working directory, environment), the JVM launch can intermittently return an
empty gateway port and fail with ``int(b'')``. Establishing the gateway at
session start sidesteps the order dependency: every later PMML test reuses the
already-running singleton instead of cold-starting it. The explicit
``pmml_runtime`` marker below keeps this protection for real PMML runs while making
focused development tests start immediately.
"""

from __future__ import annotations

from collections.abc import Iterable
from typing import Protocol

import pytest


class _CollectedItem(Protocol):
    def get_closest_marker(self, name: str): ...


def _requires_pmml_gateway(items: Iterable[_CollectedItem]) -> bool:
    """Return whether the collected selection contains a real PMML runtime test."""

    return any(item.get_closest_marker("pmml_runtime") is not None for item in items)


@pytest.fixture
def ready_contract():
    from tests.validation_builders import make_ready_contract

    return make_ready_contract()


@pytest.fixture
def pmml_contract(ready_contract):
    return ready_contract


@pytest.fixture
def direct_manifest(ready_contract):
    return ready_contract.require_pmml_manifest()


@pytest.fixture
def derived_manifest():
    from marvis.validation.input_contracts import PmmlInputManifest, StressUnit

    return PmmlInputManifest(
        schema_version="marvis.pmml_input_manifest.v1",
        raw_required_fields=("age", "income"),
        derived_fields=("age_bucket",),
        model_features=("age_bucket", "income"),
        stress_units=(
            StressUnit("age_bucket", ("age",), ("age_bucket <- age",)),
            StressUnit("income", ("income",), ()),
        ),
        unsupported_derivations=(),
        output_candidates=("probability_1",),
        algorithm="xgb",
    )


@pytest.fixture
def pipeline_settings(tmp_path):
    from docx import Document

    from marvis.db import init_db
    from marvis.pipeline import PipelineSettings
    from marvis.settings import build_settings

    app_settings = build_settings(tmp_path / "workspace")
    app_settings.report_template_path.parent.mkdir(parents=True, exist_ok=True)
    template = Document()
    template.add_paragraph("模型：{{TEXT:model_name}}")
    template.add_paragraph("{{TEXT:reproducibility_summary}}")
    template.save(app_settings.report_template_path)
    init_db(app_settings.db_path)
    return PipelineSettings(
        workspace=app_settings.workspace,
        db_path=app_settings.db_path,
        report_template_path=app_settings.report_template_path,
        pmml_scoring_chunk_size=2,
    )


@pytest.fixture
def ready_validation_task(pipeline_settings):
    from dataclasses import replace

    import pandas as pd

    from marvis.api_scan_helpers import perform_scan_task
    from marvis.db import TaskRepository
    from marvis.domain import TaskCreate, TaskStatus
    from marvis.repositories.validation_contracts import ValidationContractRepository
    from marvis.settings import build_settings
    from marvis.validation.input_confirmation import (
        validate_confirmation_against_materials,
    )
    from marvis.validation_materials import resolve_selected_validation_materials
    from tests.validation_builders import make_validation_confirmation
    from tests.validation_material_builders import write_validation_material_bundle

    app_settings = build_settings(pipeline_settings.workspace)
    bundle = write_validation_material_bundle(
        app_settings.workspace / "ready-bundle",
        notebook_source=(
            "RMC_TARGET_COL='y'\nRMC_SPLIT_COL='split'\n"
            "RMC_TIME_COL='apply_month'\n"
            "RMC_PMML_OUTPUT_FIELD='probability_1'\nRMC_MODEL_PARAMS={}\n"
        ),
        sample=pd.DataFrame(
            {
                "x1": [0.0, 1.0, 2.0, 3.0] * 3,
                "x2": [0.0, 1.0, 0.0, 1.0] * 3,
                "y": [0, 1, 0, 1] * 3,
                "split": ["train"] * 4 + ["test"] * 4 + ["oot"] * 4,
                "apply_month": ["202601"] * 4
                + ["202602"] * 4
                + ["202603"] * 4,
            }
        ),
    )
    repo = TaskRepository(app_settings.db_path)
    task = repo.create_task(
        TaskCreate(
            model_name="ready",
            model_version="v2",
            validator="pytest",
            source_dir=str(bundle.root),
            notebook_path=bundle.notebook_path.name,
            sample_path=bundle.sample_path.name,
            pmml_path=bundle.pmml_path.name,
            dictionary_path=bundle.dictionary_path.name,
        )
    )
    scan = perform_scan_task(repo, task, app_settings)
    assert scan["validation_input_contract"]["status"] == "pending_confirmation"
    contracts = ValidationContractRepository(app_settings.db_path)
    candidate = contracts.get(task.id)
    assert candidate is not None
    persisted_task = repo.get_task(task.id)
    paths = resolve_selected_validation_materials(persisted_task)
    validated = validate_confirmation_against_materials(
        contract=candidate.contract,
        sample_path=paths.sample,
        dictionary_path=paths.dictionary,
        requested=replace(make_validation_confirmation(), metadata_sheet=None),
    )
    confirmed = contracts.confirm(
        task.id,
        validated.values,
        expected_revision=candidate.revision,
        resolved_sample_schema=validated.sample_schema,
        resolved_feature_metadata=validated.feature_metadata,
    )
    assert confirmed.status == "ready"
    persisted = repo.get_task(task.id)
    assert persisted.status == TaskStatus.SCANNED
    assert persisted.validation_workflow_version == 2
    return persisted


@pytest.fixture(scope="session", autouse=True)
def _prewarm_pmml_gateway(request):
    """Start pypmml once, but only when this selection can execute PMML."""

    if _requires_pmml_gateway(request.session.items):
        try:
            from pypmml.base import PMMLContext

            PMMLContext.getOrCreate()
        except Exception:
            # If pypmml or a JVM is unavailable, individual PMML tests will surface
            # that on their own; pre-warming must never block the rest of the suite.
            pass
    yield
