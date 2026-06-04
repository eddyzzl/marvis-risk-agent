import re
from collections.abc import Sequence
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage


PLACEHOLDER_PATTERN = re.compile(r"\{\{[^{}]+\}\}")
ImageValue = Path | str | Sequence[Path | str]
REPORT_IMAGE_MAX_WIDTH_INCHES = 6.25
EMU_PER_INCH = 914400


@dataclass(frozen=True)
class TemplateTable:
    rows: list[list[str]]
    header_rows: int = 1
    style: str = "Table Grid"
    column_width_weights: list[float] | None = None
    table_width_inches: float = 6.25
    font_name: str = "微软雅黑"
    font_size_pt: int = 8
    header_fill: str = "C00000"
    header_font_color: str = "FFFFFF"
    border_color: str = "808080"


@dataclass(frozen=True)
class TemplateReportPayload:
    template_path: Path
    output_path: Path
    text_values: dict[str, str]
    image_values: dict[str, ImageValue]
    table_values: dict[str, TemplateTable] | None = None


@dataclass(frozen=True)
class TemplateReportResult:
    output_path: Path
    unresolved_placeholders: list[str]


def find_placeholders(template_path: Path) -> list[str]:
    document = Document(template_path)
    return _unique_placeholders(_document_text_parts(document))


def render_template_report(payload: TemplateReportPayload) -> TemplateReportResult:
    document = Document(payload.template_path)
    table_values = payload.table_values or {}
    max_image_width_inches = _document_image_max_width_inches(document)
    unresolved: list[str] = []

    for paragraph in document.paragraphs:
        _replace_paragraph_placeholders(
            paragraph,
            text_values=payload.text_values,
            image_values=payload.image_values,
            table_values=table_values,
            max_image_width_inches=max_image_width_inches,
            unresolved=unresolved,
        )

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_paragraph_placeholders(
                        paragraph,
                        text_values=payload.text_values,
                        image_values=payload.image_values,
                        table_values=table_values,
                        max_image_width_inches=max_image_width_inches,
                        unresolved=unresolved,
                    )

    unresolved.extend(
        placeholder
        for placeholder in _unique_placeholders(_document_text_parts(document))
        if _is_report_placeholder(placeholder)
    )
    unresolved = _unique_placeholders(unresolved)

    payload.output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(payload.output_path)
    return TemplateReportResult(
        output_path=payload.output_path,
        unresolved_placeholders=unresolved,
    )


def _replace_paragraph_placeholders(
    paragraph,
    text_values: dict[str, str],
    image_values: dict[str, ImageValue],
    table_values: dict[str, TemplateTable],
    max_image_width_inches: float,
    unresolved: list[str],
) -> None:
    text, run_ranges = _paragraph_text_and_run_ranges(paragraph)
    if "{{" not in text:
        return

    matches = list(PLACEHOLDER_PATTERN.finditer(text))
    if not matches:
        return

    _clear_paragraph(paragraph)
    cursor = 0
    anchor = paragraph._p
    for match in matches:
        _append_text_range(paragraph, text, run_ranges, cursor, match.start())

        placeholder = match.group(0)
        value_key = placeholder[2:-2]
        source_run = _run_for_placeholder(run_ranges, match.start(), match.end())
        if placeholder.startswith("{{TEXT:"):
            if value_key in text_values:
                _add_run_like(paragraph, str(text_values[value_key]), source_run)
            else:
                _append_text_range(paragraph, text, run_ranges, match.start(), match.end())
            cursor = match.end()
            continue
        if not placeholder.startswith(("{{IMAGE:", "{{TABLE:")):
            _append_text_range(paragraph, text, run_ranges, match.start(), match.end())
            cursor = match.end()
            continue

        table = table_values.get(value_key)
        image_value = image_values.get(value_key)
        if table is not None:
            anchor = _insert_table_after_paragraph(paragraph, table, anchor=anchor)
        elif placeholder.startswith("{{TABLE:"):
            unresolved.append(placeholder)
            _add_run_like(paragraph, f"【待插入表格：{value_key}】", source_run)
        elif not _image_paths_exist(image_value):
            unresolved.append(placeholder)
            _add_run_like(paragraph, f"【待插入图片：{value_key}】", source_run)
        else:
            _add_images_like(
                paragraph,
                _image_paths(image_value),
                source_run,
                max_width_inches=max_image_width_inches,
            )
        cursor = match.end()

    _append_text_range(paragraph, text, run_ranges, cursor, len(text))


def _paragraph_text_and_run_ranges(paragraph):
    text_parts: list[str] = []
    run_ranges = []
    cursor = 0
    for run in paragraph.runs:
        run_text = run.text
        if not run_text:
            continue
        text_parts.append(run_text)
        next_cursor = cursor + len(run_text)
        run_ranges.append((cursor, next_cursor, run))
        cursor = next_cursor
    return "".join(text_parts), run_ranges


def _append_text_range(paragraph, text: str, run_ranges, start: int, end: int) -> None:
    if start >= end:
        return
    appended = False
    for run_start, run_end, run in run_ranges:
        overlap_start = max(start, run_start)
        overlap_end = min(end, run_end)
        if overlap_start >= overlap_end:
            continue
        _add_run_like(paragraph, text[overlap_start:overlap_end], run)
        appended = True
    if not appended:
        paragraph.add_run(text[start:end])


def _run_at_offset(run_ranges, start: int, end: int):
    for run_start, run_end, run in run_ranges:
        if run_start <= start < run_end:
            return run
    for run_start, run_end, run in run_ranges:
        if run_start < end and start < run_end:
            return run
    return None


def _run_for_placeholder(run_ranges, start: int, end: int):
    overlapping_runs = [
        (run_start, run_end, run)
        for run_start, run_end, run in run_ranges
        if run_start < end and start < run_end
    ]
    if not overlapping_runs:
        return None
    _, _, best_run = max(
        overlapping_runs,
        key=lambda item: (
            _run_property_score(item[2]),
            min(item[1], end) - max(item[0], start),
        ),
    )
    return best_run


def _run_property_score(run) -> int:
    if run is None or run._r.rPr is None:
        return 0
    return len(list(run._r.rPr))


def _add_run_like(paragraph, text: str, source_run):
    run = paragraph.add_run(text)
    _copy_run_properties(source_run, run)
    return run


def _image_paths(value: ImageValue | None) -> list[Path]:
    if value is None:
        return []
    if isinstance(value, (str, Path)):
        return [Path(value)]
    return [Path(item) for item in value]


def _image_paths_exist(value: ImageValue | None) -> bool:
    paths = _image_paths(value)
    return bool(paths) and all(path.exists() for path in paths)


def _add_images_like(
    paragraph,
    paths: list[Path],
    source_run,
    *,
    max_width_inches: float,
) -> None:
    for index, image_path in enumerate(paths):
        if index > 0:
            _add_run_like(paragraph, "", source_run).add_break()
        _add_run_like(paragraph, "", source_run).add_picture(
            str(image_path),
            width=_report_image_width(image_path, max_width_inches=max_width_inches),
        )


def _report_image_width(image_path: Path, *, max_width_inches: float):
    natural_width = _image_natural_width_inches(image_path)
    if natural_width is None:
        return Inches(max_width_inches)
    return Inches(min(natural_width, max_width_inches))


def _document_image_max_width_inches(document) -> float:
    section_widths = [
        _section_content_width_inches(section)
        for section in document.sections
    ]
    usable_widths = [width for width in section_widths if width > 0]
    if not usable_widths:
        return REPORT_IMAGE_MAX_WIDTH_INCHES
    return min(REPORT_IMAGE_MAX_WIDTH_INCHES, min(usable_widths))


def _section_content_width_inches(section) -> float:
    return (
        section.page_width
        - section.left_margin
        - section.right_margin
    ) / EMU_PER_INCH


def _image_natural_width_inches(image_path: Path) -> float | None:
    try:
        with PILImage.open(image_path) as image:
            dpi = image.info.get("dpi")
            if not dpi:
                return None
            x_dpi = float(dpi[0])
            if x_dpi <= 1:
                return None
            rounded_dpi = round(x_dpi)
            if abs(x_dpi - rounded_dpi) < 0.1:
                x_dpi = float(rounded_dpi)
            return image.width / x_dpi
    except Exception:
        return None


def _copy_run_properties(source_run, target_run) -> None:
    if source_run is None or source_run._r.rPr is None:
        return
    if target_run._r.rPr is not None:
        target_run._r.remove(target_run._r.rPr)
    target_run._r.insert(0, deepcopy(source_run._r.rPr))


def _document_text_parts(document) -> list[str]:
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return parts


def _unique_placeholders(parts: list[str]) -> list[str]:
    placeholders: list[str] = []
    for part in parts:
        for placeholder in PLACEHOLDER_PATTERN.findall(part):
            if placeholder not in placeholders:
                placeholders.append(placeholder)
    return placeholders


def _is_report_placeholder(placeholder: str) -> bool:
    return placeholder.startswith(("{{TEXT:", "{{IMAGE:", "{{TABLE:"))


def _set_paragraph_text(paragraph, text: str) -> None:
    _clear_paragraph(paragraph)
    paragraph.add_run(text)


def _insert_table_after_paragraph(paragraph, table_data: TemplateTable, *, anchor=None):
    column_count = max((len(row) for row in table_data.rows), default=0)
    if column_count == 0:
        return paragraph._p if anchor is None else anchor

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    table_width_twips = _inches_to_twips(table_data.table_width_inches)
    column_widths = _column_widths_twips(table_data, column_count, table_width_twips)
    table = paragraph._parent.add_table(
        rows=len(table_data.rows),
        cols=column_count,
        width=Inches(table_data.table_width_inches),
    )
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    try:
        table.style = table_data.style
    except KeyError:
        pass
    _format_table_shell(table, table_width_twips, column_widths, table_data)
    for row_index, row_values in enumerate(table_data.rows):
        row = table.rows[row_index]
        for column_index in range(column_count):
            value = row_values[column_index] if column_index < len(row_values) else ""
            _format_cell(
                row.cells[column_index],
                value=str(value),
                width_twips=column_widths[column_index],
                is_header=row_index < table_data.header_rows,
                table_data=table_data,
            )
    if anchor is None:
        anchor = paragraph._p
    anchor.addnext(table._tbl)
    return table._tbl


def _clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def _format_table_shell(
    table,
    table_width_twips: int,
    column_widths: list[int],
    table_data: TemplateTable,
) -> None:
    table_properties = table._tbl.tblPr
    table_width = _get_or_add(table_properties, "w:tblW")
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(table_width_twips))

    table_layout = _get_or_add(table_properties, "w:tblLayout")
    table_layout.set(qn("w:type"), "fixed")

    table_borders = _get_or_add(table_properties, "w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = _get_or_add(table_borders, f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), table_data.border_color)

    existing_grid = table._tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        table._tbl.remove(existing_grid)
    table_grid = OxmlElement("w:tblGrid")
    for width in column_widths:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(width))
        table_grid.append(grid_column)
    table._tbl.insert(1, table_grid)


def _format_cell(
    cell,
    value: str,
    width_twips: int,
    is_header: bool,
    table_data: TemplateTable,
) -> None:
    cell.width = Inches(width_twips / 1440)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    cell_properties = cell._tc.get_or_add_tcPr()
    cell_width = _get_or_add(cell_properties, "w:tcW")
    cell_width.set(qn("w:type"), "dxa")
    cell_width.set(qn("w:w"), str(width_twips))
    vertical_alignment = _get_or_add(cell_properties, "w:vAlign")
    vertical_alignment.set(qn("w:val"), "center")
    _set_cell_margins(cell_properties)

    if is_header:
        shading = _get_or_add(cell_properties, "w:shd")
        shading.set(qn("w:fill"), table_data.header_fill)

    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(value)
    run.bold = is_header
    run.font.name = table_data.font_name
    run.font.size = Pt(table_data.font_size_pt)
    if is_header:
        run.font.color.rgb = RGBColor.from_string(table_data.header_font_color)
    run_properties = run._element.get_or_add_rPr()
    run_fonts = run_properties.get_or_add_rFonts()
    run_fonts.set(qn("w:ascii"), table_data.font_name)
    run_fonts.set(qn("w:hAnsi"), table_data.font_name)
    run_fonts.set(qn("w:eastAsia"), table_data.font_name)


def _set_cell_margins(cell_properties) -> None:
    cell_margins = _get_or_add(cell_properties, "w:tcMar")
    margin_widths = {
        "top": "40",
        "bottom": "40",
        "left": "60",
        "right": "60",
    }
    for side, width in margin_widths.items():
        margin = _get_or_add(cell_margins, f"w:{side}")
        margin.set(qn("w:w"), width)
        margin.set(qn("w:type"), "dxa")


def _column_widths_twips(
    table_data: TemplateTable,
    column_count: int,
    table_width_twips: int,
) -> list[int]:
    weights = (
        _normalized_explicit_weights(table_data.column_width_weights, column_count)
        or _content_width_weights(table_data.rows, column_count)
    )
    total_weight = sum(weights)
    widths = [
        max(360, round(table_width_twips * weight / total_weight))
        for weight in weights
    ]
    width_delta = table_width_twips - sum(widths)
    widths[-1] += width_delta
    return widths


def _normalized_explicit_weights(
    weights: list[float] | None,
    column_count: int,
) -> list[float] | None:
    if not weights:
        return None
    padded = [max(float(weight), 1.0) for weight in weights[:column_count]]
    if len(padded) < column_count:
        padded.extend([1.0] * (column_count - len(padded)))
    return padded


def _content_width_weights(rows: list[list[str]], column_count: int) -> list[float]:
    weights: list[float] = []
    for column_index in range(column_count):
        max_width = 4
        for row in rows:
            if column_index < len(row):
                max_width = max(max_width, _display_width(str(row[column_index])))
        weights.append(float(min(max_width, 22)))
    return weights


def _display_width(value: str) -> int:
    return sum(2 if ord(character) > 127 else 1 for character in value)


def _inches_to_twips(value: float) -> int:
    return round(value * 1440)


def _get_or_add(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child
